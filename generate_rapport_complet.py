#!/usr/bin/env python3
"""
Rapport d'Activite COMPLET — DATATYM™ / Marketym / H&C Executive
Toutes activites : collecte, plateforme, barometres, referentiel, livrables
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for lv in range(1, 4):
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x0F, 0x1E, 0x3A)
    hs.font.size = Pt([0, 18, 14, 12][lv])


def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.font.bold = True; r.font.size = Pt(10)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(10)
    doc.add_paragraph()


def bul(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')


def bold_bul(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        if ' : ' in it:
            parts = it.split(' : ', 1)
            r = p.add_run(parts[0] + ' : ')
            r.font.bold = True
            p.add_run(parts[1])
        else:
            p.add_run(it)


# ════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ════════════════════════════════════════════════════════════════
for _ in range(3): doc.add_paragraph()
for text, sz, bold, color in [
    ('DATATYM\u2122', 32, True, RGBColor(0x0F, 0x1E, 0x3A)),
    ('RAPPORT D\'ACTIVITE', 24, True, RGBColor(0xB5, 0x89, 0x1A)),
    ('', 12, False, None),
    ('De la collecte a la construction\ndu referentiel strategique', 16, False, RGBColor(0x6B, 0x72, 0x80)),
    ('', 12, False, None),
    ('Barometre du Marche du Talent Africain', 14, True, RGBColor(0x1A, 0x20, 0x2C)),
    ('Vague 1 — Mars-Avril 2026', 13, False, RGBColor(0x6B, 0x72, 0x80)),
    ('', 12, False, None),
    ('H&C Executive / Marketym', 12, True, RGBColor(0x1A, 0x20, 0x2C)),
    ('Document confidentiel', 10, False, RGBColor(0x9C, 0xA3, 0xAF)),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(sz); r.font.bold = bold
    if color: r.font.color.rgb = color
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# SOMMAIRE
# ════════════════════════════════════════════════════════════════
doc.add_heading('Sommaire', level=1)
for s in [
    '1. Vision et positionnement DATATYM\u2122',
    '2. Le systeme DATATYM\u2122 — Architecture',
    '3. Phase de collecte',
    '4. Developpement de la plateforme Marketym',
    '5. Construction du referentiel V3',
    '6. Barometres et analyses produits',
    '7. Indices DATATYM\u2122 — Resultats Gen Z',
    '8. Audit de coherence et calibration',
    '9. Livrables complets',
    '10. Resultats strategiques cles',
    '11. Feuille de route',
]:
    doc.add_paragraph(s, style='List Number')
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 1. VISION
# ════════════════════════════════════════════════════════════════
doc.add_heading('1. Vision et positionnement DATATYM\u2122', level=1)

doc.add_heading('1.1 Ce que DATATYM\u2122 est', level=2)
doc.add_paragraph(
    'DATATYM\u2122 n\'est pas un barometre. Ce n\'est pas une etude. '
    'C\'est un referentiel strategique — le premier systeme de notation du marche '
    'du talent africain. DATATYM\u2122 definit les standards qui servent a piloter '
    'le capital humain en Afrique.'
)

doc.add_heading('1.2 Positionnement strategique', level=2)
doc.add_paragraph(
    'DATATYM\u2122 se positionne a l\'intersection de trois modeles de reference :'
)
tbl(
    ['Modele', 'Ce qu\'il fait', 'Ce que DATATYM\u2122 fait de plus'],
    [
        ['Bloomberg', 'Decrit les prix et flux financiers', 'Impose le cadre mental, pas seulement la donnee'],
        ['McKinsey', 'Prescrit des strategies sur mesure', 'Produit des standards scalables, pas du conseil unitaire'],
        ['Statista', 'Recense et compile des statistiques', 'Prescrit des doctrines et des decisions, pas des chiffres'],
    ]
)
p = doc.add_paragraph()
r = p.add_run(
    'DATATYM\u2122 — Seule autorite produisant des indices proprietaires et des doctrines '
    'prescriptives pour le marche du talent africain francophone.'
)
r.font.bold = True; r.font.italic = True

doc.add_heading('1.3 L\'ambition', level=2)
doc.add_paragraph(
    'Objectif 2027 : DATATYM\u2122 cite dans les COMEX africains comme Bloomberg l\'est '
    'dans les salles de marche mondiales. Les dirigeants ne pensent plus AVEC DATATYM\u2122 — '
    'ils sont juges PAR DATATYM\u2122.'
)


# ════════════════════════════════════════════════════════════════
# 2. SYSTEME
# ════════════════════════════════════════════════════════════════
doc.add_heading('2. Le systeme DATATYM\u2122 — Architecture', level=1)

doc.add_heading('2.1 Les 4 piliers', level=2)
doc.add_paragraph(
    'DATATYM\u2122 est un systeme ferme a 4 verrouillages. Mesure, interpretation, '
    'usage et action sont integres — le dirigeant n\'a pas d\'espace pour penser '
    'en dehors du referentiel.'
)
tbl(
    ['Pilier', 'Nom', 'Fonction', 'Composantes'],
    [
        ['I', 'LES INDICES', 'La mesure', 'IPE\u2122, IRTA\u2122, ILUX\u2122, IGRO\u2122, ICON\u2122'],
        ['II', 'LES DOCTRINES', 'L\'interpretation', '5 verites prescriptives'],
        ['III', 'LES RITUELS', 'L\'usage recurrent', 'Barometre trimestriel, Brief CODIR, Scorecard, Diagnostic DG'],
        ['IV', 'LES DECISIONS', 'L\'action', '10 decisions obligatoires pour le CODIR'],
    ]
)

doc.add_heading('2.2 Les 4 indices proprietaires', level=2)
tbl(
    ['Indice', 'Nom complet', 'Formule', 'Ce qu\'il mesure'],
    [
        ['IPE\u2122', 'Index of Professional Experience',
         'IPE = 100 - Tension x 2\nTension = Score_interne - NPS_norm\nScore_int = SAT x0.40 + ENG x0.35 + VAL x0.25',
         'Performance employeur percue'],
        ['IRTA\u2122', 'Indice de Risque de Turnover Africain',
         'D1 + D2 + (D3 x 0.30) - deltaE',
         'Risque de depart cumule'],
        ['ILUX\u2122', 'Index of Lure and User eXperience',
         'NPS_prom x0.30 + Pref x0.25 + Reco x0.25 + Remun x0.20',
         'Capacite d\'attraction'],
        ['IGRO\u2122', 'Index of Growth, Recognition and Opportunity',
         'G1 Reconnaissance x0.54 + G2 Vision interne x0.46',
         'Developpement et retention par la croissance'],
    ]
)

doc.add_heading('2.3 Le systeme de notation DATATYM RATING\u2122', level=2)
tbl(
    ['Note', 'Label', 'Seuil', 'Interpretation'],
    [
        ['A', 'ELITE RH', '>= 80', 'Zone Verte — Organisation de reference'],
        ['B', 'STABLE', '60 - 79', 'Zone Orange moderee — Performance correcte'],
        ['C', 'RISQUE', '35 - 59', 'Zone Orange — Surveillance active, plan 90 jours'],
        ['D', 'CRITIQUE', '< 35', 'Zone Rouge — Hemorragie de talents, intervention urgente'],
    ]
)

doc.add_heading('2.4 Les 5 doctrines', level=2)
for num, titre, data, action in [
    ('1', 'L\'entreprise africaine forme ses futurs concurrents',
     '25,5% creent dans 2 ans · IGRO\u2122 = 31,6', 'Programme d\'intrapreneuriat ou disruption financee.'),
    ('2', 'La fidelite est morte, seule la progression retient',
     '73,4% en depart potentiel · 15% voient un avenir ici', 'Plan de carriere 12/24/36 mois = minimum vital.'),
    ('3', 'Le salaire est un seuil, pas un levier',
     '53% exigent salaire ET sens · 5% acceptent le sens seul', 'Culture forte sans competitivite salariale ne retient personne.'),
    ('4', 'Le NPS employeur est un KPI strategique, pas RH',
     'IPE\u2122 = 32 · NPS = -28 · 816 detracteurs actifs', 'NPS et IPE\u2122 au tableau de bord CODIR.'),
    ('5', 'Sans intrapreneuriat, vous financez votre disruption',
     '18,3% ont deja lance · 25,5% creent dans 2 ans', 'Canaliser l\'energie entrepreneuriale interne.'),
]:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f'Doctrine {num} — {titre}. ')
    r.font.bold = True
    p.add_run(f'DATA : {data}. ACTION : {action}')


# ════════════════════════════════════════════════════════════════
# 3. COLLECTE
# ════════════════════════════════════════════════════════════════
doc.add_heading('3. Phase de collecte', level=1)

doc.add_heading('3.1 Dispositif', level=2)
doc.add_paragraph(
    'La collecte a ete realisee via QuestionPro avec 8 questionnaires deployes '
    'aupres d\'enqueteurs repartis dans 16+ pays africains francophones. '
    '28 questions communes par questionnaire, structure identique validee.'
)

doc.add_heading('3.2 Chiffres de la collecte', level=2)
tbl(
    ['Indicateur', 'Valeur'],
    [
        ['Repondants totaux consolides', '1 479 (panel V3) + 1 500 Gen Z'],
        ['Questionnaires deployes', '8 surveys QuestionPro'],
        ['Pays couverts', '16+ pays africains francophones'],
        ['Gen Z pure (18-27 ans)', '57,3% du panel'],
        ['Taux de completion moyen', '~80%'],
        ['Convergence inter-surveys', '+/- 0,9% max'],
        ['Marge d\'erreur globale', '+/- 2,5% a IC 95%'],
    ]
)

doc.add_heading('3.3 Panel par questionnaire', level=2)
tbl(
    ['Questionnaire', 'ID QuestionPro', 'Repondants', 'Perimetre'],
    [
        ['PRINCIPAL', '13442593', '773', '16 pays — Panel large'],
        ['GENZ-1', '13443679', '182', 'Afrique Ouest + diaspora'],
        ['GENZ-2', '13444868', '188', 'Afrique Ouest + diaspora'],
        ['GENZ-3', '13444871', '98', 'Cameroun + Senegal'],
        ['GENZ-4', '13444872', '138', 'Senegal + Benin + CMR'],
        ['GENZ-5', '13444873', '20', 'Senegal dominant'],
        ['EN AF - V2', '13432629', '60', 'Panel Afrique senior'],
        ['EN AF - MARIE', '13432523', '20', 'Panel Afrique senior'],
        ['TOTAL', '', '1 479', '16+ pays'],
    ]
)


# ════════════════════════════════════════════════════════════════
# 4. PLATEFORME
# ════════════════════════════════════════════════════════════════
doc.add_heading('4. Developpement de la plateforme Marketym', level=1)

doc.add_heading('4.1 Description', level=2)
doc.add_paragraph(
    'Marketym est l\'application web developpee sur mesure pour piloter '
    'la collecte, suivre les enqueteurs, gerer les affectations et visualiser '
    'la progression en temps reel.'
)

doc.add_heading('4.2 Stack technique', level=2)
tbl(
    ['Composant', 'Technologie', 'Hebergement'],
    [
        ['Backend / API', 'FastAPI (Python)', 'Railway'],
        ['Frontend', 'React + Vite + TailwindCSS', 'Render'],
        ['Base de donnees', 'PostgreSQL (Supabase)', 'Supabase Cloud'],
        ['Emails', 'Brevo (ex-Sendinblue)', '-'],
        ['Sondages', 'QuestionPro API', '-'],
    ]
)

doc.add_heading('4.3 Fonctionnalites livrees', level=2)
bul([
    'Authentification securisee (JWT + OTP premiere connexion)',
    'Dashboard admin : enqueteurs, enquetes, affectations, quotas',
    'Dashboard enqueteur : suivi personnel des completions',
    'Tracking des clics par lien unique avec deduplication IP',
    'Synchronisation automatique avec QuestionPro',
    'Gestion des quotas simples et croises par segmentation',
    'Connexion PostgreSQL directe (~3ms/requete vs ~100ms REST)',
])

doc.add_heading('4.4 Evolutions techniques majeures', level=2)
tbl(
    ['Evolution', 'Probleme resolu', 'Impact'],
    [
        ['Remodelisation DB (migrations 009/010)', 'Matching texte cassait les stats', 'Fiabilite 100% par UUID FK'],
        ['DirectClient PostgreSQL', 'API REST lente (~100ms)', 'Performances x30 (~3ms)'],
        ['Batch upserts', 'N+1 queries sur sync', 'Sync 10x plus rapide'],
        ['Deduplication IP', 'Surestimation clics (993 vs 227)', 'Comptage reel'],
    ]
)


# ════════════════════════════════════════════════════════════════
# 5. REFERENTIEL V3
# ════════════════════════════════════════════════════════════════
doc.add_heading('5. Construction du referentiel V3', level=1)

doc.add_heading('5.1 Versions successives', level=2)
tbl(
    ['Version', 'Date', 'Contenu', 'Panel'],
    [
        ['V1 (estimations)', 'Mars 2026', 'IPE-26=32, IRTA=65.7, IAVA=47, IET=93.9', '1 479 resp.'],
        ['V1 Corrigee', '31 mars 2026', '8 indices calibres : IRTA=49.9, IPE-26=64.9, IET=49.9, IPS=79.4', '1 129 resp. audites'],
        ['Comparatif V1/V2', '2 avril 2026', 'Recalibration : IPE=63.9, IRTA=62.3, IAVA=54.7, IET=91.7', '1 489 resp.'],
        ['V3 (document maitre)', 'Avril 2026', '53 slides : IPE=32, IRTA=73.4, ILUX=47, IGRO=31.6, ICON=34.3', '1 479 resp.'],
        ['V3 condensee', 'Avril 2026', '25 slides pour diffusion', '1 479 resp.'],
    ]
)

doc.add_heading('5.2 Contenu du referentiel V3', level=2)
bul([
    'Positionnement strategique (Bloomberg / McKinsey / Statista)',
    'Systeme a 3 piliers : Indices, Doctrines, Rituels',
    '4 indices proprietaires avec formules detaillees',
    '5 doctrines prescriptives avec donnees de support',
    'DATATYM RATING\u2122 : systeme de notation A/B/C/D type S&P',
    'Diagnostic Verdict\u2122 : format 60 secondes pour DG',
    'Scorecard V1 : grille de pilotage Vert/Orange/Rouge',
    '10 prescriptions CODIR obligatoires',
    'Benchmark sectoriel (9 secteurs)',
    'Intervalles de confiance et validation statistique',
    'Simulateur strategique What-If avec formule ICON',
    'Feuille de route Vague 1 a Vision 2027',
    'Classements publics : Top 50 Attractivite, Ranking IRTA par secteur',
])

doc.add_heading('5.3 Scorecard V1', level=2)
tbl(
    ['Indicateur', 'Score V1', 'Rouge', 'Orange', 'Vert'],
    [
        ['IRTA\u2122 (Turnover)', '73,4', '> 60', '30-60', '< 30'],
        ['IPE\u2122 (Perf. Employeur)', '32/100', '< 30', '30-60', '> 60'],
        ['ILUX\u2122 (Attractivite)', '47/100', '< 35', '35-65', '> 65'],
        ['IGRO\u2122 (Developpement)', '31,6 (V0)', '< 20', '20-50', '> 65'],
        ['NPS Employeur', '-28', '< -10', '-10 a +20', '> +20'],
        ['Satisfaction', '68,5%', '< 55%', '55-75%', '> 75%'],
        ['Engagement', '72,9%', '< 60%', '60-80%', '> 80%'],
    ]
)


# ════════════════════════════════════════════════════════════════
# 6. BAROMETRES
# ════════════════════════════════════════════════════════════════
doc.add_heading('6. Barometres et analyses produits', level=1)

doc.add_heading('6.1 Barometre Referentiel (panel global)', level=2)
doc.add_paragraph(
    'Le barometre principal couvre l\'ensemble du panel (1 479 repondants, toutes generations) '
    'et constitue la base du referentiel DATATYM\u2122 V3. Resultats : '
    'IPE\u2122=32 (Zone Rouge), IRTA\u2122=73,4 (Zone Rouge), ILUX\u2122=47 (Zone Orange), '
    'IGRO\u2122=31,6 (Zone Orange). DATATYM RATING\u2122 = D (32,8).'
)

doc.add_heading('6.2 Barometre Gen Z', level=2)
doc.add_paragraph(
    'Le barometre Gen Z isole les 18-27 ans (1 500 repondants, 13 pays, 10+ secteurs). '
    'Analyses produites :'
)
bul([
    'Calcul des 4 indices DATATYM\u2122 sur Gen Z',
    'Matrice DATATYM\u2122 des Cases de Talents (9 cases, effectifs reels)',
    'Score de risque composite et vulnerabilite salariale',
    'Clustering organisationnel (4 profils)',
    'Correlations drivers : management vers retention (0.62), equite vers recommandation (0.78)',
    'Scenarios what-if et horizon critique par secteur',
])

doc.add_heading('6.3 Indices DATATYM\u2122 Gen Z', level=2)
tbl(
    ['Indice', 'Score Gen Z', 'Note', 'Zone', 'Panel global V3'],
    [
        ['IPE\u2122', '25,3/100', 'D', 'Rouge', '32/100'],
        ['IRTA\u2122', '68,7', 'D', 'Rouge', '73,4'],
        ['ILUX\u2122', '53,9/100', 'C', 'Orange', '47/100'],
        ['IGRO\u2122', '45,1/100', 'C', 'Orange', '31,6'],
    ]
)
doc.add_paragraph(
    'Gen Z en zone rouge sur IPE\u2122 et IRTA\u2122, mais meilleurs scores sur ILUX\u2122 et IGRO\u2122 '
    'que le panel global — potentiel non exploite.'
)

doc.add_heading('6.4 Barometre IDAT (version exploratoire)', level=2)
doc.add_paragraph(
    'Version exploratoire avec indice composite IDAT = 61,1/100 (Note B). '
    'A permis d\'explorer les donnees avant alignement V3. '
    'Audit de coherence realise (section 8).'
)

doc.add_heading('6.5 Analyses transversales', level=2)
bold_bul([
    'PCA : PC1 = Bien-etre, PC2 = Ambition',
    'K-Means : 4 profils types de talents Gen Z',
    'Matrice 9 cases : Performance (ENG) x Potentiel, effectifs reels',
    'Profils de depart : 5 profils (entrepreneur, transfuge, expatrie, indecis, fidele)',
    'Fenetre d\'intervention DG : 6 mois entre intention et depart',
    'Coefficient de productivite : 100% (engage) a 30% (depart programme)',
])


# ════════════════════════════════════════════════════════════════
# 8. AUDIT
# ════════════════════════════════════════════════════════════════
doc.add_heading('8. Audit de coherence et calibration', level=1)

doc.add_paragraph(
    'Audit (15 slides) : 11 incoherences entre les analyses Gen Z et le V3.'
)

doc.add_heading('8.1 Critiques', level=2)
bul([
    'Erreur arithmetique matrice (102% au lieu de 100%)',
    'Nomenclature sans filiation V3 (IDAT/ENG/RECO vs IPE/IRTA/ILUX/IGRO)',
    'Score contradictoire (V3 = D/32,8 vs GenZ exploratoire = C/64,5)',
])

doc.add_heading('8.2 Majeurs', level=2)
bul([
    'Architecture causale IGRO > IRTA absente',
    'Echelle de notation differente',
    'ENG comme indice autonome au lieu de composante IPE\u2122',
    '5 doctrines non integrees',
])

doc.add_heading('8.3 Mineurs', level=2)
bul([
    'Perimetre echantillon a clarifier',
    'Secteurs reduits sans justification',
    'DATATYM(TM) absent',
    'Glissement silencieux incomplet (62% du panel)',
])

doc.add_heading('8.4 Resolution', level=2)
doc.add_paragraph(
    'Les 4 indices officiels ont ete recalcules sur Gen Z avec les formules V3 (section 6.3).'
)


# ════════════════════════════════════════════════════════════════
# 9. LIVRABLES
# ════════════════════════════════════════════════════════════════
doc.add_heading('9. Livrables complets', level=1)

doc.add_heading('9.1 Referentiel et barometres', level=2)
tbl(
    ['#', 'Document', 'Format', 'Slides', 'Contenu'],
    [
        ['1', 'Referentiel V3', 'PPTX', '53', 'Document maitre complet'],
        ['2', 'Referentiel 25 slides', 'PPTX', '25', 'Version condensee'],
        ['3', 'V1 Corrigee', 'PPTX + DOCX', '17', '8 indices calibres'],
        ['4', 'Comparatif V1/V2', 'PPTX', '6', 'Recalibration'],
        ['5', 'Barometre IDAT', 'PPTX + DOCX', '11', 'Version exploratoire'],
    ]
)

doc.add_heading('9.2 Presentations Gen Z', level=2)
tbl(
    ['#', 'Document', 'Slides', 'Cible'],
    [
        ['6', 'GenZ V2', '29', 'DG'],
        ['7', 'Master Class', '23', 'CEO'],
        ['8', 'Analytique', '22', 'DRH / Data'],
        ['9', 'Presentation Strategique', '-', 'Partenaires'],
    ]
)

doc.add_heading('9.3 Audit et conformite', level=2)
tbl(
    ['#', 'Document', 'Contenu'],
    [
        ['10', 'Audit Coherence GenZ/V3', '11 incoherences + recommandations'],
        ['11', 'Audit Conformite Marketym', 'Conformite plateforme'],
        ['12', 'Analyse IPs', 'Verification adresses IP'],
        ['13', 'Dossier Confrontation', 'Elements de confrontation'],
        ['14', 'Guide Audition', 'Guide audition enqueteurs'],
        ['15', 'Checklist Objections', 'Grille objections agents'],
    ]
)

doc.add_heading('9.4 Documents financiers', level=2)
tbl(
    ['#', 'Document', 'Contenu'],
    [
        ['16', 'Situation Financiere', 'Tarification progressive, previsionnel'],
        ['17', 'Etats Financiers (x2)', 'Suivi financier'],
        ['18', 'Excedents Quotas', 'Analyse depassements'],
        ['19', 'Plan Redistribution', 'Redistribution affectations'],
        ['20', 'Frais Ecobank', 'Analyse frais bancaires'],
    ]
)

doc.add_heading('9.5 Institutionnels et donnees', level=2)
tbl(
    ['#', 'Document', 'Contenu'],
    [
        ['21', 'Lettres d\'introduction (x2)', 'DataTym + BMN'],
        ['22', 'Cahiers des charges (x2)', 'Specifications projet'],
        ['23', 'Rapports Barometres (x2)', 'Rapports narratifs'],
        ['24', 'Base exploitable Gen Z', '1 200 lignes, 56 variables'],
        ['25', 'Identifiants enqueteurs', 'Fichier reference'],
        ['26-29', 'Plateforme Marketym', 'Backend + Frontend + DB + Scripts'],
    ]
)


# ════════════════════════════════════════════════════════════════
# 10. RESULTATS
# ════════════════════════════════════════════════════════════════
doc.add_heading('10. Resultats strategiques cles', level=1)

doc.add_heading('10.1 Diagnostic panel global', level=2)
tbl(
    ['Indice', 'Score', 'Note', 'Signal'],
    [
        ['IPE\u2122', '32/100', 'D', 'Performance sous seuil de viabilite'],
        ['IRTA\u2122', '73,4', 'D', '73,4% de risque de depart cumule'],
        ['ILUX\u2122', '47/100', 'C', 'Attraction par defaut'],
        ['IGRO\u2122', '31,6', 'D', '93,9% d\'intention entrepreneuriale'],
        ['NPS', '-28', 'D', '816 detracteurs actifs'],
        ['RATING\u2122', '32,8', 'D', 'Instabilite structurelle'],
    ]
)

doc.add_heading('10.2 Constats structurels', level=2)
bold_bul([
    '73,4% en depart potentiel : actif (14,6%) + opportuniste (51,1%) + entrepreneurial (7,7%)',
    '93,9% d\'intention entrepreneuriale : l\'organisation = etape de financement',
    'NPS = -28 : marque employeur negative',
    'Paradoxe : engagement 72,9% MAIS NPS -28 — engages dans le travail, pas dans l\'organisation',
    '53% exigent salaire ET sens : double exigence non-negociable',
    '15% seulement voient un avenir chez leur employeur',
    'Gen Z : IPE=25,3 (pire) mais ILUX=53,9 et IGRO=45,1 (meilleurs) — potentiel non exploite',
])

doc.add_heading('10.3 Finding central', level=2)
p = doc.add_paragraph()
r = p.add_run(
    'Le marche africain du talent est en zone rouge structurelle. '
    'Les talents forment les futurs concurrents de leur employeur. '
    'La fenetre d\'action est maintenant.'
)
r.font.bold = True; r.font.italic = True; r.font.size = Pt(12)


# ════════════════════════════════════════════════════════════════
# 11. FEUILLE DE ROUTE
# ════════════════════════════════════════════════════════════════
doc.add_heading('11. Feuille de route', level=1)

doc.add_heading('11.1 Court terme (S17-20)', level=2)
tbl(
    ['Priorite', 'Action', 'Echeance'],
    [
        ['CRITIQUE', 'Reconstruire PPTs Gen Z avec indices V3', 'S17-18'],
        ['CRITIQUE', 'Corriger matrice (100% exact)', 'Immediat'],
        ['CRITIQUE', 'Reintegrer doctrines et architecture causale', 'S18'],
        ['MAJEUR', 'Migrations 009/010 en production', 'S18'],
        ['MAJEUR', 'Grille de notation unifiee', 'S18'],
        ['STANDARD', 'DATATYM(TM) sur tous les indices', 'S17'],
    ]
)

doc.add_heading('11.2 Vague 2 (Septembre 2026)', level=2)
bul([
    '2 500+ repondants, 20+ pays',
    'Indices mis a jour + comparatif V1/V2',
    'Variable genre ajoutee',
    'Top 50 Attractivite Africaine publie',
])

doc.add_heading('11.3 Vision 2027', level=2)
bul([
    'Barometre annuel officiel',
    'API donnees DATATYM\u2122',
    'Partenariats academiques (CESAG, HEC Abidjan)',
    'DATATYM\u2122 Summit — conference annuelle',
    'DATATYM\u2122 reference dans les COMEX africains',
])


# ── FIN ────────────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\n\nDATATYM\u2122\nH&C Executive / Marketym\nDocument confidentiel — Avril 2026')
r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

out = '/Users/Apple/Desktop/H&C/Enquetes/Suivi enquetes/enquetes/Rapport_Activite_Complet_DATATYM_202604.docx'
doc.save(out)
print(f'OK -> {out}')
