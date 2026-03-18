"""
Script pour generer le cahier de charge COMPLET en format Word (.docx)
Version 2 - avec tous les details d'implementation
"""
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, '059669')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F0FDF4')
    return table


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def h1(doc, text):
    doc.add_heading(text, level=1)

def h2(doc, text):
    doc.add_heading(text, level=2)

def h3(doc, text):
    doc.add_heading(text, level=3)

def para(doc, text):
    doc.add_paragraph(text)

def bold_para(doc, label, text=""):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    if text:
        p.add_run(text)

def bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')


def build_part1_cover_and_toc(doc):
    """Page de couverture + table des matieres"""
    for _ in range(6):
        doc.add_paragraph('')
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CAHIER DE CHARGE COMPLET')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Marketym v2 - Plateforme de Suivi d\'Enquetes')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    doc.add_paragraph('')
    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org.add_run('H&C Executive')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('Mars 2026 - Version definitive')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_page_break()
    h1(doc, 'Table des matieres')
    sections = [
        '1. Presentation generale',
        '2. Architecture technique Django',
        '3. Schema complet de base de donnees',
        '4. Application accounts - Authentification',
        '5. Application enquetes - Gestion des enquetes',
        '6. Application quotas - Quotas simples et croises',
        '7. Application sync - Synchronisation QuestionPro',
        '8. Application tracking - Suivi des clics',
        '9. Application dashboard - Tableaux de bord',
        '10. Application exports - Exports CSV',
        '11. Application notifications',
        '12. Integration QuestionPro - Specification API',
        '13. Logique metier - Excedents, calculs, edge cases',
        '14. Templates, pages et composants UI',
        '15. JavaScript minimal requis',
        '16. Emails - Templates HTML exacts',
        '17. Messages d\'erreur et codes HTTP',
        '18. Donnees initiales et migration',
        '19. Configuration deploiement',
        '20. Etapes d\'implementation detaillees',
    ]
    for s in sections:
        p = doc.add_paragraph(s)
        p.paragraph_format.space_after = Pt(3)
    doc.add_page_break()


def build_part2_presentation(doc):
    h1(doc, '1. Presentation generale')

    h2(doc, '1.1 Contexte')
    para(doc, 'H&C Executive realise des enquetes terrain via des enqueteurs. Les sondages sont heberges sur QuestionPro. Cette plateforme permet de gerer les enqueteurs, suivre leur progression, gerer les quotas et segmentations, et synchroniser les donnees avec QuestionPro.')

    h2(doc, '1.2 Utilisateurs et roles')
    add_table(doc,
        ['Role', 'Description', 'Acces'],
        [
            ['super_admin', 'Gestionnaire global', 'Tout (CRUD, suppression, roles)'],
            ['admin', 'Responsable d\'enquetes', 'CRUD enquetes/enqueteurs/affectations, sync, stats (pas de suppression)'],
            ['enqueteur', 'Realise les enquetes', 'Dashboard personnel, ses enquetes, profil'],
        ])

    h2(doc, '1.3 Stack technique')
    add_table(doc,
        ['Composant', 'Technologie'],
        [
            ['Framework', 'Django 5.x (Python 3.12+)'],
            ['Base de donnees', 'PostgreSQL (Supabase)'],
            ['Rendu', 'Django Templates + HTML/CSS (server-side rendering)'],
            ['CSS', 'TailwindCSS via CDN'],
            ['Emails', 'Brevo (ex-Sendinblue) via SDK sib_api_v3_sdk'],
            ['Sondages', 'QuestionPro API v2'],
            ['HTTP client', 'httpx (async)'],
            ['Deploiement', 'Railway (Dockerfile + gunicorn)'],
            ['Static files', 'whitenoise'],
        ])

    h2(doc, '1.4 Contraintes techniques')
    for c in [
        'Pas d\'API REST separee : tout en server-side rendering Django',
        'Pas de framework JavaScript (React, Vue, Angular)',
        'JavaScript minimal : uniquement pour modals, copier dans presse-papier, confirmations',
        'Pas de WebSocket : les mises a jour se font par rechargement de page ou sync manuelle',
        'Base PostgreSQL existante sur Supabase (migration des donnees existantes)',
        'Quota API QuestionPro : 5000 appels/mois, 300 appels/60 secondes',
    ]:
        bullet(doc, c)
    doc.add_page_break()


def build_part3_architecture(doc):
    h1(doc, '2. Architecture technique Django')

    h2(doc, '2.1 Structure du projet')
    add_code_block(doc, """marketym/
  manage.py
  marketym/                    # Projet principal
      settings.py, urls.py, wsgi.py
  accounts/                    # Auth, utilisateurs, OTP, invitations
      models.py                # Enqueteur (AbstractUser), OtpCode, InvitationToken
      views.py                 # Login, Register, Activate, ForgotPwd, ResetPwd, Profile
      forms.py                 # LoginForm, RegisterForm, OTPForm, etc.
      urls.py
      decorators.py            # @admin_required, @super_admin_required
      services/
          email_service.py     # Envoi emails Brevo (OTP, welcome, reset)
          otp_service.py       # Generate, hash, verify OTP
      templates/accounts/
  enquetes/                    # Enquetes, affectations, segmentations
      models.py                # Zone, Pays, Enquete, Affectation, Segmentation
      views.py                 # CRUD enquetes, enqueteurs (admin), affectations
      forms.py
      urls.py
      templates/enquetes/
  quotas/                      # Quotas simples et croises
      models.py                # Quota, QuotaConfig, QuotaConfigQuestion
      views.py
      forms.py
      urls.py
      templates/quotas/
  sync/                        # Synchronisation QuestionPro
      models.py                # CompletionSegment, CompletionCombination, CompletionPays, HistoriqueCompletion
      views.py                 # SyncAllView, SyncAffectationView
      services/
          questionpro.py       # Client API QuestionPro
          sync_engine.py       # Logique de synchronisation
      urls.py
      templates/sync/
  tracking/                    # Tracking clics
      models.py                # Clic
      views.py                 # TrackClickView (/r/{id}), ClicListView
      urls.py
  dashboard/                   # Dashboards
      views.py                 # AdminDashboardView, EnqueteurDashboardView
      urls.py
      templatetags/
          dashboard_tags.py    # Filtres custom (pourcentage, timeago, etc.)
      templates/dashboard/
  exports/                     # Exports CSV
      views.py
      urls.py
  notifications/               # Rappels, alertes
      views.py
      urls.py
  templates/                   # Templates globaux
      base.html
      base_admin.html
      base_enqueteur.html
      components/              # Includes reutilisables
  static/
      css/main.css
      js/main.js""")

    h2(doc, '2.2 Variables d\'environnement')
    add_table(doc,
        ['Variable', 'Obligatoire', 'Default', 'Description'],
        [
            ['DATABASE_URL', 'Oui', '-', 'URL PostgreSQL complet'],
            ['SUPABASE_URL', 'Oui', '-', 'URL du projet Supabase'],
            ['SUPABASE_KEY', 'Oui', '-', 'Cle API Supabase (service role)'],
            ['SECRET_KEY', 'Oui', '-', 'Django secret key'],
            ['QUESTIONPRO_API_KEY', 'Oui', '-', 'Cle API QuestionPro'],
            ['BREVO_API_KEY', 'Oui', '-', 'Cle API Brevo'],
            ['EMAIL_FROM', 'Non', 'noreply@example.com', 'Adresse expediteur'],
            ['EMAIL_FROM_NAME', 'Non', 'Marketym', 'Nom expediteur'],
            ['JWT_SECRET_KEY', 'Oui', '-', 'Pour tokens d\'invitation (pas auth session)'],
            ['JWT_ALGORITHM', 'Non', 'HS256', 'Algorithme JWT'],
            ['OTP_EXPIRE_MINUTES', 'Non', '5', 'Duree validite OTP'],
            ['OTP_MAX_ATTEMPTS', 'Non', '3', 'Tentatives max OTP'],
            ['FRONTEND_URL', 'Non', 'http://localhost:8000', 'URL publique du site'],
            ['SYNC_INTERVAL_MINUTES', 'Non', '30', 'Intervalle sync auto'],
            ['ALLOWED_HOSTS', 'Oui (prod)', '*', 'Hosts autorises'],
            ['DEBUG', 'Non', 'False', 'Mode debug Django'],
        ])

    h2(doc, '2.3 Dependencies Python (requirements.txt)')
    add_code_block(doc, """Django>=5.0
psycopg2-binary
bcrypt==4.1.2
PyJWT==2.8.0
sib-api-v3-sdk==7.6.0
httpx==0.27.0
python-dotenv==1.0.1
gunicorn
whitenoise""")

    h2(doc, '2.4 Authentification')
    para(doc, 'Django utilise l\'authentification par session (cookies). Le JWT n\'est utilise QUE pour les tokens d\'invitation (liens d\'activation par email). L\'authentification courante passe par request.user et le middleware SessionMiddleware.')

    doc.add_page_break()


def build_part4_database(doc):
    h1(doc, '3. Schema complet de base de donnees')

    para(doc, 'Toutes les cles primaires sont des UUID (gen_random_uuid() en PostgreSQL, uuid4 en Django). Les timestamps utilisent DateTimeField avec auto_now_add ou auto_now.')

    # Enqueteur
    h2(doc, '3.1 Enqueteur (accounts.models)')
    para(doc, 'Etend AbstractUser de Django. Le champ username n\'est pas utilise (USERNAME_FIELD = "email").')
    add_table(doc,
        ['Champ', 'Type Django', 'Type SQL', 'Contraintes', 'Description'],
        [
            ['id', 'UUIDField', 'UUID', 'PK, default=uuid4', 'Identifiant unique'],
            ['email', 'EmailField(255)', 'varchar(255)', 'UNIQUE, NOT NULL', 'Email = login'],
            ['password', 'CharField (herite)', 'varchar(128)', 'NOT NULL', 'Hash bcrypt'],
            ['nom', 'CharField(100)', 'varchar(100)', 'NOT NULL', 'Nom (.upper() a la sauvegarde)'],
            ['prenom', 'CharField(100)', 'varchar(100)', 'NOT NULL', 'Prenom (.title() a la sauvegarde)'],
            ['telephone', 'CharField(20)', 'varchar(20)', 'NULL, blank', 'Telephone'],
            ['identifiant', 'CharField(20)', 'varchar(20)', 'UNIQUE, NOT NULL', 'Code court ex: USR0001, adm00001'],
            ['token', 'CharField(20)', 'varchar(20)', 'UNIQUE, NULL', 'Token QP tracking (6 chars, sans 0OIL)'],
            ['actif', 'BooleanField', 'boolean', 'default=True', 'Compte actif'],
            ['role', 'CharField(20)', 'varchar(20)', 'default="enqueteur"', 'enqueteur|admin|super_admin'],
            ['is_admin', 'BooleanField', 'boolean', 'default=False', 'Legacy, True si role in (admin, super_admin)'],
            ['compte_configure', 'BooleanField', 'boolean', 'default=False', 'True apres validation OTP 1ere connexion'],
            ['doit_changer_mdp', 'BooleanField', 'boolean', 'default=False', 'Forcer changement MDP'],
            ['reseau_mobile', 'CharField(20)', 'varchar(20)', 'NULL, choices', 'wave|orange_money|free_money'],
            ['mode_remuneration', 'CharField(20)', 'varchar(20)', 'NULL, choices', 'virement|espece|espece_virement|cheque'],
            ['derniere_connexion', 'DateTimeField', 'timestamptz', 'NULL', 'Derniere connexion'],
            ['created_at', 'DateTimeField', 'timestamptz', 'auto_now_add', 'Date creation'],
        ])

    h3(doc, 'Generation d\'identifiant')
    para(doc, 'A la creation par admin : prefixe "adm" pour admin, "usr" pour enqueteur, suivi d\'un numero sequentiel a 5 chiffres (adm00001, usr00042). A la self-registration : prefixe "USR" + 4 chiffres aleatoires (USR3847), avec boucle d\'unicite.')

    h3(doc, 'Generation de token')
    para(doc, 'Alphabet : A-Z + 0-9 sauf 0, O, I, L (pour eviter ambiguite). Longueur : 6 caracteres. Genere avec secrets.choice(). Exemples : WK47HP, B3DNER.')

    # OtpCode
    h2(doc, '3.2 OtpCode (accounts.models)')
    add_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUIDField', 'PK', ''],
            ['email', 'EmailField', 'NOT NULL, INDEX', 'Email associe'],
            ['code_hash', 'CharField(255)', 'NOT NULL', 'Hash bcrypt du code 6 chiffres'],
            ['expires_at', 'DateTimeField', 'NOT NULL', 'now() + OTP_EXPIRE_MINUTES'],
            ['attempts', 'IntegerField', 'default=0', 'Incremente a chaque echec'],
            ['used', 'BooleanField', 'default=False', 'Marque utilise apres verification'],
            ['created_at', 'DateTimeField', 'auto_now_add', ''],
        ])

    h3(doc, 'Rate limiting OTP')
    para(doc, 'Maximum 5 codes OTP en 15 minutes par email. Verification : COUNT codes WHERE email=X AND created_at > now()-15min. Si depasse : HTTP 429 "Trop de demandes. Veuillez attendre 15 minutes."')

    # InvitationToken
    h2(doc, '3.3 InvitationToken (accounts.models)')
    add_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUIDField', 'PK', ''],
            ['enqueteur', 'ForeignKey(Enqueteur)', 'CASCADE', ''],
            ['token', 'CharField(255)', 'UNIQUE', 'secrets.token_urlsafe(48) = 64 chars'],
            ['expires_at', 'DateTimeField', 'NOT NULL', 'now() + 48 heures'],
            ['used', 'BooleanField', 'default=False', ''],
            ['used_at', 'DateTimeField', 'NULL', ''],
            ['created_at', 'DateTimeField', 'auto_now_add', ''],
        ])

    # Zone + Pays
    h2(doc, '3.4 Zone et Pays (enquetes.models)')
    para(doc, 'Zone : id (UUID), nom (varchar 100, UNIQUE). Pays : id (UUID), nom (varchar 100), code (varchar 10, UNIQUE, ISO), zone (FK Zone, NULL), quota (int, default 0, legacy), icp_pct (decimal 5,2, default 0).')
    para(doc, 'Zones predefinies : UEMOA, CEMAC. 15 pays pre-inseres (voir section 18).')

    # Enquete
    h2(doc, '3.5 Enquete (enquetes.models)')
    add_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUIDField', 'PK', ''],
            ['survey_id', 'CharField(50)', 'NOT NULL', 'ID du sondage QuestionPro'],
            ['code', 'CharField(20)', 'UNIQUE, NULL', 'Code interne (par defaut = survey_id)'],
            ['nom', 'CharField(300)', 'NOT NULL', 'Recupere depuis QP a la creation'],
            ['description', 'TextField', 'NULL, blank', ''],
            ['cible', 'TextField', 'NOT NULL', 'Description du public cible (texte libre)'],
            ['taille_echantillon', 'IntegerField', 'default=0', 'Objectif global en nombre'],
            ['statut', 'CharField(20)', 'default="en_cours"', 'brouillon|en_cours|termine|archive'],
            ['survey_url', 'TextField', 'NULL', 'URL sondage QP (cache)'],
            ['segmentation_question_id', 'CharField(50)', 'NULL', 'Legacy : ID question segmentation'],
            ['segmentation_question_text', 'TextField', 'NULL', 'Legacy : texte question segmentation'],
            ['created_at', 'DateTimeField', 'auto_now_add', ''],
        ])
    para(doc, 'Note : le champ "cible" est un texte descriptif (ex: "Dirigeants, managers"). L\'objectif numerique est dans "taille_echantillon".')

    # Affectation
    h2(doc, '3.6 Affectation (enquetes.models)')
    add_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUIDField', 'PK', ''],
            ['enqueteur', 'FK(Enqueteur)', 'CASCADE', ''],
            ['enquete', 'FK(Enquete)', 'CASCADE', ''],
            ['survey_id', 'CharField(50)', 'NOT NULL', 'Denormalise depuis enquete (ou individuel)'],
            ['lien_questionnaire', 'TextField', 'NULL', 'Lien tracking /r/{id}'],
            ['lien_direct', 'TextField', 'NULL', 'Lien QP direct avec ?custom1=token'],
            ['lien_legacy', 'TextField', 'NULL', 'Ancien lien (backup)'],
            ['objectif_total', 'IntegerField', 'default=200', ''],
            ['completions_total', 'IntegerField', 'default=0', ''],
            ['clics', 'IntegerField', 'default=0', 'Clics uniques (par IP)'],
            ['invalid_total', 'IntegerField', 'default=0', 'Toujours 0 (legacy)'],
            ['statut', 'CharField(20)', 'default="en_cours"', 'en_cours|en_retard|termine'],
            ['commentaire_admin', 'TextField', 'NULL, blank', ''],
            ['derniere_synchro', 'DateTimeField', 'NULL', ''],
            ['created_at', 'DateTimeField', 'auto_now_add', ''],
        ])
    para(doc, 'Contrainte UNIQUE(enqueteur, enquete). Trigger auto_status : si completions >= objectif -> "termine", si < 50% et > 7 jours -> "en_retard".')

    h3(doc, 'Ancien systeme vs nouveau systeme')
    para(doc, 'Ancien systeme : chaque enqueteur a son propre survey_id (copie du sondage). Aff.survey_id != Enquete.survey_id. Toutes les reponses du survey appartiennent a l\'enqueteur.')
    para(doc, 'Nouveau systeme : un seul survey partage. Aff.survey_id == Enquete.survey_id. Filtrage par custom1=token dans les reponses.')

    # Segmentation
    h2(doc, '3.7 Segmentation (enquetes.models)')
    add_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUIDField', 'PK', ''],
            ['enquete', 'FK(Enquete)', 'CASCADE', ''],
            ['question_id', 'CharField(50)', 'NOT NULL', 'ID question QP'],
            ['question_text', 'CharField(500)', 'NULL', 'Texte de la question'],
            ['nom', 'CharField(100)', 'NOT NULL', 'Nom affiche (ex: "Pays")'],
            ['answer_options', 'JSONField', 'default=list', 'Cache: [{"id":123,"text":"Senegal"},...]'],
            ['created_at', 'DateTimeField', 'auto_now_add', ''],
        ])
    para(doc, 'Contrainte UNIQUE(enquete, question_id). Les answer_options sont cachees a la creation pour eviter des appels API repetitifs.')

    # Quota
    h2(doc, '3.8 Quota (quotas.models)')
    add_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUIDField', 'PK', ''],
            ['segmentation', 'FK(Segmentation)', 'CASCADE, NULL', 'Pour quotas simples'],
            ['enquete', 'FK(Enquete)', 'via segmentation', 'Indirect'],
            ['affectation', 'FK(Affectation)', 'CASCADE, NULL', 'NULL = global, non-NULL = par enqueteur'],
            ['segment_value', 'CharField(200)', 'NOT NULL', 'Ex: "Senegal", "18-25"'],
            ['objectif', 'IntegerField', 'default=0', 'Objectif (calcule depuis pourcentage)'],
            ['completions', 'IntegerField', 'default=0', 'Mis a jour par sync'],
            ['pourcentage', 'DecimalField(5,2)', 'default=0', '% de repartition'],
            ['quota_config', 'FK(QuotaConfig)', 'CASCADE, NULL', 'Pour quotas croises'],
            ['combination', 'JSONField', 'NULL', 'Combinaison croisee ex: {"Pays":"SN"}'],
        ])
    para(doc, 'Calcul objectif depuis pourcentage : objectif = int(objectif_total_affectation * pourcentage / 100).')

    # QuotaConfig + QuotaConfigQuestion
    h2(doc, '3.9 QuotaConfig et QuotaConfigQuestion (quotas.models)')
    para(doc, 'QuotaConfig : id (UUID PK), enquete (FK CASCADE), nom (varchar 200), created_at.')
    para(doc, 'QuotaConfigQuestion : id (UUID PK), quota_config (FK CASCADE), segmentation (FK), position (int). UNIQUE(quota_config, segmentation).')

    # Completion tables
    h2(doc, '3.10 Tables de completions (sync.models)')

    h3(doc, 'CompletionSegment')
    para(doc, 'Champs : id, affectation (FK CASCADE), segmentation (FK), segment_value (varchar 200), completions (int, default 0), updated_at (auto_now). UNIQUE(affectation, segmentation, segment_value).')

    h3(doc, 'CompletionCombination')
    para(doc, 'Champs : id, affectation (FK CASCADE), quota_config (FK CASCADE), combination (JSONField), completions (int), updated_at. UNIQUE(affectation, quota_config, combination). La combination est serialisee avec json.dumps(combo, sort_keys=True) pour deduplication.')

    h3(doc, 'CompletionPays')
    para(doc, 'Champs : id, affectation (FK CASCADE), pays (FK Pays ou varchar), completions (int), objectif (int). UNIQUE(affectation, pays). A la creation d\'une affectation, une ligne est creee pour CHAQUE pays (CROSS JOIN).')

    h3(doc, 'HistoriqueCompletion')
    para(doc, 'Champs : id, date (DateField), enquete (FK CASCADE), affectation (FK CASCADE), completions (int), clics (int), created_at. UNIQUE(date, enquete_id, affectation_id). Rempli par la sync : les reponses sont groupees par date via utctimestamp.')

    # Clic
    h2(doc, '3.11 Clic (tracking.models)')
    para(doc, 'Champs : id (UUID), affectation (FK CASCADE), ip_address (varchar 45), user_agent (text, NULL, max 500 chars), created_at (auto_now_add). UNIQUE(affectation, ip_address). Deduplication par contrainte unique.')

    doc.add_page_break()


def build_part5_accounts(doc):
    h1(doc, '4. Application accounts - Authentification')

    h2(doc, '4.1 Flux d\'authentification complet')

    h3(doc, 'Inscription (self-registration)')
    para(doc, 'URL : /accounts/register (GET/POST)')
    para(doc, 'Champs : nom, prenom, email, telephone (optionnel), password (min 8 chars).')
    para(doc, 'Logique : verifier unicite email. Generer identifiant "USR{random 1000-9999}" avec boucle unicite. Sauvegarder avec nom.upper(), prenom.title(), compte_configure=False, actif=True, is_admin=False. Redirect vers /accounts/login.')
    para(doc, 'Erreur : "Cet email est deja utilise" (400).')

    h3(doc, 'Connexion')
    para(doc, 'URL : /accounts/login (GET/POST)')
    para(doc, 'Champs : email, password.')
    para(doc, 'Logique :')
    bullet(doc, 'Trouver enqueteur par email (insensible a la casse, strip). Sinon : "Email ou mot de passe incorrect".')
    bullet(doc, 'Verifier actif. Sinon : "Compte desactive" (403).')
    bullet(doc, 'Verifier password existe. Sinon : "Mot de passe non configure" (401).')
    bullet(doc, 'Verifier password avec bcrypt. Sinon : "Email ou mot de passe incorrect" (401).')
    bullet(doc, 'Si compte_configure == True : creer session Django, mettre a jour derniere_connexion, redirect vers /dashboard (enqueteur) ou /admin (admin).')
    bullet(doc, 'Si compte_configure == False : generer OTP (6 chiffres), hasher avec bcrypt, sauvegarder, envoyer par email, redirect vers /accounts/activate avec email en session.')

    h3(doc, 'Activation OTP (premiere connexion)')
    para(doc, 'URL : /accounts/activate (GET/POST)')
    para(doc, 'Champ : code (6 chiffres).')
    para(doc, 'Logique :')
    bullet(doc, 'Recuperer le dernier OTP non-utilise pour cet email, ordonne par created_at DESC, LIMIT 1.')
    bullet(doc, 'Si aucun : "Aucun code en attente" (400).')
    bullet(doc, 'Verifier expiration : si expires_at < now() : marquer used=True, "Code expire" (400).')
    bullet(doc, 'Verifier tentatives : si attempts >= 3 : marquer used=True, "Trop de tentatives. Demandez un nouveau code." (400).')
    bullet(doc, 'Verifier code avec bcrypt : si incorrect : incrementer attempts, "Code incorrect. {3-attempts} tentative(s) restante(s)." (400).')
    bullet(doc, 'Si correct : marquer used=True, mettre compte_configure=True, creer session, redirect vers dashboard.')

    h3(doc, 'Invitation par admin')
    para(doc, 'Declenchement : admin cree un enqueteur OU clique "Envoyer invitation".')
    para(doc, 'Logique :')
    bullet(doc, 'Verifier que le compte n\'est pas deja configure. Sinon : "Ce compte est deja active" (400).')
    bullet(doc, 'Invalider les anciennes invitations non-utilisees (used=True).')
    bullet(doc, 'Generer token : secrets.token_urlsafe(48) = 64 caracteres URL-safe.')
    bullet(doc, 'Sauvegarder avec expires_at = now() + 48 heures.')
    bullet(doc, 'Construire lien : {FRONTEND_URL}/activer-compte?token={token}')
    bullet(doc, 'Envoyer email de bienvenue avec ce lien.')

    h3(doc, 'Activation par lien d\'invitation')
    para(doc, 'URL : /accounts/activate/{token} (GET/POST)')
    para(doc, 'Champ : nouveau mot de passe (min 8 chars).')
    para(doc, 'Logique : trouver invitation_token par token (used=False). Verifier expiration. Hasher et sauvegarder le mot de passe. Mettre compte_configure=True, doit_changer_mdp=False. Marquer invitation used=True avec used_at. Redirect login.')

    h3(doc, 'Mot de passe oublie')
    para(doc, 'URL : /accounts/forgot-password (GET/POST)')
    para(doc, 'Champ : email.')
    para(doc, 'Logique : si utilisateur existe et actif, generer et envoyer OTP. Reponse TOUJOURS : "Si cet email existe, un code de verification a ete envoye" (securite : ne pas reveler si le compte existe). Redirect vers /accounts/reset-password.')

    h3(doc, 'Reset mot de passe')
    para(doc, 'URL : /accounts/reset-password (GET/POST)')
    para(doc, 'Champs : email, code OTP, nouveau mot de passe (min 8 chars).')
    para(doc, 'Logique : meme verification OTP que l\'activation. Hasher et sauvegarder nouveau mot de passe.')

    h3(doc, 'Changement de mot de passe (connecte)')
    para(doc, 'URL : /accounts/change-password (GET/POST)')
    para(doc, 'Champs : mot de passe actuel, nouveau mot de passe.')
    para(doc, 'Validations : mot de passe actuel correct, nouveau >= 8 chars, nouveau != ancien.')

    h3(doc, 'Modification de profil (connecte)')
    para(doc, 'URL : /accounts/profile (GET/POST)')
    para(doc, 'Flux en 2 etapes : 1) Demander OTP (POST /accounts/request-profile-otp -> envoie code). 2) Soumettre modifications avec code OTP.')
    para(doc, 'Champs modifiables : nom, prenom, email (verifier unicite), telephone, reseau_mobile (choix: wave, orange_money, free_money, vide), mode_remuneration (choix: virement, espece, espece_virement, cheque, vide).')

    h2(doc, '4.2 Decorateurs d\'acces')
    add_code_block(doc, """@login_required           # Django built-in, redirige vers /accounts/login
@admin_required           # Verifie is_admin==True OU role in ('admin','super_admin')
                          # Sinon 403 "Acces reserve aux administrateurs"
@super_admin_required     # Verifie role=='super_admin'
                          # Sinon 403 "Acces reserve aux super administrateurs" """)

    h2(doc, '4.3 Service OTP (otp_service.py)')
    add_code_block(doc, """def generate_otp(length=6) -> str:
    # Code numerique pur : secrets.choice(string.digits) x length
    return "".join(secrets.choice(string.digits) for _ in range(length))

def hash_code(code: str) -> str:
    return bcrypt.hashpw(code.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")

def verify_code(code: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(code.encode("utf-8"), hashed.encode("utf-8"))
    except Exception:
        return False

def create_otp(email: str) -> str:
    # 1. Generer code
    # 2. Hasher
    # 3. Invalider anciens codes (used=True WHERE email AND used=False)
    # 4. Sauvegarder nouveau (email, code_hash, expires_at=now+5min)
    # 5. Retourner code en clair (pour envoi email)

def verify_otp(email: str, code: str) -> tuple[bool, str]:
    # 1. Chercher dernier code non-utilise pour cet email
    # 2. Verifier expiration, tentatives, code
    # 3. Retourner (success, message_erreur)""")

    doc.add_page_break()


def build_part6_enquetes(doc):
    h1(doc, '5. Application enquetes')

    h2(doc, '5.1 CRUD Enquetes (admin)')
    add_table(doc,
        ['URL', 'Vue', 'Methode', 'Description'],
        [
            ['/admin/enquetes/', 'EnqueteListView', 'GET', 'Liste avec stats calculees'],
            ['/admin/enquetes/create/', 'EnqueteCreateView', 'GET/POST', 'Saisie survey_id -> fetch QP'],
            ['/admin/enquetes/<id>/', 'EnqueteDetailView', 'GET', 'Detail complet'],
            ['/admin/enquetes/<id>/edit/', 'EnqueteUpdateView', 'GET/POST', 'Modifier'],
            ['/admin/enquetes/<id>/delete/', 'EnqueteDeleteView', 'POST', 'Supprimer (super_admin)'],
        ])

    h3(doc, 'Creation d\'enquete - flux exact')
    para(doc, '1. Admin saisit survey_id (ex: "13445449").')
    para(doc, '2. Backend appelle GET {QP_BASE}/surveys/{survey_id} avec header api-key.')
    para(doc, '3. Parse response["response"] : extraire name, description. survey_url = shortUrl || webLink || surveyUrl || url || "".')
    para(doc, '4. Appelle GET {QP_BASE}/surveys/{survey_id}/questions pour pre-charger les questions.')
    para(doc, '5. Verifier unicite survey_id en base. Sinon : "Cette enquete existe deja" (400).')
    para(doc, '6. Inserer avec code=survey_id, statut="en_cours".')
    para(doc, '7. Retourner enquete creee + questions + survey_info.')

    h2(doc, '5.2 CRUD Enqueteurs (admin)')
    add_table(doc,
        ['URL', 'Vue', 'Methode', 'Description'],
        [
            ['/admin/enqueteurs/', 'EnqueteurListView', 'GET', 'Liste triee par identifiant'],
            ['/admin/enqueteurs/create/', 'EnqueteurCreateView', 'GET/POST', 'Creer + envoi invitation'],
            ['/admin/enqueteurs/<id>/', 'EnqueteurDetailView', 'GET', 'Dashboard enqueteur vu par admin'],
            ['/admin/enqueteurs/<id>/edit/', 'EnqueteurUpdateView', 'GET/POST', 'Modifier'],
            ['/admin/enqueteurs/<id>/delete/', 'EnqueteurDeleteView', 'POST', 'Supprimer (super_admin)'],
            ['/admin/enqueteurs/<id>/role/', 'UpdateRoleView', 'POST', 'Changer role (super_admin)'],
        ])

    h3(doc, 'Generation identifiant admin')
    para(doc, 'Si is_admin : prefixe "adm", sinon "usr". Numero sequentiel 5 chiffres base sur COUNT existant avec meme prefixe. Ex: adm00001, usr00042.')

    h2(doc, '5.3 CRUD Affectations (admin)')

    h3(doc, 'Creation d\'affectation - flux exact')
    para(doc, '1. Admin selectionne enqueteur + enquete + objectif.')
    para(doc, '2. Verifier que l\'enqueteur et l\'enquete existent.')
    para(doc, '3. Recuperer survey_url depuis QP (ou cache). Construire lien_direct :')
    add_code_block(doc, """# Si survey_url recupere depuis QP :
lien_direct = f"{survey_url}?custom1={enqueteur.token}"
# Sinon fallback :
lien_direct = f"https://hcakpo.questionpro.com/t/{survey_id}?custom1={enqueteur.token}" """)
    para(doc, '4. Inserer l\'affectation.')
    para(doc, '5. Generer lien_questionnaire (tracking) :')
    add_code_block(doc, """base_url = request.build_absolute_uri("/").rstrip("/")
# Forcer HTTPS en production (si pas localhost)
if "localhost" not in base_url and "127.0.0.1" not in base_url:
    base_url = base_url.replace("http://", "https://")
lien_questionnaire = f"{base_url}/r/{affectation.id}" """)
    para(doc, '6. Creer completions_pays : une ligne par pays (CROSS JOIN avec table pays).')
    para(doc, '7. Copier les quotas globaux en quotas individuels pour cette affectation.')

    h2(doc, '5.4 CRUD Segmentations')
    h3(doc, 'Creation de segmentation - flux exact')
    para(doc, '1. Admin selectionne enquete, entre question_id, question_text, nom.')
    para(doc, '2. Optionnel : answer_options (JSON des options de reponse QP).')
    para(doc, '3. Si answer_options est None, sauvegarder comme liste vide [].')
    para(doc, '4. Inserer avec UNIQUE(enquete, question_id).')

    h3(doc, 'Fetch des questions QP')
    para(doc, 'GET {QP_BASE}/surveys/{survey_id}/questions. Pour chaque question : extraire questionID (ou id), code, text, type. Extraire answers depuis q.answers ou q.answerChoices. Chaque answer : answerID (ou id), text. Ne retourner que les questions ayant au moins 1 option de reponse.')

    doc.add_page_break()


def build_part7_quotas(doc):
    h1(doc, '6. Application quotas')

    h2(doc, '6.1 Quotas simples')
    para(doc, 'Un quota definit un objectif par segment_value (option de reponse) pour une segmentation. Le champ "pourcentage" indique le % de l\'objectif total de l\'affectation.')
    para(doc, 'Si affectation_id IS NULL : quota global pour l\'enquete. Si non-NULL : quota par enqueteur.')
    para(doc, 'Calcul de l\'objectif effectif : objectif = int(objectif_total_affectation * pourcentage / 100).')

    h3(doc, 'Creation en masse (bulk)')
    para(doc, 'Body : enquete_id, segmentation_id, affectation_id (optionnel), quotas: [{segment_value, pourcentage}, ...]. Cree une ligne Quota par element.')

    h2(doc, '6.2 Quotas croises (cross-tabulation)')

    h3(doc, 'Processus de creation')
    para(doc, '1. Admin soumet : enquete_id, nom (ex: "Pays x Secteur"), segmentation_ids (liste d\'IDs), quotas (liste de {combination: {}, pourcentage: float}).')
    para(doc, '2. Inserer QuotaConfig(enquete_id, nom).')
    para(doc, '3. Pour chaque segmentation_id (avec position 0, 1, 2...) : inserer QuotaConfigQuestion.')
    para(doc, '4. Pour chaque quota : inserer Quota avec segment_value = " x ".join(combination.values()), combination = le dict JSON, pourcentage, quota_config_id.')

    h3(doc, 'Generation des combinaisons')
    add_code_block(doc, """from itertools import product

# Pour chaque segmentation liee au config :
# - Recuperer answer_options depuis la segmentation
# - Extraire la valeur : opt.get("text") or opt.get("value") or opt.get("label") or str(opt)
# Produit cartesien de toutes les options
combinations = [
    dict(zip(question_names, combo))
    for combo in product(*all_option_values)
]
# Retour : {"combinations": [...], "axes": [{"nom": "Pays", "values": ["SN","CI"]}, ...]}""")

    h3(doc, 'Affichage des quotas croises (admin enquete detail)')
    para(doc, 'Tableau avec colonnes dynamiques : une colonne par variable du croisement + colonnes fixes (%, Objectif, Completions, Progression). Pour chaque quota : objectif = int(taille_echantillon * pourcentage / 100). Completions = somme des completions_combinations matchant cette combination. Progression = completions/objectif*100.')

    doc.add_page_break()


def build_part8_sync(doc):
    h1(doc, '7. Application sync - Synchronisation QuestionPro')

    h2(doc, '7.1 Fonction sync_affectation - Logique complete')
    para(doc, 'Cette fonction est le coeur du systeme. Elle synchronise les donnees d\'une affectation avec QuestionPro.')

    bold_para(doc, 'Etape 1 : Charger les donnees')
    para(doc, 'Charger l\'affectation avec jointures : enqueteur (token), enquete (survey_id, segmentation_question_id). Determiner le systeme :')
    add_code_block(doc, """is_ancien_systeme = (
    aff_survey_id is not None
    and enquete_survey_id is not None
    and str(aff_survey_id) != str(enquete_survey_id)
)""")

    bold_para(doc, 'Etape 2 : Recuperer les reponses')
    para(doc, 'Ancien systeme : fetch depuis aff_survey_id. TOUTES les reponses completees appartiennent a cet enqueteur.')
    para(doc, 'Nouveau systeme : fetch depuis enquete_survey_id (fallback aff_survey_id). Filtrer par custom1 == enqueteur.token :')
    add_code_block(doc, """for resp in all_responses:
    custom_vars = resp.get("customVariables", [])
    custom1 = ""
    for cv in custom_vars:
        if cv.get("variableKey") == "custom1":
            custom1 = cv.get("variableValue", "")
            break
    if custom1 == enqueteur_token:
        enqueteur_responses.append(resp)
        if resp.get("responseStatus") == "Completed":
            completed_count += 1""")

    bold_para(doc, 'Etape 3 : Compter par segment')
    para(doc, 'Si segmentation_question_id existe (legacy) : pour chaque reponse completee, extraire la valeur du segment avec extract_segment_value_from_response() :')
    add_code_block(doc, """def extract_segment_value_from_response(response, question_id):
    for q in response.get("responseSet", []):
        if str(q.get("questionID")) == str(question_id):
            answers = q.get("answerValues", [])
            if answers:
                return answers[0].get("answerText", "Autre")
    return "Autre"  # Valeur par defaut si question non trouvee""")

    bold_para(doc, 'Etape 4 : Mettre a jour completions_segments')
    para(doc, 'DELETE tous les completions_segments pour cette affectation. Puis INSERT les nouveaux. Aucune normalisation : les valeurs sont stockees telles quelles depuis QuestionPro.')

    bold_para(doc, 'Etape 5 : Mettre a jour l\'affectation')
    add_code_block(doc, """affectation.completions_total = len(completed_responses)  # ou completed_count
affectation.clics = len(set(unique_ips))  # IPs uniques
affectation.invalid_total = 0  # Toujours 0
affectation.derniere_synchro = datetime.utcnow().isoformat()""")

    bold_para(doc, 'Etape 6 : Mettre a jour les quotas individuels')
    para(doc, 'Trouver les quotas WHERE affectation_id = cette affectation. Pour chaque quota, compter les completions_segments matchant segment_value. Mettre a jour quota.completions.')

    bold_para(doc, 'Etape 7 : Mettre a jour les quotas globaux')
    para(doc, 'Agreger completions_segments pour TOUTES les affectations de la meme enquete. Mettre a jour les quotas WHERE affectation_id IS NULL.')

    bold_para(doc, 'Etape 8 : Historique completions')
    para(doc, 'Grouper les reponses par date (utctimestamp -> datetime.utcfromtimestamp(ts).date()). Supprimer les anciens historiques pour cette affectation. Upsert avec on_conflict="date,enquete_id,affectation_id".')

    bold_para(doc, 'Etape 9 : Completions croisees (cross quotas)')
    para(doc, 'Pour chaque quota_config de l\'enquete :')
    add_code_block(doc, """# Charger les segmentations liees au config
# Pour chaque reponse completee :
#   Pour chaque segmentation du config :
#     Extraire la valeur de la reponse pour cette question
#   Construire combo = {"Pays": "Senegal", "Secteur": "Finance"}
#   combo_key = json.dumps(combo, sort_keys=True)
#   Incrementer compteur[combo_key]
# DELETE anciens completions_combinations pour (affectation, config)
# INSERT nouveaux avec les compteurs""")

    h2(doc, '7.2 Sync globale')
    para(doc, 'POST /admin/sync/ : recuperer toutes les affectations dont l\'enquete a statut="en_cours". Lancer sync_affectation pour chacune avec semaphore(5) pour parallelisme. Retourner {"synced": N, "errors": N, "results": [...]}.')

    h2(doc, '7.3 Fetch des reponses avec pagination')
    add_code_block(doc, """def fetch_all_survey_responses(survey_id):
    all_responses = []
    page = 1
    per_page = 1000  # Maximum autorise par QP
    while True:
        responses = fetch_survey_responses(survey_id, page, per_page)
        all_responses.extend(responses)
        if len(responses) < per_page:
            break
        page += 1
    return all_responses""")

    doc.add_page_break()


def build_part9_tracking(doc):
    h1(doc, '8. Application tracking')

    h2(doc, '8.1 TrackClickView')
    para(doc, 'URL : /r/<affectation_id> (GET, public, pas d\'auth)')
    add_code_block(doc, """def track_and_redirect(request, affectation_id):
    # 1. Recuperer affectation (id, lien_direct, lien_questionnaire, survey_id)
    #    Si pas trouvee : 404 "Lien invalide"

    # 2. Determiner URL de redirection
    redirect_url = affectation.lien_direct or affectation.lien_questionnaire
    #    Si aucun : 404 "Aucun lien de questionnaire configure"

    # 3. Extraire IP
    ip = request.META.get("HTTP_X_FORWARDED_FOR", "").split(",")[0].strip()
    if not ip:
        ip = request.META.get("HTTP_X_REAL_IP", "")
    if not ip:
        ip = request.META.get("REMOTE_ADDR", "unknown")

    # 4. User-Agent (tronque a 500 chars)
    ua = request.META.get("HTTP_USER_AGENT", "")[:500]

    # 5. Tenter INSERT dans clics
    try:
        Clic.objects.create(affectation=aff, ip_address=ip, user_agent=ua)
        # Si reussi : compter total clics et mettre a jour affectation.clics
        aff.clics = Clic.objects.filter(affectation=aff).count()
        aff.save()
    except IntegrityError:
        pass  # IP deja vue, ignorer silencieusement

    # 6. Redirect 302
    return redirect(redirect_url)""")

    doc.add_page_break()


def build_part10_dashboard(doc):
    h1(doc, '9. Application dashboard')

    h2(doc, '9.1 Dashboard admin principal')
    para(doc, 'URL : /admin/ (GET, @admin_required)')

    h3(doc, 'KPI globaux (4 cartes en haut)')
    add_table(doc,
        ['KPI', 'Calcul', 'Icone/Couleur'],
        [
            ['Enquetes actives', 'COUNT enquetes WHERE statut="en_cours"', 'Clipboard / Bleu #2563EB'],
            ['Enqueteurs actifs', 'COUNT enqueteurs WHERE actif=True AND is_admin=False', 'Users / Violet #7C3AED'],
            ['Completions totales', 'SUM affectations.completions_total', 'Check / Vert #059669'],
            ['Taux de completion', 'SUM(completions) / SUM(taille_echantillon) * 100', 'Target / Ambre #D97706'],
        ])
    para(doc, 'total_objectif = SUM(enquetes.taille_echantillon) et non SUM(affectations.objectif_total). Note : total_invalides est toujours 0, total_valides = total_completions (pas de calcul separe).')

    h3(doc, 'Contenu du dashboard')
    bullet(doc, 'Liste des enquetes en cours : tableau avec nom, code, cible, completions, taux, statut')
    bullet(doc, 'Enqueteurs recents : derniers connectes')

    h2(doc, '9.2 Detail enquete admin')
    para(doc, 'URL : /admin/enquetes/<id>/ (GET)')

    h3(doc, 'KPI enquete')
    add_table(doc,
        ['KPI', 'Calcul'],
        [
            ['Cible', 'enquete.taille_echantillon (ou SUM objectif_total des affectations si 0)'],
            ['Completions', 'SUM affectations.completions_total'],
            ['Valides', 'Apres logique excedent (voir section 13)'],
            ['Excedents', 'Completions - Valides'],
            ['Taux', 'Valides / Cible * 100'],
            ['Clics', 'SUM affectations.clics'],
        ])

    h3(doc, 'Tableau des affectations')
    para(doc, 'Colonnes : Enqueteur (nom prenom), Objectif, Completions, Valides, Excedents, Taux (%), Clics, Derniere sync. completions_valides = completions_total (pas de calcul excedent par affectation dans la v1 actuelle, a enrichir avec la logique excedent globale).')

    h3(doc, 'Graphiques barres par segmentation (ADMIN + ENQUETEUR)')
    para(doc, 'Pour CHAQUE segmentation de l\'enquete : un bloc avec titre (nom de la segmentation) et barres horizontales CSS. Chaque barre = un segment_value. Largeur = completions/objectif*100% (ou completions/max_completions si pas de quota). Couleur : vert si < objectif, orange si excedent.')

    h3(doc, 'Tableau croise (ADMIN UNIQUEMENT)')
    para(doc, 'Si des quota_configs existent : tableau avec colonnes dynamiques (une par variable du croisement) + %, Objectif, Completions, Progression. Pas affiche cote enqueteur.')

    h2(doc, '9.3 Dashboard enqueteur')
    para(doc, 'URL : /dashboard/ (@login_required)')

    h3(doc, 'Onglets')
    add_table(doc,
        ['Onglet', 'Contenu'],
        [
            ['Tableau de bord', 'KPI (enquetes actives, completions totales, taux moyen), liste enquetes'],
            ['Mes enquetes', 'Grille de cartes : une par affectation, avec objectif/completions/taux'],
            ['Profil', 'Infos perso, modification avec OTP, changement MDP'],
        ])

    h3(doc, 'Detail enquete enqueteur')
    para(doc, 'Au clic sur une enquete dans "Mes enquetes" :')
    bullet(doc, 'KPI : objectif, completions, valides (apres excedent global), excedents, taux')
    bullet(doc, 'Lien de collecte avec bouton copier (JS: navigator.clipboard.writeText)')
    bullet(doc, 'Graphiques barres par segmentation (MEME format que admin)')

    h2(doc, '9.4 Vue enqueteur dans admin')
    para(doc, 'URL : /admin/enqueteurs/<id>/ (@admin_required)')
    para(doc, 'Affiche le MEME contenu que le dashboard enqueteur, mais dans le layout admin (sidebar). L\'admin voit exactement ce que l\'enqueteur verrait. Donnees chargees via les memes queries (affectations de cet enqueteur, segmentations, completions).')

    h2(doc, '9.5 Graphiques barres CSS (sans JavaScript)')
    add_code_block(doc, """<!-- Template component: components/bar_chart.html -->
<!-- Inclus avec {% include "components/bar_chart.html" with segments=segments %} -->

{% for seg in segments %}
<div style="margin-bottom: 8px;">
    <div style="display: flex; justify-content: space-between; font-size: 13px; margin-bottom: 2px;">
        <span style="color: #374151;">{{ seg.label }}</span>
        <span style="color: #6B7280;">{{ seg.completions }} / {{ seg.objectif }}</span>
    </div>
    <div style="background: #E5E7EB; border-radius: 8px; height: 28px; position: relative; overflow: hidden;">
        {% if seg.percentage > 0 %}
        <div style="
            background: {% if seg.percentage > 100 %}#F59E0B{% else %}#059669{% endif %};
            width: {% if seg.percentage > 100 %}100{% else %}{{ seg.percentage }}{% endif %}%;
            height: 100%;
            border-radius: 8px;
            display: flex;
            align-items: center;
            padding: 0 8px;
        ">
            <span style="color: white; font-size: 12px; font-weight: 500;">{{ seg.percentage }}%</span>
        </div>
        {% endif %}
    </div>
</div>
{% endfor %}""")

    doc.add_page_break()


def build_part11_exports(doc):
    h1(doc, '10. Application exports')
    add_table(doc,
        ['URL', 'Export', 'Colonnes'],
        [
            ['/admin/exports/enqueteurs/', 'Enqueteurs CSV', 'id, nom, prenom, email, telephone, identifiant, actif, role, created_at'],
            ['/admin/exports/enquete/<id>/affectations/', 'Affectations CSV', 'enqueteur, objectif, completions, valides, excedents, taux, clics, derniere_synchro'],
            ['/admin/exports/enquete/<id>/quotas/', 'Quotas CSV', 'segmentation, segment_value, pourcentage, objectif, completions'],
        ])
    add_code_block(doc, """# Implementation standard Django
import csv
from django.http import HttpResponse

def export_enqueteurs(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="enqueteurs.csv"'
    writer = csv.writer(response)
    writer.writerow(['ID', 'Nom', 'Prenom', 'Email', ...])
    for e in Enqueteur.objects.filter(is_admin=False):
        writer.writerow([e.id, e.nom, e.prenom, e.email, ...])
    return response""")


def build_part12_notifications(doc):
    h1(doc, '11. Application notifications')
    add_table(doc,
        ['Type', 'Declencheur', 'Destinataire', 'Implementation'],
        [
            ['Rappel inactivite', 'Bouton admin "Envoyer rappel"', 'Enqueteur', 'Phase initiale : manuel'],
            ['Quota atteint', 'Sync detecte completions >= objectif', 'Admin', 'Alerte dans interface'],
            ['Excedent detecte', 'Sync detecte total > cible', 'Admin', 'Badge/alerte dans dashboard'],
        ])
    doc.add_page_break()


def build_part13_questionpro(doc):
    h1(doc, '12. Integration QuestionPro - Specification API')

    h2(doc, '12.1 Configuration')
    add_table(doc,
        ['Parametre', 'Valeur'],
        [
            ['Base URL', 'https://api.questionpro.com/a/api/v2'],
            ['Auth header', 'api-key: {QUESTIONPRO_API_KEY}'],
            ['Timeout GET survey', '30 secondes'],
            ['Timeout GET responses', '60 secondes'],
            ['Pagination', 'page (1-based), perPage (max 1000)'],
            ['Rate limit', '300 appels / 60 secondes'],
            ['Quota mensuel', '5000 appels gratuits'],
        ])

    h2(doc, '12.2 GET /surveys/{survey_id}')
    para(doc, 'Retourne les infos du sondage. Champs utilises :')
    add_code_block(doc, """response["response"] = {
    "name": "Nom du sondage",
    "description": "...",
    "completedResponses": 150,   # completions
    "viewedResponses": 300,      # clics/vues
    "startedResponses": 200,     # commencees
    "shortUrl": "https://...",   # URL courte (prioritaire)
    "webLink": "https://...",    # Fallback 1
    "surveyUrl": "https://...",  # Fallback 2
    "url": "https://...",        # Fallback 3
}
# survey_url = shortUrl or webLink or surveyUrl or url or "" """)

    h2(doc, '12.3 GET /surveys/{survey_id}/questions')
    para(doc, 'Retourne les questions du sondage. Parsing :')
    add_code_block(doc, """for q in response["response"]:
    question = {
        "id": q.get("questionID") or q.get("id"),
        "code": q.get("code", ""),
        "text": q.get("text", ""),
        "type": q.get("type", ""),
        "answers": []
    }
    # Reponses dans q["answers"] OU q["answerChoices"]
    raw_answers = q.get("answers", q.get("answerChoices", []))
    for a in raw_answers:
        question["answers"].append({
            "id": a.get("answerID") or a.get("id"),
            "text": a.get("text", "")
        })
    # Ne garder que les questions avec >= 1 reponse""")

    h2(doc, '12.4 GET /surveys/{survey_id}/responses')
    para(doc, 'Retourne les reponses. Structure d\'une reponse :')
    add_code_block(doc, """{
    "responseID": 12345,
    "responseStatus": "Completed",  // ou "Started", "Deleted"
    "utctimestamp": 1705312200,     // Unix timestamp
    "customVariables": [
        {"variableKey": "custom1", "variableValue": "WK47HP"}
    ],
    "responseSet": [      // OU "questionResponse" selon la version
        {
            "questionID": 67890,
            "answerValues": [
                {
                    "answerID": 111,
                    "answerText": "Senegal",   // Valeur utilisee pour segment matching
                    "value": {"text": "Senegal"}
                }
            ]
        }
    ]
}""")
    para(doc, 'IMPORTANT : le champ contenant les reponses peut s\'appeler "responseSet" ou "questionResponse" selon la version de l\'API. Le code actuel utilise "responseSet".')

    h2(doc, '12.5 Filtrage par enqueteur (custom1)')
    para(doc, 'Chaque enqueteur a un token unique (6 chars). Le lien QP inclut ?custom1={token}. A la sync, on filtre les reponses par customVariables[custom1] == token. Comparaison exacte (case-sensitive).')

    doc.add_page_break()


def build_part14_business_logic(doc):
    h1(doc, '13. Logique metier - Excedents, calculs, edge cases')

    h2(doc, '13.1 Logique des excedents (GLOBALE -> INDIVIDUELLE)')

    bold_para(doc, 'Principe fondamental :')
    para(doc, 'Les excedents sont d\'abord calcules au niveau GLOBAL de l\'enquete, puis propages aux enqueteurs individuels. Un enqueteur ne peut avoir d\'excedent que si le total global depasse la cible.')

    bold_para(doc, 'Algorithme :')
    add_code_block(doc, """def calculer_excedents(enquete, affectations):
    cible = enquete.taille_echantillon  # ou cible calculee
    total = sum(a.completions_total for a in affectations)

    if total <= cible:
        # Pas d'excedent : tout est valide
        return {
            a.id: {"valides": a.completions_total, "excedents": 0}
            for a in affectations
        }

    excedent_global = total - cible
    restant = excedent_global

    # Trier par derniere_synchro DESC (derniers = premiers a recevoir excedents)
    sorted_aff = sorted(
        affectations,
        key=lambda a: a.derniere_synchro or datetime.min,
        reverse=True
    )

    result = {}
    for aff in sorted_aff:
        if restant <= 0:
            result[aff.id] = {"valides": aff.completions_total, "excedents": 0}
        elif aff.completions_total <= restant:
            # Tout est excedent pour cet enqueteur
            result[aff.id] = {"valides": 0, "excedents": aff.completions_total}
            restant -= aff.completions_total
        else:
            # Partiellement excedent
            result[aff.id] = {
                "valides": aff.completions_total - restant,
                "excedents": restant
            }
            restant = 0

    return result""")

    bold_para(doc, 'Exemple concret :')
    add_code_block(doc, """Enquete "Satisfaction" - Cible: 100

Enqueteur A : 45 completions (sync 1er mars)
Enqueteur B : 35 completions (sync 5 mars)
Enqueteur C : 30 completions (sync 10 mars - le plus recent)
Total : 110 -> Excedent global : 10

Tri par derniere_synchro DESC : C, B, A

Enqueteur C (30 completions, restant=10) :
  30 > 10 -> valides=20, excedents=10, restant=0
Enqueteur B (35 completions, restant=0) :
  valides=35, excedents=0
Enqueteur A (45 completions, restant=0) :
  valides=45, excedents=0

Verification : 20+35+45 = 100 = cible. OK.""")

    bold_para(doc, 'Application aux segments :')
    para(doc, 'La meme logique s\'applique par segment. Pour un quota "Senegal = 50" : si le total global des completions "Senegal" (tous enqueteurs) depasse 50, les surplus sont attribues aux derniers arrivants.')

    h2(doc, '13.2 Calcul du taux')
    add_code_block(doc, """taux = round(completions_valides / objectif * 100, 1) if objectif > 0 else 0""")

    h2(doc, '13.3 Couleurs de progression')
    add_code_block(doc, """def get_progress_color(percentage):
    if percentage >= 100: return "#059669"  # vert fonce (objectif atteint)
    if percentage >= 75:  return "#10B981"  # vert
    if percentage >= 50:  return "#F59E0B"  # ambre/orange
    if percentage >= 25:  return "#F97316"  # orange
    return "#EF4444"                         # rouge""")

    h2(doc, '13.4 Edge cases geres')
    add_table(doc,
        ['Situation', 'Comportement'],
        [
            ['Enqueteur sans token', 'Token genere automatiquement (trigger DB ou code)'],
            ['Enquete sans affectation', 'Dashboard affiche 0 completions, pas d\'erreur'],
            ['Affectation sans lien_direct', 'Fallback: https://hcakpo.questionpro.com/t/{survey_id}?custom1={token}'],
            ['Affectation sans lien_questionnaire', 'Fallback vers lien_direct'],
            ['QP retourne 0 reponses', 'completions_total = 0, pas d\'erreur'],
            ['QP retourne erreur non-200', 'fetch_survey_stats retourne None, erreur propagee'],
            ['QP rate limit (429)', 'Pas de retry automatique, erreur retournee'],
            ['Question non trouvee dans reponse', 'segment_value = "Autre"'],
            ['Duplicate clic (meme IP)', 'IntegrityError catchee, ignoree silencieusement'],
            ['OTP expire', 'Marque used=True, message "Code expire"'],
            ['OTP 3 echecs', 'Marque used=True, message "Trop de tentatives"'],
            ['objectif = 0', 'taux = 0 (pas de division par zero)'],
            ['Email inexistant (forgot-password)', 'Reponse identique (securite)'],
            ['Mot de passe non configure', 'Message specifique "Mot de passe non configure"'],
        ])

    doc.add_page_break()


def build_part15_templates(doc):
    h1(doc, '14. Templates, pages et composants UI')

    h2(doc, '14.1 Charte graphique')
    add_table(doc,
        ['Element', 'Valeur'],
        [
            ['Couleur primaire', '#059669 (emeraude), hover #047857'],
            ['Couleur secondaire', '#111827'],
            ['Excedent', '#F59E0B (orange/ambre)'],
            ['Danger', '#EF4444 (rouge), soft #FEF2F2'],
            ['Info', '#2563EB (bleu), soft #EFF6FF'],
            ['Violet', '#7C3AED, soft #F5F3FF'],
            ['Ambre', '#D97706, soft #FFFBEB'],
            ['Fond page', '#F9FAFB'],
            ['Fond carte', '#FFFFFF, ombre shadow-sm'],
            ['Texte principal', '#111827'],
            ['Texte secondaire', '#6B7280'],
            ['Texte tertiaire', '#9CA3AF'],
            ['Bordures', '#E5E7EB, #D1D5DB'],
            ['Fond gris', '#F3F4F6'],
            ['Vert soft', '#ECFDF5, #D1FAE5, #A7F3D0'],
            ['Police texte', 'Inter, -apple-system, BlinkMacSystemFont, sans-serif'],
            ['Police code', 'SF Mono, Fira Code, Consolas, monospace'],
            ['Border radius cartes', '12px (rounded-xl)'],
            ['Border radius boutons', '8px (rounded-lg)'],
            ['Border radius modals', '16px (rounded-2xl)'],
            ['Nom application', 'Marketym'],
            ['Favicon', 'SVG inline: rect vert 10B981, ligne + courbe blanche'],
        ])

    h2(doc, '14.2 Layout admin (base_admin.html)')
    para(doc, 'Sidebar a gauche (260px, collapsible a 68px). Menu : Dashboard, Enquetes, Enqueteurs, Mes enquetes (si admin est aussi enqueteur), Profil, Deconnexion. Bouton sync en bas de sidebar. Zone de contenu a droite avec padding.')

    h2(doc, '14.3 Layout enqueteur (base_enqueteur.html)')
    para(doc, 'Header en haut avec logo + onglets (Tableau de bord, Mes enquetes, Profil) + avatar/deconnexion. Zone de contenu en dessous.')

    h2(doc, '14.4 Composants reutilisables')
    add_table(doc,
        ['Composant', 'Props/Variables', 'Description'],
        [
            ['stat_card.html', 'title, value, sub_value, color, icon_svg', 'Carte KPI avec icone et couleur de fond'],
            ['progress_bar.html', 'percentage, color, label', 'Barre progression avec fond gris E5E7EB'],
            ['bar_chart.html', 'segments (list of {label,completions,objectif,percentage})', 'Barres horizontales CSS'],
            ['cross_table.html', 'config, quotas, axes', 'Tableau croise pour quotas croises'],
            ['badge.html', 'text, variant (success/warning/danger/info)', 'Badge statut'],
            ['modal.html', 'id, title', 'Modal avec overlay, close button'],
            ['table.html', 'headers, rows', 'Tableau avec en-tetes'],
            ['alert.html', 'type (success/error/info/warning), message', 'Alerte dismissible'],
            ['copy_button.html', 'text_to_copy', 'Bouton copier avec JS clipboard'],
            ['empty_state.html', 'icon_svg, title, message', 'Etat vide'],
        ])

    h2(doc, '14.5 Liste complete des pages')
    add_table(doc,
        ['Page', 'Template', 'Layout', 'Acces'],
        [
            ['Connexion', 'accounts/login.html', 'base.html (minimal)', 'Public'],
            ['Inscription', 'accounts/register.html', 'base.html', 'Public'],
            ['Activation OTP', 'accounts/activate.html', 'base.html', 'Public'],
            ['Activation invitation', 'accounts/activate_token.html', 'base.html', 'Public'],
            ['Mot de passe oublie', 'accounts/forgot_password.html', 'base.html', 'Public'],
            ['Reset mot de passe', 'accounts/reset_password.html', 'base.html', 'Public'],
            ['Dashboard admin', 'dashboard/admin_dashboard.html', 'base_admin.html', 'Admin'],
            ['Liste enquetes', 'enquetes/enquete_list.html', 'base_admin.html', 'Admin'],
            ['Detail enquete', 'enquetes/enquete_detail.html', 'base_admin.html', 'Admin'],
            ['Form enquete', 'enquetes/enquete_form.html', 'base_admin.html', 'Admin'],
            ['Liste enqueteurs', 'enquetes/enqueteur_list.html', 'base_admin.html', 'Admin'],
            ['Detail enqueteur', 'dashboard/enqueteur_detail_admin.html', 'base_admin.html', 'Admin'],
            ['Form enqueteur', 'enquetes/enqueteur_form.html', 'base_admin.html', 'Admin'],
            ['Form affectation', 'enquetes/affectation_form.html', 'base_admin.html', 'Admin'],
            ['Form segmentation', 'enquetes/segmentation_form.html', 'base_admin.html', 'Admin'],
            ['Quotas enquete', 'quotas/quota_list.html', 'base_admin.html', 'Admin'],
            ['Form quota', 'quotas/quota_form.html', 'base_admin.html', 'Admin'],
            ['Form quota bulk', 'quotas/quota_bulk_form.html', 'base_admin.html', 'Admin'],
            ['Form quota croise', 'quotas/quota_config_form.html', 'base_admin.html', 'Admin'],
            ['Detail quota croise', 'quotas/quota_config_detail.html', 'base_admin.html', 'Admin'],
            ['Panel sync', 'sync/sync_panel.html', 'base_admin.html', 'Admin'],
            ['Statistiques', 'dashboard/stats.html', 'base_admin.html', 'Admin'],
            ['Exports', 'exports/export_panel.html', 'base_admin.html', 'Admin'],
            ['Dashboard enqueteur', 'dashboard/enqueteur_dashboard.html', 'base_enqueteur.html', 'Enqueteur'],
            ['Detail enquete enq.', 'dashboard/enqueteur_enquete_detail.html', 'base_enqueteur.html', 'Enqueteur'],
            ['Profil enqueteur', 'dashboard/enqueteur_profil.html', 'base_enqueteur.html', 'Enqueteur'],
        ])

    doc.add_page_break()


def build_part16_javascript(doc):
    h1(doc, '15. JavaScript minimal requis')

    para(doc, 'Pas de framework JS. Uniquement du vanilla JS pour les interactions suivantes :')

    h2(doc, '15.1 Modals')
    add_code_block(doc, """// Ouvrir/fermer un modal
function openModal(id) {
    document.getElementById(id).classList.remove('hidden');
    document.body.style.overflow = 'hidden';
}
function closeModal(id) {
    document.getElementById(id).classList.add('hidden');
    document.body.style.overflow = '';
}
// Fermer en cliquant sur l'overlay
document.querySelectorAll('.modal-overlay').forEach(el => {
    el.addEventListener('click', e => {
        if (e.target === el) closeModal(el.id);
    });
});""")

    h2(doc, '15.2 Copier dans le presse-papier')
    add_code_block(doc, """function copyToClipboard(text, buttonEl) {
    navigator.clipboard.writeText(text).then(() => {
        const originalText = buttonEl.textContent;
        buttonEl.textContent = 'Copie !';
        buttonEl.style.background = '#059669';
        buttonEl.style.color = 'white';
        setTimeout(() => {
            buttonEl.textContent = originalText;
            buttonEl.style.background = '';
            buttonEl.style.color = '';
        }, 2000);
    });
}""")

    h2(doc, '15.3 Confirmation de suppression')
    add_code_block(doc, """function confirmDelete(formId, itemName) {
    if (confirm('Supprimer "' + itemName + '" ? Cette action est irreversible.')) {
        document.getElementById(formId).submit();
    }
}""")

    h2(doc, '15.4 Toggle sidebar (admin)')
    add_code_block(doc, """function toggleSidebar() {
    const sidebar = document.getElementById('sidebar');
    const content = document.getElementById('main-content');
    sidebar.classList.toggle('collapsed');  // 260px <-> 68px
    content.classList.toggle('sidebar-collapsed');
}""")

    h2(doc, '15.5 Formulaire multi-etapes (quotas croises)')
    para(doc, 'Le formulaire de creation de quotas croises comporte 2 etapes gerees par JavaScript :')
    add_code_block(doc, """// Etape 1: Selection des segmentations (checkboxes)
// Etape 2: Saisie des pourcentages par combinaison
// Bouton "Repartition equitable": 100 / nb_combinaisons pour chaque input

function goToStep(step) {
    document.querySelectorAll('.step').forEach(el => el.classList.add('hidden'));
    document.getElementById('step-' + step).classList.remove('hidden');
}

function equalDistribution() {
    const inputs = document.querySelectorAll('.percentage-input');
    const equalPct = (100 / inputs.length).toFixed(2);
    inputs.forEach(input => input.value = equalPct);
}""")

    doc.add_page_break()


def build_part17_emails(doc):
    h1(doc, '16. Emails - Templates HTML exacts')

    h2(doc, '16.1 Style commun a tous les emails')
    add_code_block(doc, """body {
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    line-height: 1.6;
    color: #333;
}
.container { max-width: 480px; margin: 0 auto; padding: 40px 20px; }
.button {
    display: inline-block; background: #059669; color: white;
    padding: 14px 28px; border-radius: 8px;
    text-decoration: none; font-weight: 600; margin: 24px 0;
}
.info { background: #F3F4F6; border-radius: 12px; padding: 20px; margin: 24px 0; }
.warning {
    background: #FEF3C7; border-radius: 8px;
    padding: 12px 16px; margin: 16px 0; color: #92400E; font-size: 14px;
}
.footer {
    color: #9CA3AF; font-size: 12px; text-align: center;
    margin-top: 40px; padding-top: 20px; border-top: 1px solid #E5E7EB;
}""")

    h2(doc, '16.2 Email OTP (send_otp_email)')
    para(doc, 'Sujet : "Votre code de verification : {code}" (code inclus dans le sujet !)')
    para(doc, 'Contenu :')
    bullet(doc, 'Titre h2 centre : "Code de verification"')
    bullet(doc, 'Greeting : "Bonjour {prenom}," si prenom, sinon "Bonjour,"')
    bullet(doc, '"Voici votre code de verification :"')
    bullet(doc, 'Code dans boite grise : font-size 32px, font-weight bold, letter-spacing 8px, color #059669, font-family monospace')
    bullet(doc, 'Sous le code : "Ce code expire dans {OTP_EXPIRE_MINUTES} minutes" en gris 14px')
    bullet(doc, '"Si vous n\'avez pas demande ce code, vous pouvez ignorer cet email."')
    bullet(doc, 'Footer : "Marketym - Plateforme de suivi d\'enquetes"')
    para(doc, 'Version texte incluse.')

    h2(doc, '16.3 Email de bienvenue (send_welcome_email)')
    para(doc, 'Sujet : "Bienvenue sur Marketym - Activez votre compte"')
    para(doc, 'Contenu :')
    bullet(doc, 'Titre h2 centre : "Bienvenue sur Marketym !"')
    bullet(doc, '"Bonjour {prenom},"')
    bullet(doc, '"Votre compte a ete cree avec succes. Pour activer votre compte, cliquez sur le bouton ci-dessous pour definir votre mot de passe :"')
    bullet(doc, 'Bouton vert centre : "Activer mon compte" -> {activation_link}')
    bullet(doc, 'Boite info grise : "Votre email de connexion :" + email en monospace')
    bullet(doc, 'Warning jaune : "Important : Ce lien expire dans 48 heures."')
    bullet(doc, 'Texte petit : lien en clair sous le bouton (si le bouton ne fonctionne pas)')
    bullet(doc, 'Footer')
    para(doc, 'Version texte incluse.')

    h2(doc, '16.4 Email reset password (send_password_reset_email)')
    para(doc, 'Sujet : "Reinitialisation de votre mot de passe Marketym"')
    para(doc, 'Legacy (plus utilise en production). Affiche le nouveau mot de passe en 24px bold monospace vert dans une boite grise.')

    doc.add_page_break()


def build_part18_errors(doc):
    h1(doc, '17. Messages d\'erreur et codes HTTP')

    h2(doc, '17.1 Authentification')
    add_table(doc,
        ['Code', 'Message exact', 'Contexte'],
        [
            ['400', 'Cet email est deja utilise', 'Register, update profil'],
            ['400', 'Le mot de passe doit contenir au moins 8 caracteres', 'Register, reset, change'],
            ['400', 'Le nouveau mot de passe doit etre different de l\'ancien', 'Change password'],
            ['400', 'Aucun code en attente', 'Verify OTP sans code actif'],
            ['400', 'Code expire', 'OTP expire'],
            ['400', 'Trop de tentatives. Demandez un nouveau code.', 'OTP >= 3 echecs'],
            ['400', 'Code incorrect. {N} tentative(s) restante(s).', 'OTP mauvais code'],
            ['400', 'Ce compte est deja active', 'Send invitation deja configure'],
            ['400', 'Lien d\'activation invalide ou expire', 'Setup password token invalide'],
            ['400', 'Lien d\'activation expire. Contactez l\'administrateur.', 'Token expire'],
            ['400', 'Aucune modification a effectuer', 'Update profil sans changement'],
            ['400', 'Reseau mobile invalide', 'Valeur hors choix'],
            ['400', 'Mode de remuneration invalide', 'Valeur hors choix'],
            ['401', 'Email ou mot de passe incorrect', 'Login mauvais email ou mdp'],
            ['401', 'Mot de passe non configure', 'Login sans mot de passe'],
            ['401', 'Mot de passe actuel incorrect', 'Change password'],
            ['401', 'Token manquant', 'Header Authorization absent'],
            ['401', 'Token invalide ou expire', 'JWT invalide'],
            ['403', 'Compte desactive', 'Login compte inactif'],
            ['403', 'Acces reserve aux administrateurs', 'Non-admin sur route admin'],
            ['403', 'Acces reserve aux super administrateurs', 'Non-super sur route super'],
            ['404', 'Aucun compte associe a cet email', 'Request OTP email inexistant'],
            ['404', 'Enqueteur introuvable', 'ID invalide'],
            ['429', 'Trop de demandes. Veuillez attendre 15 minutes.', 'Rate limit OTP'],
        ])

    h2(doc, '17.2 Enquetes et affectations')
    add_table(doc,
        ['Code', 'Message', 'Contexte'],
        [
            ['400', 'Cette enquete existe deja', 'survey_id en double'],
            ['400', 'Role invalide. Valeurs autorisees: enqueteur, admin, super_admin', 'Changement role'],
            ['404', 'Enqueteur introuvable', ''],
            ['404', 'Enquete introuvable', ''],
            ['404', 'Survey QuestionPro introuvable', 'fetch_survey_stats retourne None'],
            ['404', 'Lien invalide', 'Track clic affectation inexistante'],
            ['404', 'Aucun lien de questionnaire configure', 'Track clic sans lien'],
        ])

    doc.add_page_break()


def build_part19_initial_data(doc):
    h1(doc, '18. Donnees initiales et migration')

    h2(doc, '18.1 Zones')
    add_table(doc,
        ['Code', 'Nom'],
        [
            ['UEMOA', 'Union Economique et Monetaire Ouest Africaine'],
            ['CEMAC', 'Communaute Economique et Monetaire de l\'Afrique Centrale'],
        ])

    h2(doc, '18.2 Pays (15 pays)')
    add_table(doc,
        ['Zone', 'Code', 'Pays', 'Quota', 'ICP %'],
        [
            ['UEMOA', 'CI', 'Cote d\'Ivoire', '29', '14.36'],
            ['UEMOA', 'SN', 'Senegal', '18', '9.23'],
            ['UEMOA', 'ML', 'Mali', '17', '8.34'],
            ['UEMOA', 'BF', 'Burkina Faso', '15', '7.33'],
            ['UEMOA', 'NE', 'Niger', '16', '7.87'],
            ['UEMOA', 'TG', 'Togo', '13', '6.41'],
            ['UEMOA', 'BJ', 'Benin', '16', '7.91'],
            ['UEMOA', 'GW', 'Guinee-Bissau', '8', '4.14'],
            ['UEMOA', 'MR', 'Mauritanie', '10', '5.00'],
            ['CEMAC', 'CM', 'Cameroun', '22', '11.11'],
            ['CEMAC', 'GA', 'Gabon', '14', '7.18'],
            ['CEMAC', 'CG', 'Congo', '8', '4.08'],
            ['CEMAC', 'TD', 'Tchad', '11', '5.57'],
            ['CEMAC', 'CF', 'RCA', '6', '3.09'],
            ['CEMAC', 'GQ', 'Guinee Equatoriale', '7', '3.38'],
        ])

    h2(doc, '18.3 Compte admin par defaut')
    para(doc, 'Super Admin : wilfredkouadjo006@gmail.com')
    para(doc, 'Email expediteur : marketym@hcexecutive.net')

    h2(doc, '18.4 Migration des donnees existantes')
    para(doc, 'Script de migration pour transferer les donnees de l\'ancienne base Supabase (FastAPI) vers les modeles Django :')
    bullet(doc, 'Enqueteurs : mapper tous les champs, conserver les UUIDs')
    bullet(doc, 'Enquetes : mapper, conserver survey_id et UUIDs')
    bullet(doc, 'Affectations : mapper avec FK vers enqueteurs et enquetes')
    bullet(doc, 'Segmentations : mapper avec FK vers enquetes')
    bullet(doc, 'Quotas : mapper avec FK vers segmentations et affectations')
    bullet(doc, 'CompletionSegments, CompletionPays, HistoriqueCompletions : mapper')
    bullet(doc, 'Clics : mapper')
    bullet(doc, 'OtpCodes, InvitationTokens : mapper')
    para(doc, 'Important : les UUIDs doivent etre preserves pour ne pas casser les liens existants (tracking /r/{id}, etc.).')

    doc.add_page_break()


def build_part20_deployment(doc):
    h1(doc, '19. Configuration deploiement')

    h2(doc, '19.1 Dockerfile')
    add_code_block(doc, """FROM python:3.12-slim

WORKDIR /app
COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt

COPY . .
RUN python manage.py collectstatic --noinput

ENV PORT=8080
EXPOSE 8080

CMD gunicorn marketym.wsgi --bind 0.0.0.0:${PORT:-8080}""")

    h2(doc, '19.2 Settings production')
    add_code_block(doc, """# En production :
DEBUG = False
ALLOWED_HOSTS = [os.environ.get("ALLOWED_HOSTS", "*")]
CSRF_TRUSTED_ORIGINS = [os.environ.get("FRONTEND_URL", "")]

# Static files avec whitenoise
MIDDLEWARE = [
    'django.middleware.security.SecurityMiddleware',
    'whitenoise.middleware.WhiteNoiseMiddleware',  # Apres SecurityMiddleware
    ...
]
STATIC_ROOT = BASE_DIR / 'staticfiles'
STATICFILES_STORAGE = 'whitenoise.storage.CompressedManifestStaticFilesStorage'

# Securite
SECURE_SSL_REDIRECT = True
SESSION_COOKIE_SECURE = True
CSRF_COOKIE_SECURE = True
SECURE_HSTS_SECONDS = 31536000""")

    h2(doc, '19.3 Variables Railway')
    add_table(doc,
        ['Variable', 'Valeur'],
        [
            ['DATABASE_URL', 'Fourni par Supabase'],
            ['SECRET_KEY', 'Generer avec django.core.management.utils.get_random_secret_key()'],
            ['SUPABASE_URL', 'URL du projet'],
            ['SUPABASE_KEY', 'Service role key'],
            ['QUESTIONPRO_API_KEY', 'Cle API'],
            ['BREVO_API_KEY', 'Cle API Brevo'],
            ['EMAIL_FROM', 'marketym@hcexecutive.net'],
            ['FRONTEND_URL', 'URL publique du site'],
            ['ALLOWED_HOSTS', 'Domaine du site'],
            ['DEBUG', 'False'],
        ])

    doc.add_page_break()


def build_part21_implementation(doc):
    h1(doc, '20. Etapes d\'implementation detaillees')

    phases = [
        ('Phase 0 : Setup projet', [
            ('0.1', 'Initialisation Django', 'django-admin startproject marketym. Creer les 8 apps (accounts, enquetes, quotas, sync, tracking, dashboard, exports, notifications). Configurer settings.py : DATABASE_URL, INSTALLED_APPS, MIDDLEWARE, TEMPLATES, STATIC_FILES, AUTH_USER_MODEL. Configurer urls.py principal. Installer dependencies. Creer requirements.txt.'),
        ]),
        ('Phase 1 : Modeles', [
            ('1.1', 'Tous les modeles', 'Definir chaque modele dans son app respective selon le schema section 3. Enqueteur extends AbstractUser avec USERNAME_FIELD="email". Configurer les Meta (unique_together, ordering). makemigrations + migrate. Creer superuser.'),
        ]),
        ('Phase 2 : Authentification', [
            ('2.1', 'Services OTP et Email', 'otp_service.py : generate_otp, hash_code, verify_code, create_otp, verify_otp. email_service.py : configuration Brevo, send_email, send_otp_email, send_welcome_email (avec les templates HTML exacts de la section 16).'),
            ('2.2', 'Vues auth', 'LoginView (avec logique OTP premiere connexion), RegisterView (avec generation identifiant), ActivateView, ForgotPasswordView, ResetPasswordView, ChangePasswordView, ProfileView. Decorateurs admin_required, super_admin_required.'),
            ('2.3', 'Templates auth', 'base.html (TailwindCSS CDN, structure), login.html, register.html, activate.html, forgot_password.html, reset_password.html. Style : page centree, carte blanche avec ombre, input arrondis.'),
        ]),
        ('Phase 3 : Layout et composants', [
            ('3.1', 'Templates de base', 'base.html, base_admin.html (sidebar collapsible 260px/68px), base_enqueteur.html (header avec onglets). Composants includes : stat_card, progress_bar, badge, alert, modal, table, empty_state, bar_chart, copy_button. static/css/main.css, static/js/main.js (modals, clipboard, sidebar toggle, confirmations).'),
        ]),
        ('Phase 4 : Enquetes', [
            ('4.1', 'CRUD enquetes', 'Vues + templates : liste (grille de cartes), create (fetch QP par survey_id), detail (KPI + onglets affectations/segmentations), edit, delete.'),
            ('4.2', 'CRUD enqueteurs', 'Vues + templates : liste (tableau triable avec recherche), create (+ envoi invitation), detail (= dashboard enqueteur), edit, delete. Generation identifiant (adm/usr + seq).'),
            ('4.3', 'CRUD affectations', 'Create (selection enqueteur + enquete, generation liens), edit (objectif, statut, commentaire), delete. Generation lien_direct (QP + custom1) et lien_questionnaire (tracking /r/{id}).'),
        ]),
        ('Phase 5 : Segmentations et quotas', [
            ('5.1', 'Segmentations', 'Create (fetch questions QP, selection, cache answer_options en JSON), edit, delete.'),
            ('5.2', 'Quotas simples', 'Create (simple + bulk), edit pourcentage, delete. Calcul objectif depuis pourcentage.'),
            ('5.3', 'Quotas croises', 'Form multi-etapes (JS steps) : 1) checkboxes segmentations 2) saisie pourcentages par combinaison (produit cartesien). Bouton repartition equitable. Detail avec tableau croise.'),
        ]),
        ('Phase 6 : Synchronisation', [
            ('6.1', 'Client QuestionPro', 'questionpro.py : QuestionProClient avec get_survey, get_questions, get_responses (pagination auto 1000/page), extract_segment_value.'),
            ('6.2', 'Moteur de sync', 'sync_engine.py : sync_affectation (logique complete section 7), sync_enquete, sync_all (semaphore 5). Vues SyncAllView, SyncAffectationView. Template sync_panel.html avec bouton et historique.'),
        ]),
        ('Phase 7 : Tracking', [
            ('7.1', 'TrackClickView', '/r/<id> : get IP (X-Forwarded-For), insert clic (IntegrityError = ignore), update count, redirect 302.'),
        ]),
        ('Phase 8 : Dashboards', [
            ('8.1', 'Dashboard admin', 'KPI globaux (4 cartes), liste enquetes en cours, enqueteurs recents.'),
            ('8.2', 'Detail enquete admin', 'KPI enquete, tableau affectations, graphiques barres CSS par segmentation, tableau croise si quota_configs. Logique excedents globale->individuelle.'),
            ('8.3', 'Dashboard enqueteur', '3 onglets (tableau de bord, mes enquetes, profil). Detail enquete avec KPI, lien copiable, graphiques barres. Profil avec modification OTP.'),
            ('8.4', 'Vue enqueteur dans admin', 'Reutiliser les memes donnees/templates que le dashboard enqueteur mais dans base_admin.html.'),
        ]),
        ('Phase 9 : Exports', [
            ('9.1', 'CSV exports', 'Enqueteurs, affectations par enquete, quotas. Boutons dans pages admin.'),
        ]),
        ('Phase 10 : Stats', [
            ('10.1', 'Pages statistiques', 'Stats par pays (completions_pays), par segment (completions_segments), globales. Graphiques barres CSS.'),
        ]),
        ('Phase 11 : Notifications', [
            ('11.1', 'Base', 'Bouton "Envoyer rappel" par enqueteur. Badge alerte quota atteint.'),
        ]),
        ('Phase 12 : Deploiement', [
            ('12.1', 'Config', 'Dockerfile, gunicorn, whitenoise, collectstatic, variables env Railway. HTTPS, HSTS, CSRF_TRUSTED_ORIGINS.'),
        ]),
        ('Phase 13 : Migration donnees', [
            ('13.1', 'Script migration', 'Management command Django pour migrer les donnees Supabase existantes. Preserver les UUIDs. Migrer : enqueteurs, enquetes, affectations, segmentations, quotas, completions, clics, historique, OTP, invitations.'),
        ]),
    ]

    for phase_title, etapes in phases:
        h2(doc, phase_title)
        for num, titre, desc in etapes:
            bold_para(doc, f'Etape {num} - {titre}')
            para(doc, desc)

    doc.add_page_break()

    # Annexe finale
    h1(doc, 'Annexe - Checklist de verification')
    checks = [
        'Inscription -> login -> OTP -> dashboard enqueteur',
        'Admin cree enqueteur -> email invitation recu -> activation par lien',
        'Admin cree enquete (survey_id QP) -> nom/URL recuperes',
        'Admin cree affectation -> liens generes (tracking + direct)',
        'Admin cree segmentation -> questions QP fetchees, answer_options cachees',
        'Admin cree quotas simples -> pourcentages, objectifs calcules',
        'Admin cree quotas croises -> combinaisons generees, pourcentages saisis',
        'Sync -> completions, segments, combos mis a jour',
        'Clic sur lien tracking -> IP deduplication -> redirect QP',
        'Dashboard admin -> KPI corrects, graphiques barres',
        'Detail enquete admin -> tableau croise, excedents calcules',
        'Dashboard enqueteur -> graphiques barres par segmentation',
        'Admin voit dashboard enqueteur -> meme vue',
        'Export CSV -> fichier telecharge avec bonnes donnees',
        'Mot de passe oublie -> OTP -> reset',
        'Profil -> modification avec OTP -> sauvegarde',
        'Build production -> collectstatic OK, gunicorn OK',
        'HTTPS + CSRF OK en production',
    ]
    for c in checks:
        bullet(doc, c)


def main():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
    doc.styles['Heading 1'].font.size = Pt(20)
    doc.styles['Heading 2'].font.size = Pt(14)
    doc.styles['Heading 3'].font.size = Pt(12)

    # Build all parts
    build_part1_cover_and_toc(doc)
    build_part2_presentation(doc)
    build_part3_architecture(doc)
    build_part4_database(doc)
    build_part5_accounts(doc)
    build_part6_enquetes(doc)
    build_part7_quotas(doc)
    build_part8_sync(doc)
    build_part9_tracking(doc)
    build_part10_dashboard(doc)
    build_part11_exports(doc)
    build_part12_notifications(doc)
    build_part13_questionpro(doc)
    build_part14_business_logic(doc)
    build_part15_templates(doc)
    build_part16_javascript(doc)
    build_part17_emails(doc)
    build_part18_errors(doc)
    build_part19_initial_data(doc)
    build_part20_deployment(doc)
    build_part21_implementation(doc)

    output = '/Users/Apple/Desktop/H&C/Enquetes/Suivi enquetes/enquetes/CAHIER_DE_CHARGE_COMPLET.docx'
    doc.save(output)
    print(f'Document genere : {output}')


if __name__ == '__main__':
    main()
