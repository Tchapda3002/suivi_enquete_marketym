#!/usr/bin/env python3
"""
GETITHERE — Document Strategique
Projet DATATYM™ · Plateforme de Business Intelligence Africaine
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(5)
style.paragraph_format.line_spacing = 1.15

for lv in range(1, 4):
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x0F, 0x1E, 0x3A)
    hs.font.size = Pt([0, 18, 14, 12][lv])

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.font.bold = True; r.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()

def bul(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

def quote(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = True; r.font.italic = True
    r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x0F, 0x1E, 0x3A)

def note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.italic = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


# ════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ════════════════════════════════════════════════════════════════
for _ in range(3): doc.add_paragraph()
for text, sz, bold, col in [
    ('DATATYM\u2122', 28, True, RGBColor(0x0F, 0x1E, 0x3A)),
    ('', 6, False, None),
    ('GETITHERE', 40, True, RGBColor(0xB5, 0x89, 0x1A)),
    ('', 10, False, None),
    ('Document strategique', 18, False, RGBColor(0x6B, 0x72, 0x80)),
    ('', 6, False, None),
    ('La plateforme de Business Intelligence\nqui manque a l\'Afrique', 15, True, RGBColor(0x1A, 0x20, 0x2C)),
    ('', 14, False, None),
    ('H&C Executive — Avril 2026', 12, False, RGBColor(0x6B, 0x72, 0x80)),
    ('Document confidentiel — Usage Direction Generale', 10, False, RGBColor(0x9C, 0xA3, 0xAF)),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(sz); r.font.bold = bold
    if col: r.font.color.rgb = col
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 1. LA THESE
# ════════════════════════════════════════════════════════════════
doc.add_heading('1. La these', level=1)

doc.add_paragraph(
    'Les decideurs africains prennent chaque jour des decisions a plusieurs milliards '
    'de FCFA sans donnees fiables. Pas parce qu\'ils sont incompetents — parce que '
    'l\'infrastructure informationnelle qui permettrait de decider n\'existe pas.'
)
doc.add_paragraph(
    'En Europe ou aux Etats-Unis, un DG qui veut savoir combien coute le m2 dans un quartier, '
    'quelle est la performance financiere de ses concurrents, ou dans quoi son secteur '
    'investit l\'annee prochaine — il ouvre Bloomberg, Statista, ou appelle son consultant '
    'qui a la donnee sous 48h.'
)
doc.add_paragraph(
    'En Afrique francophone, ce meme DG appelle un ami. Envoie un stagiaire faire le tour '
    'des agences. Attend 3 mois un rapport qui sera obsolete a la livraison. Ou — le plus '
    'souvent — il decide sans. A l\'intuition. Au feeling. A l\'experience.'
)

quote(
    'GETITHERE ne vend pas de la donnee.\n'
    'GETITHERE vend ce qu\'aucun tableau Excel, aucun rapport PDF, '
    'aucun consultant ne peut livrer en 30 secondes : '
    'une aide a la decision fondee sur de la donnee terrain fiable.'
)

doc.add_paragraph(
    'GETITHERE est un projet de l\'ecosysteme DATATYM\u2122. '
    'Il herite de l\'ADN DATATYM\u2122 : des indices proprietaires, des doctrines prescriptives, '
    'une posture qui ne decrit pas le marche mais qui dit au dirigeant quoi faire. '
    'DATATYM\u2122 a prouve le modele sur le talent (IPE\u2122, IRTA\u2122, 1 500 repondants, 16+ pays). '
    'GETITHERE l\'etend a toutes les decisions strategiques.'
)


# ════════════════════════════════════════════════════════════════
# 2. LE PROBLEME
# ════════════════════════════════════════════════════════════════
doc.add_heading('2. Le probleme', level=1)

doc.add_heading('2.1 Les decisions impossibles', level=2)
doc.add_paragraph(
    'Le probleme n\'est pas le manque de donnees. C\'est le manque de decisions fondees. '
    'Chaque jour, des DG, des investisseurs, des DAF africains font face a des questions '
    'strategiques pour lesquelles aucune reponse fiable n\'est disponible :'
)

tbl(
    ['La decision', 'Ce qu\'il faudrait pour decider', 'Ce qui existe aujourd\'hui'],
    [
        ['Je rachete cette entreprise', 'CA reel, effectifs, charges, tendance 5 ans',
         'Un PDF de 2019 ou un ami qui connait le DG'],
        ['J\'ouvre un bureau a Douala', 'Prix m2, cout salarial, bassin de talents',
         'Des estimations verbales d\'un agent'],
        ['J\'alloue mon budget IT 2027', 'Ce que font les autres dans mon secteur',
         'Son intuition + le budget de l\'an dernier +5%'],
        ['Je retiens mes meilleurs talents', 'Risque de depart, drivers reels',
         'Rien — jusqu\'a DATATYM\u2122 Talent'],
        ['J\'entre sur le marche senegalais', 'Taille marche, concurrents, reglementation',
         'Un rapport Deloitte a 30 000$ qui parle de "l\'Afrique de l\'Ouest"'],
        ['Je negocie un bail commercial', 'Prix de reference du quartier, tendance',
         'Le prix que le proprietaire annonce'],
        ['Je dimensionne mon equipe', 'Effectifs comparables dans mon secteur/pays',
         'Aucune donnee comparative'],
    ]
)

doc.add_heading('2.2 Le cout de l\'aveuglement', level=2)
doc.add_paragraph(
    'Ce brouillard informationnel a un cout enorme — mais invisible, '
    'parce que personne ne le mesure :'
)
bul([
    'Investissements mal calibres : des millions engages sur des estimations',
    'Talents perdus : faute de diagnostic, les meilleurs partent sans qu\'on comprenne pourquoi',
    'Opportunites ratees : des marches non explores parce que la donnee n\'etait pas la',
    'Budgets au doigt mouille : des allocations basees sur l\'habitude, pas sur l\'evidence',
    'Negociations perdues : sans reference de marche, c\'est le vendeur qui fixe le prix',
    'Retard strategique : pendant que les decisions trainent, les concurrents agissent',
])

quote(
    'En Afrique francophone, la donnee n\'est pas un luxe. '
    'C\'est l\'infrastructure manquante de la decision. '
    'GETITHERE construit cette infrastructure.'
)

doc.add_heading('2.3 Pourquoi le probleme persiste', level=2)
doc.add_paragraph(
    'Ce n\'est pas que la donnee n\'existe pas du tout. '
    'C\'est qu\'elle est dans un etat qui la rend inutilisable pour decider :'
)
tbl(
    ['Etat de la donnee', 'Consequence pour le decideur'],
    [
        ['Dispersee (PDF, sites regulateurs, tableurs informels)',
         'Il faudrait 3 semaines pour rassembler ce qu\'il faut'],
        ['Non structuree (pas de format standard, pas de series)',
         'Impossible de comparer, de tracer une tendance'],
        ['Inaccessible (murs payants occidentaux ou introuvable)',
         'Payer 25 000$/an pour Bloomberg ou ne rien avoir'],
        ['Non fiable (estimations, donnees obsoletes)',
         'Decider sur une donnee fausse est pire que decider sans'],
        ['Non locale (rapports qui traitent "l\'Afrique" comme un bloc)',
         '"L\'Afrique de l\'Ouest" ne dit rien sur Abidjan vs Cotonou'],
    ]
)

doc.add_heading('2.4 Le vrai concurrent de GETITHERE', level=2)
doc.add_paragraph(
    'Le concurrent principal de GETITHERE n\'est pas Bloomberg, ni Statista, ni Deloitte. '
    'Le concurrent principal, c\'est l\'appel a un ami. C\'est le reseau personnel. '
    'C\'est le WhatsApp envoye a un contact qui "connait quelqu\'un dans le secteur". '
    'GETITHERE ne remplace pas Bloomberg — Bloomberg ne sert pas ce marche. '
    'GETITHERE remplace le systeme D.'
)


# ════════════════════════════════════════════════════════════════
# 3. LA SOLUTION
# ════════════════════════════════════════════════════════════════
doc.add_heading('3. La solution', level=1)

doc.add_heading('3.1 Ce que GETITHERE fait', level=2)
doc.add_paragraph(
    'GETITHERE est une plateforme de business intelligence qui permet aux decideurs '
    'africains de prendre des decisions fondees sur de la donnee terrain fiable. '
    'Pas un moteur de recherche de donnees — un outil de decision.'
)

doc.add_paragraph('Le parcours utilisateur est simple :')
bul([
    'Un DG a une question strategique',
    'Il va sur datatym.ai',
    'Il trouve en 30 secondes la donnee structuree, fiable, locale dont il a besoin',
    'Il comprend la donnee (indices, benchmarks, tendances, alertes)',
    'Il decide — informe, pas a l\'aveugle',
])

doc.add_heading('3.2 Ce qui rend GETITHERE different', level=2)
tbl(
    ['Dimension', 'Les autres', 'GETITHERE'],
    [
        ['Source des donnees', 'Agregation, scraping, estimation',
         'Collecte terrain proprietaire + partenariats locaux + donnees publiques structurees'],
        ['Granularite', '"L\'Afrique de l\'Ouest"',
         'Par pays, par ville, par quartier, par secteur, par entreprise'],
        ['Posture', 'Decrire le marche',
         'Dire au decideur quoi faire (ADN DATATYM\u2122)'],
        ['Prix', '25 000$/an (Bloomberg) ou 30 000$/rapport (Big 4)',
         'Freemium + abonnements accessibles'],
        ['Fiabilite', 'Estimations, donnees secondaires',
         'Sources primaires, methodologie transparente, indices proprietaires'],
        ['Recurrence', 'Rapport ponctuel',
         'Plateforme vivante, donnees mises a jour, alertes automatiques'],
    ]
)

doc.add_heading('3.3 Les verticales', level=2)
doc.add_paragraph(
    'GETITHERE est concu pour accueillir n\'importe quelle verticale ou la donnee '
    'a de la valeur pour une decision. Les premieres verticales :'
)

tbl(
    ['Verticale', 'La decision qu\'elle sert', 'Donnee cle', 'Source'],
    [
        ['Finance &\nPerformance', 'Faut-il investir dans ce secteur ?\nQuelle est la sante de mes concurrents ?',
         'CA, effectifs, charges, ratios\ndes banques, IMF, telcos, miniers',
         'Rapports de gestion publics,\nregulateurs (BCEAO, BEAC, ARTP)'],
        ['Immobilier', 'Combien ca coute ici ?\nOu installer mon bureau ?',
         'Prix m2 locatif et acquisitif\npar ville, quartier, type de bien',
         'Agences immobilieres partenaires,\nnotaires, transactions reelles'],
        ['Budgets &\nDepenses', 'Dans quoi investir ?\nSuis-je aligne avec mon secteur ?',
         'Budgets previsionnels par secteur,\npostes de depense, tendances',
         'DAF, controleurs de gestion\n(enquetes structurees)'],
        ['Talent\n(existe deja)', 'Vais-je perdre mes meilleurs ?\nMon management est-il un risque ?',
         'Risque turnover, performance\nemployeur, attractivite, NPS',
         'Barometres terrain DATATYM\u2122,\nenqueteurs'],
    ]
)

note(
    'Chaque verticale est independante mais interconnectee : un DG qui consulte les donnees '
    'Finance decouvre les donnees Talent. Un investisseur qui regarde l\'Immobilier '
    'voit aussi la performance des entreprises du quartier. La plateforme cree des ponts.'
)

doc.add_heading('3.4 Les produits', level=2)
tbl(
    ['Produit', 'Pour qui', 'Ce qu\'il contient', 'Modele tarifaire'],
    [
        ['GETITHERE Free', 'Tout le monde',
         'Donnees agregees par pays/secteur,\nindicateurs macro, tendances',
         'Gratuit (acquisition + credibilite)'],
        ['GETITHERE Pro', 'DG, DAF, consultants',
         'Donnees detaillees, historiques 5 ans,\nbenchmarks, export Excel/PDF',
         'Abonnement 50-200 $/mois'],
        ['GETITHERE Report', 'Investisseurs, CODIR',
         'Rapport sectoriel sur mesure\navec analyse + recommandations DATATYM\u2122',
         '2 000-10 000 $ par rapport'],
        ['GETITHERE Enterprise', 'Banques, institutions, grandes entreprises',
         'API complete, donnees temps reel,\ndashboard dedie, alertes',
         'Abonnement 500-2 000 $/mois'],
        ['GETITHERE Index', 'Marche (publication)',
         'Indices proprietaires DATATYM\u2122\npar verticale (IPFI\u2122, IPIM\u2122...)',
         'Publication trimestrielle\n(inclus dans Pro/Enterprise)'],
    ]
)


# ════════════════════════════════════════════════════════════════
# 4. ARCHITECTURE DATATYM
# ════════════════════════════════════════════════════════════════
doc.add_heading('4. GETITHERE dans l\'ecosysteme DATATYM\u2122', level=1)

doc.add_paragraph(
    'DATATYM\u2122 est la marque mere — l\'autorite. '
    'GETITHERE est un de ses projets — le vehicule qui rend la donnee '
    'DATATYM\u2122 (et au-dela) exploitable pour le marche.'
)

tbl(
    ['Niveau', 'Composante', 'Role'],
    [
        ['Autorite', 'DATATYM\u2122', 'La marque. La credibilite. Les indices. Les doctrines.'],
        ['Projet 1', 'DATATYM\u2122 Talent', 'Barometre RH — IPE\u2122/IRTA\u2122/ILUX\u2122/IGRO\u2122'],
        ['Projet 2', 'GETITHERE', 'Plateforme BI multi-verticales'],
        ['Projet N', '(futurs)', 'Sante, Education, Agriculture...'],
        ['Plateforme', 'datatym.ai', 'Point d\'acces unique — tous les projets, une seule porte'],
    ]
)

doc.add_paragraph(
    'L\'avantage strategique : chaque projet renforce la marque mere. '
    'La credibilite gagnee sur le Talent profite a la Finance. '
    'La donnee Immobilier rend la plateforme indispensable pour les investisseurs '
    'qui decouvrent ensuite le Talent. L\'ecosysteme se nourrit de lui-meme.'
)

doc.add_heading('4.1 Le principe DATATYM\u2122 applique a chaque verticale', level=2)
doc.add_paragraph(
    'Chaque verticale GETITHERE herite du systeme DATATYM\u2122 :'
)
tbl(
    ['Pilier DATATYM\u2122', 'Application GETITHERE'],
    [
        ['I. Indices', 'Chaque verticale a son indice proprietaire (IPFI\u2122, IPIM\u2122, IBUD\u2122...)'],
        ['II. Doctrines', 'Chaque verticale produit ses verites prescriptives fondees sur les donnees'],
        ['III. Rituels', 'Publication trimestrielle, alertes, brief sectoriel'],
        ['IV. Decisions', 'Chaque donnee mene a une decision — pas a un constat'],
    ]
)


# ════════════════════════════════════════════════════════════════
# 5. LE MARCHE
# ════════════════════════════════════════════════════════════════
doc.add_heading('5. Le marche', level=1)

doc.add_heading('5.1 Taille et potentiel', level=2)
doc.add_paragraph(
    'Le marche de la business intelligence et de la donnee economique en Afrique '
    'est estime entre 300 et 800 M$ annuels (donnees financieres, immobilieres, '
    'rapports sectoriels, conseil data). Ce marche est domine par des acteurs '
    'internationaux qui ne servent pas — ou mal — le segment francophone.'
)
doc.add_paragraph(
    'GETITHERE ne cible pas l\'ensemble de ce marche. Il cible le segment '
    '"donnees structurees accessibles en self-service pour decideurs africains" — '
    'un segment ou il n\'y a aujourd\'hui aucun acteur dominant.'
)

doc.add_heading('5.2 Les clients', level=2)
tbl(
    ['Segment', 'Besoin', 'Pourquoi il paye', 'Taille estimee'],
    [
        ['DG / CODIR', 'Piloter avec de la donnee',
         'Eviter les decisions a l\'aveugle', '5 000+ entreprises (top 10 pays)'],
        ['Investisseurs (PE, VC, DFI)', 'Due diligence, sizing',
         'Valider une these d\'investissement', '200+ fonds actifs'],
        ['Banques / Institutions', 'Benchmark, risk scoring',
         'Reduire le risque credit / investissement', '100+ banques UEMOA/CEMAC'],
        ['Cabinets conseil locaux', 'Donnees pour missions',
         'Produire plus vite, facturer plus', '500+ cabinets'],
        ['Organisations internationales', 'Donnees pour programmes',
         'Justifier les allocations', 'BM, BAD, AFD, UE...'],
        ['Gouvernements / Regulateurs', 'Tableaux de bord',
         'Piloter les politiques publiques', '16+ pays'],
        ['Promoteurs immobiliers', 'Prix de reference',
         'Calibrer les investissements', '1 000+ operateurs'],
    ]
)

doc.add_heading('5.3 La concurrence', level=2)
tbl(
    ['', 'Bloomberg', 'Statista', 'Big 4\n(Deloitte/PwC)', 'GETITHERE'],
    [
        ['Donnees proprietaires', 'Oui', 'Non (agregation)', 'Partiellement', 'Oui (terrain)'],
        ['Afrique francophone', 'Marginal', 'Marginal', 'Oui mais cher', 'Coeur de cible'],
        ['Granularite locale', 'Non', 'Non', 'Non', 'Oui (ville/quartier)'],
        ['Prix accessible', 'Non\n(25K$/an)', 'Non\n(1K$/mois)', 'Non\n(30K$/mandat)', 'Oui\n(freemium)'],
        ['Aide a la decision', 'Terminal', 'Compilation', 'Conseil', 'Prescription\n(ADN DATATYM)'],
        ['Mise a jour', 'Temps reel', 'Annuelle', 'Ponctuelle', 'Continue\n(reseau terrain)'],
    ]
)

quote(
    'Le concurrent de GETITHERE n\'est pas Bloomberg. '
    'C\'est l\'appel a un ami. C\'est le systeme D. '
    'C\'est le WhatsApp envoye a un contact qui "connait quelqu\'un". '
    'GETITHERE professionnalise ce qui se fait de maniere informelle '
    'depuis toujours — et le rend scalable.'
)


# ════════════════════════════════════════════════════════════════
# 6. LE MODELE ECONOMIQUE
# ════════════════════════════════════════════════════════════════
doc.add_heading('6. Le modele economique', level=1)

doc.add_heading('6.1 Comment GETITHERE gagne de l\'argent', level=2)
tbl(
    ['Source de revenus', 'Description', 'Part a maturite'],
    [
        ['Abonnements\n(Pro + Enterprise)', 'Acces mensuel/annuel aux donnees detaillees,\nbenchmarks, historiques, exports, API',
         '40-50%'],
        ['Rapports sur mesure', 'Analyses sectorielles commandees\navec recommandations DATATYM\u2122',
         '20-25%'],
        ['API & Data Licensing', 'Acces programmatique pour integration\ndans les SI des banques, fonds, institutions',
         '15-20%'],
        ['Indices & Rankings', 'Publication de classements (Top 50 Attractivite,\nRanking sectoriel) — visibilite sponsorisee',
         '5-10%'],
        ['Publicite premium', 'Visibilite pour entreprises sur la plateforme\n(recrutement, marque employeur, vitrine)',
         '5-10%'],
    ]
)

doc.add_heading('6.2 Comment GETITHERE depense', level=2)
tbl(
    ['Poste de cout', 'Description', 'Part des charges'],
    [
        ['Collecte terrain', 'Enqueteurs, partenariats, deplacements,\nincentives fournisseurs de donnees', '35-40%'],
        ['Equipe', 'Analystes, data engineers, commerciaux,\nmanagement projet', '25-30%'],
        ['Technologie', 'Plateforme web, hebergement, API, securite,\nmaintenance', '20-25%'],
        ['Marketing &\nCommercial', 'Acquisition clients, evenements,\npartenariats strategiques', '10-15%'],
    ]
)

doc.add_heading('6.3 Le cercle vertueux', level=2)
doc.add_paragraph(
    'GETITHERE cree de la valeur par un mecanisme auto-renforçant :'
)
bul([
    'Plus de donnees collectees -> plus de valeur pour les utilisateurs -> plus d\'abonnes',
    'Plus d\'abonnes -> plus de revenus -> plus de capacite de collecte -> plus de donnees',
    'Plus de verticales -> plus de raisons de venir -> plus d\'utilisateurs',
    'Les fournisseurs de donnees (agences, regulateurs) gagnent en visibilite -> ils contribuent volontairement',
    'Les indices DATATYM\u2122 creent de la dependance — un DG qui pilote avec l\'IRTA\u2122 ne revient pas en arriere',
    'Chaque nouveau client genere des donnees (comportement, besoins) qui enrichissent la plateforme',
])

doc.add_heading('6.4 Monetisation immediate vs long terme', level=2)
tbl(
    ['Horizon', 'Source principale', 'Logique'],
    [
        ['Immediat (0-6 mois)', 'Rapports sur mesure',
         'Revenus de consulting — capitalise sur le savoir-faire H&C existant.\nPas besoin d\'attendre la plateforme.'],
        ['Court terme (6-18 mois)', 'Abonnements Pro',
         'La plateforme attire les premiers utilisateurs.\nConversion free -> payant.'],
        ['Moyen terme (18-36 mois)', 'Enterprise + API',
         'Les institutions integrent GETITHERE dans leurs processus.\nRevenus recurrents a forte valeur.'],
        ['Long terme (36+ mois)', 'Indices + Rankings',
         'GETITHERE devient la reference citee.\nLes entreprises payent pour y figurer.'],
    ]
)


# ════════════════════════════════════════════════════════════════
# 7. LA COLLECTE
# ════════════════════════════════════════════════════════════════
doc.add_heading('7. Le modele de collecte — Le coeur de la valeur', level=1)

doc.add_paragraph(
    'La collecte est ce qui rend GETITHERE inreplicable. '
    'N\'importe qui peut construire un site web. Personne ne peut repliquer '
    'un reseau de collecte terrain dans 16 pays africains. '
    'C\'est le moat — la barriere a l\'entree.'
)

doc.add_heading('7.1 Trois canaux complementaires', level=2)

doc.add_heading('Canal 1 — Donnees publiques structurees', level=3)
doc.add_paragraph('Effort faible. Impact immediat. Premiere source a activer.')
tbl(
    ['Source', 'Donnees', 'Couverture'],
    [
        ['BCEAO / BEAC', 'Bilans bancaires, donnees monetaires', '8 pays UEMOA + 6 CEMAC'],
        ['ARTP / ARCEP / regulateurs telecoms', 'Penetration mobile, CA operateurs', 'Par pays'],
        ['Rapports annuels (entreprises cotees)', 'CA, effectifs, charges', 'BRVM, DSX, etc.'],
        ['Instituts statistiques nationaux', 'PIB, inflation, demographie', '16+ pays'],
        ['Banques centrales', 'Taux, liquidite, credit', 'UEMOA + CEMAC'],
    ]
)

doc.add_heading('Canal 2 — Partenariats de donnees', level=3)
doc.add_paragraph('Effort moyen. Valeur elevee. Cree la differenciation.')
tbl(
    ['Partenaire', 'Ce qu\'il apporte', 'Ce qu\'il recoit en echange'],
    [
        ['Agences immobilieres', 'Prix reels, transactions, tendances',
         'Visibilite sur datatym.ai, leads clients, indice de reference'],
        ['Cabinets comptables', 'Donnees financieres PME (anonymisees)',
         'Acces gratuit aux benchmarks — valeur pour leurs clients'],
        ['Associations professionnelles', 'Donnees de branche, effectifs, tendances',
         'Rapport sectoriel annuel offert — leur donne de la valeur'],
        ['Universites / Ecoles', 'Donnees emploi, insertion, salaires',
         'Co-branding recherche, acces donnees pour publications'],
    ]
)

doc.add_heading('Canal 3 — Collecte terrain proprietaire', level=3)
doc.add_paragraph('Effort eleve. Valeur maximale. C\'est la signature DATATYM\u2122.')
tbl(
    ['Enquete', 'Cible', 'Frequence', 'Methode'],
    [
        ['Barometre Talent', 'Salaries (toutes generations)', 'Semestriel', 'QuestionPro + enqueteurs'],
        ['Barometre Budget', 'DAF / Controleurs de gestion', 'Annuel', 'Enquete ciblee + interviews'],
        ['Barometre Immobilier', 'Agences + notaires + terrain', 'Trimestriel', 'Collecte terrain + partenariats'],
        ['Barometre Sectoriel', 'DG / Directeurs par secteur', 'Annuel', 'Enquete + donnees publiques'],
    ]
)

quote(
    'GETITHERE va chercher la donnee la ou personne ne va — '
    'sur le terrain, dans les regulateurs, dans les agences, dans les entreprises — '
    'et la transforme en decision.'
)


# ════════════════════════════════════════════════════════════════
# 8. LA PLATEFORME
# ════════════════════════════════════════════════════════════════
doc.add_heading('8. La plateforme — datatym.ai', level=1)

doc.add_heading('8.1 Experience utilisateur', level=2)
doc.add_paragraph(
    'L\'utilisateur arrive sur datatym.ai avec une question. '
    'Il repart avec une reponse exploitable. Le parcours :'
)
bul([
    'Recherche : par pays, secteur, indicateur, ou question en langage naturel',
    'Consultation : donnees structurees, graphiques, tendances, comparatifs',
    'Comprehension : indices DATATYM\u2122, benchmarks, alertes, interpretation',
    'Action : recommandations, export, integration API, prise de decision',
])

doc.add_heading('8.2 Stack technique', level=2)
tbl(
    ['Composant', 'Technologie', 'Raison'],
    [
        ['Frontend', 'React / Next.js', 'SEO + performance + savoir-faire existant'],
        ['Backend', 'FastAPI (Python)', 'Stack Marketym reutilisable'],
        ['Base de donnees', 'PostgreSQL', 'Supabase existant, requetes complexes'],
        ['Recherche', 'Meilisearch', 'Recherche rapide multi-verticales'],
        ['Visualisation', 'Charts.js / Plotly', 'Dashboards interactifs'],
        ['API', 'REST + GraphQL', 'Integration clients Enterprise'],
    ]
)

doc.add_heading('8.3 Reutilisation de Marketym', level=2)
doc.add_paragraph(
    'L\'infrastructure Marketym (auth, API, base de donnees, sync) sert de fondation. '
    'GETITHERE n\'est pas un projet from scratch — c\'est l\'evolution naturelle de '
    'l\'outil de collecte vers un outil de diffusion et de decision.'
)


# ════════════════════════════════════════════════════════════════
# 9. FAISABILITE & RISQUES
# ════════════════════════════════════════════════════════════════
doc.add_heading('9. Faisabilite et risques', level=1)

doc.add_heading('9.1 Ce qui est faisable maintenant', level=2)
bul([
    'Verticale Talent : operationnelle, indices calcules, barometres publies',
    'Verticale Finance : rapports publics accessibles (BCEAO, rapports annuels)',
    'Plateforme MVP : stack Marketym reutilisable, equipe technique en place',
    'Marque DATATYM\u2122 : positionnee, premiers livrables reconnus',
    'Premiers rapports sur mesure : vendables immediatement (monetisation J+0)',
])

doc.add_heading('9.2 Ce qui necessite du temps', level=2)
bul([
    'Partenariats agences immobilieres : negociation, confiance (6-12 mois)',
    'Reseau de DAF pour le Barometre Budget : acces, credibilite (12+ mois)',
    'Couverture multi-pays : deploiement pays par pays, pas en bloc',
    'API Enterprise : developpement technique (6 mois apres MVP)',
])

doc.add_heading('9.3 Risques', level=2)
tbl(
    ['Risque', 'Probabilite', 'Impact', 'Comment on le gere'],
    [
        ['Dispersion\n(trop de verticales trop tot)', 'Elevee', 'Critique',
         'Regle : 1 nouvelle verticale tous les 6 mois, pas avant.\nSequencer est non-negociable.'],
        ['Qualite donnees\ninsuffisante', 'Moyenne', 'Critique',
         'Sources multiples, processus de validation,\ntransparence methodologique.\nMieux vaut moins de donnees fiables que beaucoup de bruit.'],
        ['Adoption lente\n(les gens ne viennent pas)', 'Moyenne', 'Eleve',
         'Contenu gratuit attractif en porte d\'entree.\nPartenariats institutionnels.\nEvenements (DATATYM Summit).'],
        ['Collecte terrain\ndifficile', 'Elevee', 'Eleve',
         'Commencer par les donnees publiques (Canal 1).\nAjouter les partenariats (Canal 2).\nTerrain proprietaire (Canal 3) en dernier.'],
        ['Un concurrent\ns\'installe avant', 'Faible\n(court terme)', 'Moyen',
         'Avantage terrain local + vitesse d\'execution.\nLe reseau de collecte est un moat.'],
        ['Financement\ninsuffisant', 'Moyenne', 'Critique',
         'MVP lean (reutilisation Marketym).\nMonetisation rapide par les rapports.\nAutofinancement avant croissance.'],
    ]
)


# ════════════════════════════════════════════════════════════════
# 10. PLAN DE DEPLOIEMENT
# ════════════════════════════════════════════════════════════════
doc.add_heading('10. Plan de deploiement', level=1)

tbl(
    ['Phase', 'Periode', 'Objectifs cles', 'Verticales', 'Livrable'],
    [
        ['Phase 0\nFondation\n(actuel)', 'Jan-Avr 2026',
         'DATATYM Talent operationnel\nReferentiel V3 livre\nMarketym en production\nPremiers rapports sur mesure',
         'Talent', 'Referentiel V3\nBarometres Gen Z\nPlateforme Marketym'],
        ['Phase 1\nMVP', 'Mai-Oct 2026',
         'datatym.ai en ligne (MVP)\nVerticale Finance lancee (3 pays)\nPremiers abonnes Pro\n100 utilisateurs inscrits\n5 rapports sur mesure vendus',
         'Talent\n+ Finance', 'datatym.ai MVP\nPremiers indices Finance\nBase donnees 3 pays'],
        ['Phase 2\nCroissance', 'Nov 2026-Avr 2027',
         'Verticale Immobilier (5 villes)\nAPI Enterprise disponible\n500 utilisateurs\nPartenariats agences signes\n20+ rapports vendus',
         'Talent\n+ Finance\n+ Immobilier', 'Indices IPIM\u2122\nAPI publique\nPremier DATATYM Summit'],
        ['Phase 3\nScale', 'Mai-Dec 2027',
         'Verticale Budget\n10 pays couverts\n2 000 utilisateurs\nRentabilite operationnelle\nTop 50 Attractivite publie',
         'Toutes\nverticales', 'Classement public\nApp mobile\nRentabilite'],
    ]
)

note(
    'Regle de gouvernance : aucune nouvelle verticale n\'est lancee tant que la precedente '
    'n\'a pas atteint ses objectifs Phase 1 (MVP + premiers utilisateurs + premiere monetisation).'
)


# ════════════════════════════════════════════════════════════════
# 11. PROJECTIONS FINANCIERES
# ════════════════════════════════════════════════════════════════
doc.add_heading('11. Projections financieres', level=1)

note('Projections indicatives sur 3 ans. Hypotheses conservatrices. Marche : Afrique francophone.')

tbl(
    ['', 'Annee 1\n(S2 2026)', 'Annee 2\n(2027)', 'Annee 3\n(2028)'],
    [
        ['Utilisateurs inscrits', '200', '2 000', '8 000'],
        ['Abonnes payants (Pro+Enterprise)', '10', '150', '600'],
        ['Rapports sur mesure vendus', '5', '30', '80'],
        ['', '', '', ''],
        ['Revenus abonnements', '6 K$', '108 K$', '432 K$'],
        ['Revenus rapports', '25 K$', '150 K$', '400 K$'],
        ['Revenus API/Enterprise', '-', '36 K$', '180 K$'],
        ['Autres (indices, pub)', '-', '-', '50 K$'],
        ['TOTAL REVENUS', '31 K$', '294 K$', '1 062 K$'],
        ['', '', '', ''],
        ['Couts operationnels', '80 K$', '200 K$', '450 K$'],
        ['RESULTAT', '-49 K$', '+94 K$', '+612 K$'],
    ]
)

doc.add_paragraph(
    'Point mort estime : milieu d\'Annee 2 (mi-2027). '
    'La cle de la Phase 1 est la vente de rapports sur mesure (revenus immediats) '
    'qui financent le developpement de la plateforme.'
)


# ════════════════════════════════════════════════════════════════
# 12. EQUIPE
# ════════════════════════════════════════════════════════════════
doc.add_heading('12. Equipe', level=1)

tbl(
    ['Role', 'Phase 1\n(6 pers.)', 'Phase 2\n(13 pers.)', 'Phase 3\n(21+ pers.)'],
    [
        ['Direction projet', '1', '1', '1'],
        ['Data analyst / Ingenieur donnees', '1', '2', '3'],
        ['Developpeur plateforme', '1', '2', '2'],
        ['Responsable collecte & partenariats', '1', '2', '3'],
        ['Collecteurs terrain', '2', '5', '10+'],
        ['Commercial / Growth', '-', '1', '2'],
    ]
)

doc.add_paragraph(
    'L\'equipe Phase 1 est lean : 6 personnes. '
    'Le Directeur de projet rapporte au DG. '
    'Chaque verticale a un Responsable (collecte + qualite + publication). '
    'Comite de pilotage mensuel : DG + Directeur projet + Leads verticales.'
)


# ════════════════════════════════════════════════════════════════
# 13. CONCLUSION
# ════════════════════════════════════════════════════════════════
doc.add_heading('13. La decision', level=1)

doc.add_paragraph(
    'GETITHERE n\'est pas une idee. C\'est une extension naturelle de ce que '
    'DATATYM\u2122 fait deja — transformer de la donnee terrain en decisions. '
    'Le savoir-faire existe. L\'infrastructure existe. Le marche attend.'
)

doc.add_paragraph(
    'La question n\'est pas de savoir s\'il faut le faire. '
    'La question est : qui le fera en premier — nous ou quelqu\'un d\'autre ?'
)

tbl(
    ['Decision', 'Recommandation'],
    [
        ['Lancer GETITHERE ?', 'GO — Phase 1 (MVP) des mai 2026'],
        ['Premiere verticale ?', 'Finance (donnees publiques accessibles, monetisation rapide)'],
        ['Premiers pays ?', 'Cote d\'Ivoire, Senegal, Cameroun (hubs DATATYM existants)'],
        ['Budget Phase 1 (6 mois)', '40-80 K$ selon ambition (lean si reutilisation Marketym)'],
        ['Equipe Phase 1', '6 personnes'],
        ['MVP datatym.ai', 'Octobre 2026'],
        ['Premiere monetisation', 'Rapports sur mesure — vendables immediatement'],
    ]
)

quote(
    'En Afrique, la donnee n\'est pas un luxe.\n'
    'C\'est l\'infrastructure manquante de la decision.\n'
    'GETITHERE la construit.\n\n'
    'Le premier qui l\'occupe verrouille le marche.'
)


# ── FIN ────────────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    '\n\nDATATYM\u2122\nGETITHERE\n\n'
    'Document strategique — Avril 2026\n'
    'H&C Executive — Confidentiel'
)
r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

out = '/Users/Apple/Desktop/H&C/Enquetes/Suivi enquetes/enquetes/GETITHERE_Document_Strategique_2026.docx'
doc.save(out)
print(f'OK -> {out}')
