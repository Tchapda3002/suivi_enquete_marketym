#!/usr/bin/env python3
"""
Rapport d'Activite SYNTHETIQUE — DATATYM™ / H&C Executive
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for lv in range(1, 4):
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x0F, 0x1E, 0x3A)
    hs.font.size = Pt([0, 16, 13, 11][lv])


def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.font.bold = True; r.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()


# ── PAGE DE GARDE ──────────────────────────────────────────────
for _ in range(4): doc.add_paragraph()
for text, sz, bold, col in [
    ('DATATYM\u2122', 28, True, RGBColor(0x0F, 0x1E, 0x3A)),
    ('Rapport d\'activite synthetique', 18, True, RGBColor(0xB5, 0x89, 0x1A)),
    ('', 10, False, None),
    ('Vague 1 — Mars-Avril 2026', 13, False, RGBColor(0x6B, 0x72, 0x80)),
    ('H&C Executive / Marketym — Confidentiel', 11, False, RGBColor(0x9C, 0xA3, 0xAF)),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(sz); r.font.bold = bold
    if col: r.font.color.rgb = col
doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# 1. DATATYM EN UNE PAGE
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. DATATYM\u2122 — Ce que c\'est', level=1)
doc.add_paragraph(
    'DATATYM\u2122 est le premier referentiel de notation du marche du talent africain. '
    'Il ne decrit pas l\'Afrique — il structure la maniere dont l\'Afrique se comprend elle-meme. '
    'Indices proprietaires, doctrines prescriptives, systeme de notation type S&P : '
    'les dirigeants citent DATATYM\u2122 comme reference de pilotage.'
)
tbl(
    ['Pilier', 'Fonction', 'Contenu'],
    [
        ['I. Indices', 'Mesurer', 'IPE\u2122 · IRTA\u2122 · ILUX\u2122 · IGRO\u2122'],
        ['II. Doctrines', 'Interpreter', '5 verites prescriptives'],
        ['III. Rituels', 'Ancrer dans les routines', 'Barometre trim. · Brief CODIR · Scorecard · Diagnostic DG'],
        ['IV. Decisions', 'Agir', '10 prescriptions CODIR obligatoires'],
    ]
)
doc.add_paragraph(
    'Notation : A (>=80, Elite) · B (60-79, Stable) · C (35-59, Risque) · D (<35, Critique). '
    'Ambition 2027 : DATATYM\u2122 cite dans les COMEX africains comme Bloomberg dans les salles de marche.'
)


# ══════════════════════════════════════════════════════════════
# 2. CE QUI A ETE FAIT
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. Travaux realises', level=1)

doc.add_heading('2.1 Collecte', level=2)
doc.add_paragraph(
    '1 479 repondants consolides · 8 questionnaires · 16+ pays · 28 questions communes · '
    'Taux de completion ~80% · Convergence inter-surveys +/-0,9% · Panel Gen Z filtre : 1 500 repondants.'
)

doc.add_heading('2.2 Plateforme Marketym', level=2)
doc.add_paragraph(
    'Application web (FastAPI + React + Supabase) deployee sur Railway/Render. '
    'Auth JWT/OTP, dashboards admin/enqueteur, tracking clics IP unique, '
    'sync QuestionPro, quotas/segmentations, connexion PostgreSQL directe (~3ms/req).'
)

doc.add_heading('2.3 Referentiel DATATYM\u2122', level=2)
tbl(
    ['Version', 'Date', 'Contenu cle'],
    [
        ['V1 (estimations)', 'Mars 2026', 'IPE=32, IRTA=65.7, IAVA=47, IET=93.9'],
        ['V1 Corrigee', '31 mars', '8 indices calibres sur 1 129 reponses auditees'],
        ['Comparatif V1/V2', '2 avril', 'Recalibration donnees brutes : IPE=63.9, IRTA=62.3'],
        ['V3 (document maitre)', 'Avril', '53 slides : indices, doctrines, rituels, rating, benchmarks, what-if'],
        ['V3 condensee', 'Avril', '25 slides pour diffusion'],
    ]
)

doc.add_heading('2.4 Barometres et analyses Gen Z', level=2)
doc.add_paragraph(
    '3 presentations produites (V2=29 slides, Master Class=23 slides, Analytique=22 slides). '
    'Analyses : matrice talents 9 cases, clustering K-Means, PCA, score de risque composite, '
    'vulnerabilite salariale, profils d\'organisation, scenarios what-if, '
    'correlations drivers (management vers retention = 0.62, equite vers recommandation = 0.78).'
)

doc.add_heading('2.5 Audit de coherence', level=2)
doc.add_paragraph(
    '11 incoherences identifiees entre les livrables Gen Z et le referentiel V3. '
    '3 critiques (erreur arithmetique, nomenclature, score contradictoire), '
    '4 majeurs (architecture causale, grille, doctrines), 4 mineurs. '
    'Indices V3 recalcules sur Gen Z : IPE=25.3 (D), IRTA=68.7 (D), ILUX=53.9 (C), IGRO=45.1 (C).'
)


# ══════════════════════════════════════════════════════════════
# 3. LIVRABLES
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. Inventaire des livrables', level=1)

tbl(
    ['Cat.', '#', 'Document'],
    [
        ['Referentiel', '5', 'V3 53sl. · V3 25sl. · V1 Corrigee · Comparatif V1/V2 · Barometre IDAT'],
        ['Gen Z', '4', 'PPT V2 (29sl.) · Master Class (23sl.) · Analytique (22sl.) · Presentation Strategique'],
        ['Audit', '6', 'Coherence GenZ/V3 · Conformite Marketym · Analyse IPs · Confrontation · Guide Audition · Checklist'],
        ['Finance', '5', 'Situation Financiere · Etats Financiers (x2) · Excedents Quotas · Frais Ecobank'],
        ['Instit.', '4', 'Lettres d\'intro (x2) · Cahiers des charges (x2) · Rapports Barometres (x2)'],
        ['Data', '4', 'Base GENZ (1200 lignes) · Identifiants enq. · Plan Redistribution · Plateforme Marketym'],
        ['TOTAL', '29', 'livrables produits'],
    ]
)


# ══════════════════════════════════════════════════════════════
# 4. DIAGNOSTIC
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. Diagnostic DATATYM\u2122 — Verdict', level=1)

tbl(
    ['Indice', 'Panel global (1 479)', 'Gen Z (1 500)', 'Zone'],
    [
        ['IPE\u2122 (Perf. Employeur)', '32/100 — D', '25,3/100 — D', 'ROUGE'],
        ['IRTA\u2122 (Turnover)', '73,4 — D', '68,7 — D', 'ROUGE'],
        ['ILUX\u2122 (Attractivite)', '47/100 — C', '53,9/100 — C', 'ORANGE'],
        ['IGRO\u2122 (Developpement)', '31,6 — D', '45,1 — C', 'ORANGE'],
        ['NPS Employeur', '-28', '-29', 'ROUGE'],
        ['DATATYM RATING\u2122', 'D (32,8)', 'D', 'ROUGE'],
    ]
)

p = doc.add_paragraph()
r = p.add_run(
    'Verdict : Le marche africain du talent est en zone rouge structurelle. '
    'Les talents sont engages dans leur travail (72,9%) mais pas dans leur organisation (NPS -28). '
    '93,9% veulent entreprendre. 15% seulement voient un avenir chez leur employeur. '
    'La fenetre d\'action est maintenant.'
)
r.font.bold = True
r.font.size = Pt(10)


# ══════════════════════════════════════════════════════════════
# 5. PROCHAINES ETAPES
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. Prochaines etapes', level=1)

tbl(
    ['Horizon', 'Actions'],
    [
        ['Immediat (S17-18)', 'Reconstruire PPTs Gen Z avec indices V3 · Corriger matrice 100% · '
         'Reintegrer doctrines et architecture causale · Appliquer DATATYM(TM) partout'],
        ['Court terme (S19-20)', 'Appliquer migrations 009/010 en production · '
         'Harmoniser grille de notation unique · Integrer situation financiere dans Marketym'],
        ['Vague 2 (Sept. 2026)', '2 500+ repondants · 20+ pays · Indices mis a jour · '
         'Variable genre ajoutee · Classement Top 50 Attractivite'],
        ['Vision 2027', 'Barometre annuel officiel · Benchmarks sectoriels publies · '
         'API DATATYM\u2122 · Partenariats academiques · DATATYM\u2122 Summit'],
    ]
)


# ── FIN ────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DATATYM\u2122 · H&C Executive · Confidentiel · Avril 2026')
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

out = '/Users/Apple/Desktop/H&C/Enquetes/Suivi enquetes/enquetes/Rapport_Activite_DATATYM_202604.docx'
doc.save(out)
print(f'OK -> {out}')
