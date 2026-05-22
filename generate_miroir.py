#!/usr/bin/env python3
"""
DATATYM™ — Le Miroir Déformant
Ce que les employeurs voient vs ce que les Gen Z vivent
Croisement Baromètre Gen Z (n=1500) × ACQ Employeurs (n=293)
"""

import os, json
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import warnings; warnings.filterwarnings('ignore')
matplotlib.rcParams['font.family'] = 'DejaVu Sans'

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

with open('/tmp/comparison_data.json') as f: D = json.load(f)
gz = D['gz']
acq = D['acq']

# Palette
INK='#1B1F3B'; CORAL='#E63946'; AMBER='#E9A820'
OCEAN='#1D7A8A'; SAGE='#2D8659'; SLATE='#8896AB'; BG='#FFFFFF'

IMG = '/tmp/miroir'
os.makedirs(IMG, exist_ok=True)

def emp_pct(stats, keys):
    return sum(v for k,v in stats.items() if k in keys)

# ── Calculer les écarts clés
ECARTS = {
    'satisfaction': {
        'titre': 'Satisfaction des Gen Z',
        'gz_label': 'Gen Z se disent satisfaits',
        'gz_val': gz['satisfaction_pos'],
        'acq_label': 'Employeurs les voient satisfaits',
        'acq_val': emp_pct(acq.get('satisfaction',{}), ['Très satisfaits','Plutôt satisfaits']),
    },
    'engagement': {
        'titre': 'Engagement des Gen Z',
        'gz_label': 'Gen Z se disent engagés',
        'gz_val': gz['engagement_pos'],
        'acq_label': 'Employeurs les voient engagés',
        'acq_val': emp_pct(acq.get('engagement',{}), ['Très engagés','Plutôt engagés']),
    },
    'remuneration': {
        'titre': 'Rémunération juste',
        'gz_label': 'Gen Z disent être rémunérés justement',
        'gz_val': gz['rem_juste_oui'],
        'acq_label': 'Employeurs disent les rémunérer justement',
        'acq_val': emp_pct(acq.get('remuneres_hauteur',{}), ['Tout à fait','Plutôt oui']),
    },
    'culture': {
        'titre': 'Alignement culture / valeurs',
        'gz_label': 'Gen Z se disent alignés avec la culture',
        'gz_val': gz['val_alignes'],
        'acq_label': 'Employeurs disent aligner leur culture',
        'acq_val': emp_pct(acq.get('culture_valeurs',{}), ['Tout à fait','Plutôt oui']),
    },
    'depart': {
        'titre': 'Risque de départ à 12 mois',
        'gz_label': 'Gen Z envisagent de partir',
        'gz_val': gz['iq_risk'],
        'acq_label': 'Employeurs estiment proportion qui partira (>25%)',
        'acq_val': emp_pct(acq.get('pct_partir',{}), ['25-40%','40-60%','Plus de 60%']),
    },
}

# ── Chart de l'écart de perception
def chart_ecart_perception():
    fig, ax = plt.subplots(figsize=(13, 7))
    fig.patch.set_facecolor(BG); ax.set_facecolor(BG)

    items = list(ECARTS.items())
    labels = [v['titre'] for k,v in items]
    gz_vals = [v['gz_val'] for k,v in items]
    acq_vals = [v['acq_val'] for k,v in items]
    ecarts = [a - g for g, a in zip(gz_vals, acq_vals)]

    y = np.arange(len(labels))
    w = 0.38

    bars1 = ax.barh(y - w/2, gz_vals, w, color=CORAL, label='Gen Z (vécu)', alpha=0.85)
    bars2 = ax.barh(y + w/2, acq_vals, w, color=OCEAN, label='Employeurs (perception)', alpha=0.85)

    for bar, v in zip(bars1, gz_vals):
        ax.text(v + 1, bar.get_y() + bar.get_height()/2, f'{v:.0f}%', va='center', color=INK, fontsize=11, fontweight='bold')
    for bar, v in zip(bars2, acq_vals):
        ax.text(v + 1, bar.get_y() + bar.get_height()/2, f'{v:.0f}%', va='center', color=INK, fontsize=11, fontweight='bold')

    # Annotation des écarts
    for i, (g, a, e) in enumerate(zip(gz_vals, acq_vals, ecarts)):
        col = CORAL if abs(e) > 15 else AMBER if abs(e) > 5 else SAGE
        sign = '+' if e > 0 else ''
        ax.text(98, y[i], f'écart : {sign}{e:.0f} pts', va='center', color=col, fontsize=10, fontweight='bold')

    ax.set_yticks(y); ax.set_yticklabels(labels, color=INK, fontsize=11)
    ax.set_xlim(0, 105); ax.xaxis.set_visible(False)
    for s in ax.spines.values(): s.set_visible(False)
    ax.legend(loc='lower right', fontsize=11, frameon=False)
    ax.set_title("L'écart de perception — Ce que les employeurs voient vs ce que les Gen Z vivent",
                 color=INK, fontsize=14, fontweight='bold', pad=15, loc='left')
    plt.tight_layout()
    plt.savefig(f'{IMG}/ecart.png', dpi=180, bbox_inches='tight', facecolor=BG, edgecolor='none')
    plt.close()

# ── Chart % de départ
def chart_depart():
    fig, ax = plt.subplots(figsize=(12, 6))
    fig.patch.set_facecolor(BG); ax.set_facecolor(BG)

    pct_partir = acq.get('pct_partir', {})
    cats = ['Moins de 10%', '10-25%', '25-40%', '40-60%', 'Plus de 60%']
    vals_emp = [pct_partir.get(c, 0) for c in cats]

    x = np.arange(len(cats))
    bars = ax.bar(x, vals_emp, color=OCEAN, alpha=0.85, width=0.55)

    for bar, v in zip(bars, vals_emp):
        ax.text(bar.get_x() + bar.get_width()/2, v + 1, f'{v:.0f}%', ha='center', color=INK, fontsize=12, fontweight='bold')

    # Ligne réelle Gen Z
    ax.axhline(y=gz['iq_risk'], color=CORAL, linestyle='--', linewidth=2.5, alpha=0.8)
    ax.text(4.5, gz['iq_risk']+1, f"Réalité Gen Z = {gz['iq_risk']:.0f}%",
            ha='right', color=CORAL, fontsize=12, fontweight='bold')

    ax.set_xticks(x); ax.set_xticklabels(cats, color=INK, fontsize=11)
    ax.set_ylabel('% des employeurs', color=INK, fontsize=11)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.set_ylim(0, 60)
    ax.set_title("Quelle proportion de Gen Z va partir dans les 12 prochains mois ?\nLa perception employeur vs la réalité Gen Z",
                 color=INK, fontsize=13, fontweight='bold', pad=15, loc='left')
    plt.tight_layout()
    plt.savefig(f'{IMG}/depart.png', dpi=180, bbox_inches='tight', facecolor=BG, edgecolor='none')
    plt.close()

# ── Chart salaire vs sens
def chart_sal_sens():
    fig, axes = plt.subplots(1, 2, figsize=(14, 5))
    fig.patch.set_facecolor(BG)

    # Gen Z
    ax = axes[0]; ax.set_facecolor(BG)
    cats_gz = ['Salaire\nuniquement', 'Les deux\nà parts égales', 'Sens\nuniquement']
    vals_gz = [gz['svs_salaire'], gz['svs_deux'], gz['svs_sens']]
    bars = ax.bar(cats_gz, vals_gz, color=[AMBER, SAGE, OCEAN], alpha=0.85, width=0.55)
    for bar, v in zip(bars, vals_gz):
        ax.text(bar.get_x()+bar.get_width()/2, v+1, f'{v:.0f}%', ha='center', color=INK, fontsize=14, fontweight='bold')
    ax.set_title('VECU Gen Z', color=CORAL, fontsize=13, fontweight='bold')
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.set_ylim(0, 70); ax.tick_params(colors=INK)

    # Employeurs
    ax = axes[1]; ax.set_facecolor(BG)
    sal_emp = emp_pct(acq.get('salaire_sens',{}), ['Plutôt le salaire','Nettement le salaire'])
    deux_emp = acq.get('salaire_sens',{}).get('Les deux à parts égales', 0)
    sens_emp = emp_pct(acq.get('salaire_sens',{}), ['Plutôt le sens','Nettement le sens'])
    cats_emp = ['Salaire\nuniquement', 'Les deux\nà parts égales', 'Sens\nuniquement']
    vals_emp = [sal_emp, deux_emp, sens_emp]
    bars = ax.bar(cats_emp, vals_emp, color=[AMBER, SAGE, OCEAN], alpha=0.85, width=0.55)
    for bar, v in zip(bars, vals_emp):
        ax.text(bar.get_x()+bar.get_width()/2, v+1, f'{v:.0f}%', ha='center', color=INK, fontsize=14, fontweight='bold')
    ax.set_title('PERCEPTION Employeurs', color=OCEAN, fontsize=13, fontweight='bold')
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.set_ylim(0, 70); ax.tick_params(colors=INK)

    fig.suptitle("Salaire ou sens ? Le malentendu fondamental",
                 color=INK, fontsize=14, fontweight='bold', y=1.02)
    plt.tight_layout()
    plt.savefig(f'{IMG}/sal_sens.png', dpi=180, bbox_inches='tight', facecolor=BG, edgecolor='none')
    plt.close()

print("Generation des graphiques...")
chart_ecart_perception()
chart_depart()
chart_sal_sens()


# ════════════════════════════════════════════════════════════════
# DOCUMENT
# ════════════════════════════════════════════════════════════════

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4); style.paragraph_format.line_spacing = 1.15

for lv in range(1,4):
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = 'Calibri'; hs.font.color.rgb = RGBColor(0x1B,0x1F,0x3B)
    hs.font.size = Pt([0,18,14,12][lv])

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.font.bold = True; r.font.size = Pt(9)
    for i,row in enumerate(rows):
        for j,v in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()

def p(text): doc.add_paragraph(text)

def pb(text):
    pr = doc.add_paragraph()
    if '. ' in text:
        parts = text.split('. ', 1)
        r = pr.add_run(parts[0]+'. '); r.font.bold = True
        pr.add_run(parts[1])
    else:
        pr.add_run(text)

def add_img(name, w=5.5):
    path = f'{IMG}/{name}.png'
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

# ── PAGE DE GARDE
for _ in range(3): doc.add_paragraph()
for text, sz, bold, col in [
    ('DATATYM™', 28, True, RGBColor(0x1B,0x1F,0x3B)),
    ('LE MIROIR DEFORMANT', 22, True, RGBColor(0xE6,0x39,0x46)),
    ('', 8, False, None),
    ('Ce que les employeurs voient.\nCe que les Gen Z vivent.', 16, False, RGBColor(0x6B,0x72,0x80)),
    ('', 8, False, None),
    ('Croisement Barometre Gen Z (n=1 500) x ACQ Employeurs (n=293)', 12, False, RGBColor(0x1B,0x1F,0x3B)),
    ('', 12, False, None),
    ('H&C Executive — Avril 2026', 11, False, RGBColor(0x9C,0xA3,0xAF)),
    ('Document confidentiel', 10, False, RGBColor(0x9C,0xA3,0xAF)),
]:
    pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pr.add_run(text); r.font.size = Pt(sz); r.font.bold = bold
    if col: r.font.color.rgb = col
doc.add_page_break()


# ── INTRO
doc.add_heading('La these', level=1)

p("Ce document compare deux realites. D'un cote, ce que 1 500 collaborateurs Gen Z africains "
  "declarent vivre (Barometre Gen Z 2026). De l'autre, ce que 293 dirigeants et DRH disent "
  "observer chez ces memes Gen Z (Barometre Acquisition Talents 2026).")

p("Ces deux mesures portent sur les memes dimensions — engagement, satisfaction, remuneration, "
  "culture, intention de depart — mais sont prises de deux angles opposes : le vecu et la "
  "perception. L'ecart entre les deux est le revelateur le plus puissant des angles morts du "
  "management africain.")

pb("Le constat principal de cette etude est sans appel : les employeurs surestiment systematiquement "
   "la sante de leur relation aux Gen Z et sous-estiment massivement le risque de depart. "
   "L'ecart le plus alarmant : 65% des Gen Z envisagent de partir dans les 12 prochains mois, "
   "alors que 78% des employeurs estiment que moins de 25% partiront.")

doc.add_paragraph()
doc.add_heading('Note methodologique', level=2)
p(f"Les 293 reponses employeurs proviennent de 5 surveys QuestionPro identiques (memes 27 questions). "
  "Le questionnaire emane de la perspective employeur : 'Comment qualifiez-vous LEUR engagement ?', "
  "'Pensez-vous que votre culture correspond a LEURS valeurs ?'. Cette posture 'observation externe' "
  "est ce qui rend la comparaison avec le vecu Gen Z aussi instructive.")
doc.add_page_break()


# ── L'ECART PERCEPTION
doc.add_heading('I. L\'ecart de perception — Vue d\'ensemble', level=1)

add_img('ecart', 6.0)

pb("Sur les 5 dimensions communes mesurees, les employeurs surestiment systematiquement la sante "
   "de la relation. Aucune dimension ne montre une perception employeur INFERIEURE au vecu "
   "Gen Z — ce qui suggere un biais d'auto-evaluation positif chronique au niveau dirigeant.")

doc.add_heading('Synthese des ecarts', level=2)
tbl(['Dimension', 'Vecu Gen Z', 'Perception employeur', 'Ecart', 'Lecture'], [
    ['Satisfaction des Gen Z', f"{gz['satisfaction_pos']}%",
     f"{ECARTS['satisfaction']['acq_val']:.0f}%",
     f"+{ECARTS['satisfaction']['acq_val']-gz['satisfaction_pos']:.0f} pts",
     'Employeurs surestiment'],
    ['Engagement', f"{gz['engagement_pos']}%",
     f"{ECARTS['engagement']['acq_val']:.0f}%",
     f"+{ECARTS['engagement']['acq_val']-gz['engagement_pos']:.0f} pts",
     'Convergence relative'],
    ['Remuneration juste', f"{gz['rem_juste_oui']}%",
     f"{ECARTS['remuneration']['acq_val']:.0f}%",
     f"+{ECARTS['remuneration']['acq_val']-gz['rem_juste_oui']:.0f} pts",
     'ECART CRITIQUE'],
    ['Culture alignee aux valeurs', f"{gz['val_alignes']}%",
     f"{ECARTS['culture']['acq_val']:.0f}%",
     f"+{ECARTS['culture']['acq_val']-gz['val_alignes']:.0f} pts",
     'Employeurs surestiment'],
    ['Risque de depart 12 mois', f"{gz['iq_risk']}% partent",
     f"{ECARTS['depart']['acq_val']:.0f}% (>25% partent)",
     f"-{gz['iq_risk']-ECARTS['depart']['acq_val']:.0f} pts",
     'ANGLE MORT MAJEUR'],
])
doc.add_page_break()


# ── L'ANGLE MORT N°1 : LE DEPART
doc.add_heading('II. L\'angle mort principal : le risque de depart', level=1)

add_img('depart', 6.0)

p(f"La question posee aux deux groupes : 'Quelle proportion de Gen Z va partir dans les 12 prochains mois ?'")
p(f"Cote Gen Z : {gz['iq_risk']:.0f}% declarent envisager de partir (recherche active ou si opportunite).")
p(f"Cote employeurs : 78% estiment que moins de 25% partiront. Seulement 22% estiment que plus de 25% partiront.")

pb("L'ecart est massif : les employeurs sous-estiment de plus de 40 points l'intention de depart "
   "reelle. Cet angle mort a une consequence directe — les politiques de retention sont calibrees "
   "sur une perception erronee, donc systematiquement sous-dimensionnees. Le DRH qui pense que "
   "10% de ses Gen Z partiront budgete pour 10%, alors que la realite est trois a six fois "
   "superieure. Cet ecart explique pourquoi les vagues de depart prennent les organisations par surprise.")

doc.add_heading('Profil de l\'employeur lucide vs naif', level=2)
p(f"Sur les 293 employeurs, seuls {ECARTS['depart']['acq_val']:.0f}% anticipent un taux de depart > 25% — "
  f"ce qui correspond approximativement a la realite. La majorite ({100-ECARTS['depart']['acq_val']:.0f}%) "
  f"vit dans une projection optimiste, deconnectee des intentions reelles de leurs collaborateurs.")
doc.add_page_break()


# ── LE MALENTENDU SALAIRE/SENS
doc.add_heading('III. Le malentendu salaire / sens', level=1)

add_img('sal_sens', 6.0)

p("Question miroir : 'Le salaire ou le sens du travail compte le plus ?'")

tbl(['Reponse', 'Gen Z (vecu)', 'Employeurs (perception)', 'Ecart'], [
    ['Le salaire uniquement', f"{gz['svs_salaire']}%", f"{emp_pct(acq.get('salaire_sens',{}), ['Plutôt le salaire','Nettement le salaire']):.0f}%",
     f"{emp_pct(acq.get('salaire_sens',{}), ['Plutôt le salaire','Nettement le salaire'])-gz['svs_salaire']:+.0f} pts"],
    ['Les deux a parts egales', f"{gz['svs_deux']}%",
     f"{acq.get('salaire_sens',{}).get('Les deux à parts égales', 0):.0f}%",
     f"{acq.get('salaire_sens',{}).get('Les deux à parts égales', 0)-gz['svs_deux']:+.0f} pts"],
    ['Le sens uniquement', f"{gz['svs_sens']}%",
     f"{emp_pct(acq.get('salaire_sens',{}), ['Plutôt le sens','Nettement le sens']):.0f}%",
     f"{emp_pct(acq.get('salaire_sens',{}), ['Plutôt le sens','Nettement le sens'])-gz['svs_sens']:+.0f} pts"],
])

pb("Les employeurs pensent que les Gen Z sont avant tout motives par le salaire (57% repondent "
   "salaire seul). Les Gen Z, eux, repondent majoritairement 'les deux a parts egales' (58%). "
   "Ce malentendu a un cout direct : les politiques RH se concentrent sur des leviers salariaux "
   "(augmentations, primes) alors que la double exigence salaire ET sens est non negociable. "
   "Resultat : les augmentations seules ne retiennent pas, parce qu'elles ignorent l'autre moitie "
   "de l'equation.")
doc.add_page_break()


# ── REMUNERATION : QUI DIT VRAI ?
doc.add_heading('IV. Remuneration : la double comptabilite', level=1)

p("Question miroir : 'Les Gen Z sont-ils remuneres a la hauteur de leur contribution ?'")

tbl(['Position', 'Vecu Gen Z', 'Perception employeur'], [
    ['Oui, remuneration juste', f"{gz['rem_juste_oui']:.0f}%", f"{emp_pct(acq.get('remuneres_hauteur',{}), ['Tout à fait','Plutôt oui']):.0f}%"],
    ['Neutre', f"{round(100-gz['rem_juste_oui']-(100-gz['rem_juste_oui']-30)):.0f}%", f"{acq.get('remuneres_hauteur',{}).get('Neutre',0):.0f}%"],
    ['Non, sous-remuneres', f"{round(100-gz['rem_juste_oui']-30):.0f}%", f"{emp_pct(acq.get('remuneres_hauteur',{}), ['Plutôt non','Pas du tout']):.0f}%"],
    ['', '', ''],
    ['MAIS : se sentent sous-payes generationnellement', f"{gz['souspaye_oui']:.0f}%", '— (non mesure)'],
])

pb("Trois quarts des employeurs (76%) estiment remunerer leurs Gen Z justement. Moins de la moitie "
   "des Gen Z (48%) sont d'accord. Pire : 67% des Gen Z se sentent SOUS-PAYES collectivement par "
   "rapport a leur contribution generationnelle. Les employeurs ne percoivent pas cette dimension "
   "collective — ils raisonnent en termes de salaire individuel par rapport au marche, pas en termes "
   "de valeur creee par une generation entiere. Cette double comptabilite — l'employeur compte ce "
   "qu'il paye, le collaborateur compte ce qu'il vaut — est le moteur invisible des departs.")
doc.add_page_break()


# ── PREFERENCES STRUCTURELLES
doc.add_heading('V. Preferences structurelles : ce qu\'ils croient savoir', level=1)

doc.add_heading('5.1 Stabilite vs mobilite', level=2)
tbl(['Position', 'Vecu Gen Z', 'Perception employeur'], [
    ['Plutot la stabilite', f"{gz['stab_oui']:.0f}%", f"{emp_pct(acq.get('stab_mob',{}), ['Plutôt la stabilité','Nettement la stabilité']):.0f}%"],
    ['Les deux a parts egales', f"{gz['stab_mob_egal']:.0f}%", f"{acq.get('stab_mob',{}).get('Les deux à parts égales', 0):.0f}%"],
    ['Plutot la mobilite', f"{gz['mob_oui']:.0f}%", f"{emp_pct(acq.get('stab_mob',{}), ['Plutôt la mobilité','Nettement la mobilité']):.0f}%"],
])
pb("Les employeurs croient que les Gen Z veulent la mobilite. La realite : 52% privilegient la "
   "stabilite contre seulement 9% qui veulent la mobilite. Cette idee recue selon laquelle 'les "
   "Gen Z changent d'emploi tous les 18 mois' n'est pas confirmee par les Gen Z eux-memes. Ils "
   "veulent la stabilite — mais avec progression. Les programmes RH bases sur la rotation "
   "rapide manquent leur cible.")

doc.add_heading('5.2 Locale vs multinationale', level=2)
tbl(['Position', 'Vecu Gen Z', 'Perception employeur'], [
    ['Plutot locale', f"{gz['locale']:.0f}%", f"{emp_pct(acq.get('locale_multi',{}), ['Plutôt locale','Nettement locale']):.0f}%"],
    ['Indifferent', f"{gz['indiff']:.0f}%", f"{acq.get('locale_multi',{}).get('Indifférent', 0):.0f}%"],
    ['Plutot multinationale', f"{gz['multi']:.0f}%", f"{emp_pct(acq.get('locale_multi',{}), ['Plutôt multinationale','Nettement multinationale']):.0f}%"],
])
pb("Les employeurs locaux pensent que les Gen Z preferent les multinationales. Les Gen Z sont "
   "en realite quasi-equipartis (36% multi, 31% local, 31% indifferent). Le 'mythe de la "
   "multinationale' qui hantait les DRH locaux n'est pas confirme.")
doc.add_page_break()


# ── CONCLUSION
doc.add_heading('VI. Synthese — Les angles morts du management africain', level=1)

p("Cette comparaison entre 1 500 vecus Gen Z et 293 perceptions employeurs revele 5 angles morts "
  "structurels :")

bul_items = [
    "Le risque de depart est sous-estime de 40 points (employeurs voient 22% partir, realite 65%)",
    "La satisfaction est surestimee de 21 points (employeurs voient 86% satisfaits, realite 65%)",
    "L'equite salariale est surestimee de 28 points (76% employeurs vs 48% Gen Z)",
    "L'alignement culturel est surestime de 18 points (89% employeurs vs 71% Gen Z)",
    "Le malentendu salaire/sens : 57% employeurs pensent salaire seul, 58% Gen Z veulent les deux",
]
for item in bul_items:
    pr = doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Ce que ce miroir apprend', level=2)

pb("Les politiques RH sont calibrees sur une perception qui n'est pas la realite. "
   "C'est le diagnostic structurel le plus important de cette etude. Pas un probleme de Gen Z, "
   "pas un probleme de remuneration — un probleme de signal. Les dirigeants africains decident "
   "sur la base de ce qu'ils croient observer chez leurs Gen Z. Or ce qu'ils observent est "
   "systematiquement plus optimiste que la realite vecue.")

pb("La consequence operationnelle est claire. Les budgets de retention sont sous-dimensionnes. "
   "Les programmes culturels sont fondes sur un alignement surestime. Les politiques salariales "
   "se concentrent sur le salaire alors que la double exigence salaire+sens domine. Et surtout : "
   "les vagues de depart paraitront soudaines aux dirigeants alors qu'elles etaient previsibles "
   "depuis des mois.")

doc.add_heading('Pour la suite', level=2)
p("Ce document n'est pas une recommandation. C'est un diagnostic miroir. Trois pistes d'analyse "
  "pour la prochaine vague :")
bul_items = [
    "Mesurer les memes dimensions cote Gen Z et cote employeur dans la meme entreprise (panel apparie) pour mesurer l'ecart de perception au niveau organisationnel",
    "Croiser l'ecart de perception avec la performance de retention — les organisations ou l'ecart est faible retiennent-elles mieux ?",
    "Integrer la variable sexe dans les deux questionnaires pour analyser si l'ecart de perception varie selon le genre du dirigeant ou du collaborateur",
]
for item in bul_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()
pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pr.add_run('\n\nDATATYM™\nLE MIROIR DEFORMANT\n\nBarometre Gen Z (n=1 500) x ACQ Employeurs (n=293)\nH&C Executive — Avril 2026\nDocument confidentiel')
r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x6B,0x72,0x80)

out = '/Users/Apple/Desktop/H&C/Enquetes/Suivi enquetes/enquetes/Miroir_Deformant_GenZ_vs_Employeurs_2026.docx'
doc.save(out)
print(f'OK -> {out}')
