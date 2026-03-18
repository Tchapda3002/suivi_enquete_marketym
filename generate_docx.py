"""
Script pour generer le cahier de charge en format Word (.docx)
"""
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, '059669')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F0FDF4')

    return table


def create_document():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

    doc.styles['Heading 1'].font.size = Pt(22)
    doc.styles['Heading 2'].font.size = Pt(16)
    doc.styles['Heading 3'].font.size = Pt(13)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    for _ in range(6):
        doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CAHIER DE CHARGE')
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Marketym v2 - Plateforme de Suivi d\'Enquetes')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    doc.add_paragraph('')

    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org.add_run('H&C Executive')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph('')

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('Mars 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_page_break()

    # =========================================================================
    # TABLE DES MATIERES
    # =========================================================================
    doc.add_heading('Table des matieres', level=1)

    toc_items = [
        '1. Presentation generale',
        '2. Architecture technique',
        '3. Schema de base de donnees',
        '4. Application accounts',
        '5. Application enquetes',
        '6. Application quotas',
        '7. Application sync',
        '8. Application tracking',
        '9. Application dashboard',
        '10. Application exports',
        '11. Application notifications',
        '12. Integration QuestionPro',
        '13. Logique metier detaillee',
        '14. Templates et pages',
        '15. Etapes d\'implementation',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # =========================================================================
    # 1. PRESENTATION GENERALE
    # =========================================================================
    doc.add_heading('1. Presentation generale', level=1)

    doc.add_heading('1.1 Contexte', level=2)
    doc.add_paragraph(
        'H&C Executive realise des enquetes terrain via des enqueteurs. Les sondages sont '
        'heberges sur QuestionPro. Cette plateforme permet de :'
    )
    for item in [
        'Gerer les enqueteurs (comptes, affectations, suivi)',
        'Gerer les enquetes (creation, configuration, quotas)',
        'Suivre la progression en temps reel (completions, quotas, segmentations)',
        'Synchroniser les donnees avec QuestionPro',
        'Visualiser les statistiques (dashboards admin et enqueteur)',
        'Gerer les quotas simples et croises',
        'Tracker les clics sur les liens de collecte',
        'Exporter les donnees',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('1.2 Utilisateurs', level=2)
    add_styled_table(doc,
        ['Role', 'Description', 'Acces'],
        [
            ['Super Admin', 'Gestionnaire de la plateforme', 'Tout'],
            ['Admin', 'Responsable d\'enquetes', 'Gestion enquetes, enqueteurs, affectations'],
            ['Enqueteur', 'Realise les enquetes terrain', 'Dashboard personnel, ses enquetes'],
        ]
    )

    doc.add_heading('1.3 Stack technique', level=2)
    add_styled_table(doc,
        ['Composant', 'Technologie'],
        [
            ['Framework', 'Django 5.x'],
            ['Base de donnees', 'PostgreSQL (Supabase)'],
            ['Templates', 'Django Templates + HTML/CSS'],
            ['CSS', 'TailwindCSS (via CDN ou build)'],
            ['Emails', 'Brevo (ex-Sendinblue) via API'],
            ['Sondages', 'QuestionPro API v2'],
            ['Deploiement', 'Railway / Render'],
            ['Taches planifiees', 'django-crontab ou Celery (optionnel)'],
        ]
    )

    doc.add_heading('1.4 Contraintes', level=2)
    for item in [
        'Pas d\'API REST separee : tout en server-side rendering',
        'Pas de framework JavaScript (React, Vue, etc.)',
        'JavaScript minimal uniquement pour interactions UI (modals, toggles, AJAX ponctuel)',
        'Hebergement sur Railway/Render',
        'Base de donnees PostgreSQL existante sur Supabase (migration des donnees)',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # 2. ARCHITECTURE TECHNIQUE
    # =========================================================================
    doc.add_heading('2. Architecture technique', level=1)

    doc.add_heading('2.1 Structure du projet Django', level=2)
    doc.add_paragraph(
        'Le projet Django est organise en 8 applications independantes dans un meme projet :'
    )

    add_styled_table(doc,
        ['Application', 'Responsabilite', 'Fichiers cles'],
        [
            ['accounts', 'Authentification, utilisateurs, OTP, invitations', 'models.py, views.py, services/email_service.py, services/otp_service.py'],
            ['enquetes', 'Enquetes, affectations, segmentations', 'models.py, views.py, forms.py'],
            ['quotas', 'Quotas simples et croises', 'models.py, views.py, forms.py'],
            ['sync', 'Synchronisation QuestionPro', 'services/questionpro.py, services/sync_engine.py'],
            ['tracking', 'Tracking clics avec deduplication IP', 'models.py, views.py'],
            ['dashboard', 'Dashboards admin et enqueteur', 'views.py, templatetags/'],
            ['exports', 'Export CSV/Excel', 'views.py, services/'],
            ['notifications', 'Emails, rappels, alertes', 'models.py, services/'],
        ]
    )

    doc.add_heading('2.2 Variables d\'environnement', level=2)
    add_styled_table(doc,
        ['Variable', 'Description', 'Exemple'],
        [
            ['DATABASE_URL', 'URL PostgreSQL (Supabase)', 'postgres://user:pass@host:5432/db'],
            ['SUPABASE_URL', 'URL du projet Supabase', 'https://xxx.supabase.co'],
            ['SUPABASE_KEY', 'Cle API Supabase', 'eyJ...'],
            ['QUESTIONPRO_API_KEY', 'Cle API QuestionPro', 'abc123...'],
            ['BREVO_API_KEY', 'Cle API Brevo', 'xkeysib-...'],
            ['EMAIL_FROM', 'Adresse expediteur', 'marketym@hcexecutive.net'],
            ['EMAIL_FROM_NAME', 'Nom expediteur', 'Marketym'],
            ['SECRET_KEY', 'Django secret key', '(genere)'],
            ['JWT_SECRET_KEY', 'Pour tokens d\'invitation', '(genere)'],
            ['OTP_EXPIRE_MINUTES', 'Duree validite OTP', '5'],
            ['OTP_MAX_ATTEMPTS', 'Tentatives max OTP', '3'],
            ['FRONTEND_URL', 'URL publique du site', 'https://marketym.com'],
            ['SYNC_INTERVAL_MINUTES', 'Intervalle sync auto', '30'],
        ]
    )

    doc.add_heading('2.3 Middleware et securite', level=2)
    for item in [
        'SessionMiddleware : authentification par session Django (pas de JWT pour l\'auth)',
        'AuthenticationMiddleware : gestion utilisateur connecte',
        'CsrfViewMiddleware : protection CSRF sur tous les formulaires',
        'SecurityMiddleware : HTTPS, HSTS en production',
        'Pas besoin de CORS (pas d\'API separee, tout est server-side)',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # 3. SCHEMA DE BASE DE DONNEES
    # =========================================================================
    doc.add_heading('3. Schema de base de donnees', level=1)

    doc.add_heading('3.1 Vue d\'ensemble des tables', level=2)
    doc.add_paragraph(
        'Le schema comporte 15 tables reparties entre les applications Django. '
        'Toutes les cles primaires sont des UUID generes automatiquement.'
    )

    # Enqueteur
    doc.add_heading('3.2 Enqueteur (accounts)', level=2)
    doc.add_paragraph('Etend AbstractUser de Django. Represente un utilisateur de la plateforme.')
    add_styled_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUID', 'PK, auto', 'Identifiant unique'],
            ['email', 'varchar(255)', 'UNIQUE, NOT NULL', 'Email (= username)'],
            ['nom', 'varchar(100)', 'NOT NULL', 'Nom de famille'],
            ['prenom', 'varchar(100)', 'NOT NULL', 'Prenom'],
            ['telephone', 'varchar(20)', 'NULL', 'Telephone'],
            ['identifiant', 'varchar(20)', 'UNIQUE, NOT NULL', 'Code court (ex: "WK47HP")'],
            ['token', 'varchar(20)', 'UNIQUE, NULL', 'Token QuestionPro tracking'],
            ['actif', 'boolean', 'default TRUE', 'Compte actif'],
            ['role', 'varchar(20)', 'default "enqueteur"', 'super_admin, admin, enqueteur'],
            ['compte_configure', 'boolean', 'default FALSE', 'OTP premiere connexion valide'],
            ['doit_changer_mdp', 'boolean', 'default FALSE', 'Forcer changement MDP'],
            ['derniere_connexion', 'datetime', 'NULL', 'Derniere connexion'],
            ['created_at', 'datetime', 'auto', 'Date de creation'],
        ]
    )

    # OtpCode
    doc.add_heading('3.3 OtpCode (accounts)', level=2)
    add_styled_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUID', 'PK', ''],
            ['email', 'varchar(255)', 'NOT NULL, INDEX', 'Email associe'],
            ['code_hash', 'varchar(255)', 'NOT NULL', 'Hash bcrypt du code'],
            ['expires_at', 'datetime', 'NOT NULL', 'Expiration'],
            ['attempts', 'int', 'default 0', 'Tentatives echouees'],
            ['used', 'boolean', 'default FALSE', 'Deja utilise'],
            ['created_at', 'datetime', 'auto', ''],
        ]
    )

    # InvitationToken
    doc.add_heading('3.4 InvitationToken (accounts)', level=2)
    add_styled_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUID', 'PK', ''],
            ['enqueteur', 'FK(Enqueteur)', 'NOT NULL', 'Enqueteur invite'],
            ['token', 'varchar(255)', 'UNIQUE, NOT NULL', 'Token securise (64 chars)'],
            ['expires_at', 'datetime', 'NOT NULL', 'Expiration (48h)'],
            ['used', 'boolean', 'default FALSE', 'Deja utilise'],
            ['used_at', 'datetime', 'NULL', 'Date d\'utilisation'],
            ['created_at', 'datetime', 'auto', ''],
        ]
    )

    # Enquete
    doc.add_heading('3.5 Enquete (enquetes)', level=2)
    add_styled_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUID', 'PK', ''],
            ['survey_id', 'varchar(50)', 'NOT NULL', 'ID QuestionPro'],
            ['code', 'varchar(20)', 'UNIQUE, NULL', 'Code interne'],
            ['nom', 'varchar(300)', 'NOT NULL', 'Nom (depuis QuestionPro)'],
            ['description', 'text', 'NULL', 'Description'],
            ['cible', 'int', 'default 0', 'Objectif global'],
            ['statut', 'varchar(20)', 'default "en_cours"', 'en_cours, termine, archive'],
            ['survey_url', 'text', 'NULL', 'URL du sondage QuestionPro'],
            ['created_at', 'datetime', 'auto', ''],
        ]
    )

    # Affectation
    doc.add_heading('3.6 Affectation (enquetes)', level=2)
    add_styled_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUID', 'PK', ''],
            ['enqueteur', 'FK(Enqueteur)', 'NOT NULL, CASCADE', ''],
            ['enquete', 'FK(Enquete)', 'NOT NULL, CASCADE', ''],
            ['survey_id', 'varchar(50)', 'NOT NULL', 'ID QP (denormalise)'],
            ['lien_questionnaire', 'text', 'NULL', 'Lien de tracking /r/{id}'],
            ['lien_direct', 'text', 'NULL', 'Lien QuestionPro direct'],
            ['objectif_total', 'int', 'default 0', 'Objectif pour cet enqueteur'],
            ['completions_total', 'int', 'default 0', 'Total completions'],
            ['clics', 'int', 'default 0', 'Nombre de clics uniques'],
            ['statut', 'varchar(20)', 'default "en_cours"', ''],
            ['commentaire_admin', 'text', 'NULL', 'Note admin'],
            ['derniere_synchro', 'datetime', 'NULL', 'Derniere sync QP'],
            ['created_at', 'datetime', 'auto', ''],
        ]
    )
    doc.add_paragraph('Contrainte : UNIQUE(enqueteur, enquete)')

    # Segmentation
    doc.add_heading('3.7 Segmentation (enquetes)', level=2)
    add_styled_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUID', 'PK', ''],
            ['enquete', 'FK(Enquete)', 'NOT NULL, CASCADE', ''],
            ['question_id', 'varchar(50)', 'NOT NULL', 'ID question QuestionPro'],
            ['question_text', 'varchar(500)', 'NOT NULL', 'Texte de la question'],
            ['answer_options', 'JSON', 'default []', 'Options de reponse cachees depuis QP'],
            ['created_at', 'datetime', 'auto', ''],
        ]
    )
    doc.add_paragraph('Contrainte : UNIQUE(enquete, question_id)')

    # Quota
    doc.add_heading('3.8 Quota (quotas)', level=2)
    add_styled_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUID', 'PK', ''],
            ['enquete', 'FK(Enquete)', 'NOT NULL, CASCADE', ''],
            ['affectation', 'FK(Affectation)', 'NULL, CASCADE', 'NULL = quota global'],
            ['segmentation', 'FK(Segmentation)', 'NULL, CASCADE', 'Pour quotas simples'],
            ['segment_value', 'varchar(200)', 'NULL', 'Valeur du segment'],
            ['objectif', 'int', 'default 0', 'Objectif'],
            ['quota_config', 'FK(QuotaConfig)', 'NULL, CASCADE', 'Pour quotas croises'],
            ['combination', 'JSON', 'NULL', 'Combinaison croisee'],
            ['pourcentage', 'decimal(5,2)', 'default 0', '% de repartition'],
        ]
    )

    # QuotaConfig
    doc.add_heading('3.9 QuotaConfig (quotas)', level=2)
    add_styled_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUID', 'PK', ''],
            ['enquete', 'FK(Enquete)', 'NOT NULL, CASCADE', ''],
            ['nom', 'varchar(200)', 'NOT NULL', 'Nom du croisement'],
            ['created_at', 'datetime', 'auto', ''],
        ]
    )

    # QuotaConfigQuestion
    doc.add_heading('3.10 QuotaConfigQuestion (quotas)', level=2)
    add_styled_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUID', 'PK', ''],
            ['quota_config', 'FK(QuotaConfig)', 'NOT NULL, CASCADE', ''],
            ['segmentation', 'FK(Segmentation)', 'NOT NULL', ''],
            ['position', 'int', 'default 0', 'Ordre'],
        ]
    )
    doc.add_paragraph('Contrainte : UNIQUE(quota_config, segmentation)')

    # CompletionSegment
    doc.add_heading('3.11 CompletionSegment (sync)', level=2)
    add_styled_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUID', 'PK', ''],
            ['affectation', 'FK(Affectation)', 'NOT NULL, CASCADE', ''],
            ['segmentation', 'FK(Segmentation)', 'NOT NULL', ''],
            ['segment_value', 'varchar(200)', 'NOT NULL', 'Valeur exacte QP'],
            ['completions', 'int', 'default 0', ''],
            ['updated_at', 'datetime', 'auto', ''],
        ]
    )
    doc.add_paragraph('Contrainte : UNIQUE(affectation, segmentation, segment_value)')

    # CompletionCombination
    doc.add_heading('3.12 CompletionCombination (sync)', level=2)
    add_styled_table(doc,
        ['Champ', 'Type', 'Contraintes', 'Description'],
        [
            ['id', 'UUID', 'PK', ''],
            ['affectation', 'FK(Affectation)', 'NOT NULL, CASCADE', ''],
            ['quota_config', 'FK(QuotaConfig)', 'NOT NULL, CASCADE', ''],
            ['combination', 'JSON', 'NOT NULL', 'Ex: {"Pays":"Senegal","Secteur":"Finance"}'],
            ['completions', 'int', 'default 0', ''],
            ['updated_at', 'datetime', 'auto', ''],
        ]
    )
    doc.add_paragraph('Contrainte : UNIQUE(affectation, quota_config, combination)')

    # Autres tables
    doc.add_heading('3.13 Autres tables', level=2)

    doc.add_paragraph('CompletionPays (sync) : completions par pays et par affectation. '
                       'Contrainte UNIQUE(affectation, pays).')
    doc.add_paragraph('HistoriqueCompletion (sync) : snapshots des completions dans le temps '
                       'pour chaque affectation.')
    doc.add_paragraph('Clic (tracking) : clics sur les liens de collecte. '
                       'Contrainte UNIQUE(affectation, ip_address) pour deduplication.')
    doc.add_paragraph('Zone (enquetes) : zones geographiques.')
    doc.add_paragraph('Pays (enquetes) : pays avec FK vers Zone.')

    doc.add_page_break()

    # =========================================================================
    # 4. APPLICATION ACCOUNTS
    # =========================================================================
    doc.add_heading('4. Application accounts', level=1)

    doc.add_heading('4.1 Fonctionnalites', level=2)
    for item in [
        'Inscription (self-registration) avec email, nom, prenom, telephone, mot de passe',
        'Connexion par email + mot de passe',
        'OTP premiere connexion (code 6 chiffres envoye par email, expire en 5 min)',
        'Mot de passe oublie (envoi OTP par email, puis reset)',
        'Changement de mot de passe (utilisateur connecte)',
        'Profil enqueteur (modification avec OTP de confirmation)',
        'Invitation par admin (email avec lien d\'activation, token 48h)',
        'Gestion des roles : super_admin, admin, enqueteur',
        'Creation d\'enqueteur par admin (avec envoi invitation automatique)',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.2 Flux d\'authentification', level=2)

    doc.add_paragraph('Inscription :', style='List Bullet')
    doc.add_paragraph('L\'utilisateur remplit le formulaire (nom, prenom, email, telephone, mot de passe). '
                       'Le compte est cree avec compte_configure=False. Redirect vers la page de connexion.')

    doc.add_paragraph('Premiere connexion :', style='List Bullet')
    doc.add_paragraph('Email + mot de passe corrects. Si compte_configure=False : un code OTP est envoye '
                       'par email et l\'utilisateur est redirige vers /accounts/activate. '
                       'Si compte_configure=True : session creee, redirect vers /dashboard.')

    doc.add_paragraph('Activation OTP :', style='List Bullet')
    doc.add_paragraph('L\'utilisateur saisit le code 6 chiffres recu par email. '
                       'Si correct : compte_configure=True, session creee, redirect vers /dashboard. '
                       'Max 3 tentatives. Code expire en 5 minutes.')

    doc.add_paragraph('Invitation par admin :', style='List Bullet')
    doc.add_paragraph('L\'admin cree un enqueteur. Un email est envoye avec un lien d\'activation '
                       'contenant un token securise (64 caracteres, expire en 48h). '
                       'L\'enqueteur clique et definit son mot de passe.')

    doc.add_paragraph('Mot de passe oublie :', style='List Bullet')
    doc.add_paragraph('L\'utilisateur saisit son email. Un OTP est envoye. '
                       'Il saisit ensuite l\'email + le code + le nouveau mot de passe.')

    doc.add_heading('4.3 Vues', level=2)
    add_styled_table(doc,
        ['URL', 'Vue', 'Methode', 'Description'],
        [
            ['/accounts/login', 'LoginView', 'GET/POST', 'Formulaire de connexion'],
            ['/accounts/register', 'RegisterView', 'GET/POST', 'Formulaire d\'inscription'],
            ['/accounts/activate', 'ActivateView', 'GET/POST', 'Saisie code OTP'],
            ['/accounts/activate/{token}', 'ActivateTokenView', 'GET/POST', 'Activation par lien d\'invitation'],
            ['/accounts/forgot-password', 'ForgotPasswordView', 'GET/POST', 'Demande reset'],
            ['/accounts/reset-password', 'ResetPasswordView', 'GET/POST', 'Reset avec OTP'],
            ['/accounts/change-password', 'ChangePasswordView', 'GET/POST', 'Changer MDP (authentifie)'],
            ['/accounts/profile', 'ProfileView', 'GET/POST', 'Modifier profil (avec OTP)'],
            ['/accounts/logout', 'LogoutView', 'POST', 'Deconnexion'],
        ]
    )

    doc.add_heading('4.4 Securite', level=2)
    for item in [
        'Mots de passe hashes avec bcrypt (via Django make_password/check_password)',
        'Sessions Django (cookie securise, HttpOnly, SameSite)',
        'Protection CSRF sur tous les formulaires POST',
        'OTP : max 3 tentatives, expiration 5 minutes, code a usage unique',
        'Tokens d\'invitation : 64 caracteres securises (secrets.token_urlsafe(48)), expiration 48h',
        'Identifiants enqueteurs : 6 caracteres alphanumeriques sans ambiguite (sans 0, O, I, L)',
        'Decorateurs d\'acces : @login_required, @admin_required, @super_admin_required',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.5 Service email (Brevo)', level=2)
    doc.add_paragraph(
        'L\'envoi d\'emails utilise l\'API Brevo (ex-Sendinblue) via le SDK Python sib_api_v3_sdk. '
        'Les templates HTML sont integres dans le code Python avec un style inline : '
        'fond blanc, bouton vert #059669, police system, max-width 480px.'
    )
    doc.add_paragraph('Types d\'emails envoyes :')
    for item in [
        'Email OTP : code 6 chiffres affiche en grand, indication d\'expiration',
        'Email de bienvenue/invitation : lien d\'activation, email de connexion rappele',
        'Email reset password : code OTP pour reinitialisation',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # 5. APPLICATION ENQUETES
    # =========================================================================
    doc.add_heading('5. Application enquetes', level=1)

    doc.add_heading('5.1 Fonctionnalites', level=2)
    for item in [
        'CRUD complet des enquetes (admin)',
        'CRUD complet des affectations (admin)',
        'CRUD complet des segmentations (admin)',
        'Generation automatique des liens de collecte (tracking + direct)',
        'Vue detail enquete avec affectations, segmentations, quotas',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5.2 Creation d\'enquete', level=2)
    doc.add_paragraph(
        'A la creation, l\'admin saisit uniquement le survey_id QuestionPro. Le systeme :'
    )
    doc.add_paragraph('1. Appelle GET /surveys/{survey_id} de l\'API QuestionPro')
    doc.add_paragraph('2. Recupere : nom, survey_url, status')
    doc.add_paragraph('3. Sauvegarde en base avec statut = "en_cours"')
    doc.add_paragraph('L\'admin peut ensuite ajouter une description et definir la cible (objectif global).')

    doc.add_heading('5.3 Creation d\'affectation', level=2)
    doc.add_paragraph('Lors de la creation d\'une affectation :')
    doc.add_paragraph('1. L\'admin selectionne un enqueteur et une enquete', style='List Bullet')
    doc.add_paragraph('2. Saisit l\'objectif total', style='List Bullet')
    doc.add_paragraph('3. Le systeme genere automatiquement :', style='List Bullet')
    doc.add_paragraph(
        '   - lien_direct : URL QuestionPro avec custom variable custom1={token_enqueteur}\n'
        '   - lien_questionnaire : URL de tracking interne /r/{affectation_id} qui redirige vers lien_direct'
    )

    doc.add_heading('5.4 Creation de segmentation', level=2)
    doc.add_paragraph('Lors de la creation d\'une segmentation :')
    doc.add_paragraph('1. L\'admin selectionne une enquete', style='List Bullet')
    doc.add_paragraph('2. Le systeme fetch les questions du sondage via l\'API QuestionPro', style='List Bullet')
    doc.add_paragraph('3. L\'admin selectionne une question (type choix unique/multiple)', style='List Bullet')
    doc.add_paragraph('4. Le systeme sauvegarde :', style='List Bullet')
    doc.add_paragraph(
        '   - question_id, question_text\n'
        '   - answer_options : JSON des options de reponse [{"id": 123, "text": "Senegal"}, ...]\n'
        '   Les answer_options sont cachees en base pour eviter des appels API repetitifs.'
    )

    doc.add_heading('5.5 Vues admin', level=2)
    add_styled_table(doc,
        ['URL', 'Vue', 'Description'],
        [
            ['/admin/enquetes/', 'EnqueteListView', 'Liste des enquetes'],
            ['/admin/enquetes/create/', 'EnqueteCreateView', 'Creer enquete (survey_id)'],
            ['/admin/enquetes/{id}/', 'EnqueteDetailView', 'Detail avec affectations, segmentations, quotas'],
            ['/admin/enquetes/{id}/edit/', 'EnqueteUpdateView', 'Modifier enquete'],
            ['/admin/enquetes/{id}/delete/', 'EnqueteDeleteView', 'Supprimer enquete'],
            ['/admin/enqueteurs/', 'EnqueteurListView', 'Liste des enqueteurs'],
            ['/admin/enqueteurs/create/', 'EnqueteurCreateView', 'Creer + envoi invitation'],
            ['/admin/enqueteurs/{id}/', 'EnqueteurDetailView', 'Dashboard enqueteur vu par admin'],
            ['/admin/enqueteurs/{id}/edit/', 'EnqueteurUpdateView', 'Modifier enqueteur'],
            ['/admin/affectations/create/', 'AffectationCreateView', 'Creer affectation'],
            ['/admin/affectations/{id}/edit/', 'AffectationUpdateView', 'Modifier affectation'],
            ['/admin/affectations/{id}/delete/', 'AffectationDeleteView', 'Supprimer affectation'],
        ]
    )

    doc.add_heading('5.6 Statuts', level=2)
    doc.add_paragraph('Statuts des enquetes :')
    add_styled_table(doc,
        ['Statut', 'Description', 'Actions possibles'],
        [
            ['en_cours', 'Active, sync possible', 'Sync, editer, affecter'],
            ['termine', 'Terminee, pas de sync', 'Voir stats, exporter'],
            ['archive', 'Archivee', 'Voir seulement'],
        ]
    )

    doc.add_page_break()

    # =========================================================================
    # 6. APPLICATION QUOTAS
    # =========================================================================
    doc.add_heading('6. Application quotas', level=1)

    doc.add_heading('6.1 Quotas simples', level=2)
    doc.add_paragraph(
        'Un quota simple est defini par une enquete, une segmentation (question QP), '
        'un segment_value (option de reponse) et un objectif. '
        'Si affectation = NULL : quota global. Si affectation != NULL : quota par enqueteur.'
    )
    doc.add_paragraph(
        'L\'admin peut creer des quotas en masse : pour une segmentation donnee, '
        'definir un objectif par option de reponse. Le systeme cree une ligne Quota par option.'
    )

    doc.add_heading('6.2 Quotas croises (cross-tabulation)', level=2)
    doc.add_paragraph(
        'Un quota croise combine N segmentations (ex: Pays x Secteur x Taille) '
        'pour definir des objectifs par combinaison.'
    )

    doc.add_paragraph('Structure :', style='List Bullet')
    doc.add_paragraph(
        'QuotaConfig contient le nom et l\'enquete.\n'
        'QuotaConfigQuestion lie N segmentations au QuotaConfig (avec position/ordre).\n'
        'Quota contient chaque combinaison avec son objectif et son pourcentage.'
    )

    doc.add_paragraph('Processus de creation :', style='List Bullet')
    doc.add_paragraph(
        '1. L\'admin clique "Creer quota croise"\n'
        '2. Etape 1 : selectionne N segmentations (checkboxes)\n'
        '3. Etape 2 : le systeme genere le produit cartesien des answer_options\n'
        '4. Etape 3 : l\'admin saisit un pourcentage par combinaison\n'
        '   (bouton "Repartition equitable" disponible)\n'
        '5. Les objectifs sont calcules : objectif_combo = pourcentage * cible_enquete / 100\n'
        '6. Le systeme cree : 1 QuotaConfig + N QuotaConfigQuestion + P Quota'
    )

    doc.add_heading('6.3 Generation des combinaisons', level=2)
    doc.add_paragraph(
        'Le produit cartesien est genere a partir des answer_options de chaque segmentation. '
        'Par exemple, si on croise Pays (Senegal, Cote d\'Ivoire) et Secteur (Finance, Telecom), '
        'on obtient 4 combinaisons :'
    )
    add_styled_table(doc,
        ['Pays', 'Secteur'],
        [
            ['Senegal', 'Finance'],
            ['Senegal', 'Telecom'],
            ['Cote d\'Ivoire', 'Finance'],
            ['Cote d\'Ivoire', 'Telecom'],
        ]
    )
    doc.add_paragraph(
        'Chaque combinaison est stockee en JSON : {"Pays": "Senegal", "Secteur": "Finance"}. '
        'La serialisation utilise json.dumps(combo, sort_keys=True) pour garantir l\'unicite.'
    )

    doc.add_heading('6.4 Vues', level=2)
    add_styled_table(doc,
        ['URL', 'Vue', 'Description'],
        [
            ['/admin/quotas/enquete/{id}/', 'QuotasByEnqueteView', 'Quotas d\'une enquete'],
            ['/admin/quotas/create/', 'QuotaCreateView', 'Creer quota simple'],
            ['/admin/quotas/bulk-create/', 'QuotaBulkCreateView', 'Creer quotas en masse'],
            ['/admin/quotas/{id}/edit/', 'QuotaUpdateView', 'Modifier quota'],
            ['/admin/quotas/{id}/delete/', 'QuotaDeleteView', 'Supprimer quota'],
            ['/admin/quota-configs/create/', 'QuotaConfigCreateView', 'Creer quota croise'],
            ['/admin/quota-configs/{id}/', 'QuotaConfigDetailView', 'Detail quota croise'],
            ['/admin/quota-configs/{id}/delete/', 'QuotaConfigDeleteView', 'Supprimer quota croise'],
        ]
    )

    doc.add_page_break()

    # =========================================================================
    # 7. APPLICATION SYNC
    # =========================================================================
    doc.add_heading('7. Application sync', level=1)

    doc.add_heading('7.1 Processus de synchronisation', level=2)
    doc.add_paragraph('Pour chaque affectation d\'une enquete en_cours :')
    doc.add_paragraph(
        '1. Appeler GET /surveys/{survey_id}/responses?page=1&perPage=1000\n'
        '   Filtrer par custom1 = {enqueteur.token}. Paginer si necessaire.\n\n'
        '2. Pour chaque reponse :\n'
        '   a. Compter les completions totales pour cet enqueteur\n'
        '   b. Pour chaque segmentation : trouver la reponse a la question, '
        'extraire le segment_value, upsert dans completions_segments\n'
        '   c. Pour chaque quota_config : construire la combinaison JSON, '
        'upsert dans completions_combinations\n\n'
        '3. Mettre a jour affectation.completions_total\n'
        '4. Mettre a jour affectation.derniere_synchro\n'
        '5. Inserer dans historique_completions (snapshot)'
    )

    doc.add_heading('7.2 Correspondance des reponses', level=2)
    doc.add_paragraph(
        'Les valeurs de segment sont stockees TELLES QUELLES depuis QuestionPro. '
        'Aucune normalisation n\'est effectuee. Le matching se fait par comparaison '
        'exacte entre answer_options[].text et la valeur retournee par l\'API.'
    )

    doc.add_heading('7.3 Optimisations', level=2)
    for item in [
        'per_page = 1000 (max autorise par QP) au lieu de 100',
        'Sync uniquement les enquetes statut = "en_cours"',
        'answer_options cachees en base (pas de re-fetch des questions a chaque sync)',
        'Rate limiting : 300 appels / 60 secondes sur QP',
        'Quota API : 5000 appels gratuits / mois',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.4 Vues', level=2)
    add_styled_table(doc,
        ['URL', 'Vue', 'Description'],
        [
            ['/admin/sync/', 'SyncAllView', 'POST : lancer sync complete'],
            ['/admin/sync/{affectation_id}/', 'SyncAffectationView', 'POST : sync une affectation'],
            ['/enqueteur/{id}/sync/', 'EnqueteurSyncView', 'POST : sync enqueteur'],
        ]
    )

    doc.add_page_break()

    # =========================================================================
    # 8. APPLICATION TRACKING
    # =========================================================================
    doc.add_heading('8. Application tracking', level=1)

    doc.add_heading('8.1 Fonctionnement', level=2)
    doc.add_paragraph(
        'L\'enqueteur partage le lien de tracking (ex: https://marketym.com/r/{affectation_id}). '
        'Quand un visiteur clique :'
    )
    doc.add_paragraph(
        '1. Le serveur recupere l\'IP du visiteur (X-Forwarded-For ou REMOTE_ADDR)\n'
        '2. Tente un INSERT dans la table clics (affectation_id, ip_address, user_agent)\n'
        '   - Si UNIQUE violation (IP deja vue) : ignore silencieusement\n'
        '   - Si OK : incremente affectation.clics += 1\n'
        '3. Redirect 302 vers affectation.lien_direct (URL QuestionPro)'
    )
    doc.add_paragraph(
        'La deduplication est assuree par la contrainte UNIQUE(affectation, ip_address). '
        'Un seul clic est compte par IP unique par affectation.'
    )

    doc.add_heading('8.2 Vues', level=2)
    add_styled_table(doc,
        ['URL', 'Vue', 'Description'],
        [
            ['/r/{affectation_id}', 'TrackClickView', 'GET : track + redirect'],
            ['/admin/affectations/{id}/clics/', 'ClicListView', 'Liste des clics (admin)'],
        ]
    )

    doc.add_page_break()

    # =========================================================================
    # 9. APPLICATION DASHBOARD
    # =========================================================================
    doc.add_heading('9. Application dashboard', level=1)

    doc.add_heading('9.1 Dashboard admin', level=2)

    doc.add_paragraph('Page principale (/admin/) :', style='List Bullet')
    doc.add_paragraph(
        'KPI globaux (cartes en haut) : enquetes actives, enqueteurs actifs, '
        'completions totales, taux de completion moyen.\n'
        'Liste des enquetes en cours : tableau avec nom, cible, completions, taux, statut.\n'
        'Enqueteurs recents : derniers connectes.'
    )

    doc.add_paragraph('Detail enquete (/admin/enquetes/{id}/) :', style='List Bullet')
    doc.add_paragraph(
        'KPI enquete : cible, completions totales, completions valides (apres logique excedent), '
        'excedents, taux de completion, clics totaux.\n'
        'Tableau des affectations : enqueteur, objectif, completions, valides, excedents, taux, clics.\n'
        'Graphiques barres par segmentation (voir section 9.3).\n'
        'Tableau croise si quota_configs existent (colonnes par variable + %, Objectif, Completions, Progression).'
    )

    doc.add_paragraph('Detail enqueteur (/admin/enqueteurs/{id}/) :', style='List Bullet')
    doc.add_paragraph(
        'Affiche le MEME dashboard que l\'enqueteur verrait, mais dans le layout admin. '
        'Permet a l\'admin de voir exactement ce que l\'enqueteur voit : '
        'KPI personnels, liste des enquetes, graphiques par segmentation.'
    )

    doc.add_heading('9.2 Dashboard enqueteur', level=2)

    doc.add_paragraph('Onglet Tableau de bord :', style='List Bullet')
    doc.add_paragraph(
        'KPI : enquetes actives, completions totales, taux moyen.\n'
        'Liste des enquetes avec progression.'
    )

    doc.add_paragraph('Onglet Mes enquetes :', style='List Bullet')
    doc.add_paragraph(
        'Liste des enquetes affectees. Pour chaque enquete : objectif, completions, taux. '
        'Clic sur une enquete ouvre le detail.'
    )

    doc.add_paragraph('Detail enquete enqueteur :', style='List Bullet')
    doc.add_paragraph(
        'KPI : objectif, completions, completions valides, excedents, taux.\n'
        'Lien de collecte avec bouton copier.\n'
        'Graphiques barres par segmentation (meme format que l\'admin).'
    )

    doc.add_paragraph('Onglet Profil :', style='List Bullet')
    doc.add_paragraph(
        'Infos personnelles (nom, prenom, email, telephone, identifiant).\n'
        'Modification avec OTP de confirmation.\n'
        'Changement de mot de passe.'
    )

    doc.add_heading('9.3 Graphiques barres CSS', level=2)
    doc.add_paragraph(
        'Pas de librairie JavaScript. Barres horizontales en CSS pur. '
        'Pour chaque segmentation, un graphique affiche une barre par option de reponse '
        'avec objectif (si defini) vs completions.'
    )
    doc.add_paragraph('Couleurs :')
    add_styled_table(doc,
        ['Couleur', 'Code', 'Usage'],
        [
            ['Vert', '#059669', 'En cours, normal'],
            ['Orange', '#F59E0B', 'Excedent'],
            ['Gris', '#E5E7EB', 'Fond de barre'],
        ]
    )
    doc.add_paragraph(
        'Important : les graphiques barres sont affiches dans BOTH le dashboard admin '
        'ET le dashboard enqueteur. Le tableau croise est affiche UNIQUEMENT dans le '
        'detail enquete admin (trop detaille pour les enqueteurs).'
    )

    doc.add_page_break()

    # =========================================================================
    # 10. APPLICATION EXPORTS
    # =========================================================================
    doc.add_heading('10. Application exports', level=1)

    doc.add_heading('10.1 Exports disponibles', level=2)
    add_styled_table(doc,
        ['Export', 'Contenu', 'Format'],
        [
            ['Enqueteurs', 'id, nom, prenom, email, telephone, identifiant, actif, role, date creation', 'CSV'],
            ['Affectations par enquete', 'enqueteur, objectif, completions, valides, excedents, taux, clics', 'CSV'],
            ['Quotas par enquete', 'segment, objectif, completions, valides, excedents, %', 'CSV'],
            ['Historique completions', 'date, enqueteur, enquete, completions', 'CSV'],
        ]
    )

    doc.add_heading('10.2 Implementation', level=2)
    doc.add_paragraph(
        'Les exports utilisent le module csv de Python avec HttpResponse(content_type="text/csv"). '
        'Des boutons d\'export sont places dans les pages admin concernees.'
    )

    doc.add_page_break()

    # =========================================================================
    # 11. APPLICATION NOTIFICATIONS
    # =========================================================================
    doc.add_heading('11. Application notifications', level=1)

    doc.add_heading('11.1 Types de notifications', level=2)
    add_styled_table(doc,
        ['Type', 'Declencheur', 'Destinataire'],
        [
            ['Rappel inactivite', 'Enqueteur inactif > X jours', 'Enqueteur'],
            ['Quota atteint', 'Completions >= objectif sur un segment', 'Admin'],
            ['Excedent detecte', 'Completions > cible globale', 'Admin'],
            ['Recapitulatif', 'Periodique (hebdomadaire)', 'Admin'],
        ]
    )

    doc.add_paragraph(
        'Phase initiale : manuellement via bouton admin ("Envoyer rappel"). '
        'Phase ulterieure : tache Celery/cron pour automatiser.'
    )

    doc.add_page_break()

    # =========================================================================
    # 12. INTEGRATION QUESTIONPRO
    # =========================================================================
    doc.add_heading('12. Integration QuestionPro', level=1)

    doc.add_heading('12.1 API QuestionPro v2', level=2)
    add_styled_table(doc,
        ['Parametre', 'Valeur'],
        [
            ['Base URL', 'https://api.questionpro.com/a/api/v2'],
            ['Authentification', 'Header api-key: {QUESTIONPRO_API_KEY}'],
            ['Limites', '5000 appels gratuits/mois, 300 appels/60s'],
            ['Pagination', 'page (1-based) + perPage (max 1000)'],
        ]
    )

    doc.add_heading('12.2 Endpoints utilises', level=2)
    add_styled_table(doc,
        ['Endpoint', 'Usage', 'Quand'],
        [
            ['GET /surveys/{id}', 'Info sondage (nom, URL, statut)', 'Creation enquete'],
            ['GET /surveys/{id}/questions', 'Questions + options reponse', 'Creation segmentation'],
            ['GET /surveys/{id}/responses', 'Reponses des participants', 'Synchronisation'],
        ]
    )

    doc.add_heading('12.3 Filtrage par enqueteur', level=2)
    doc.add_paragraph(
        'Chaque enqueteur a un token unique (6 caracteres). A la creation d\'une affectation, '
        'le lien QuestionPro est genere avec ?custom1={token}. Lors de la sync, on filtre '
        'les reponses dont customVariables[custom1] == token.'
    )

    doc.add_heading('12.4 Structure de reponse', level=2)
    doc.add_paragraph(
        'Les reponses QuestionPro contiennent :\n'
        '- responseID : identifiant unique de la reponse\n'
        '- timestamp : date/heure de la reponse\n'
        '- customVariables : [{"variableKey": "custom1", "variableValue": "WK47HP"}]\n'
        '- questionResponse : [{"questionID": 67890, "answerValues": [{"answerID": 111, "value": {"text": "Senegal"}}]}]'
    )

    doc.add_page_break()

    # =========================================================================
    # 13. LOGIQUE METIER
    # =========================================================================
    doc.add_heading('13. Logique metier detaillee', level=1)

    doc.add_heading('13.1 Logique des excedents', level=2)
    doc.add_paragraph(
        'Les excedents sont calcules au niveau GLOBAL de l\'enquete d\'abord, '
        'puis propages aux enqueteurs individuels.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Principe :')
    run.bold = True

    doc.add_paragraph(
        '1. Une enquete a une cible (objectif global), par exemple 100\n'
        '2. Plusieurs enqueteurs y sont affectes\n'
        '3. Tant que le total des completions de TOUS les enqueteurs <= 100, '
        'TOUTES les completions sont valides\n'
        '4. Des que le total depasse 100, les completions en surplus deviennent des excedents\n'
        '5. Les excedents sont attribues chronologiquement : les DERNIERES completions sont les excedents'
    )

    p = doc.add_paragraph()
    run = p.add_run('Exemple :')
    run.bold = True

    doc.add_paragraph(
        'Enquete "Satisfaction Client" - Cible: 100\n\n'
        'Enqueteur A : 45 completions (arrivees en premier)\n'
        'Enqueteur B : 35 completions\n'
        'Enqueteur C : 30 completions (arrivees en dernier)\n'
        'Total : 110 -> Excedent global : 10\n\n'
        'Repartition (chronologiquement, derniers = excedents) :\n'
        '- Enqueteur C : 30 completions, dont 10 excedent -> 20 valides, 10 excedents\n'
        '- Enqueteur B : 35 completions -> 35 valides, 0 excedent\n'
        '- Enqueteur A : 45 completions -> 45 valides, 0 excedent'
    )

    p = doc.add_paragraph()
    run = p.add_run('Algorithme :')
    run.bold = True

    doc.add_paragraph(
        '1. Calculer excedent_global = sum(completions) - cible\n'
        '2. Si excedent_global <= 0 : tout est valide\n'
        '3. Sinon : trier les affectations par derniere_synchro DESC\n'
        '4. Pour chaque affectation (du plus recent au plus ancien) :\n'
        '   - Si restant_excedent >= completions : tout est excedent\n'
        '   - Sinon : excedent = restant_excedent, valides = completions - excedent\n'
        '   - restant_excedent -= min(completions, restant_excedent)'
    )

    doc.add_paragraph(
        'La meme logique s\'applique au niveau des segments : pour un quota "Senegal = 50", '
        'si le total global des completions "Senegal" depasse 50, les surplus sont des excedents '
        'attribues aux derniers arrivants.'
    )

    doc.add_heading('13.2 Calcul du taux de completion', level=2)
    doc.add_paragraph('taux = (completions_valides / objectif) * 100')
    doc.add_paragraph('Si objectif = 0, le taux est 0.')

    doc.add_heading('13.3 Generation d\'identifiant', level=2)
    doc.add_paragraph(
        'Code de 6 caracteres alphanumeriques : lettres majuscules + chiffres, '
        'sans les caracteres ambigus (0, O, I, L). Genere avec secrets.choice(). '
        'Unicite verifiee en base. Exemples : WK47HP, B3DNER.'
    )

    doc.add_heading('13.4 Statuts', level=2)
    doc.add_paragraph('Enquetes : en_cours, termine, archive')
    doc.add_paragraph('Affectations : en_cours, termine, suspendu')

    doc.add_page_break()

    # =========================================================================
    # 14. TEMPLATES ET PAGES
    # =========================================================================
    doc.add_heading('14. Templates et pages', level=1)

    doc.add_heading('14.1 Layout general', level=2)
    doc.add_paragraph(
        'base.html : layout principal avec TailwindCSS (CDN), meta tags, favicon, '
        'block title, block content, main.js.\n'
        'base_admin.html : extends base.html, ajoute sidebar et breadcrumb.\n'
        'base_enqueteur.html : extends base.html, ajoute navbar avec onglets.'
    )

    doc.add_heading('14.2 Pages publiques (accounts/)', level=2)
    add_styled_table(doc,
        ['Page', 'Description', 'Elements'],
        [
            ['login.html', 'Connexion', 'Formulaire email + password, liens inscription/oublie'],
            ['register.html', 'Inscription', 'Formulaire nom, prenom, email, telephone, password'],
            ['activate.html', 'OTP', 'Champ code 6 chiffres, timer expiration'],
            ['activate_token.html', 'Invitation', 'Formulaire nouveau mot de passe'],
            ['forgot_password.html', 'Oublie', 'Champ email'],
            ['reset_password.html', 'Reset', 'Champs email + code + nouveau MDP'],
        ]
    )

    doc.add_heading('14.3 Pages admin', level=2)
    add_styled_table(doc,
        ['Page', 'Description'],
        [
            ['admin_dashboard.html', 'KPI globaux + liste enquetes + enqueteurs recents'],
            ['enquete_list.html', 'Tableau des enquetes'],
            ['enquete_detail.html', 'KPI + affectations + graphiques segmentations + tableau croise'],
            ['enquete_form.html', 'Formulaire creation/edition enquete'],
            ['enqueteur_list.html', 'Tableau des enqueteurs'],
            ['enqueteur_detail.html', 'Dashboard enqueteur vu par admin'],
            ['enqueteur_form.html', 'Formulaire creation/edition enqueteur'],
            ['affectation_form.html', 'Formulaire creation/edition affectation'],
            ['segmentation_form.html', 'Formulaire creation segmentation'],
            ['quota_list.html', 'Quotas d\'une enquete (simples + croises)'],
            ['quota_form.html', 'Formulaire creation quota simple'],
            ['quota_bulk_form.html', 'Formulaire creation quotas en masse'],
            ['quota_config_form.html', 'Multi-etapes : selection, generation, saisie %'],
            ['quota_config_detail.html', 'Tableau croise avec progressions'],
            ['sync_panel.html', 'Bouton sync + historique'],
            ['stats.html', 'Statistiques globales'],
            ['export_panel.html', 'Boutons export CSV'],
        ]
    )

    doc.add_heading('14.4 Pages enqueteur', level=2)
    add_styled_table(doc,
        ['Page', 'Description'],
        [
            ['enqueteur_dashboard.html', 'KPI personnels + liste enquetes'],
            ['enqueteur_enquete_detail.html', 'Detail : KPI, lien, graphiques barres'],
            ['enqueteur_profil.html', 'Infos personnelles, modification, changement MDP'],
        ]
    )

    doc.add_heading('14.5 Composants reutilisables (includes)', level=2)
    add_styled_table(doc,
        ['Composant', 'Usage'],
        [
            ['stat_card.html', 'Carte KPI (titre, valeur, icone, couleur)'],
            ['progress_bar.html', 'Barre de progression (%, couleur)'],
            ['bar_chart.html', 'Graphique barres horizontales CSS'],
            ['cross_table.html', 'Tableau croise quotas'],
            ['modal.html', 'Modal generique (confirmation, formulaire)'],
            ['table.html', 'Tableau avec tri et pagination'],
            ['badge.html', 'Badge statut (en_cours, termine, etc.)'],
            ['alert.html', 'Alerte (succes, erreur, info, warning)'],
            ['copy_button.html', 'Bouton copier dans le presse-papier'],
            ['empty_state.html', 'Etat vide (icone + message)'],
        ]
    )

    doc.add_heading('14.6 Charte graphique', level=2)
    add_styled_table(doc,
        ['Element', 'Valeur'],
        [
            ['Couleur primaire', '#059669 (vert emeraude)'],
            ['Couleur secondaire', '#111827 (gris fonce)'],
            ['Couleur excedent', '#F59E0B (orange)'],
            ['Couleur danger', '#EF4444 (rouge)'],
            ['Couleur info', '#3B82F6 (bleu)'],
            ['Couleur fond', '#F9FAFB'],
            ['Couleur texte', '#111827'],
            ['Couleur texte secondaire', '#6B7280'],
            ['Couleur bordures', '#E5E7EB'],
            ['Police', 'system-ui, -apple-system, sans-serif'],
            ['Rayon bordures', '8px (cartes), 6px (boutons), 12px (modals)'],
            ['Nom de l\'application', 'Marketym'],
        ]
    )

    doc.add_page_break()

    # =========================================================================
    # 15. ETAPES D'IMPLEMENTATION
    # =========================================================================
    doc.add_heading('15. Etapes d\'implementation', level=1)

    phases = [
        {
            'titre': 'Phase 0 : Initialisation du projet',
            'etapes': [
                ('0.1 - Setup Django',
                 'django-admin startproject marketym. Creer les 8 apps. '
                 'Configurer settings.py (database PostgreSQL/Supabase, installed apps, middleware, '
                 'templates, static files). Configurer urls.py principal. '
                 'Variables d\'environnement avec python-dotenv. '
                 'requirements.txt : Django, psycopg2-binary, bcrypt, sib-api-v3-sdk, httpx, '
                 'python-dotenv, gunicorn, whitenoise.'),
            ]
        },
        {
            'titre': 'Phase 1 : Modeles et base de donnees',
            'etapes': [
                ('1.1 - Definir tous les modeles',
                 'accounts/models.py (Enqueteur, OtpCode, InvitationToken). '
                 'enquetes/models.py (Zone, Pays, Enquete, Affectation, Segmentation). '
                 'quotas/models.py (Quota, QuotaConfig, QuotaConfigQuestion). '
                 'sync/models.py (CompletionSegment, CompletionCombination, CompletionPays, HistoriqueCompletion). '
                 'tracking/models.py (Clic). '
                 'makemigrations + migrate.'),
            ]
        },
        {
            'titre': 'Phase 2 : Authentification',
            'etapes': [
                ('2.1 - Services de base',
                 'otp_service.py : generate_otp, hash_code, verify_code, create_otp, verify_otp. '
                 'email_service.py : send_email (Brevo), send_otp_email, send_welcome_email.'),
                ('2.2 - Vues d\'authentification',
                 'LoginView, RegisterView, LogoutView, ActivateView (OTP), '
                 'ForgotPasswordView, ResetPasswordView. '
                 'Decorateurs : @admin_required, @super_admin_required.'),
                ('2.3 - Templates auth',
                 'login.html, register.html, activate.html, forgot_password.html, reset_password.html. '
                 'base.html (layout minimal).'),
            ]
        },
        {
            'titre': 'Phase 3 : Layout et templates de base',
            'etapes': [
                ('3.1 - Templates globaux',
                 'base.html (TailwindCSS, structure). base_admin.html (sidebar, breadcrumb). '
                 'base_enqueteur.html (navbar onglets). '
                 'Composants : stat_card, progress_bar, badge, alert, modal, table, empty_state. '
                 'static/css/main.css, static/js/main.js.'),
            ]
        },
        {
            'titre': 'Phase 4 : Gestion des enquetes',
            'etapes': [
                ('4.1 - CRUD enquetes',
                 'EnqueteListView, CreateView, DetailView, UpdateView, DeleteView. '
                 'Templates. Integration QuestionPro pour creation.'),
                ('4.2 - CRUD enqueteurs (admin)',
                 'EnqueteurListView, CreateView, DetailView, UpdateView, DeleteView. '
                 'Envoi invitation email a la creation.'),
                ('4.3 - CRUD affectations',
                 'AffectationCreateView, UpdateView, DeleteView. '
                 'Generation automatique des liens (tracking + direct).'),
            ]
        },
        {
            'titre': 'Phase 5 : Segmentations et quotas',
            'etapes': [
                ('5.1 - Segmentations',
                 'SegmentationCreateView (fetch questions QP, cache answer_options). '
                 'UpdateView, DeleteView.'),
                ('5.2 - Quotas simples',
                 'QuotaCreateView, BulkCreateView, UpdateView, DeleteView.'),
                ('5.3 - Quotas croises',
                 'QuotaConfigCreateView (multi-etapes). Generation combinaisons. '
                 'DetailView, DeleteView. Composant cross_table.html.'),
            ]
        },
        {
            'titre': 'Phase 6 : Synchronisation QuestionPro',
            'etapes': [
                ('6.1 - Client QuestionPro',
                 'questionpro.py : QuestionProClient (get_survey, get_questions, get_responses, '
                 'get_all_responses). Gestion pagination, rate limiting, erreurs.'),
                ('6.2 - Moteur de sync',
                 'sync_engine.py : sync_affectation, sync_enquete, sync_all. '
                 'Remplissage completions. Vues SyncAllView, SyncAffectationView.'),
            ]
        },
        {
            'titre': 'Phase 7 : Tracking des clics',
            'etapes': [
                ('7.1 - Tracking',
                 'TrackClickView : GET /r/{id} -> track IP + redirect. '
                 'Deduplication UNIQUE(affectation, ip_address). ClicListView (admin).'),
            ]
        },
        {
            'titre': 'Phase 8 : Dashboards',
            'etapes': [
                ('8.1 - Dashboard admin',
                 'AdminDashboardView : KPI globaux, enquetes en cours, enqueteurs recents. '
                 'Composants graphiques barres CSS.'),
                ('8.2 - Detail enquete admin',
                 'Vue enrichie avec graphiques par segmentation. '
                 'Tableau croise pour quota_configs. Logique excedents (global -> individuel).'),
                ('8.3 - Dashboard enqueteur',
                 'KPI personnels, liste enquetes, detail avec graphiques barres. '
                 'ProfilView avec modification + changement MDP.'),
                ('8.4 - Vue enqueteur dans admin',
                 'EnqueteurDetailView (admin) : reutilise les memes donnees que le dashboard enqueteur.'),
            ]
        },
        {
            'titre': 'Phase 9 : Exports',
            'etapes': [
                ('9.1 - Exports CSV',
                 'ExportEnqueteursView, ExportAffectationsView, ExportQuotasView. '
                 'Boutons dans les pages admin.'),
            ]
        },
        {
            'titre': 'Phase 10 : Statistiques avancees',
            'etapes': [
                ('10.1 - Pages stats',
                 'Stats par pays, par segment, globales. Graphiques barres CSS.'),
            ]
        },
        {
            'titre': 'Phase 11 : Notifications',
            'etapes': [
                ('11.1 - Notifications de base',
                 'Bouton "Envoyer rappel" par enqueteur. '
                 'Alerte quota atteint. Templates email rappel.'),
            ]
        },
        {
            'titre': 'Phase 12 : Deploiement',
            'etapes': [
                ('12.1 - Configuration deploiement',
                 'Procfile : gunicorn marketym.wsgi. whitenoise pour static files. '
                 'Variables d\'environnement Railway/Render. '
                 'ALLOWED_HOSTS, CSRF_TRUSTED_ORIGINS. collectstatic dans le build.'),
            ]
        },
        {
            'titre': 'Phase 13 : Migration des donnees',
            'etapes': [
                ('13.1 - Script de migration',
                 'Script Python pour migrer les donnees existantes de Supabase vers les nouveaux '
                 'modeles Django. Verifier correspondance UUIDs. Migrer : enqueteurs, enquetes, '
                 'affectations, segmentations, quotas, completions, clics, historique.'),
            ]
        },
    ]

    for phase in phases:
        doc.add_heading(phase['titre'], level=2)
        for etape_titre, etape_desc in phase['etapes']:
            p = doc.add_paragraph()
            run = p.add_run(etape_titre)
            run.bold = True
            doc.add_paragraph(etape_desc)

    doc.add_page_break()

    # =========================================================================
    # ANNEXES
    # =========================================================================
    doc.add_heading('Annexes', level=1)

    doc.add_heading('A. Comptes de test', level=2)
    doc.add_paragraph('Super Admin : wilfredkouadjo006@gmail.com')
    doc.add_paragraph('Email expediteur : marketym@hcexecutive.net')

    doc.add_heading('B. Contraintes de performance', level=2)
    for item in [
        'Sync : max 1000 reponses par page QP',
        'Sync uniquement enquetes en_cours',
        'Cache answer_options en base',
        '5000 appels QP/mois maximum',
        '300 appels/60s rate limit',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('C. Securite', level=2)
    for item in [
        'HTTPS obligatoire en production',
        'CSRF protection sur tous les POST',
        'Sessions securisees (HttpOnly, SameSite=Lax)',
        'Mots de passe hashes bcrypt',
        'OTP a usage unique, expire en 5 min, max 3 tentatives',
        'Tokens d\'invitation expire en 48h',
        'Pas de secrets en dur (tout en variables d\'environnement)',
        'Validation de toutes les entrees via Django Forms',
        'Protection XSS via auto-escaping Django templates',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('D. Dependencies Python', level=2)
    for dep in [
        'Django >= 5.0',
        'psycopg2-binary',
        'bcrypt',
        'sib-api-v3-sdk',
        'httpx',
        'python-dotenv',
        'gunicorn',
        'whitenoise',
        'django-environ',
    ]:
        doc.add_paragraph(dep, style='List Bullet')

    return doc


if __name__ == '__main__':
    doc = create_document()
    output_path = '/Users/Apple/Desktop/H&C/Enquetes/Suivi enquetes/enquetes/CAHIER_DE_CHARGE.docx'
    doc.save(output_path)
    print(f'Document genere : {output_path}')
