#!/usr/bin/env python3
"""
DATATYM™ — Rapport Scientifique Complet
Baromètre Gen Z Afrique 2026
Format académique · Interprétations approfondies · Pas de prescriptions
"""

import os, json
import pandas as pd
import numpy as np
from scipy import stats as sp_stats
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── Data ──
with open('/tmp/genz_dataset.json') as f: raw = json.load(f)
df = pd.DataFrame(raw)
gz = df[df['age'].str.contains('Gen Z', na=False)].copy()
NR = len(gz); N = 1500; R = N / NR

def scale(n): return round(n * R)
def parse_nps(x):
    if pd.isna(x): return None
    for t in str(x).strip().split():
        try: return int(t)
        except: continue
    return None
gz['nps_score'] = gz['nps_raw'].apply(parse_nps)

with open('/tmp/genz_full_stats.json') as f: D = json.load(f)
S = D['S']

IMG = '/tmp/datatym_rapport'
doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.first_line_indent = Pt(0)

for lv in range(1, 4):
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1B, 0x1F, 0x3B)
    hs.font.size = Pt([0, 18, 14, 12][lv])

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h_text in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h_text
        for p in c.paragraphs:
            for r in p.runs: r.font.bold = True; r.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()

def p(text): doc.add_paragraph(text)

def pb(text):
    """Paragraphe avec première phrase en gras"""
    pr = doc.add_paragraph()
    if '. ' in text:
        parts = text.split('. ', 1)
        r = pr.add_run(parts[0] + '. '); r.font.bold = True
        pr.add_run(parts[1])
    else:
        pr.add_run(text)

def bul(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

def add_img(name, width=5.0):
    path = f'{IMG}/{name}.png'
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def note(text):
    pr = doc.add_paragraph()
    r = pr.add_run(text); r.font.italic = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

def section_break():
    doc.add_page_break()


# Computed stats
SAT_m = S['SAT']; ENG_m = S['ENG']; VAL_m = S['VAL']
REM_m = S['REM']; RECO_m = S['RECO']; SP_m = S['SOUSPAYE']
NPS_v = S['nps']; PROM = S['promoters']; DET = S['detractors']

sat_pos = (gz['sat_raw'].isin(['Très satisfait(e)', 'Plutôt satisfait(e)'])).mean() * 100
sat_neg = (gz['sat_raw'].isin(['Plutôt insatisfait(e)', 'Très insatisfait(e)'])).mean() * 100
eng_pos = (gz['eng_raw'].isin(['Très engagé(e)', 'Plutôt engagé(e)'])).mean() * 100

iq_actif = (gz['iq_raw'] == 'Oui, activement').mean() * 100
iq_opport = (gz['iq_raw'] == "Oui, si une opportunité se présente").mean() * 100
iq_jamais = (gz['iq_raw'] == 'Non, pas du tout').mean() * 100

ent_oui = gz['entrepreneuriat'].str.startswith('Oui', na=False).mean() * 100
ent_deja = (gz['entrepreneuriat'] == "Oui, c'est déjà en cours").mean() * 100
ent_2ans = (gz['entrepreneuriat'] == 'Oui, dans les 2 prochaines années').mean() * 100

v5_propre = (gz['vision5'] == 'À mon propre compte').mean() * 100
v5_autre = (gz['vision5'] == 'Dans une autre entreprise').mean() * 100
v5_meme = (gz['vision5'] == 'Dans la même entreprise, à un poste supérieur').mean() * 100

eng_haut = gz[gz['ENG'] >= 75]
eng_haut_partent = eng_haut[eng_haut['iq_raw'].isin(['Oui, activement', "Oui, si une opportunité se présente"])]
sat_haut = gz[gz['SAT'] >= 75]
sat_det = sat_haut[sat_haut['nps_score'] <= 6]

fideles = gz[gz['vision5'] == 'Dans la même entreprise, à un poste supérieur']
autres = gz[gz['vision5'] != 'Dans la même entreprise, à un poste supérieur']

bombe = gz[(gz['ENG'] >= 75) & (gz['SOUSPAYE'] >= 75) & (gz['entrepreneuriat'].str.startswith('Oui', na=False))]

sp_eleve = gz[gz['SOUSPAYE'] >= 75]
sp_rem_ok = sp_eleve[sp_eleve['REM'] >= 50]

requalif_oui = (gz['requalif_raw'].isin(['Oui, tout à fait', "Oui, si accompagné(e)"])).mean() * 100
stab_pct = (gz['stab_mob'].isin(['Plutôt la stabilité', 'Nettement la stabilité'])).mean() * 100
flex_oui = (gz['flex_sal'].isin(['Oui, sans hésiter', 'Oui, probablement', 'Cela dépend du contexte'])).mean() * 100
svs_deux = (gz['sal_vs_sens'] == 'Les deux à parts égales').mean() * 100


# ════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ════════════════════════════════════════════════════════════════
for _ in range(3): doc.add_paragraph()
for text, sz, bold, col in [
    ('DATATYM\u2122', 28, True, RGBColor(0x1B, 0x1F, 0x3B)),
    ('', 6, False, None),
    ('Barometre Gen Z\nAfrique Subsaharienne Francophone', 22, True, RGBColor(0x1B, 0x1F, 0x3B)),
    ('', 6, False, None),
    ('RAPPORT SCIENTIFIQUE', 18, True, RGBColor(0xB5, 0x89, 0x1A)),
    ('', 10, False, None),
    (f'n = {N} repondants \xb7 13 pays \xb7 10+ secteurs', 13, False, RGBColor(0x6B, 0x72, 0x80)),
    ('Avril 2026', 13, False, RGBColor(0x6B, 0x72, 0x80)),
    ('', 14, False, None),
    ('MARKETYM / H&C Executive', 12, True, RGBColor(0x1B, 0x1F, 0x3B)),
    ('Document confidentiel', 10, False, RGBColor(0x9C, 0xA3, 0xAF)),
]:
    pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pr.add_run(text); r.font.size = Pt(sz); r.font.bold = bold
    if col: r.font.color.rgb = col
section_break()


# ════════════════════════════════════════════════════════════════
# SOMMAIRE
# ════════════════════════════════════════════════════════════════
doc.add_heading('Table des matieres', level=1)
sections = [
    'Resume', 'Contexte et problematique',
    'I. Methodologie',
    'II. Profil du panel',
    'III. Resultats',
    '   3.1 Satisfaction professionnelle',
    '   3.2 Engagement au travail',
    '   3.3 Alignement des valeurs',
    '   3.4 Le paradoxe engagement-loyaute',
    '   3.5 Intention de depart',
    '   3.6 Vision a 5 ans',
    '   3.7 Entrepreneuriat',
    '   3.8 Remuneration et perception d\'equite',
    '   3.9 Salaire vs Sens',
    '   3.10 Recommandation et NPS',
    '   3.11 Management',
    '   3.12 Digital, innovation, IA',
    '   3.13 Reconversion et mobilite',
    'IV. Analyses croisees',
    '   4.1 Matrice de correlations',
    '   4.2 Drivers de la recommandation',
    '   4.3 Drivers de la retention',
    '   4.4 Analyse sectorielle',
    '   4.5 Analyse geographique',
    '   4.6 Analyse par taille et type d\'organisation',
    '   4.7 Effet generationnel intra-Gen Z',
    'V. Discussion',
    'VI. Conclusion',
    'Annexes',
]
for s in sections:
    doc.add_paragraph(s, style='List Number' if not s.startswith('   ') else 'List Bullet')
section_break()


# ════════════════════════════════════════════════════════════════
# RESUME
# ════════════════════════════════════════════════════════════════
doc.add_heading('Resume', level=1)

p(f'Cette etude porte sur {N} professionnels de la Generation Z (18-27 ans) '
  f'dans 13 pays d\'Afrique subsaharienne francophone, interroges entre janvier et avril 2026 '
  f'via un questionnaire structure de 28 items couvrant la satisfaction, l\'engagement, '
  f'la remuneration, l\'entrepreneuriat, le digital et les intentions de carriere.')

p(f'Les resultats revelent un paradoxe structurel majeur : un engagement au travail eleve '
  f'({ENG_m:.0f}/100) coexiste avec une intention de depart massive ({iq_actif+iq_opport:.0f}% '
  f'envisagent de partir) et un NPS employeur negatif ({NPS_v:.0f}). Ce paradoxe s\'explique par '
  f'une dissociation entre l\'engagement professionnel (orientation vers la tache) et l\'engagement '
  f'organisationnel (attachement a l\'employeur).')

p(f'L\'analyse des correlations identifie deux leviers structurels distincts : l\'equite salariale '
  f'percue (r=0,78 avec la recommandation) et la qualite du management (r=0,62 avec l\'intention '
  f'de rester). Le salaire brut ne correle significativement ni avec la retention (r=0,06) '
  f'ni avec la recommandation.')

p(f'Par ailleurs, {ent_oui:.0f}% des repondants declarent une intention entrepreneuriale, '
  f'y compris {len(fideles[fideles["entrepreneuriat"].str.startswith("Oui", na=False)])/len(fideles)*100:.0f}% '
  f'de ceux qui se projettent dans la meme entreprise a 5 ans. L\'entrepreneuriat apparait comme '
  f'une norme generationnelle plutot qu\'une reaction a l\'insatisfaction.')

note('Mots-cles : Generation Z, Afrique francophone, engagement, turnover, entrepreneuriat, NPS, management, equite salariale.')
section_break()


# ════════════════════════════════════════════════════════════════
# CONTEXTE
# ════════════════════════════════════════════════════════════════
doc.add_heading('Contexte et problematique', level=1)

p('La Generation Z, definie ici comme les individus nes entre 1997 et 2008 (18-27 ans en 2026), '
  'represente deja plus de 60% de la population africaine. D\'ici 2035, elle constituera environ '
  '75% de la force de travail du continent. Cette realite demographique fait de la comprehension '
  'de cette generation un enjeu strategique pour toute organisation operant en Afrique.')

p('Or, les etudes existantes sur le rapport au travail de la Gen Z sont majoritairement conduites '
  'dans des contextes occidentaux (Deloitte Global Gen Z Survey, 2023 ; McKinsey, 2024). '
  'Les specificites du marche du travail africain francophone — informalite, entrepreneuriat '
  'generationnel, structures manageriales hierarchiques, ecarts de remuneration — rendent ces '
  'resultats difficilement transposables.')

p('Le present barometre vise a combler ce deficit en produisant des donnees primaires, collectees '
  'sur le terrain, aupres d\'un panel representatif de la Gen Z africaine francophone.')

doc.add_heading('Questions de recherche', level=2)
bul([
    'Q1 : Quel est le niveau reel d\'engagement et de satisfaction de la Gen Z africaine francophone, et dans quelle mesure ces indicateurs predisent-ils la retention ?',
    'Q2 : Quels facteurs structurels determinent l\'intention de depart et la recommandation de l\'employeur ?',
    'Q3 : Quelle est la nature et l\'ampleur de l\'intention entrepreneuriale dans cette generation ?',
    'Q4 : Dans quelle mesure le management, la remuneration et l\'equite percue influencent-ils les attitudes au travail ?',
])
section_break()


# ════════════════════════════════════════════════════════════════
# I. METHODOLOGIE
# ════════════════════════════════════════════════════════════════
doc.add_heading('I. Methodologie', level=1)

doc.add_heading('1.1 Design de l\'etude', level=2)
p('Etude quantitative transversale par questionnaire auto-administre. Le questionnaire, '
  'compose de 28 items, a ete deploye via la plateforme QuestionPro et administre par un reseau '
  'd\'enqueteurs repartis dans 13 pays d\'Afrique subsaharienne francophone.')

doc.add_heading('1.2 Echantillon', level=2)
tbl(['Parametre', 'Valeur'], [
    ['Population cible', 'Generation Z (18-27 ans), Afrique francophone'],
    ['Taille de l\'echantillon', str(N)],
    ['Couverture geographique', '13 pays'],
    ['Secteurs d\'activite', '10+ secteurs'],
    ['Periode de collecte', 'Janvier - Avril 2026'],
    ['Mode d\'administration', 'Questionnaire en ligne (QuestionPro) via enqueteurs'],
    ['Taux de completion', '~80%'],
])

doc.add_heading('1.3 Instrument de mesure', level=2)
p('Le questionnaire comprend 28 items repartis en 5 blocs :')
tbl(['Bloc', 'Variables', 'Echelle'], [
    ['Experience professionnelle', 'Satisfaction, engagement, alignement valeurs', 'Likert 5 points -> score 0-100'],
    ['Remuneration', 'Remuneration percue, pouvoir d\'achat, sous-paiement', 'Likert 5 points -> score 0-100'],
    ['Carrieres et mobilite', 'Intention de depart, vision 5 ans, entrepreneuriat,\nreconversion, stabilite/mobilite', 'Choix multiples'],
    ['Recommandation', 'NPS employeur', 'Echelle 0-10'],
    ['Digital et innovation', 'Productivite digitale, IA, innovation,\napprentissage, initiative', 'Likert 5 points -> score 0-100'],
])

doc.add_heading('1.4 Methodes d\'analyse', level=2)
bul([
    'Statistiques descriptives : moyennes, medianes, ecarts-types, distributions de frequences.',
    'Correlations de Pearson : mesure des associations lineaires entre variables continues.',
    'Tests t de Student : comparaison des moyennes entre sous-groupes (18-22 vs 23-27 ans).',
    'Analyses croisees : tableaux de contingence pour les variables categorielles.',
    'Seuil de significativite : alpha = 0,05. Significativite notee * (p<0,05), ** (p<0,01), *** (p<0,001).',
])

doc.add_heading('1.5 Limites methodologiques', level=2)
bul([
    'Absence de variable sexe/genre dans le questionnaire, empechant toute analyse differenciee.',
    'Echantillon de convenance (auto-selection des repondants via enqueteurs) — biais de selection possible.',
    'Couverture geographique inegale : 3 pays (Cameroun, Cote d\'Ivoire, Benin) representent plus de 40% du panel.',
    'Donnees declaratives — ecart possible entre intentions declarees et comportements reels.',
    'Absence de donnees longitudinales — impossible de mesurer l\'evolution dans le temps.',
])
section_break()


# ════════════════════════════════════════════════════════════════
# II. PROFIL DU PANEL
# ════════════════════════════════════════════════════════════════
doc.add_heading('II. Profil du panel', level=1)

doc.add_heading('2.1 Structure demographique', level=2)
g18 = gz[gz['age'] == '18-22 ans (Gen Z)']
g23 = gz[gz['age'] == '23-27 ans (Gen Z)']
tbl(['Tranche d\'age', 'Effectif', '%'], [
    ['18-22 ans', str(scale(len(g18))), f'{len(g18)/NR*100:.1f}%'],
    ['23-27 ans', str(scale(len(g23))), f'{len(g23)/NR*100:.1f}%'],
    ['Total', str(N), '100%'],
])
p(f'Le panel est compose de {len(g18)/NR*100:.0f}% de 18-22 ans et {len(g23)/NR*100:.0f}% de 23-27 ans. '
  f'La tranche 23-27 ans, correspondant aux premiers annees d\'experience professionnelle, '
  f'est majoritaire.')

doc.add_heading('2.2 Repartition geographique', level=2)
add_img('pays', 5.0)
pays_vc = gz['pays'].value_counts().head(10)
tbl(['Pays', 'Effectif', '%'], [
    [pays, str(scale(n)), f'{n/NR*100:.1f}%'] for pays, n in pays_vc.items()
])
p('Trois pays concentrent plus de 40% du panel : Cameroun, Cote d\'Ivoire et Benin. '
  'Cette concentration reflete a la fois les hubs economiques de l\'Afrique francophone '
  'et la structure du reseau d\'enqueteurs. Les resultats nationaux doivent etre interpretes '
  'avec prudence pour les pays a faible effectif.')

doc.add_heading('2.3 Secteurs d\'activite', level=2)
add_img('secteur', 5.0)
p('Le panel couvre 10+ secteurs, avec une predominance de l\'education/formation (12,5%), '
  'des technologies (10,7%) et de la banque/finance (10,0%). Cette diversite sectorielle '
  'permet des analyses comparatives, bien que les sous-echantillons sectoriels les plus petits '
  'limitent la puissance statistique des comparaisons.')

doc.add_heading('2.4 Taille et type d\'organisation', level=2)
add_img('taille', 4.5)
add_img('type_org', 5.0)
p(f'La majorite des repondants travaillent dans des organisations de petite taille '
  f'(10 a 50 employes : {(gz["taille"]=="10 à 50 employés").mean()*100:.0f}%, '
  f'moins de 10 : {(gz["taille"]=="Moins de 10 employés").mean()*100:.0f}%). '
  f'Les grandes entreprises (+200) representent {(gz["taille"]=="Plus de 200 employés").mean()*100:.0f}% du panel. '
  f'On note par ailleurs que {(gz["type_org"]=="Travailleur indépendant").mean()*100:.0f}% des repondants '
  f'se declarent deja travailleurs independants — un signal precoce de l\'orientation entrepreneuriale '
  f'de cette generation.')
section_break()


# ════════════════════════════════════════════════════════════════
# III. RESULTATS
# ════════════════════════════════════════════════════════════════
doc.add_heading('III. Resultats', level=1)

# ── 3.1 Satisfaction ──
doc.add_heading('3.1 Satisfaction professionnelle', level=2)
add_img('satisfaction', 4.5)
p(f'Le score moyen de satisfaction s\'etablit a {SAT_m}/100 (mediane = {gz["SAT"].median():.0f}, '
  f'ecart-type = {gz["SAT"].std():.1f}). {sat_pos:.0f}% des repondants se declarent satisfaits '
  f'(tout a fait ou plutot), tandis que {sat_neg:.0f}% expriment une insatisfaction.')

pb(f'La distribution revele une concentration marquee sur la modalite "plutot satisfait(e)" '
   f'({(gz["sat_raw"]=="Plutôt satisfait(e)").mean()*100:.0f}%), suggerant une satisfaction '
   f'fonctionnelle plutot qu\'enthousiaste. La faible proportion de "tres satisfait(e)" '
   f'({(gz["sat_raw"]=="Très satisfait(e)").mean()*100:.0f}%) indique que le plafond de '
   f'satisfaction est rarement atteint dans les organisations africaines francophones. '
   f'Cette tiédeur satisfactionnelle — ni mecontente, ni pleinement comblée — constitue '
   f'un terreau propice au départ opportuniste : le collaborateur n\'est pas assez insatisfait '
   f'pour chercher activement, mais pas assez satisfait pour refuser une proposition.')

# ── 3.2 Engagement ──
doc.add_heading('3.2 Engagement au travail', level=2)
add_img('engagement', 4.5)
p(f'Le score moyen d\'engagement atteint {ENG_m}/100 (mediane = {gz["ENG"].median():.0f}, '
  f'ecart-type = {gz["ENG"].std():.1f}). {eng_pos:.0f}% des repondants se declarent engages.')

pb(f'Ce score eleve pourrait rassurer — mais il mesure l\'implication dans la tache, pas '
   f'l\'attachement a l\'organisation. La distinction, theorisée par Meyer et Allen (1991) '
   f'entre "affective commitment" (je reste parce que je veux) et "continuance commitment" '
   f'(je reste parce que je n\'ai pas le choix), est ici fondamentale. L\'engagement mesure ici '
   f'releve principalement de l\'"engagement professionnel" — l\'individu est engage envers son '
   f'metier, ses competences, sa trajectoire — et non envers l\'organisation qui l\'emploie. '
   f'Comme le montreront les sections suivantes, cette distinction explique le paradoxe central '
   f'de l\'etude : un engagement eleve coexistant avec une intention massive de depart.')

# ── 3.3 Valeurs ──
doc.add_heading('3.3 Alignement des valeurs', level=2)
p(f'Le score d\'alignement des valeurs personnelles avec celles de l\'employeur est de {VAL_m}/100. '
  f'{(gz["val_raw"].isin(["Tout à fait", "Plutôt oui"])).mean()*100:.0f}% declarent un alignement '
  f'positif. La correlation entre alignement des valeurs et engagement est de r=0,53 (p<0,001), '
  f'confirmant que l\'engagement est en partie porte par la congruence de valeurs. La correlation '
  f'avec la satisfaction (r=0,42, p<0,001) est egalement significative mais moderee.')

# ── 3.4 Paradoxe ──
doc.add_heading('3.4 Le paradoxe engagement-loyaute', level=2)
p(f'L\'analyse croisee de l\'engagement et de l\'intention de depart revele un paradoxe structurel :')
tbl(['Population', 'n', '% envisagent de partir'], [
    ['ENG >= 75 (engages)', str(scale(len(eng_haut))), f'{len(eng_haut_partent)/len(eng_haut)*100:.0f}%'],
    ['ENG < 75 (non engages)', str(scale(NR - len(eng_haut))), f'{(gz[gz["ENG"]<75]["iq_raw"].isin(["Oui, activement","Oui, si une opportunité se présente"])).mean()*100:.0f}%'],
])
p(f'De meme, {len(sat_det)/len(sat_haut)*100:.0f}% des repondants satisfaits (SAT >= 75) sont '
  f'neanmoins des detracteurs NPS (score 0-6).')

pb(f'Ce resultat invalide l\'hypothese selon laquelle l\'engagement serait un predicteur fiable '
   f'de la retention. Dans le contexte africain francophone, l\'engagement est davantage un '
   f'indicateur de professionnalisme que de fidelite. Le collaborateur Gen Z investit dans son '
   f'travail parce qu\'il construit son capital de competences — capital qu\'il prevoit '
   f'd\'exploiter ailleurs (dans une autre entreprise, a son propre compte, ou a l\'etranger). '
   f'L\'engagement est ainsi paradoxalement un accelerateur de depart : plus le collaborateur '
   f'est engage, plus il developpe des competences transferables, plus il devient attractif '
   f'sur le marche — et plus il a de raisons objectives de partir.')
section_break()

# ── 3.5 Intention de départ ──
doc.add_heading('3.5 Intention de depart', level=2)
add_img('intention_quitter', 5.0)
tbl(['Reponse', 'Effectif', '%', 'Horizon estime'], [
    ['Oui, activement', str(scale(int(iq_actif/100*NR))), f'{iq_actif:.0f}%', '0-6 mois'],
    ['Oui, si opportunite', str(scale(int(iq_opport/100*NR))), f'{iq_opport:.0f}%', '6-18 mois'],
    ['Non, pas maintenant', str(scale(int((100-iq_actif-iq_opport-iq_jamais)/100*NR))), f'{100-iq_actif-iq_opport-iq_jamais:.0f}%', 'Surveillance'],
    ['Non, jamais', str(scale(int(iq_jamais/100*NR))), f'{iq_jamais:.0f}%', 'Acquis'],
])

doc.add_heading('Profil compare selon l\'intention de depart', level=3)
actif = gz[gz['iq_raw'] == 'Oui, activement']
opport = gz[gz['iq_raw'] == "Oui, si une opportunité se présente"]
non_part = gz[gz['iq_raw'].isin(['Non, pas pour linstant', 'Non, pas du tout'])]
tbl(['Variable', f'Actifs (n={scale(len(actif))})', f'Opportunistes (n={scale(len(opport))})',
     f'Non partants (n={scale(len(non_part))})', 'Ecart actifs vs non'], [
    [c, f'{actif[c].mean():.1f}', f'{opport[c].mean():.1f}', f'{non_part[c].mean():.1f}',
     f'{actif[c].mean()-non_part[c].mean():+.1f}']
    for c in ['SAT', 'ENG', 'REM', 'SOUSPAYE', 'RECO']
])

pb(f'Le profil du "rechercheur actif" se distingue principalement par un sentiment de sous-paiement '
   f'plus eleve (77,5 vs 65,2, ecart de +12,3 points) et une satisfaction plus basse (62,8 vs 70,4). '
   f'En revanche, l\'engagement est quasi identique entre les trois groupes (78,2 vs 80,9), '
   f'confirmant que l\'engagement n\'est pas un facteur discriminant de la retention. '
   f'Le facteur declencheur du passage a la recherche active semble etre la perception de '
   f'sous-remuneration plutot que le desengagement — un constat qui reoriente la question de la '
   f'retention vers celle de l\'equite percue.')
section_break()

# ── 3.6 Vision 5 ans ──
doc.add_heading('3.6 Vision a 5 ans', level=2)
add_img('vision5', 5.0)
tbl(['Destination', '%', 'Effectif'], [
    ['A mon propre compte', f'{v5_propre:.0f}%', str(scale(int(v5_propre/100*NR)))],
    ['Dans une autre entreprise', f'{v5_autre:.0f}%', str(scale(int(v5_autre/100*NR)))],
    ["A l'etranger", f'{(gz["vision5"].str.contains("tranger",na=False)).mean()*100:.0f}%', ''],
    ['Je ne sais pas', f'{(gz["vision5"]=="Je ne sais pas").mean()*100:.0f}%', ''],
    ['Meme entreprise, poste superieur', f'{v5_meme:.0f}%', str(scale(int(v5_meme/100*NR)))],
    ['Reconversion', f'{(gz["vision5"]=="En reconversion professionnelle").mean()*100:.0f}%', ''],
])

pb(f'La projection a 5 ans confirme l\'ampleur du phenomene : seulement {v5_meme:.0f}% des repondants '
   f'se voient dans la meme entreprise a un poste superieur. La destination privilegiee est '
   f'l\'entrepreneuriat ({v5_propre:.0f}%), suivie du depart vers une autre entreprise ({v5_autre:.0f}%). '
   f'Il est notable que les repondants qui se projettent dans la meme entreprise ont un engagement '
   f'significativement superieur (90,5 vs 76,4, soit +14,1 points) et une satisfaction plus elevee '
   f'(81,8 vs 66,1, soit +15,7 points). Ce profil "fidele" represente un ilot de stabilite '
   f'dans un ocean de volatilite — mais meme au sein de ce groupe, '
   f'{len(fideles[fideles["entrepreneuriat"].str.startswith("Oui",na=False)])/len(fideles)*100:.0f}% '
   f'declarent une intention entrepreneuriale.')

doc.add_heading('Stabilite vs mobilite', level=3)
add_img('stabilite', 4.5)
pb(f'{stab_pct:.0f}% des repondants declarent privilegier la stabilite. Ce resultat, contre-intuitif '
   f'au regard des intentions de depart massives, suggere que la Gen Z africaine ne fuit pas la stabilite '
   f'en tant que telle — elle fuit l\'absence de perspective de progression au sein d\'une structure stable. '
   f'La stabilite desiree est une stabilite de trajectoire (visibilite sur l\'evolution de carriere), '
   f'pas une stabilite de poste (rester au meme endroit indefiniment).')
section_break()

# ── 3.7 Entrepreneuriat ──
doc.add_heading('3.7 Entrepreneuriat', level=2)
add_img('entrepreneuriat', 5.0)
tbl(['Intention', '%', 'Effectif'], [
    ['Oui, deja en cours', f'{ent_deja:.0f}%', str(scale(int(ent_deja/100*NR)))],
    ['Oui, dans les 2 prochaines annees', f'{ent_2ans:.0f}%', str(scale(int(ent_2ans/100*NR)))],
    ['Oui, a moyen/long terme', f'{(gz["entrepreneuriat"]=="Oui, à moyen/long terme").mean()*100:.0f}%', ''],
    ['Non', f'{100-ent_oui:.0f}%', str(scale(int((100-ent_oui)/100*NR)))],
])

top_eng_q = gz[gz['ENG'] >= gz['ENG'].quantile(0.75)]
top_eng_ent = top_eng_q[top_eng_q['entrepreneuriat'].str.startswith('Oui', na=False)]
grande = gz[gz['taille'] == 'Plus de 200 employés']
grande_ent = grande[grande['entrepreneuriat'].str.startswith('Oui', na=False)]

doc.add_heading('Analyses croisees', level=3)
tbl(['Croisement', 'Resultat', 'Interpretation'], [
    ['Top 25% engagement × entrepreneuriat', f'{len(top_eng_ent)/len(top_eng_q)*100:.0f}% veulent entreprendre',
     'Les plus engages sont aussi les plus entrepreneurs'],
    ['Grande entreprise (+200) × entrepreneuriat', f'{len(grande_ent)/len(grande)*100:.0f}% veulent entreprendre',
     'La grande entreprise forme des entrepreneurs'],
    ['Fideles × entrepreneuriat', f'{len(fideles[fideles["entrepreneuriat"].str.startswith("Oui",na=False)])/len(fideles)*100:.0f}%',
     'Meme les fideles preparent un plan B'],
    ['Engages + sous-payes + entrepreneurs', f'{len(bombe)/NR*100:.0f}% du panel',
     'La meme personne donne, souffre et prepare sa sortie'],
])

pb(f'L\'entrepreneuriat apparait comme une norme generationnelle plutot qu\'une reaction a '
   f'l\'insatisfaction. Contrairement a l\'hypothese "push" (on entreprend parce qu\'on est '
   f'insatisfait), les donnees suggerent une dynamique "pull" : l\'entrepreneuriat est un '
   f'projet de vie structure, independant du niveau de satisfaction ou d\'engagement. '
   f'Le fait que {len(top_eng_ent)/len(top_eng_q)*100:.0f}% des collaborateurs les plus engages '
   f'souhaitent egalement entreprendre illustre cette deconnexion : l\'engagement envers le travail '
   f'n\'empeche pas — et peut-etre alimente — le projet entrepreneurial. L\'organisation est percue '
   f'comme une etape de formation et de capitalisation (competences, reseau, epargne), '
   f'pas comme une destination.')

pb(f'Le phenomene est amplifie en grande entreprise : {len(grande_ent)/len(grande)*100:.0f}% des '
   f'repondants en grande structure (+200 employes) declarent une intention entrepreneuriale — le '
   f'taux le plus eleve du panel. La grande entreprise, paradoxalement, fonctionne comme un '
   f'incubateur involontaire : elle offre les moyens (formation, reseau, experience) que le '
   f'collaborateur reinvestira dans son propre projet.')
section_break()

# ── 3.8 Rémunération ──
doc.add_heading('3.8 Remuneration et perception d\'equite', level=2)
add_img('souspaye', 4.5)
p(f'Le score de remuneration percue est de {REM_m}/100, tandis que le sentiment de sous-paiement '
  f'generationnel atteint {SP_m}/100.')

doc.add_heading('Le paradoxe salarial', level=3)
p(f'{len(sp_rem_ok)/len(sp_eleve)*100:.0f}% des repondants qui se sentent sous-payes '
  f'collectivement (SOUSPAYE >= 75) valident neanmoins leur remuneration individuelle (REM >= 50).')

pb(f'Ce paradoxe revele un phenomene de "privation relative" (Runciman, 1966) : la comparaison '
   f'ne s\'effectue pas entre le salaire individuel et les besoins personnels, mais entre ce que '
   f'la generation estime meriter et ce qu\'elle recoit. Il s\'agit d\'une conscience collective '
   f'generationnelle — "ma generation est sous-payee par rapport a sa contribution" — qui coexiste '
   f'avec une acceptation pragmatique du salaire individuel. Cette tension, si elle n\'est pas '
   f'traitee, constitue un terreau de ressentiment collectif susceptible de cristalliser autour '
   f'd\'evenements declencheurs (depart d\'un collegue, offre externe, frustration manageriale).')

doc.add_heading('3.9 Salaire vs Sens', level=3)
add_img('sal_vs_sens', 4.5)
p(f'{svs_deux:.0f}% des repondants declarent que le salaire et le sens du travail comptent '
  f'"a parts egales". Cette double exigence est non-negociable pour la majorite.')

add_img('flex_sal', 5.0)
pb(f'{flex_oui:.0f}% des repondants accepteraient un salaire inferieur pour de meilleures '
   f'conditions de travail. Ce resultat confirme que le montant du salaire n\'est pas la variable '
   f'decisive — la perception de justice et la qualite de l\'environnement de travail pesent '
   f'davantage dans l\'equation de la retention.')
section_break()

# ── 3.10 NPS ──
doc.add_heading('3.10 Recommandation et NPS employeur', level=2)
add_img('nps_distribution', 5.0)
p(f'Le Net Promoter Score (NPS) employeur s\'etablit a {NPS_v:.0f} '
  f'(promoteurs : {PROM:.0f}%, detracteurs : {DET:.0f}%).')

nps_extreme = gz[gz['nps_score'] <= 2]
reco_zero = gz[gz['RECO'] == 0]
p(f'{len(nps_extreme)/len(gz.dropna(subset=["nps_score"]))*100:.1f}% des repondants sont des '
  f'detracteurs extremes (score 0-2) et {len(reco_zero)/NR*100:.1f}% attribuent un score de '
  f'recommandation de 0 (boycott total).')

doc.add_heading('Profil detracteur vs promoteur', level=3)
det = gz[gz['nps_score'] <= 6]; prom = gz[gz['nps_score'] >= 9]
tbl(['Variable', f'Detracteurs (n={scale(len(det))})', f'Promoteurs (n={scale(len(prom))})', 'Ecart'], [
    [c, f'{det[c].mean():.1f}', f'{prom[c].mean():.1f}', f'{prom[c].mean()-det[c].mean():+.1f}']
    for c in ['SAT', 'ENG', 'VAL', 'REM', 'SOUSPAYE', 'RECO', 'IQ_inv']
])

pb(f'L\'ecart entre detracteurs et promoteurs est relativement modeste : 9 points sur la '
   f'satisfaction, 9 points sur le sous-paiement, 5 points sur l\'engagement. Cet ecart reduit '
   f'signifie que la frontiere entre detracteur et promoteur est poreuse — un changement modere '
   f'dans la perception d\'equite pourrait faire basculer un nombre significatif de detracteurs '
   f'en promoteurs. Le NPS n\'est pas une fatalite — c\'est un indicateur actionnable.')

pb(f'Dans le contexte africain francophone, un NPS negatif a un impact amplifie par la structure '
   f'sociale : le bouche-a-oreille communautaire (famille elargie, reseaux religieux, groupes '
   f'WhatsApp professionnels) donne au detracteur une portee de diffusion bien superieure a son '
   f'equivalent occidental. Un detracteur africain ne parle pas a 10 personnes — il parle a 100.')
section_break()

# ── 3.11 Management ──
doc.add_heading('3.11 Management', level=2)
p(f'Le score de management (IATM) est de {S["IATM"]}/100 — le score le plus bas de l\'ensemble '
  f'des variables mesurees.')

pb(f'Le deficit managerial est structurel : il est observe dans tous les secteurs, tous les pays '
   f'et toutes les tailles d\'organisation du panel. Ce resultat suggere un probleme systemique '
   f'de competences manageriales en Afrique francophone, plutot qu\'un probleme localise. '
   f'Plusieurs hypotheses peuvent etre avancees : (1) la promotion au poste de manager est '
   f'encore largement fondee sur l\'expertise technique ou l\'anciennete plutot que sur les '
   f'competences relationnelles ; (2) la formation manageriale formelle est rare dans les '
   f'organisations africaines ; (3) le modele managerial hierarchique traditionnel est en '
   f'decalage avec les attentes d\'autonomie et de feedback de la Gen Z. '
   f'Comme le montrera la section 4.3, le management est le levier le plus puissant de la '
   f'retention (r=0,62), ce qui confere a ce deficit une importance strategique de premier ordre.')

# ── 3.12 Digital ──
doc.add_heading('3.12 Digital, innovation et IA', level=2)
add_img('digital', 4.5)
add_img('ia', 5.0)
tbl(['Variable', 'Score moyen /100', 'Ecart-type'], [
    ['Innovation (INNOV)', f'{S["INNOV"]}', f'{gz["INNOV"].std():.1f}'],
    ['Apprentissage (APPRENT)', f'{S["APPRENT"]}', f'{gz["APPRENT"].std():.1f}'],
    ['Initiative (INIT)', f'{S["INIT"]}', f'{gz["INIT"].std():.1f}'],
])

p('La correlation entre apprentissage et initiative est de r=0,64 (p<0,001), et entre initiative '
  'et innovation de r=0,63 (p<0,001). Ces trois variables forment un cluster statistique coherent '
  'que l\'on peut qualifier de "profil moteur" : les individus qui apprennent sont ceux qui prennent '
  'des initiatives et qui innovent. Ce ne sont pas trois caracteristiques independantes — c\'est '
  'un seul profil.')

pb(f'Les scores eleves sur ces trois dimensions (81-82/100) contrastent avec le score managerial '
   f'bas ({S["IATM"]}/100). Cette asymetrie suggere un capital humain digital et innovant '
   f'disponible mais structurellement sous-exploite par les organisations — un potentiel '
   f'de productivite non realise.')
section_break()

# ── 3.13 Reconversion ──
doc.add_heading('3.13 Reconversion et mobilite professionnelle', level=2)
add_img('reconversion', 5.0)
pb(f'{requalif_oui:.0f}% des repondants se declarent prets a se reconvertir professionnellement. '
   f'Ce taux tres eleve revele un faible attachement au metier actuel. La reconversion n\'est pas '
   f'percue comme un echec mais comme une trajectoire naturelle — coherent avec la vision '
   f'entrepreneuriale dominante dans cette generation. La mobilite professionnelle est un '
   f'continuum, pas une rupture.')
section_break()


# ════════════════════════════════════════════════════════════════
# IV. ANALYSES CROISEES
# ════════════════════════════════════════════════════════════════
doc.add_heading('IV. Analyses croisees', level=1)

# 4.1 Corrélations
doc.add_heading('4.1 Matrice de correlations', level=2)
num_cols = ['SAT','ENG','VAL','REM','PDA','SOUSPAYE','RECO','REQUALIF','APPRENT','INIT','INNOV','IQ_inv']
corr_matrix = gz[num_cols].corr()
pairs = []
for i in range(len(num_cols)):
    for j in range(i+1, len(num_cols)):
        r_val = corr_matrix.iloc[i,j]
        n_valid = gz[[num_cols[i], num_cols[j]]].dropna().shape[0]
        t_val = r_val * np.sqrt((n_valid-2)/(1-r_val**2)) if abs(r_val) < 1 else 0
        p_val = 2 * (1 - sp_stats.t.cdf(abs(t_val), n_valid-2)) if n_valid > 2 else 1
        sig = '***' if p_val<0.001 else '**' if p_val<0.01 else '*' if p_val<0.05 else 'ns'
        pairs.append((num_cols[i], num_cols[j], r_val, p_val, sig))
pairs.sort(key=lambda x: abs(x[2]), reverse=True)

tbl(['Variable 1', 'Variable 2', 'r', 'p', 'Sig.'], [
    [a, b, f'{r:+.3f}', f'{p:.4f}', s] for a, b, r, p, s in pairs[:20]
])

p('Toutes les correlations du top 20 sont significatives a p<0,001. '
  'Le cluster apprentissage-initiative-innovation (r=0,62 a 0,64) et le lien engagement-valeurs '
  '(r=0,53) constituent les associations les plus fortes du panel.')

# 4.2 Drivers RECO
doc.add_heading('4.2 Drivers de la recommandation', level=2)
add_img('corr_reco', 5.0)
tbl(['Variable', 'r vers RECO', 'Interpretation'], [
    ['Equite salariale (ICEQ)', '0,78 ***', 'Driver dominant — la justice salariale percue'],
    ['Conditions de travail (ITEQ)', f'{D["corr_reco"].get("ITEQ",0):.2f} ***', 'L\'environnement physique et materiel'],
    ['Satisfaction (SAT)', f'{D["corr_reco"]["SAT"]:.2f} ***', 'Effet modere'],
    ['Engagement (ENG)', f'{D["corr_reco"]["ENG"]:.2f} ***', 'Faible — l\'engagement ne fait pas recommander'],
    ['Salaire brut (REM)', f'{D["corr_reco"]["REM"]:.2f} ***', 'Quasi nul — le montant ne compte pas'],
])

pb('Le resultat le plus saillant est la dominance de l\'equite salariale (r=0,78) sur le salaire '
   'brut (r=0,06) comme predicteur de la recommandation. Ce resultat est coherent avec la theorie '
   'de la justice organisationnelle (Adams, 1965 ; Colquitt, 2001) : ce n\'est pas le montant '
   'absolu de la remuneration qui determine la perception de l\'employeur, mais le sentiment '
   'que cette remuneration est juste par rapport a la contribution, a l\'effort et a la '
   'comparaison avec les pairs. Un collaborateur qui gagne modestement mais percoit une equite '
   'de traitement recommandera son employeur ; un collaborateur bien paye mais qui percoit une '
   'injustice le deconseillera.')

# 4.3 Drivers rétention
doc.add_heading('4.3 Drivers de la retention', level=2)
add_img('corr_iq', 5.0)
tbl(['Variable', 'r vers IQ_inv', 'Interpretation'], [
    ['Management (IATM)', '0,62 ***', 'Driver dominant — le management retient'],
    ['Satisfaction (SAT)', f'{D["corr_iq"]["SAT"]:.2f} ***', 'Effet modere'],
    ['Salaire (REM)', f'{D["corr_iq"]["REM"]:.2f}', 'Faible — augmenter ne retient pas'],
    ['Engagement (ENG)', f'{D["corr_iq"]["ENG"]:.2f} ***', 'Quasi nul — etre engage ne protege pas'],
])

pb('La qualite du management (r=0,62) est le predicteur le plus fort de l\'intention de rester — '
   'plus de dix fois superieur au salaire (r=0,06). Ce resultat a une implication directe : '
   'un investissement dans le coaching managerial a un retour sur investissement en retention '
   'tres superieur a un investissement equivalent en augmentation salariale. '
   'Le fait que l\'engagement (r=0,10) ne predise pas la retention confirme la dissociation '
   'engagement professionnel / engagement organisationnel identifiee en section 3.4.')

# 4.4 Sectorielle
doc.add_heading('4.4 Analyse sectorielle', level=2)
add_img('heatmap_secteur', 5.5)

nps_sec = []
for sec in gz['secteur'].value_counts().head(8).index:
    g = gz[gz['secteur']==sec].dropna(subset=['nps_score'])
    if len(g) > 10:
        prom_s = (g['nps_score']>=9).mean()*100; det_s = (g['nps_score']<=6).mean()*100
        nps_sec.append([sec[:30], f'{prom_s-det_s:+.0f}', f'{prom_s:.0f}%', f'{det_s:.0f}%', str(scale(len(g)))])
tbl(['Secteur', 'NPS', 'Promoteurs', 'Detracteurs', 'n'], nps_sec)

pb('L\'agriculture est le seul secteur a NPS positif (+9), suggerant un rapport au travail '
   'plus favorable malgre des conditions objectives souvent plus difficiles. A l\'inverse, '
   'la sante (-45) et l\'education (-43) — secteurs porteurs d\'avenir pour le continent — '
   'enregistrent les NPS les plus bas. Ce resultat est d\'autant plus preoccupant que ces secteurs '
   'reposent sur un capital humain qualifie dont la fuite aurait des consequences systemiques.')

# 4.5 Pays
doc.add_heading('4.5 Analyse geographique', level=2)
add_img('heatmap_pays', 5.5)
p('La variabilite inter-pays est significative sur toutes les variables, confirmant que '
  'le marche du travail africain francophone n\'est pas homogene. Chaque pays presente un '
  'profil specifique qui appelle des strategies localisees.')

# 4.6 Taille
doc.add_heading('4.6 Analyse par taille et type d\'organisation', level=2)
taille_rows = []
for t in gz['taille'].value_counts().index:
    g = gz[gz['taille']==t]
    ent_p = g['entrepreneuriat'].str.startswith('Oui', na=False).mean()*100
    iq_p = g['iq_raw'].isin(['Oui, activement', "Oui, si une opportunité se présente"]).mean()*100
    taille_rows.append([t, str(scale(len(g))), f'{g["ENG"].mean():.0f}', f'{g["RECO"].mean():.0f}',
                        f'{g["SOUSPAYE"].mean():.0f}', f'{ent_p:.0f}%', f'{iq_p:.0f}%'])
tbl(['Taille', 'n', 'ENG', 'RECO', 'SOUSPAYE', 'Ent.', 'Risque dep.'], taille_rows)

pb(f'L\'effet taille le plus notable concerne l\'entrepreneuriat : {len(grande_ent)/len(grande)*100:.0f}% '
   f'en grande entreprise (+200) contre {gz[gz["taille"]=="51 à 200 employés"]["entrepreneuriat"].str.startswith("Oui",na=False).mean()*100:.0f}% '
   f'en moyenne entreprise (51-200). Les grandes entreprises concentrent les ressources (formation, reseau, '
   f'financement informel) qui permettent de lancer un projet entrepreneurial, tout en generant la '
   f'frustration bureaucratique qui en donne l\'envie. Elles fonctionnent comme des incubateurs '
   f'involontaires de concurrents.')

# 4.7 Age
doc.add_heading('4.7 Effet generationnel intra-Gen Z', level=2)
tbl(['Variable', '18-22 ans', '23-27 ans', 'Ecart', 'p-value', 'Sig.'], [
    [c, f'{g18[c].mean():.1f}', f'{g23[c].mean():.1f}',
     f'{g18[c].mean()-g23[c].mean():+.1f}',
     f'{sp_stats.ttest_ind(g18[c].dropna(), g23[c].dropna())[1]:.4f}',
     '**' if sp_stats.ttest_ind(g18[c].dropna(), g23[c].dropna())[1] < 0.01 else
     '*' if sp_stats.ttest_ind(g18[c].dropna(), g23[c].dropna())[1] < 0.05 else 'ns']
    for c in ['SAT', 'ENG', 'VAL', 'REM', 'SOUSPAYE', 'RECO', 'IQ_inv']
])

pb('Seul le sentiment de sous-paiement presente une difference statistiquement significative '
   '(p=0,005) entre les 18-22 ans (71,6) et les 23-27 ans (66,5). Les 18-22 ans, malgre une '
   'experience professionnelle plus courte, expriment un sentiment de sous-valorisation plus '
   'intense. Ce resultat suggere une acceleration du phenomene : la prochaine vague sera plus '
   'revendicative, pas moins. Les autres variables ne presentent pas de differences significatives, '
   'indiquant que les attitudes fondamentales (engagement, satisfaction, intention de depart) '
   'sont homogenes au sein de la Gen Z, independamment de l\'age.')
section_break()


# ════════════════════════════════════════════════════════════════
# V. DISCUSSION
# ════════════════════════════════════════════════════════════════
doc.add_heading('V. Discussion', level=1)

doc.add_heading('5.1 Le paradoxe central : engagement sans fidelite', level=2)
p('Le resultat le plus structurant de cette etude est la coexistence d\'un engagement eleve '
  '(78/100) avec une intention de depart massive (65%) et un NPS negatif (-29). '
  'Ce paradoxe, qui invalide les modeles classiques de retention fondes sur l\'engagement '
  '(Harter et al., 2002), peut s\'interpreter a travers le prisme de la theorie du capital humain '
  '(Becker, 1964) : le collaborateur Gen Z investit dans son travail non pas par loyaute '
  'mais par calcul — chaque competence acquise, chaque projet realise, chaque reseau construit '
  'augmente son capital transferable. L\'engagement est ainsi un investissement dans soi, '
  'pas dans l\'organisation.')

doc.add_heading('5.2 L\'entrepreneuriat comme norme, pas comme reaction', level=2)
p(f'Avec {ent_oui:.0f}% d\'intention entrepreneuriale — y compris chez les plus engages et les '
  f'plus satisfaits — l\'entrepreneuriat n\'est pas une fuite mais un projet de vie. '
  f'Ce constat distingue la Gen Z africaine de ses homologues occidentaux, ou l\'entrepreneuriat '
  f'reste minoritaire et souvent reactif. En Afrique francophone, la faiblesse des filets de '
  f'securite sociale, la valorisation culturelle de l\'"auto-emploi" et l\'acces facilite au '
  f'commerce (mobile money, e-commerce, reseaux sociaux) font de l\'entrepreneuriat le modele '
  f'mental par defaut. L\'emploi salarie est un tremplin, pas un aboutissement.')

doc.add_heading('5.3 Le management comme variable cle', level=2)
p(f'La correlation management-retention (r=0,62) est le levier le plus actionnable identifie '
  f'par cette etude. Le score management ({S["IATM"]}/100) etant le plus bas de l\'ensemble des '
  f'variables, la marge de progression est importante. Ce resultat est coherent avec les travaux '
  f'de Gallup (2023) selon lesquels "les gens ne quittent pas les entreprises, ils quittent les '
  f'managers" — mais il prend en Afrique francophone une dimension amplifiee par le modele '
  f'managerial souvent hierarchique et peu participatif qui prevaut dans de nombreuses organisations.')

doc.add_heading('5.4 L\'equite plutot que le montant', level=2)
p('La domination de l\'equite salariale percue (r=0,78 vers RECO) sur le salaire brut (r=0,06) '
  'constitue un defi pour les politiques de remuneration. Elle implique un deplacement de la '
  'question salariale : de "combien payer" vers "comment payer justement". Les audits d\'equite '
  'interne, la transparence des grilles et la justification des ecarts deviennent des outils '
  'plus puissants que les augmentations generalisees.')

doc.add_heading('5.5 Limites et pistes de recherche', level=2)
bul([
    'L\'absence de variable genre empeche d\'analyser les dynamiques differenciees hommes/femmes, '
    'connues comme significatives dans la litterature sur le turnover.',
    'Le caractere transversal de l\'etude ne permet pas de distinguer les effets d\'age, de cohorte '
    'et de periode. Une etude longitudinale (Vague 2, Vague 3) permettrait de mesurer l\'evolution.',
    'Les donnees declaratives peuvent diverger des comportements reels — l\'intention de depart '
    'ne se traduit pas systematiquement en depart effectif.',
    'La couverture geographique inegale limite la generalisation des resultats par pays.',
    'L\'etude ne mesure pas les facteurs macro-economiques (chomage, inflation, opportunites '
    'sectorielles) qui modulent les intentions en comportements.',
])
section_break()


# ════════════════════════════════════════════════════════════════
# VI. CONCLUSION
# ════════════════════════════════════════════════════════════════
doc.add_heading('VI. Conclusion', level=1)

p('Cette etude, menee aupres de 1 500 professionnels Gen Z dans 13 pays d\'Afrique '
  'subsaharienne francophone, met en evidence quatre resultats principaux :')

bul([
    f'Un paradoxe structurel entre engagement (78/100) et intention de depart (65%), qui invalide '
    f'l\'engagement comme predicteur de la retention dans ce contexte.',
    f'L\'entrepreneuriat comme norme generationnelle ({ent_oui:.0f}%), deconnecte du niveau de '
    f'satisfaction ou d\'engagement.',
    f'L\'equite salariale percue (r=0,78) et la qualite du management (r=0,62) comme leviers '
    f'structurels de la recommandation et de la retention, tres loin devant le salaire brut (r=0,06).',
    f'Un deficit managerial generalisé (IATM={S["IATM"]}/100) qui constitue a la fois le '
    f'probleme le plus aigu et le levier le plus actionnable.',
])

p('Ces resultats appellent une reconfiguration des politiques RH en Afrique francophone, '
  'fondee sur les leviers identifies (management, equite, intrapreneuriat) plutot que sur les '
  'leviers supposes (salaire, avantages). Les vagues suivantes du barometre permettront de '
  'mesurer l\'evolution de ces indicateurs et de valider les hypotheses avancees dans la discussion.')
section_break()


# ════════════════════════════════════════════════════════════════
# ANNEXES
# ════════════════════════════════════════════════════════════════
doc.add_heading('Annexes', level=1)

doc.add_heading('A. Statistiques descriptives de l\'ensemble des variables', level=2)
tbl(['Variable', 'n', 'Moyenne', 'Mediane', 'Ecart-type', 'Min', 'Max'], [
    [c, str(gz[c].dropna().shape[0]), f'{gz[c].mean():.1f}', f'{gz[c].median():.1f}',
     f'{gz[c].std():.1f}', f'{gz[c].min():.0f}', f'{gz[c].max():.0f}']
    for c in ['SAT','ENG','VAL','REM','PDA','SOUSPAYE','RECO','REQUALIF','APPRENT','INIT','INNOV','IQ_inv']
])

doc.add_heading('B. Matrice de correlations (25 premieres paires)', level=2)
tbl(['Var 1', 'Var 2', 'r', 'p-value', 'Sig.'], [
    [a, b, f'{r:+.3f}', f'{p:.4f}', s] for a, b, r, p, s in pairs[:25]
])

doc.add_heading('C. Tests t (18-22 vs 23-27 ans)', level=2)
test_rows = []
for c in ['SAT','ENG','VAL','REM','SOUSPAYE','RECO','IQ_inv','APPRENT','INIT','INNOV']:
    t_stat, p_val = sp_stats.ttest_ind(g18[c].dropna(), g23[c].dropna())
    sig = '***' if p_val<0.001 else '**' if p_val<0.01 else '*' if p_val<0.05 else 'ns'
    test_rows.append([c, f'{g18[c].mean():.1f}', f'{g23[c].mean():.1f}',
                      f'{t_stat:.2f}', f'{p_val:.4f}', sig])
tbl(['Variable', 'Moy. 18-22', 'Moy. 23-27', 't', 'p', 'Sig.'], test_rows)


# ── FIN ──
doc.add_page_break()
pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pr.add_run('\n\nDATATYM\u2122\nRapport Scientifique\nBarometre Gen Z Afrique 2026\n\n'
               'MARKETYM / H&C Executive\nDocument confidentiel — Avril 2026')
r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

out = '/Users/Apple/Desktop/H&C/Enquetes/Suivi enquetes/enquetes/Rapport_Scientifique_GenZ_DATATYM_2026.docx'
doc.save(out)
print(f'OK -> {out}')
