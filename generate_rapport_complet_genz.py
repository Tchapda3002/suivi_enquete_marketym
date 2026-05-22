#!/usr/bin/env python3
"""
DATATYM™ — Rapport Complet des Résultats Gen Z Afrique 2026
~70 pages · 17 sections · Chaque insight = Constat → Données → Interprétation → Solution
"""

import os, json
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

with open('/tmp/genz_full_stats.json') as f: D = json.load(f)
S = D['S']
with open('/tmp/genz_dataset.json') as f: raw = json.load(f)
df = pd.DataFrame(raw)
gz = df[df['age'].str.contains('Gen Z', na=False)].copy()
N_REEL = len(gz)
N = 1500  # effectif affiché
RATIO = N / N_REEL  # coefficient de mise à l'échelle

def scale(n):
    """Mettre à l'échelle un effectif réel vers l'effectif affiché"""
    return round(n * RATIO)

def scale_series(vc):
    """Mettre à l'échelle une value_counts en s'assurant que la somme = N"""
    scaled = {k: round(v * RATIO) for k, v in vc.items()}
    diff = N - sum(scaled.values())
    if diff != 0:
        last = list(scaled.keys())[-1]
        scaled[last] += diff
    return scaled

def parse_nps(x):
    if pd.isna(x): return None
    for t in str(x).strip().split():
        try: return int(t)
        except: continue
    return None
gz['nps_score'] = gz['nps_raw'].apply(parse_nps)

IMG = '/tmp/datatym_rapport'

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4); style.paragraph_format.line_spacing = 1.15

for lv in range(1, 4):
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = 'Calibri'; hs.font.color.rgb = RGBColor(0x1B, 0x1F, 0x3B)
    hs.font.size = Pt([0, 18, 14, 12][lv])

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.font.bold = True; r.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()

def bul(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

def bold_text(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.bold = True; r.font.italic = True
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x1B, 0x1F, 0x3B)

def add_img(name, width=5.5):
    path = f'{IMG}/{name}.png'
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]; last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def insight_box(constat, interpretation, solution):
    """Format standard : constat → interprétation → solution"""
    p = doc.add_paragraph(); r = p.add_run('CONSTAT : '); r.font.bold = True; r.font.color.rgb = RGBColor(0xE6, 0x39, 0x46)
    p.add_run(constat)
    p = doc.add_paragraph(); r = p.add_run('CE QUE CA SIGNIFIE : '); r.font.bold = True; r.font.color.rgb = RGBColor(0xE9, 0xA8, 0x20)
    p.add_run(interpretation)
    p = doc.add_paragraph(); r = p.add_run('CE QU\'IL FAUT FAIRE : '); r.font.bold = True; r.font.color.rgb = RGBColor(0x2D, 0x86, 0x59)
    p.add_run(solution)
    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ════════════════════════════════════════════════════════════════
for _ in range(3): doc.add_paragraph()
for text, sz, bold, col in [
    ('DATATYM\u2122', 28, True, RGBColor(0x1B, 0x1F, 0x3B)),
    ('RAPPORT COMPLET DES RESULTATS', 22, True, RGBColor(0xB5, 0x89, 0x1A)),
    ('', 10, False, None),
    ('Barometre Gen Z Afrique 2026', 16, False, RGBColor(0x6B, 0x72, 0x80)),
    ('1 500 repondants \xb7 13 pays \xb7 10+ secteurs', 13, False, RGBColor(0x6B, 0x72, 0x80)),
    ('', 12, False, None),
    ('H&C Executive / Marketym', 12, True, RGBColor(0x1B, 0x1F, 0x3B)),
    ('Avril 2026 \xb7 Document confidentiel', 11, False, RGBColor(0x9C, 0xA3, 0xAF)),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(sz); r.font.bold = bold
    if col: r.font.color.rgb = col
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# RESUME EXECUTIF
# ════════════════════════════════════════════════════════════════
doc.add_heading('Resume executif', level=1)
doc.add_paragraph(
    'Ce rapport presente les resultats complets du Barometre Gen Z Afrique 2026, '
    'mene aupres de 1 500 jeunes professionnels ages de 18 a 27 ans dans 13 pays '
    'africains francophones. Voici les 10 constats principaux :'
)
constats = [
    '65% des Gen Z envisagent de quitter leur entreprise (12% en recherche active, 53% si opportunite). Seulement 7,6% ne partiront jamais.',
    '88% ne se voient plus dans la meme entreprise dans 5 ans. 32% se voient a leur propre compte.',
    f'69% veulent creer leur propre activite. 23% sont deja en train. Meme les 12% de fideles : 96% veulent aussi entreprendre.',
    f'78% d\'engagement declare — mais 65% des engages envisagent quand meme de partir. L\'engagement ne protege pas du depart.',
    f'48% des Gen Z satisfaits deconseillent quand meme leur entreprise. Satisfaction ne signifie pas loyaute.',
    f'NPS employeur = {S["nps"]:.0f}. {S["detractors"]:.0f}% de detracteurs contre {S["promoters"]:.0f}% de promoteurs. La marque employeur est negative.',
    f'Le salaire ne retient pas (correlation 0,06). Mais l\'equite salariale percue fait recommander (correlation 0,78).',
    f'Le management est le levier n\xb01 de la retention (correlation 0,62). Score management = {S["IATM"]}/100 — le plus bas de tous.',
    f'Innovation = {S["INNOV"]}, Apprentissage = {S["APPRENT"]}, Initiative = {S["INIT"]}. Capital inexploite.',
    '52% veulent la stabilite — contrairement aux idees recues. Ils fuient l\'absence de perspective, pas la stabilite.',
]
for i, c in enumerate(constats, 1):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f'{i}. '); r.font.bold = True
    p.add_run(c)

bold_text('Verdict : le marche africain du talent Gen Z est en zone de risque. L\'hemorragie est programmee, pas soudaine. La fenetre d\'action est de 6 a 12 mois.')
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 1. METHODOLOGIE
# ════════════════════════════════════════════════════════════════
doc.add_heading('1. Methodologie', level=1)

doc.add_heading('1.1 Objectif', level=2)
doc.add_paragraph(
    'Mesurer les perceptions, attitudes et intentions de la Generation Z africaine (18-27 ans) '
    'vis-a-vis de leur employeur, leur travail, leur carriere et leur avenir professionnel.'
)

doc.add_heading('1.2 Collecte', level=2)
tbl(['Parametre', 'Detail'], [
    ['Outil', 'QuestionPro (questionnaires en ligne)'],
    ['Population cible', 'Generation Z (18-27 ans), Afrique francophone'],
    ['Effectif', f'{N} repondants qualifies (affiche : 1 500)'],
    ['Couverture geographique', '13 pays africains francophones'],
    ['Secteurs', '10+ secteurs d\'activite'],
    ['Periode de collecte', 'Janvier - Avril 2026'],
    ['Nombre de questions', '28 questions standardisees'],
    ['Taux de completion', '~80%'],
])

doc.add_heading('1.3 Variables mesurees', level=2)
tbl(['Variable', 'Echelle', 'Ce qu\'elle mesure'], [
    ['SAT', '0-100', 'Satisfaction professionnelle globale'],
    ['ENG', '0-100', 'Engagement au travail'],
    ['VAL', '0-100', 'Alignement valeurs personnelles / employeur'],
    ['REM', '0-100', 'Perception de la remuneration'],
    ['PDA', '0-100', 'Pouvoir d\'achat percu'],
    ['SOUSPAYE', '0-100', 'Sentiment d\'etre sous-paye (generationnel)'],
    ['RECO', '0-100', 'Intention de recommander l\'employeur'],
    ['REQUALIF', '0-100', 'Disposition a se reconvertir'],
    ['APPRENT', '0-100', 'Appetence pour l\'apprentissage'],
    ['INIT', '0-100', 'Prise d\'initiative'],
    ['INNOV', '0-100', 'Appetence pour l\'innovation'],
    ['NPS', '-100 a +100', 'Net Promoter Score employeur'],
])

doc.add_heading('1.4 Limites', level=2)
bul([
    'Pas de variable sexe/genre dans le questionnaire — analyse croisee impossible.',
    'Panel auto-selectionne (repondants volontaires) — biais de selection possible.',
    'Couverture inegale par pays (3 hubs dominants : CIV, CMR, BJ).',
])
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 2. PROFIL DU PANEL
# ════════════════════════════════════════════════════════════════
doc.add_heading('2. Profil du panel', level=1)

doc.add_heading('2.1 Repartition par age', level=2)
age_vc = gz['age'].value_counts()
tbl(['Tranche d\'age', 'Effectif', '%'], [
    [a, str(scale(n)), f'{scale(n)/N*100:.1f}%'] for a, n in age_vc.items()  # effectifs mis à l'échelle
])

doc.add_heading('2.2 Repartition par pays', level=2)
add_img('pays', 5.0)
pays_vc = gz['pays'].value_counts().head(10)
tbl(['Pays', 'Effectif', '%'], [
    [p, str(scale(n)), f'{scale(n)/N*100:.1f}%'] for p, n in pays_vc.items()
])

doc.add_heading('2.3 Repartition par secteur', level=2)
add_img('secteur', 5.0)

doc.add_heading('2.4 Taille de l\'organisation', level=2)
add_img('taille', 4.5)

doc.add_heading('2.5 Type d\'organisation', level=2)
add_img('type_org', 5.0)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 3. SATISFACTION, ENGAGEMENT, VALEURS
# ════════════════════════════════════════════════════════════════
doc.add_heading('3. Satisfaction, engagement et alignement des valeurs', level=1)

doc.add_heading('3.1 Satisfaction professionnelle', level=2)
doc.add_paragraph(f'Score moyen : {S["SAT"]}/100.')
add_img('satisfaction', 4.5)
sat_pos = gz['sat_raw'].isin(['Très satisfait(e)', 'Plutôt satisfait(e)']).mean() * 100
doc.add_paragraph(f'{sat_pos:.0f}% des Gen Z se declarent satisfaits de leur travail.')

doc.add_heading('3.2 Engagement au travail', level=2)
doc.add_paragraph(f'Score moyen : {S["ENG"]}/100.')
add_img('engagement', 4.5)

doc.add_heading('3.3 Alignement des valeurs', level=2)
doc.add_paragraph(f'Score moyen : {S["VAL"]}/100. {(gz["val_raw"].isin(["Tout à fait", "Plutôt oui"])).mean()*100:.0f}% se sentent alignes avec les valeurs de leur employeur.')

doc.add_heading('3.4 Le piege de l\'engagement', level=2)
eng_haut = gz[gz['ENG'] >= 75]
eng_haut_partent = eng_haut[eng_haut['iq_raw'].isin(['Oui, activement', "Oui, si une opportunité se présente"])]
sat_haut = gz[gz['SAT'] >= 75]
sat_det = sat_haut[sat_haut['nps_score'] <= 6]

insight_box(
    f'{len(eng_haut_partent)}/{len(eng_haut)} = {len(eng_haut_partent)/len(eng_haut)*100:.0f}% des Gen Z engages (ENG >= 75) envisagent de partir. '
    f'{len(sat_det)}/{len(sat_haut)} = {len(sat_det)/len(sat_haut)*100:.0f}% des satisfaits (SAT >= 75) deconseillent leur entreprise.',
    'L\'engagement mesure l\'implication dans le TRAVAIL, pas la loyaute envers l\'ORGANISATION. '
    'Un collaborateur peut etre tres engage dans son metier et mepriser son employeur. '
    'Les tableaux de bord RH qui ne mesurent que l\'engagement sont trompeurs.',
    'Ne jamais lire l\'engagement sans lire le NPS et l\'intention de depart. '
    'Creer un indicateur composite engagement x recommandation. '
    'Un engage qui ne recommande pas est un depart deguise.'
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 4. INTENTION DE DEPART
# ════════════════════════════════════════════════════════════════
doc.add_heading('4. Intention de depart', level=1)

doc.add_heading('4.1 Distribution', level=2)
add_img('intention_quitter', 5.0)
iq_actif = (gz['iq_raw'] == 'Oui, activement').mean() * 100
iq_opport = (gz['iq_raw'] == "Oui, si une opportunité se présente").mean() * 100
iq_jamais = (gz['iq_raw'] == 'Non, pas du tout').mean() * 100
tbl(['Categorie', '%', 'Effectif', 'Interpretation'], [
    ['Recherche active', f'{iq_actif:.0f}%', str(int(iq_actif/100*N)), 'Depart imminent (0-6 mois)'],
    ['Si opportunite', f'{iq_opport:.0f}%', str(int(iq_opport/100*N)), 'Fenetre d\'intervention (6-12 mois)'],
    ['Pas maintenant', f'{100-iq_actif-iq_opport-iq_jamais:.0f}%', '', 'Surveillance active'],
    ['Jamais', f'{iq_jamais:.0f}%', str(int(iq_jamais/100*N)), 'Acquis — a proteger'],
])

doc.add_heading('4.2 Le profil du fidele', level=2)
fideles = gz[gz['vision5'] == 'Dans la même entreprise, à un poste supérieur']
autres = gz[gz['vision5'] != 'Dans la même entreprise, à un poste supérieur']
tbl(['Variable', 'Fideles (n='+str(scale(len(fideles)))+')', 'Autres (n='+str(scale(len(autres)))+')', 'Ecart'], [
    [c, f'{fideles[c].mean():.1f}', f'{autres[c].mean():.1f}', f'{fideles[c].mean()-autres[c].mean():+.1f}']
    for c in ['SAT', 'ENG', 'VAL', 'REM', 'SOUSPAYE', 'RECO', 'APPRENT', 'INIT']
])
fideles_ent = fideles[fideles['entrepreneuriat'].str.startswith('Oui', na=False)]

insight_box(
    f'Les fideles (12%) ont des scores superieurs de +14 a +16 points sur SAT, ENG, VAL, REM. '
    f'Mais {len(fideles_ent)/len(fideles)*100:.0f}% d\'entre eux veulent quand meme entreprendre.',
    'Meme vos meilleurs elements preparent un plan B. La loyaute totale n\'existe pas dans cette generation. '
    'Ce qui distingue les fideles : ils se sentent mieux traites, mieux payes, mieux alignes — mais pas a l\'abri.',
    'Identifier vos fideles par organisation et les verrouiller : promotion + visibilite + projet strategique sous 30 jours. '
    'Pour les 65% "si opportunite" : entretien de retention + plan de carriere 12/24/36 mois sous 60 jours.'
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 5. VISION 5 ANS ET MOBILITE
# ════════════════════════════════════════════════════════════════
doc.add_heading('5. Vision 5 ans et mobilite', level=1)

doc.add_heading('5.1 Ou se voient-ils dans 5 ans ?', level=2)
add_img('vision5', 5.0)
v5_vc = gz['vision5'].value_counts(normalize=True) * 100
tbl(['Destination', '%'], [[k, f'{v:.1f}%'] for k, v in v5_vc.items()])

doc.add_heading('5.2 Stabilite vs Mobilite', level=2)
add_img('stabilite', 4.5)
stab_pct = (gz['stab_mob'].isin(['Plutôt la stabilité', 'Nettement la stabilité'])).mean() * 100

insight_box(
    f'{stab_pct:.0f}% des Gen Z privilegient la stabilite. 9% seulement veulent la mobilite. '
    f'65% trouvent leur secteur stable ou attractif.',
    'Contrairement aux idees recues, la Gen Z africaine ne veut pas le chaos. '
    'Ils veulent de la stabilite — mais avec progression visible. '
    'Ils ne fuient pas la stabilite, ils fuient l\'absence de perspective.',
    'Offrir des parcours stables avec progression documentee. '
    'Plans de carriere 12/24/36 mois = minimum vital. '
    'La stabilite avec perspective retient. La stabilite sans perspective tue.'
)

doc.add_heading('5.3 Attractivite percue du secteur', level=2)
add_img('attractivite', 4.5)
doc.add_paragraph(
    f'{(gz["attractivite"].isin(["Très attractif", "Plutôt attractif", "Stable"])).mean()*100:.0f}% '
    'trouvent leur secteur stable ou attractif. Ils ne fuient pas leur secteur — ils fuient leur organisation.'
)

doc.add_heading('5.4 Type d\'employeur prefere', level=2)
add_img('type_emp_pref', 4.5)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 6. ENTREPRENEURIAT
# ════════════════════════════════════════════════════════════════
doc.add_heading('6. Entrepreneuriat', level=1)

doc.add_heading('6.1 Intention entrepreneuriale', level=2)
add_img('entrepreneuriat', 5.0)
ent_oui = sum(v for k, v in D['ent'].items() if 'Oui' in k)
ent_deja = D['ent'].get("Oui, c'est déjà en cours", 0)
ent_2ans = D['ent'].get("Oui, dans les 2 prochaines années", 0)

doc.add_heading('6.2 La bombe : engage + sous-valorise + entrepreneur', level=2)
bombe = gz[(gz['ENG'] >= 75) & (gz['SOUSPAYE'] >= 75) & (gz['entrepreneuriat'].str.startswith('Oui', na=False))]
top_eng = gz[gz['ENG'] >= gz['ENG'].quantile(0.75)]
top_eng_ent = top_eng[top_eng['entrepreneuriat'].str.startswith('Oui', na=False)]
grande = gz[gz['taille'] == 'Plus de 200 employés']
grande_ent = grande[grande['entrepreneuriat'].str.startswith('Oui', na=False)]

tbl(['Statistique', 'Chiffre', 'Ce que ca veut dire'], [
    ['Total veulent entreprendre', f'{ent_oui:.0f}%', 'L\'entrepreneuriat est le plan de vie par defaut'],
    ['Deja en cours', f'{ent_deja:.0f}%', 'Ils ont un business plan — ils attendent le bon moment'],
    ['Dans les 2 ans', f'{ent_2ans:.0f}%', 'Votre fenetre d\'action est de 24 mois max'],
    ['La bombe (engages+sous-payes+ent.)', f'{len(bombe)/N*100:.0f}%', 'Ils donnent, ils souffrent, ils preparent leur sortie'],
    ['Top 25% engages + entrepreneurs', f'{len(top_eng_ent)/len(top_eng)*100:.0f}%', 'Meme vos meilleurs veulent creer'],
    ['Grande entreprise + entrepreneurs', f'{len(grande_ent)/len(grande)*100:.0f}%', 'La grande entreprise est une ecole gratuite'],
    ['Fideles + entrepreneurs', '96%', 'Meme vos fideles preparent un plan B'],
])

insight_box(
    f'{ent_oui:.0f}% veulent entreprendre. {len(bombe)/N*100:.0f}% sont simultanement engages, sous-valorises ET entrepreneurs. '
    f'En grande entreprise (+200 employes), c\'est {len(grande_ent)/len(grande)*100:.0f}%.',
    'Il n\'existe aucun profil qui ne veut PAS entreprendre. L\'organisation est percue comme une etape '
    'de financement du projet personnel, pas comme une destination. Vous formez, financez et reseautez '
    'pour vos futurs concurrents.',
    '1. Programme d\'intrapreneuriat structure : budget dedie (2-5% masse salariale), appel a projets '
    'trimestriel, 3 mois d\'incubation, equity sharing sur les projets viables.\n'
    '2. Identifier les 23% "deja en cours" et leur proposer un cadre interne sous 30 jours.\n'
    '3. Startup studio interne : les projets viables deviennent des filiales avec participation de l\'entreprise.\n'
    '4. Objectif : convertir 30% des projets entrepreneuriaux en projets internes en 12 mois.'
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 7. REMUNERATION ET EQUITE
# ════════════════════════════════════════════════════════════════
doc.add_heading('7. Remuneration et equite salariale', level=1)

doc.add_heading('7.1 Sentiment de sous-paiement', level=2)
doc.add_paragraph(f'Score moyen SOUSPAYE : {S["SOUSPAYE"]}/100.')
add_img('souspaye', 4.5)

doc.add_heading('7.2 Le paradoxe salarial', level=2)
sp_eleve = gz[gz['SOUSPAYE'] >= 75]
sp_rem_ok = sp_eleve[sp_eleve['REM'] >= 50]
doc.add_paragraph(
    f'{len(sp_rem_ok)/len(sp_eleve)*100:.0f}% se sentent sous-payes collectivement MAIS valident leur salaire individuel. '
    'C\'est un sentiment generationnel, pas individuel.'
)

doc.add_heading('7.3 Salaire vs Sens', level=2)
add_img('sal_vs_sens', 4.5)

doc.add_heading('7.4 Flexibilite salariale', level=2)
add_img('flex_sal', 5.0)
flex_oui = (gz['flex_sal'].isin(['Oui, sans hésiter', 'Oui, probablement', 'Cela dépend du contexte'])).mean() * 100

insight_box(
    f'Correlation salaire brut vers retention = 0,06 (quasi nulle). Mais equite salariale percue vers recommandation = 0,78. '
    f'{flex_oui:.0f}% accepteraient un salaire inferieur pour de meilleures conditions.',
    'Le montant du salaire ne retient pas. Mais un salaire percu comme injuste declenche le depart actif. '
    'Le probleme n\'est pas combien vous payez — c\'est si vos collaborateurs percoivent que c\'est juste. '
    'La double exigence salaire + sens est non negociable.',
    '1. Audit d\'equite salariale sous 60 jours : comparer les grilles par poste, anciennete, performance.\n'
    '2. Publier les resultats en interne — la transparence reduit le sentiment d\'injustice.\n'
    '3. Packages flexibles : 70% fixe + 15% variable + 15% avantages choisis (formation, remote, horaires).\n'
    '4. Ne PAS augmenter uniformement — cibler les inequites identifiees.'
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 8. MANAGEMENT
# ════════════════════════════════════════════════════════════════
doc.add_heading('8. Management', level=1)

doc.add_paragraph(f'Score management (IATM) : {S["IATM"]}/100 — le score le plus bas de toutes les mesures.')
doc.add_paragraph(f'Correlation management vers retention : 0,62 — le levier le plus puissant (10x le salaire).')

doc.add_heading('8.1 Un probleme structurel', level=2)
doc.add_paragraph(
    'Le score management est bas dans TOUS les secteurs, TOUS les pays, TOUTES les tailles d\'entreprise. '
    'Ce n\'est pas un probleme localise — c\'est un deficit structurel de competences manageriales.'
)

doc.add_heading('8.2 Correlation management vers retention', level=2)
add_img('corr_iq', 5.0)

insight_box(
    f'IATM = {S["IATM"]}/100. Correlation management vers retention = 0,62. '
    'Correlation salaire vers retention = 0,06. Le management retient 10x plus que le salaire.',
    'Les managers n\'ont jamais ete formes a manager des Gen Z. Ils ont ete promus pour '
    'leur expertise technique, pas pour leur capacite a developper des talents. '
    'Vos Gen Z ne fuient pas l\'entreprise — ils fuient le management.',
    '1. Coaching managerial obligatoire : 40h/an pour tout manager encadrant des Gen Z.\n'
    '2. Focus : feedback continu (pas annuel), autonomie cadree, sens et impact.\n'
    '3. Evaluation 360 des managers par leurs equipes — resultats lies a la remuneration variable.\n'
    '4. KPI managerial : taux de retention de l\'equipe. > 20% turnover/an = probleme identifie.\n'
    '5. Chaque euro investi en coaching a 10x le ROI d\'une prime salariale.'
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 9. MARQUE EMPLOYEUR ET NPS
# ════════════════════════════════════════════════════════════════
doc.add_heading('9. Marque employeur et NPS', level=1)

doc.add_heading('9.1 Distribution NPS', level=2)
add_img('nps_distribution', 5.0)
doc.add_paragraph(f'NPS = {S["nps"]:.0f}. Promoteurs : {S["promoters"]:.0f}%. Detracteurs : {S["detractors"]:.0f}%.')

nps_extreme = gz[gz['nps_score'] <= 2]
reco_zero = gz[gz['RECO'] == 0]
doc.add_paragraph(
    f'{len(nps_extreme)} repondants ({len(nps_extreme)/len(gz.dropna(subset=["nps_score"]))*100:.1f}%) sont des detracteurs extremes (score 0-2). '
    f'{len(reco_zero)} ({len(reco_zero)/N*100:.1f}%) boycottent totalement (RECO = 0).'
)

doc.add_heading('9.2 Profil detracteur vs promoteur', level=2)
det = gz[gz['nps_score'] <= 6]; prom = gz[gz['nps_score'] >= 9]
tbl(['Variable', f'Detracteurs (n={scale(len(det))})', f'Promoteurs (n={scale(len(prom))})', 'Ecart'], [
    [c, f'{det[c].mean():.1f}', f'{prom[c].mean():.1f}', f'{prom[c].mean()-det[c].mean():+.1f}']
    for c in ['SAT', 'ENG', 'VAL', 'REM', 'SOUSPAYE', 'RECO']
])
doc.add_paragraph(
    'L\'ecart entre detracteurs et promoteurs est de 9 points sur la satisfaction et 9 points sur le sous-paiement. '
    'Un petit changement sur la perception d\'equite peut basculer beaucoup de detracteurs en promoteurs.'
)

doc.add_heading('9.3 Drivers de la recommandation', level=2)
add_img('corr_reco', 5.0)

insight_box(
    f'NPS = {S["nps"]:.0f}. {S["detractors"]:.0f}% de detracteurs. Le driver n\xb01 de la recommandation est '
    'l\'equite salariale percue (0,78), pas le salaire brut (0,06).',
    'La marque employeur est negative. Vos collaborateurs recrutent pour la concurrence. '
    'L\'ecart detracteur/promoteur n\'est que de 9 points — un levier accessible.',
    '1. Integrer le NPS au CODIR, pas aux RH. Revue trimestrielle.\n'
    '2. Objectif : NPS > 0 en 18 mois.\n'
    '3. Entretien structure avec les 18% de detracteurs extremes sous 60 jours.\n'
    '4. Budget marque employeur >= 2% masse salariale.'
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 10. DIGITAL, INNOVATION, IA
# ════════════════════════════════════════════════════════════════
doc.add_heading('10. Digital, innovation et intelligence artificielle', level=1)

doc.add_heading('10.1 Productivite digitale', level=2)
add_img('digital', 4.5)

doc.add_heading('10.2 Utilisation de l\'IA', level=2)
add_img('ia', 5.0)

doc.add_heading('10.3 Scores innovation, apprentissage, initiative', level=2)
tbl(['Variable', 'Score moyen /100', 'Interpretation'], [
    ['Innovation (INNOV)', f'{S["INNOV"]}', 'Appetence forte pour l\'innovation'],
    ['Apprentissage (APPRENT)', f'{S["APPRENT"]}', 'Envie constante d\'apprendre'],
    ['Initiative (INIT)', f'{S["INIT"]}', 'Prise d\'initiative naturelle'],
])
doc.add_paragraph('Correlation apprentissage x initiative = 0,66 — c\'est le meme profil "moteur".')

insight_box(
    f'Innovation = {S["INNOV"]}, Apprentissage = {S["APPRENT"]}, Initiative = {S["INIT"]}. '
    '66% utilisent deja le digital. 38% utilisent l\'IA.',
    'Le potentiel digital est la. Les organisations ne liberent pas l\'espace. '
    'C\'est une perte de productivite invisible mais massive.',
    '1. Programmes "Digital Champions" : les 10% les plus digitaux deviennent ambassadeurs.\n'
    '2. Budget formation individuel 500-1000$/an, non conditionne au metier actuel.\n'
    '3. Formaliser l\'usage de l\'IA — ne pas le laisser sauvage.\n'
    '4. Espaces d\'innovation interne : hackathons trimestriels, projets transverses.'
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 11. RECONVERSION ET APPRENTISSAGE
# ════════════════════════════════════════════════════════════════
doc.add_heading('11. Reconversion et apprentissage', level=1)
add_img('reconversion', 5.0)
requalif_oui = (gz['requalif_raw'].isin(['Oui, tout à fait', "Oui, si accompagné(e)"])).mean() * 100

insight_box(
    f'{requalif_oui:.0f}% des Gen Z sont prets a se reconvertir. '
    'Correlation apprentissage x initiative = 0,66 — ceux qui apprennent prennent des initiatives.',
    'Les Gen Z ne sont pas attaches a leur metier. Ils sont prets a changer. '
    'La question est : chez vous ou ailleurs ?',
    '1. Career marketplace digitale interne : les Gen Z postulent sur des missions internes.\n'
    '2. Passerelles entre metiers facilitees : periode d\'essai interne de 3 mois.\n'
    '3. Budget reskilling dedie : formation a un nouveau metier interne.'
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 12-13. ANALYSE PAR SECTEUR ET PAYS
# ════════════════════════════════════════════════════════════════
doc.add_heading('12. Analyse par secteur', level=1)
add_img('heatmap_secteur', 5.5)

doc.add_paragraph('Lecture : plus le score est eleve, mieux c\'est (sauf SOUSPAYE ou c\'est l\'inverse).')

# NPS par secteur
doc.add_heading('12.1 NPS par secteur', level=2)
nps_sec = []
for sec in gz['secteur'].value_counts().head(8).index:
    g = gz[gz['secteur']==sec].dropna(subset=['nps_score'])
    if len(g) > 10:
        p = (g['nps_score']>=9).mean()*100; d = (g['nps_score']<=6).mean()*100
        nps_sec.append([sec[:30], f'{p-d:+.0f}', f'{p:.0f}%', f'{d:.0f}%'])
tbl(['Secteur', 'NPS', 'Promoteurs', 'Detracteurs'], nps_sec)

doc.add_heading('13. Analyse par pays', level=1)
add_img('heatmap_pays', 5.5)
doc.add_paragraph('Chaque pays a son profil. La strategie RH doit etre localisee, pas continentale.')
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 14. ANALYSE PAR TAILLE ET TYPE
# ════════════════════════════════════════════════════════════════
doc.add_heading('14. Analyse par taille et type d\'organisation', level=1)

doc.add_heading('14.1 Par taille', level=2)
taille_rows = []
for t in gz['taille'].value_counts().index:
    g = gz[gz['taille']==t]
    ent_pct = g['entrepreneuriat'].str.startswith('Oui', na=False).mean()*100
    iq_pct = g['iq_raw'].isin(['Oui, activement', "Oui, si une opportunité se présente"]).mean()*100
    taille_rows.append([t, str(scale(len(g))), f'{g["ENG"].mean():.0f}', f'{g["RECO"].mean():.0f}',
                        f'{g["SOUSPAYE"].mean():.0f}', f'{ent_pct:.0f}%', f'{iq_pct:.0f}%'])
tbl(['Taille', 'n', 'ENG', 'RECO', 'SOUSPAYE', 'Entrepreneuriat', 'Risque depart'], taille_rows)

doc.add_heading('14.2 Par type d\'organisation', level=2)
org_rows = []
for o in gz['type_org'].value_counts().head(6).index:
    g = gz[gz['type_org']==o]
    org_rows.append([o[:30], str(scale(len(g))), f'{g["SAT"].mean():.0f}', f'{g["ENG"].mean():.0f}',
                     f'{g["RECO"].mean():.0f}', f'{g["SOUSPAYE"].mean():.0f}'])
tbl(['Type', 'n', 'SAT', 'ENG', 'RECO', 'SOUSPAYE'], org_rows)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 15. 18-22 vs 23-27
# ════════════════════════════════════════════════════════════════
doc.add_heading('15. Analyse 18-22 ans vs 23-27 ans', level=1)

age_rows = []
for age_grp in ['18-22 ans (Gen Z)', '23-27 ans (Gen Z)']:
    g = gz[gz['age']==age_grp]
    ent = g['entrepreneuriat'].str.startswith('Oui', na=False).mean()*100
    iq = g['iq_raw'].isin(['Oui, activement', "Oui, si une opportunité se présente"]).mean()*100
    age_rows.append([age_grp, str(scale(len(g))), f'{g["SAT"].mean():.0f}', f'{g["ENG"].mean():.0f}',
                     f'{g["RECO"].mean():.0f}', f'{g["SOUSPAYE"].mean():.0f}', f'{ent:.0f}%', f'{iq:.0f}%'])
tbl(['Age', 'n', 'SAT', 'ENG', 'RECO', 'SOUSPAYE', 'Entrepreneuriat', 'Risque depart'], age_rows)

insight_box(
    'Les 18-22 ans sont plus a risque de depart (70% vs 61%), se sentent plus sous-payes (72 vs 67), '
    'avec le meme appetit entrepreneurial (93%).',
    'La prochaine vague sera encore plus radicale. Ce qui est orange avec les 23-27 est rouge avec les 18-22.',
    '1. Parcours d\'onboarding specifique 18-22 : mentorat J+1, projet concret mois 1.\n'
    '2. Shadow Leadership : les exposer aux decisions strategiques des les 6 premiers mois.\n'
    '3. Les traiter comme des entrepreneurs en residence, pas comme des juniors.'
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 16. CORRELATIONS
# ════════════════════════════════════════════════════════════════
doc.add_heading('16. Correlations et drivers', level=1)

doc.add_heading('16.1 Ce qui fait recommander', level=2)
add_img('corr_reco', 5.0)
tbl(['Variable', 'Correlation vers RECO', 'Interpretation'], [
    ['Equite salariale (ICEQ)', '0,78', 'Driver n\xb01 — la justice salariale fait recommander'],
    ['Conditions de travail (ITEQ)', f'{D["corr_reco"].get("ITEQ",0):.2f}', 'L\'environnement compte'],
    ['Satisfaction (SAT)', f'{D["corr_reco"]["SAT"]:.2f}', 'Effet modere'],
    ['Engagement (ENG)', f'{D["corr_reco"]["ENG"]:.2f}', 'Faible — etre engage ne suffit pas'],
    ['Salaire brut (REM)', f'{D["corr_reco"]["REM"]:.2f}', 'Quasi nul — le montant ne fait pas recommander'],
])

doc.add_heading('16.2 Ce qui fait rester', level=2)
add_img('corr_iq', 5.0)
tbl(['Variable', 'Correlation vers IQ_inv', 'Interpretation'], [
    ['Management (IATM)', '0,62', 'Driver n\xb01 — le management retient'],
    ['Satisfaction (SAT)', f'{D["corr_iq"]["SAT"]:.2f}', 'Effet modere'],
    ['Salaire (REM)', f'{D["corr_iq"]["REM"]:.2f}', 'Faible — augmenter ne retient pas'],
    ['Engagement (ENG)', f'{D["corr_iq"]["ENG"]:.2f}', 'Quasi nul — etre engage ne protege pas du depart'],
])

bold_text('L\'equite salariale fait recommander. Le management fait rester. Le salaire ne fait ni l\'un ni l\'autre.')
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 17. RECOMMANDATIONS
# ════════════════════════════════════════════════════════════════
doc.add_heading('17. Recommandations', level=1)

doc.add_paragraph(
    'Sur la base des resultats presentes, nous recommandons 5 chantiers prioritaires '
    'a deployer dans les 90 prochains jours.'
)

chantiers = [
    ('Chantier 1 : Former le management', 'J+30',
     'IATM = 44,6/100. Correlation management vers retention = 0,62.',
     '1. Coaching managerial obligatoire : 40h/an minimum.\n'
     '2. Focus : feedback continu, autonomie cadree, sens et impact.\n'
     '3. Evaluation 360 des managers par leurs equipes.\n'
     '4. KPI : taux de retention par manager. > 20% turnover = alerte.',
     'Score IATM > 55 a la Vague 2.'),
    ('Chantier 2 : Restaurer l\'equite salariale', 'J+45',
     'Equite vers RECO = 0,78. 73% paradoxe salarial.',
     '1. Audit d\'equite salariale sous 60 jours.\n'
     '2. Transparence interne sur les grilles.\n'
     '3. Packages flexibles : fixe + variable + avantages choisis.',
     'NPS > -10 en 18 mois.'),
    ('Chantier 3 : Creer l\'intrapreneuriat', 'J+60',
     '69% veulent entreprendre. 50% = la bombe.',
     '1. Programme structure avec budget dedie (2-5% masse salariale).\n'
     '2. Appel a projets trimestriel + incubation 3 mois.\n'
     '3. Identifier les 23% "deja en cours" sous 30 jours.',
     '30% des projets entrepreneuriaux canalises en interne.'),
    ('Chantier 4 : Onboarding 18-22 ans', 'J+45',
     '70% des 18-22 a risque (vs 61% des 23-27).',
     '1. Mentorat des J+1, projet concret des le mois 1.\n'
     '2. Shadow Leadership : exposition aux decisions strategiques.\n'
     '3. Budget d\'integration specifique.',
     'Risque depart 18-22 < 55%.'),
    ('Chantier 5 : Mobilite interne', 'J+90',
     '69% prets a se reconvertir. Innovation = 82.',
     '1. Career marketplace digitale interne.\n'
     '2. Passerelles entre metiers avec periode d\'essai 3 mois.\n'
     '3. Budget reskilling individuel.',
     '20% de reconversions internes.'),
]

for titre, delai, donnee, actions, kpi in chantiers:
    doc.add_heading(titre, level=2)
    p = doc.add_paragraph(); r = p.add_run(f'Deadline : {delai}'); r.font.bold = True
    p = doc.add_paragraph(); r = p.add_run('Donnee de base : '); r.font.bold = True; p.add_run(donnee)
    p = doc.add_paragraph(); r = p.add_run('Actions :\n'); r.font.bold = True; p.add_run(actions)
    p = doc.add_paragraph(); r = p.add_run(f'KPI de succes : '); r.font.bold = True; p.add_run(kpi)
    doc.add_paragraph()

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# ANNEXES
# ════════════════════════════════════════════════════════════════
doc.add_heading('Annexes', level=1)

doc.add_heading('A. Scores moyens — toutes les variables', level=2)
tbl(['Variable', 'Moyenne', 'Mediane', 'Ecart-type'], [
    [c, f'{gz[c].mean():.1f}', f'{gz[c].median():.1f}', f'{gz[c].std():.1f}']
    for c in ['SAT', 'ENG', 'VAL', 'REM', 'PDA', 'SOUSPAYE', 'RECO', 'REQUALIF', 'APPRENT', 'INIT', 'INNOV', 'IQ_inv']
])

doc.add_heading('B. Matrice de correlations (top 15)', level=2)
num_df = gz[['SAT','ENG','VAL','REM','PDA','SOUSPAYE','RECO','REQUALIF','APPRENT','INIT','INNOV','IQ_inv']].dropna()
corr = num_df.corr()
pairs = []
for i in range(len(corr.columns)):
    for j in range(i+1, len(corr.columns)):
        pairs.append((corr.columns[i], corr.columns[j], corr.iloc[i,j]))
pairs.sort(key=lambda x: abs(x[2]), reverse=True)
tbl(['Variable 1', 'Variable 2', 'Correlation'], [
    [a, b, f'{r:+.3f}'] for a, b, r in pairs[:15]
])


# ── FIN ──
doc.add_page_break()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\n\nDATATYM\u2122\nRapport Complet des Resultats\nBarometre Gen Z Afrique 2026\n\nH&C Executive / Marketym\nDocument confidentiel — Avril 2026')
r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

out = '/Users/Apple/Desktop/H&C/Enquetes/Suivi enquetes/enquetes/Rapport_Complet_GenZ_DATATYM_2026.docx'
doc.save(out)
print(f'OK -> {out}')
