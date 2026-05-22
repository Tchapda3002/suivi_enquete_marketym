#!/usr/bin/env python3
"""
DATATYM™ — Benchmark Plateformes Data & BI Afrique
Éléments pertinents pour DATATYM AI
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4); style.paragraph_format.line_spacing = 1.12

for lv in range(1,4):
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = 'Calibri'; hs.font.color.rgb = RGBColor(0x1B,0x1F,0x3B)
    hs.font.size = Pt([0,16,13,11][lv])

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=h
        for p in c.paragraphs:
            for r in p.runs: r.font.bold=True; r.font.size=Pt(8)
    for i,row in enumerate(rows):
        for j,v in enumerate(row):
            c=t.rows[i+1].cells[j]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(8)
    doc.add_paragraph()

def bul(items):
    for it in items: doc.add_paragraph(it, style='List Bullet')

def bold_p(text):
    p=doc.add_paragraph(); r=p.add_run(text); r.font.bold=True; r.font.italic=True
    r.font.size=Pt(10); r.font.color.rgb=RGBColor(0x1B,0x1F,0x3B)

# ═══ COVER ═══
for _ in range(3): doc.add_paragraph()
for text,sz,bold,col in [
    ('DATATYM\u2122',28,True,RGBColor(0x1B,0x1F,0x3B)),
    ('BENCHMARK',22,True,RGBColor(0xB5,0x89,0x1A)),
    ('',8,False,None),
    ('Plateformes Data & Business Intelligence\nen Afrique',15,False,RGBColor(0x6B,0x72,0x80)),
    ('',8,False,None),
    ('Elements pertinents pour DATATYM AI',13,True,RGBColor(0x1B,0x1F,0x3B)),
    ('',10,False,None),
    ('H&C Executive \xb7 Avril 2026 \xb7 Confidentiel',11,False,RGBColor(0x9C,0xA3,0xAF)),
]:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.font.size=Pt(sz); r.font.bold=bold
    if col: r.font.color.rgb=col
doc.add_page_break()


# ═══ SYNTHESE ═══
doc.add_heading('Synthese strategique', level=1)

doc.add_paragraph(
    'Ce benchmark couvre 15 plateformes operant sur le marche africain de la donnee, '
    'de la business intelligence et des etudes de marche. L\'objectif est d\'identifier '
    'les elements pertinents pour le positionnement de DATATYM AI.')

doc.add_heading('Carte du paysage concurrentiel', level=2)
tbl(['Segment','Acteurs','Position DATATYM'], [
    ['Etudes terrain premium','Sagaci, Kantar, Ipsos, Ask Afrika',
     'DATATYM a les enqueteurs (Marketym) mais vise le\nself-service, pas le sur-mesure cher'],
    ['BI / Intelligence Africa','Stears, Asoko Insight',
     'ZONE DE COMPETITION DIRECTE\nSe differencier par le francophone, les indices,\nla couverture multi-sectorielle'],
    ['Mapping ecosysteme tech','Briter Bridges, Partech',
     'Adjacent, pas prioritaire'],
    ['Statistiques generales','Statista, Open Data BAD, Numbeo',
     'DATATYM = "Statista africain"\navec donnees primaires'],
    ['Media / influence decideurs','Jeune Afrique, Africa CEO Forum',
     'Partenaires de distribution, pas concurrents'],
    ['Collecte mobile','GeoPoll',
     'Partenaire potentiel pour la collecte'],
])

doc.add_heading('5 avantages concurrentiels a construire', level=2)
bul([
    'LE "STATISTA AFRICAIN" AVEC DONNEES PRIMAIRES — Aucun acteur ne combine donnees terrain propres + plateforme self-service + indices proprietaires. C\'est le positionnement unique.',
    'COUVERTURE FRANCOPHONE — Stears est anglophone, Sagaci n\'a pas de plateforme self-service. Le marche francophone (UEMOA, CEMAC) est sous-desservi en BI digitale.',
    'INDICES PROPRIETAIRES PUBLICS — Des indices de reference (IPE, IRTA, ILUX, IGRO) qui deviennent des standards cites par la presse et les decideurs. Comme Bloomberg a ses indices.',
    'PRIX ADAPTE — Statista a 2 400$/an, Sagaci a 10-50K$/etude, Kantar a 100K+. DATATYM peut viser 500-2 000$/an pour les PME et ETI africaines.',
    'IA PREDICTIVE — Aucun acteur du benchmark ne propose d\'IA integree pour des recommandations personnalisees. C\'est le facteur de differenciation le plus fort a moyen terme.',
])

doc.add_heading('Partenariats strategiques prioritaires', level=2)
tbl(['Partenaire','Apport pour DATATYM','Type'], [
    ['Jeune Afrique / Africa CEO Forum','Distribution vers les DG africains.\nAudience captive de decideurs.','Distribution'],
    ['GeoPoll','Collecte complementaire dans les pays\nsans equipe Marketym. 3$/complete.','Collecte'],
    ['Open Data BAD','Couche de donnees macro gratuite\na integrer comme socle.','Donnees'],
    ['Partech Africa','Credibilite ecosysteme tech.\nModele de rapport gratuit a repliquer.','Credibilite'],
])
doc.add_page_break()


# ═══ FICHES DETAILLEES ═══
doc.add_heading('Fiches detaillees par plateforme', level=1)

fiches = [
    {
        'nom': 'STEARS', 'url': 'stears.co', 'pays': 'Nigeria',
        'cat': 'CONCURRENT DIRECT',
        'desc': 'Intelligence financiere et economique sur l\'Afrique. A pivote du B2C (media/data) vers le B2B pur. Offre : market sizing, previsions, indices consommateurs, datasets macroeconomiques.',
        'couv': '5 marches cles (Nigeria, Kenya, Ghana, Egypte, Afrique du Sud).',
        'prix': 'Anciennement 10$/mois (B2C). Maintenant B2B enterprise sur devis. Le B2B genere 75% du CA. Leve 3.3M$ (Serena Ventures, Mac VC).',
        'forces': ['Marque forte au Nigeria', 'Donnees financieres profondes', 'Equipe editoriale solide', 'Backed par des VC de renom'],
        'faiblesses': ['Couverture limitee a 5 pays', 'A abandonne le B2C (perte de communaute)', 'Pas de terrain/enquetes propres', 'Zero Afrique francophone'],
        'pertinence': 'Concurrent le plus direct. Meme positionnement (BI africaine pour decideurs). DATATYM se differencie par : (a) couverture francophone, (b) indices sectoriels proprietaires, (c) donnees terrain via Marketym, (d) plus de pays. Le pivot B2B de Stears valide le modele.',
    },
    {
        'nom': 'SAGACI RESEARCH', 'url': 'sagaciresearch.com', 'pays': 'Multi-pays',
        'cat': 'CONCURRENT INDIRECT',
        'desc': 'Etudes de marche terrain en Afrique. Produits proprietaires : SagaTracker (audit retail), SagaBrand (tracking marque, 25 KPIs, 25 marches), SagaCube (analyse conso). CA ~16M$.',
        'couv': '54 pays africains, panels dans 34 pays. Bureaux Abidjan, Accra, Lagos, Nairobi, Johannesburg, Le Caire, Casablanca, Paris.',
        'prix': 'B2B sur devis. 10-50K$ par etude. Abonnements trackers.',
        'forces': ['Couverture panafricaine la plus complete', 'Equipes terrain propres', 'Methodologie cross-pays', 'Fondes par ex-BCG'],
        'faiblesses': ['Prix eleves (inaccessible PME)', 'Focus FMCG/retail', 'Pas de plateforme self-service', 'Pas d\'indices publics'],
        'pertinence': 'DATATYM se differencie par le self-service, les indices accessibles en ligne, un prix adapte. Sagaci vend du "sur mesure cher" ; DATATYM vend du "pret-a-porter intelligent".',
    },
    {
        'nom': 'STATISTA (Afrique)', 'url': 'statista.com', 'pays': 'Mondial',
        'cat': 'BENCHMARK DE REFERENCE',
        'desc': 'Portail mondial de statistiques. Section Afrique avec +60 000 sujets. Agregation de sources.',
        'couv': 'Mondiale, section Afrique presente mais moins profonde.',
        'prix': 'A partir de 2 388$/an. Comptes entreprise sur devis.',
        'forces': ['Marque mondiale reconnue', 'Interface excellente', 'Volume de donnees enorme', 'Visualisations pretes a l\'emploi'],
        'faiblesses': ['Donnees Afrique superficielles/secondaires', 'Pas de collecte terrain', 'Pas d\'expertise contextuelle africaine', 'Prix eleve pour les entreprises africaines'],
        'pertinence': 'Talon d\'Achille en Afrique = profondeur et fraicheur. DATATYM = "Statista africain" avec donnees primaires et indices proprietaires. C\'est le pitch le plus vendeur.',
    },
    {
        'nom': 'ASOKO INSIGHT', 'url': 'asokoinsight.com', 'pays': 'Multi-pays',
        'cat': 'COMPLEMENTAIRE',
        'desc': 'Base de +200 000 profils d\'entreprises africaines, historiques de deals PE/dette. Cible : fonds d\'investissement, banques.',
        'couv': 'Panafricaine (focus marches actifs en investissement).',
        'prix': 'Abonnement + "company snapshots" unitaires. Leve 3.6M$. Estime 5-20K$/an.',
        'forces': ['Database corporate la plus profonde', 'Utile pour le sourcing de deals', 'Recherche bespoke'],
        'faiblesses': ['Focus investisseurs/M&A uniquement', 'Pas d\'indices', 'Interface fonctionnelle mais pas sexy', 'Pas d\'IA'],
        'pertinence': 'Complementaire. DATATYM vise les DG (decision operationnelle) ; Asoko vise les investisseurs (deal-making). Possibilite de partenariat.',
    },
    {
        'nom': 'ASK AFRIKA', 'url': 'askafrika.co.za', 'pays': 'Afrique du Sud',
        'cat': 'MODELE A ETUDIER',
        'desc': 'Plus grande societe d\'etudes independante d\'Afrique du Sud (depuis 1995). Indices proprietaires : Target Group Index, Orange Index.',
        'couv': 'Terrain dans 95% de l\'Afrique subsaharienne, base en Afrique du Sud.',
        'prix': 'Etudes sur devis. Indices syndiques (Orange Index par abonnement).',
        'forces': ['30 ans d\'experience', 'Capacite terrain massive', 'Indices reconnus localement', 'Independance'],
        'faiblesses': ['Tres centree Afrique du Sud', 'Pas de plateforme digitale', 'Methodologies classiques', 'Marque peu connue hors Afrique australe'],
        'pertinence': 'Modele a etudier pour les indices proprietaires. DATATYM s\'inspire de l\'approche "indices de reference" en version digitale et panafricaine.',
    },
    {
        'nom': 'XRI', 'url': 'xri-africa.com', 'pays': 'Zimbabwe',
        'cat': 'CONCURRENT LOCAL LIMITE',
        'desc': 'Cabinet d\'etudes de marche. Communaute online "The Lounge" pour recrutement de repondants.',
        'couv': 'Principalement Zimbabwe, expansion regionale.',
        'prix': 'Etudes sur devis + panel remunere.',
        'forces': ['Ancrage local fort', 'Panel communautaire en ligne'],
        'faiblesses': ['1 seul pays', 'Pas de techno differenciante', 'Pas de plateforme data', 'Marche limite'],
        'pertinence': 'Preoccupant seulement s\'il scale. Le modele panel communautaire est a observer. DATATYM est deja bien au-dela.',
    },
    {
        'nom': 'GEOPOLL', 'url': 'geopoll.com', 'pays': 'Multi-pays',
        'cat': 'PARTENAIRE POTENTIEL',
        'desc': 'Enquetes mobiles (SMS, app, CATI). 53M+ interviews. Panel de 5M repondants. Acces a 300M+ personnes via operateurs telecom.',
        'couv': 'Afrique (Ghana, Kenya, Nigeria, Afrique du Sud) + Asie + Amerique Latine.',
        'prix': 'A partir de 3$/complete. Etudes custom sur devis.',
        'forces': ['Infrastructure mobile unique', 'Rapidite', 'Acces populations non connectees', 'Prix accessibles'],
        'faiblesses': ['Collecte seulement, pas d\'analyse', 'Pas de plateforme BI', 'Qualite variable (SMS)', 'Pas d\'indices'],
        'pertinence': 'PARTENAIRE IDEAL pour la collecte. DATATYM peut utiliser GeoPoll dans les pays sans enqueteurs Marketym. Prix benchmark : 3$/complete.',
    },
    {
        'nom': 'JEUNE AFRIQUE / AFRICA CEO FORUM', 'url': 'jeuneafrique.com / theafricaceoforum.com', 'pays': 'Panafricain',
        'cat': 'PARTENAIRE STRATEGIQUE',
        'desc': 'Ecosysteme media + evenements. Africa CEO Forum = 2 000 participants, 70 pays, 1 000+ CEOs. The Africa Report = 1M lecteurs/mois. JA = 35 000 abonnes.',
        'couv': 'Panafricaine, bilingue FR/EN.',
        'prix': 'Abonnements media + billetterie + sponsoring.',
        'forces': ['Marque la plus puissante aupres des decideurs africains', 'Reseau CEO inegale', 'FR + EN', 'Legitimite historique (60+ ans)'],
        'faiblesses': ['Pas une plateforme data/BI', 'Pas de datasets exploitables', 'Contenu editorial, pas analytique'],
        'pertinence': 'PARTENAIRE ideal. DATATYM fournit la data structuree ; JAMG fournit l\'audience. Co-branding de rapports. Africa CEO Forum = canal de distribution parfait.',
    },
    {
        'nom': 'BRITER BRIDGES', 'url': 'briter.co', 'pays': 'Multi-emergents',
        'cat': 'BENCHMARK PRODUIT',
        'desc': 'Plateforme BI pour marches emergents. Cartographie startups, investisseurs, deals. Dashboards interactifs. Clients : Mastercard Foundation, Banque Mondiale.',
        'couv': 'Afrique + Amerique Latine + Asie-Pacifique.',
        'prix': 'Freemium + abonnement payant. Finance par FMO.',
        'forces': ['Interface moderne', 'Dashboards interactifs', 'Multi-emergents'],
        'faiblesses': ['Focus tech/startup uniquement', 'Pas de donnees conso/sectorielle classique'],
        'pertinence': 'Bon benchmark UX/produit. Le modele freemium + dashboards est pertinent pour DATATYM.',
    },
    {
        'nom': 'PARTECH AFRICA', 'url': 'partechpartners.com', 'pays': 'Panafricain',
        'cat': 'MODELE DE CONTENU',
        'desc': 'Fonds VC avec rapport annuel de reference sur le financement tech en Afrique. 2025 : 4.1Mds$ (+25% YoY).',
        'couv': 'Panafricaine (tech/startups).',
        'prix': 'Rapport gratuit (visibilite pour le fonds).',
        'forces': ['Reference absolue pour le VC tech', 'Grande visibilite mediatique', 'Donnees fiables'],
        'faiblesses': ['Uniquement tech/VC', '1 rapport/an', 'Pas de plateforme'],
        'pertinence': 'Modele "rapport gratuit = acquisition notoriete" a repliquer. DATATYM pourrait produire un equivalent mensuel et multi-sectoriel.',
    },
    {
        'nom': 'KANTAR / IPSOS (Afrique)', 'url': 'kantar.com / ipsos.com', 'pays': 'Multi-pays',
        'cat': 'GEANTS A DISRUPTER',
        'desc': 'Geants mondiaux des etudes. En Afrique : panels FMCG, tracking marques, etudes ad hoc. Kantar : 16 pays, 6 500 menages panel. Ipsos : ex-Synovate.',
        'couv': 'Kantar : 16 pays (78% PIB africain). Ipsos : Nigeria, Kenya, Ghana, CIV, Maroc, Afrique du Sud.',
        'prix': 'Panels : centaines de milliers $/an. Etudes : 50-200K$.',
        'forces': ['Panels les plus fiables', 'Methodologie mondiale', 'Marque premium'],
        'faiblesses': ['Prix prohibitifs', 'Lent', 'Focus FMCG', 'Pas de self-service abordable'],
        'pertinence': 'Les clients qui trouvent Kantar/Ipsos trop cher sont un vivier pour DATATYM. Disruption par le self-service et les indices continus a prix accessible.',
    },
    {
        'nom': 'NUMBEO', 'url': 'numbeo.com', 'pays': 'Mondial',
        'cat': 'MODELE D\'INDICES',
        'desc': 'Plus grande base mondiale cout de la vie, immobilier, criminalite. Donnees crowdsourcees. ~40 villes africaines.',
        'couv': 'Mondiale, couverture Afrique incomplete.',
        'prix': 'Gratuit + API par abonnement.',
        'forces': ['Indices comparatifs inter-villes', 'Gratuit', 'SEO massif'],
        'faiblesses': ['Crowdsource (fiabilite variable)', 'Couverture Afrique incomplete', 'Pas d\'analyse business'],
        'pertinence': 'Modele d\'indices comparatifs a repliquer. DATATYM peut creer des indices "attractivite d\'investissement" par ville avec des donnees primaires plus fiables.',
    },
    {
        'nom': 'OPEN DATA BAD', 'url': 'dataportal.opendataforafrica.org', 'pays': '54 pays',
        'cat': 'SOURCE DE DONNEES',
        'desc': 'Plateforme open data de la Banque Africaine de Developpement. 54 pays, donnees macro, sante, genre, ODD.',
        'couv': '54 pays africains (couverture complete).',
        'prix': 'Entierement gratuit.',
        'forces': ['Couverture continentale', 'Donnees officielles', 'Outils de visualisation'],
        'faiblesses': ['Donnees datees (2-3 ans)', 'Interface archaique', 'Zero donnees privees/sectorielles', 'Zero intelligence'],
        'pertinence': 'Source de donnees macro a integrer comme couche de base. DATATYM ajoute la couche "intelligence" que la BAD ne fournit pas.',
    },
]

for f in fiches:
    doc.add_heading(f'{f["nom"]}', level=2)

    # Catégorie
    p = doc.add_paragraph()
    r = p.add_run(f['cat']); r.font.bold = True; r.font.size = Pt(10)
    cat_colors = {'CONCURRENT DIRECT': RGBColor(0xE6,0x39,0x46),
                  'CONCURRENT INDIRECT': RGBColor(0xE9,0xA8,0x20),
                  'BENCHMARK DE REFERENCE': RGBColor(0x1D,0x7A,0x8A),
                  'COMPLEMENTAIRE': RGBColor(0x2D,0x86,0x59),
                  'PARTENAIRE POTENTIEL': RGBColor(0x2D,0x86,0x59),
                  'PARTENAIRE STRATEGIQUE': RGBColor(0x2D,0x86,0x59),
                  'MODELE A ETUDIER': RGBColor(0x1D,0x7A,0x8A),
                  'BENCHMARK PRODUIT': RGBColor(0x1D,0x7A,0x8A),
                  'MODELE DE CONTENU': RGBColor(0x1D,0x7A,0x8A),
                  'MODELE D\'INDICES': RGBColor(0x1D,0x7A,0x8A),
                  'GEANTS A DISRUPTER': RGBColor(0xE6,0x39,0x46),
                  'CONCURRENT LOCAL LIMITE': RGBColor(0x88,0x96,0xAB),
                  'SOURCE DE DONNEES': RGBColor(0x2D,0x86,0x59)}
    r.font.color.rgb = cat_colors.get(f['cat'], RGBColor(0x6B,0x72,0x80))

    tbl(['','Detail'], [
        ['URL', f['url']],
        ['Origine', f['pays']],
        ['Description', f['desc']],
        ['Couverture', f['couv']],
        ['Prix', f['prix']],
    ])

    p = doc.add_paragraph()
    r = p.add_run('Forces : '); r.font.bold = True; r.font.color.rgb = RGBColor(0x2D,0x86,0x59)
    p.add_run(' | '.join(f['forces']))

    p = doc.add_paragraph()
    r = p.add_run('Faiblesses : '); r.font.bold = True; r.font.color.rgb = RGBColor(0xE6,0x39,0x46)
    p.add_run(' | '.join(f['faiblesses']))

    p = doc.add_paragraph()
    r = p.add_run('Pertinence DATATYM : '); r.font.bold = True; r.font.color.rgb = RGBColor(0xB5,0x89,0x1A)
    p.add_run(f['pertinence'])

    doc.add_paragraph()


# ═══ MATRICE COMPARATIVE ═══
doc.add_page_break()
doc.add_heading('Matrice comparative', level=1)

tbl(['Plateforme','Donnees\nprimaires','Self-\nservice','Indices\nproprio.','Afrique\nfranc.','Prix\naccessible','IA /\nPredictif'], [
    ['DATATYM AI','OUI\n(Marketym)','OUI\n(datatym.ai)','OUI\n(IPE,IRTA...)','COEUR\nde cible','OUI\n(500-2K$/an)','EN DEV'],
    ['Stears','NON','OUI','Partiel','NON','Moyen','NON'],
    ['Sagaci','OUI','NON','NON','Partiel','NON\n(10-50K$)','NON'],
    ['Statista','NON','OUI','NON','Marginal','NON\n(2.4K$)','NON'],
    ['Asoko','Partiel','OUI','NON','NON','NON\n(5-20K$)','NON'],
    ['Ask Afrika','OUI','NON','OUI','NON','NON','NON'],
    ['GeoPoll','OUI','OUI','NON','NON','OUI\n(3$/compl.)','NON'],
    ['Kantar/Ipsos','OUI','NON','OUI','Partiel','NON\n(100K$+)','NON'],
    ['Numbeo','Crowd','OUI','OUI','Marginal','Gratuit','NON'],
    ['Briter','NON','OUI','NON','NON','Freemium','NON'],
])

bold_p('DATATYM AI est la seule plateforme qui combine : donnees primaires terrain + self-service + indices proprietaires + couverture francophone + prix accessible. C\'est un ocean bleu.')

# ═══ CE QUE DATATYM DOIT PRENDRE DE CHACUN ═══
doc.add_heading('Ce que DATATYM doit prendre de chaque acteur', level=1)

tbl(['Acteur','Ce qu\'on prend','Concretement'], [
    ['Statista','Interface + UX + SEO','Visualisations pretes a l\'emploi.\nMoteur de recherche de donnees.\nPages sectorielles indexees Google.'],
    ['Stears','Positionnement BI + pivot B2B','Se vendre comme BI, pas comme "etudes".\nLe B2B genere 75% du CA chez eux.'],
    ['Sagaci','Rigueur terrain + cross-pays','Methodologie documentee et repetable.\nComparabilite inter-pays.'],
    ['Partech','Rapport gratuit = notoriete','Un rapport annuel phare, gratuit,\nqui genere du buzz et des leads.'],
    ['Numbeo','Indices comparatifs + API','Indices par ville/pays vendus via API.\nSEO sur "cout de [X] en [pays]".'],
    ['Briter','Dashboards interactifs','Interface moderne avec filtres\npays/secteur/periode.'],
    ['GeoPoll','Canal de collecte mobile','3$/complete dans les pays sans\nenqueteurs Marketym.'],
    ['Ask Afrika','Indices reconnus localement','Faire de l\'IPE/IRTA des references\ncitees par la presse.'],
    ['JA/Africa CEO Forum','Distribution decideurs','Co-branding de rapports.\nPresence au Forum annuel.'],
])

# ═══ FIN ═══
doc.add_page_break()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('\n\nDATATYM\u2122\nBenchmark Plateformes Data & BI Afrique\n\nH&C Executive \xb7 Confidentiel \xb7 Avril 2026')
r.font.size=Pt(12); r.font.color.rgb=RGBColor(0x6B,0x72,0x80)

out='/Users/Apple/Desktop/H&C/Enquetes/Suivi enquetes/enquetes/Benchmark_Plateformes_DATATYM_2026.docx'
doc.save(out); print(f'OK -> {out}')
