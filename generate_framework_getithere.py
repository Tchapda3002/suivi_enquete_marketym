#!/usr/bin/env python3
"""
GETITHERE — Framework Universel de Questions Sectorielles V2
Version améliorée : tables remplies, priorités, sources concrètes, lien décision
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.12

for lv in range(1, 4):
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x0F, 0x1E, 0x3A)
    hs.font.size = Pt([0, 16, 13, 11][lv])
    hs.paragraph_format.space_before = Pt(12)


def tbl(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.font.bold = True; r.font.size = Pt(8)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(8)
    doc.add_paragraph()

def bul(items, size=10):
    for it in items:
        p = doc.add_paragraph(it, style='List Bullet')
        for r in p.runs: r.font.size = Pt(size)

def quote(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = True; r.font.italic = True
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x0F, 0x1E, 0x3A)

def note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

def section_break():
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ════════════════════════════════════════════════════════════════
for _ in range(2): doc.add_paragraph()
for text, sz, bold, col in [
    ('GETITHERE by datatym.ai', 12, True, RGBColor(0x00, 0x8F, 0x82)),
    ('Plateforme de Business Intelligence | Afrique', 11, False, RGBColor(0x6B, 0x72, 0x80)),
    ('', 8, False, None),
    ('FRAMEWORK UNIVERSEL', 32, True, RGBColor(0x0F, 0x1E, 0x3A)),
    ('Plan de Questions Sectorielles', 18, False, RGBColor(0xB5, 0x89, 0x1A)),
    ('', 6, False, None),
    ('7 dimensions  x  5 questions  =  35 questions universelles', 13, False, RGBColor(0x1A, 0x20, 0x2C)),
    ('adaptees a 8 secteurs  x  16+ pays africains', 13, False, RGBColor(0x1A, 0x20, 0x2C)),
    ('', 12, False, None),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(sz); r.font.bold = bold
    if col: r.font.color.rgb = col

# Chiffres clés
tbl(
    ['7', '35', '8', '280', '3'],
    [['Dimensions\nuniverselles', 'Questions\ncles', 'Secteurs\ncouverts',
      'Adaptations\nsectorielles', 'Niveaux de\npriorite']],
)

note('Document confidentiel — Usage interne GETITHERE / H&C Executive — V2 — Avril 2026')
section_break()


# ════════════════════════════════════════════════════════════════
# SOMMAIRE
# ════════════════════════════════════════════════════════════════
doc.add_heading('Sommaire', level=1)
for s in [
    'Introduction & Principes directeurs',
    'Mode d\'emploi operationnel',
    'Systeme de priorites & Scoring',
    'DIMENSION 01 — Performance & Resultats economiques',
    'DIMENSION 02 — Structure du marche & Concurrence',
    'DIMENSION 03 — Demande & Comportement des clients',
    'DIMENSION 04 — Couts, Prix & Rentabilite',
    'DIMENSION 05 — Regulation, Conformite & Risques',
    'DIMENSION 06 — Ressources Humaines & Capacites',
    'DIMENSION 07 — Perspectives & Decisions d\'Investissement',
    'Matrice de collecte par pays',
    'KPIs de collecte & Scoring de maturite',
    'Recapitulatif des 35 questions',
]:
    doc.add_paragraph(s, style='List Number')
section_break()


# ════════════════════════════════════════════════════════════════
# INTRODUCTION
# ════════════════════════════════════════════════════════════════
doc.add_heading('Introduction & Principes directeurs', level=1)

doc.add_heading('Pourquoi ce framework', level=2)
doc.add_paragraph(
    'GETITHERE ne vend pas de la donnee — il vend la capacite a decider. '
    'Ce framework garantit que pour chaque secteur couvert, on cherche systematiquement '
    'les memes 7 categories d\'information, dans le meme ordre, avec la meme rigueur. '
    'C\'est la colonne vertebrale de toute la strategie de collecte.'
)

doc.add_heading('Les 4 principes GETITHERE', level=2)
tbl(
    ['Principe', 'Regle', 'Pourquoi'],
    [
        ['Pas de donnee\nsans decision', 'Chaque question collectee doit servir\nune decision concrete d\'un decideur.',
         'La donnee sans usage est du bruit.\nOn ne collecte que ce qui sert a decider.'],
        ['Triangulation\nobligatoire', 'Minimum 3 sources croisees\npar reponse.',
         'Une source unique = une opinion.\nTrois sources croisees = un fait.'],
        ['Terrain d\'abord', 'Prioriser les sources primaires\n(dirigeants, regulateurs) sur les secondaires.',
         'Les rapports publics sont en retard de 2 ans.\nLe terrain donne le present.'],
        ['Granularite locale', 'Toujours descendre au niveau\npays > ville > quartier.',
         '"L\'Afrique de l\'Ouest" ne sert personne.\nUn DG a Abidjan veut Abidjan.'],
    ]
)

doc.add_heading('Les 7 dimensions', level=2)
tbl(
    ['#', 'Dimension', 'Ce qu\'elle mesure', 'La decision qu\'elle sert'],
    [
        ['D01', 'Performance &\nResultats economiques', 'Ce que le secteur produit\ncomme chiffres cles',
         'Faut-il investir dans ce secteur ?\nQuelle est sa trajectoire ?'],
        ['D02', 'Structure du marche\n& Concurrence', 'Qui joue, comment,\navec quelle part',
         'Ou me positionner ?\nQui sont mes vrais concurrents ?'],
        ['D03', 'Demande &\nComportement clients', 'Qui achete, combien,\npourquoi',
         'Mon marche est-il solvable ?\nA quel prix ?'],
        ['D04', 'Couts, Prix\n& Rentabilite', 'La mecanique financiere\nreelle du secteur',
         'Est-ce rentable ?\nQuel est mon point mort ?'],
        ['D05', 'Regulation,\nConformite & Risques', 'Le cadre qui contraint\net faconne le secteur',
         'Quels risques je prends ?\nQu\'est-ce qui va changer ?'],
        ['D06', 'Ressources Humaines\n& Capacites', 'Le capital humain et\norganisationnel',
         'Vais-je trouver les competences ?\nA quel cout ?'],
        ['D07', 'Perspectives &\nDecisions d\'Investissement', 'Ou va le secteur,\nqui investit, combien',
         'Ou va l\'argent ?\nQu\'est-ce qui va disruper ?'],
    ]
)

doc.add_heading('Les 8 secteurs couverts', level=2)
tbl(
    ['#', 'Secteur', 'Regulateur principal (UEMOA)', 'Priorite de lancement'],
    [
        ['01', 'Banques & IMF', 'BCEAO, Commission Bancaire', 'P1 — Immediat'],
        ['02', 'Telecoms', 'ARTP/ARCEP par pays', 'P1 — Immediat'],
        ['03', 'Immobilier', 'Ordres des notaires, cadastre', 'P1 — Immediat'],
        ['04', 'Mines & Ressources naturelles', 'Ministeres des Mines, cadastres miniers', 'P2 — Court terme'],
        ['05', 'Energie', 'Autorites de regulation energie', 'P2 — Court terme'],
        ['06', 'Fintech', 'BCEAO (instructions), banques centrales', 'P2 — Court terme'],
        ['07', 'FMCG & Distribution', 'Directions du commerce, douanes', 'P3 — Moyen terme'],
        ['08', 'Agriculture', 'Ministeres agriculture, ECOWAS/CEDEAO', 'P3 — Moyen terme'],
    ]
)
section_break()


# ════════════════════════════════════════════════════════════════
# MODE D'EMPLOI
# ════════════════════════════════════════════════════════════════
doc.add_heading('Mode d\'emploi operationnel', level=1)

doc.add_heading('Les 5 etapes de collecte', level=2)
tbl(
    ['Etape', 'Action', 'Livrable', 'Duree'],
    [
        ['1. Identifier\nles cibles', 'Pour chaque dimension, lister 3 profils\nde detenteurs par secteur et par pays.\nPrioriser sources primaires.',
         'Liste de contacts\npar dimension/secteur/pays', '1 semaine\n/ secteur / pays'],
        ['2. Adapter le\nvocabulaire', 'Utiliser les reformulations sectorielles\nde ce framework. "CA" devient\n"encours" pour les banques, etc.',
         'Grille de questions\nadaptee au secteur', '1 jour\n/ secteur'],
        ['3. Prioriser par\nvaleur commerciale', 'Commencer par D01 (performance)\net D04 (couts & prix) — les dimensions\nque les decideurs payent le plus cher.',
         'Plan de collecte\npriorise', '1 jour'],
        ['4. Collecter\n& trianguler', 'Minimum 3 sources par reponse.\nRegulateur + acteur prive + source\nacademique/media.',
         'Fiches de donnees\nbrutes par question', '2-4 semaines\n/ secteur / pays'],
        ['5. Structurer\n& publier', 'Transformer les donnees brutes en\nindicateurs exploitables sur datatym.ai.\nCalculer les indices GETITHERE.',
         'Donnees structurees\n+ indices sur plateforme', '1 semaine\n/ secteur'],
    ]
)

doc.add_heading('Regles de qualite', level=2)
bul([
    'JAMAIS de donnee sans source identifiee — chaque chiffre a un auteur et une date',
    'JAMAIS de source unique — minimum 3 sources croisees par reponse',
    'TOUJOURS dater la donnee — une donnee de 2023 n\'a pas la meme valeur qu\'une de 2026',
    'TOUJOURS indiquer la methodologie — comment le chiffre a ete obtenu',
    'PREFERER les donnees primaires aux estimations — le terrain bat toujours le modele',
], size=9)
section_break()


# ════════════════════════════════════════════════════════════════
# SYSTEME DE PRIORITES
# ════════════════════════════════════════════════════════════════
doc.add_heading('Systeme de priorites & Scoring', level=1)

doc.add_heading('3 niveaux de priorite', level=2)
tbl(
    ['Niveau', 'Label', 'Critere', 'Delai de collecte'],
    [
        ['P1', 'CRITIQUE', 'Donnee pour laquelle un decideur\npayerait immediatement.\nValeur commerciale maximale.',
         'Collecte en priorite.\nDoit etre disponible\nau lancement du secteur.'],
        ['P2', 'IMPORTANT', 'Donnee qui enrichit la decision\nmais n\'est pas bloquante.\nValeur contextuelle.',
         'A collecter dans les\n3 mois suivant le\nlancement du secteur.'],
        ['P3', 'UTILE', 'Donnee de fond, tendancielle.\nEnrichit la plateforme mais\nn\'est pas urgente.',
         'A collecter dans les\n6 mois. Peut attendre\nla vague 2.'],
    ]
)

doc.add_heading('Priorite par dimension (recommandation)', level=2)
tbl(
    ['Dimension', 'Priorite', 'Justification'],
    [
        ['D01 Performance & Resultats', 'P1', 'C\'est la premiere question de tout decideur : "ca vaut combien ?"'],
        ['D04 Couts, Prix & Rentabilite', 'P1', 'La deuxieme question : "c\'est rentable ?"'],
        ['D02 Structure & Concurrence', 'P1', '"Qui sont mes concurrents et quelle est leur part ?"'],
        ['D05 Regulation & Risques', 'P2', 'Critique pour les investisseurs, moins urgent pour les operateurs'],
        ['D07 Perspectives & Investissement', 'P2', '"Ou va l\'argent ?" — question d\'investisseur'],
        ['D03 Demande & Comportement', 'P2', 'Important mais plus dur a collecter — necessite du terrain'],
        ['D06 Ressources Humaines', 'P3', 'Couvert en partie par DATATYM Talent. Complementaire.'],
    ]
)

doc.add_heading('Scoring de maturite par secteur/pays', level=2)
doc.add_paragraph('Pour chaque couple secteur x pays, un score de maturite de collecte :')
tbl(
    ['Score', 'Label', 'Critere', 'Couverture'],
    [
        ['0', 'VIERGE', 'Aucune donnee collectee', '0 / 35 questions'],
        ['1', 'AMORCE', 'D01 et D04 partiellement couvertes\n(performance + prix)', '1-10 / 35'],
        ['2', 'EMERGENT', 'D01, D02, D04 couvertes\n+ sources identifiees pour le reste', '11-20 / 35'],
        ['3', 'OPERATIONNEL', 'Toutes les P1 et P2 couvertes\nTriangulation validee', '21-30 / 35'],
        ['4', 'MATURE', 'Les 35 questions couvertes\nMise a jour trimestrielle active', '31-35 / 35'],
    ]
)
section_break()


# ════════════════════════════════════════════════════════════════
# DIMENSIONS 01-07
# ════════════════════════════════════════════════════════════════

dimensions = [
    {
        'num': '01', 'titre': 'Performance & Resultats economiques',
        'sous': 'Ce que le secteur produit comme chiffres cles',
        'decision': 'Faut-il investir ? Le secteur croit-il ? A quelle vitesse ?',
        'priorite': 'P1 — CRITIQUE',
        'questions': [
            ('Q01', 'Quel est le chiffre d\'affaires moyen des acteurs de ce secteur, et comment se repartit-il entre les leaders et les acteurs secondaires ?'),
            ('Q02', 'Quelle est la marge operationnelle reelle (pas declaree) du secteur, et quels sont les principaux postes de charge ?'),
            ('Q03', 'Quelle est la trajectoire de croissance du secteur sur les 3 dernieres annees — volume, valeur, nombre d\'acteurs ?'),
            ('Q04', 'Quels sont les ratios de productivite cles du secteur — revenu par employe, rotation des actifs, rendement du capital investi ?'),
            ('Q05', 'Quelle est la saisonnalite de l\'activite — quels mois ou trimestres concentrent l\'essentiel du chiffre d\'affaires et pourquoi ?'),
        ],
        'detenteurs': 'Directions financieres des acteurs majeurs, regulateurs sectoriels, ordres professionnels, rapports annuels publies, ministeres de tutelle.',
        'adaptations': [
            ('Banques & IMF', 'PNB (Produit Net Bancaire), encours\nde credits, total bilan, coefficient\nd\'exploitation', 'BCEAO, Commission Bancaire,\nrapports annuels banques', '1'),
            ('Telecoms', 'CA par operateur, ARPU, nombre\nd\'abonnes, revenus data vs voix', 'ARTP/ARCEP, rapports operateurs,\nGSMA Intelligence', '1'),
            ('Immobilier', 'Volume de transactions, prix moyen\nau m2, nombre de permis delivres,\nloyers moyens', 'Ordres des notaires, cadastre,\nagences immobilieres partenaires', '2'),
            ('Mines', 'Production en tonnes/onces, valeur\nexportee, redevances versees,\ncours des matieres', 'Cadastres miniers, ministeres,\nrapports EITI, operateurs', '2'),
            ('Energie', 'Capacite installee (MW), production\n(GWh), taux d\'acces, pertes\ntechniques', 'Autorites de regulation, societes\nnationales (EDC, CIE, SENELEC)', '2'),
            ('Fintech', 'Volume de transactions, nombre\nd\'utilisateurs actifs, valeur des\ntransferts, GMV', 'BCEAO (GIM-UEMOA), operateurs,\nrapports investisseurs', '2'),
            ('FMCG & Distribution', 'CA par categorie, volumes vendus,\nparts de marche par marque,\ntaux de penetration', 'Nielsen/Kantar (si dispo), grandes\nsurfaces, importateurs', '3'),
            ('Agriculture', 'Production par culture (tonnes),\nsuperficies cultivees, rendements,\nexportations', 'Ministeres agriculture, FAO,\ncooperatives, bourses matieres', '2'),
        ],
    },
    {
        'num': '02', 'titre': 'Structure du marche & Concurrence',
        'sous': 'Qui joue, comment, avec quelle part',
        'decision': 'Ou me positionner ? Qui sont mes vrais concurrents ? Le marche se consolide-t-il ?',
        'priorite': 'P1 — CRITIQUE',
        'questions': [
            ('Q06', 'Quelle est la concentration du marche — les 3 premiers acteurs representent quelle part du total, et cette concentration evolue-t-elle ?'),
            ('Q07', 'Quelles sont les barrieres a l\'entree reelles — capital, licence, reseau, technologie — et comment ont-elles evolue ces 3 dernieres annees ?'),
            ('Q08', 'Quels sont les modeles economiques qui gagnent vs ceux qui peinent, et pourquoi dans ce contexte africain specifique ?'),
            ('Q09', 'Quelle est la geographie reelle du marche — quelle part de l\'activite est concentree dans les capitales vs les villes secondaires vs le rural ?'),
            ('Q10', 'Quels sont les acteurs regionaux qui s\'etendent — qui gagne des parts de marche trans-pays et avec quelle strategie ?'),
        ],
        'detenteurs': 'Associations sectorielles, regulateurs, cabinets de conseil, dirigeants des acteurs leaders, presse economique.',
        'adaptations': [
            ('Banques & IMF', 'Part de marche par total bilan,\nHHI index, nombre d\'agences,\ntaux de bancarisation', 'BCEAO classement, rapports\nannuels, Commission Bancaire', '1'),
            ('Telecoms', 'Parts de marche par abonnes et CA,\nconcentration HHI, strategies MVNO,\ninfrastructure partagee', 'ARTP/ARCEP, rapports trimestriels\noperateurs', '1'),
            ('Immobilier', 'Nombre d\'agences par ville,\npromoteurs dominants, part formel\nvs informel, entree de groupes\ninternationaux', 'Chambre des notaires, registres\ndu commerce, enquete terrain', '2'),
            ('Mines', 'Nombre de permis actifs vs en\nexploration, part des juniors vs\nmajors, acteurs chinois/turcs', 'Cadastre minier, ministere,\nrapports EITI', '2'),
            ('Energie', 'Monopole public vs IPP vs off-grid,\npart du solaire, nombre de\nlicences privees', 'Autorite regulation, ministere\nenergie, IRENA', '2'),
            ('Fintech', 'Mobile money vs neobanques vs\nagreateurs, parts par volume,\nnombre de licences EME', 'BCEAO, GIM-UEMOA, Partech\nAfrica reports', '2'),
            ('FMCG & Distribution', 'Part moderne vs traditionnel vs\ninformel, couverture geographique\ndes distributeurs', 'Nielsen, grandes surfaces,\nimportateurs', '3'),
            ('Agriculture', 'Part cooperatives vs traders vs\nexport direct, concentration\npar filiere', 'Ministere agriculture, bourses,\ninterprofessions', '2'),
        ],
    },
    {
        'num': '03', 'titre': 'Demande & Comportement des clients',
        'sous': 'Qui achete, combien, pourquoi',
        'decision': 'Mon marche est-il solvable ? A quel prix la demande decroche ?',
        'priorite': 'P2 — IMPORTANT',
        'questions': [
            ('Q11', 'Quelle est la taille reelle de la demande solvable — combien de clients potentiels ont la capacite financiere d\'acheter ou de souscrire ?'),
            ('Q12', 'Quelle est la structure de la decision d\'achat — qui decide, en combien de temps, avec quels criteres de choix prioritaires ?'),
            ('Q13', 'Quels sont les principaux motifs d\'insatisfaction ou d\'abandon des clients actuels du secteur ?'),
            ('Q14', 'Quelle est l\'elasticite prix — au-dela de quel niveau de prix la demande se contracte-t-elle significativement ?'),
            ('Q15', 'Quelle est la part du marche informel ou parallele qui capte de la demande que le secteur formel ne sert pas ?'),
        ],
        'detenteurs': 'Directions commerciales, etudes consommateurs (Nielsen, Kantar), regulateurs, INS/ANSD, associations de consommateurs, panels qualitatifs terrain.',
        'adaptations': [
            ('Banques & IMF', 'Taux de bancarisation, demande de\ncredit non servie, nombre de comptes\nactifs vs inactifs', 'BCEAO, Global Findex, enquetes\nterrain', '2'),
            ('Telecoms', 'Penetration mobile, nombre de SIM\nactives, ARPU par segment,\nchurn rate', 'ARTP/ARCEP, enquetes conso,\nrapports operateurs', '1'),
            ('Immobilier', 'Deficit logement par ville, demande\nlocative vs acquisitive, budget moyen\npar segment', 'INS, enquetes terrain, agences,\nbanques (credits immo)', '3'),
            ('Mines', 'Demande internationale (cours),\ncontrats long terme vs spot,\nclient final vs trader', 'LME, bourses matieres, contrats\npublies, operateurs', '2'),
            ('Energie', 'Demande non servie (gap acces),\nconsommation par segment\n(residentiel/industriel/commercial)', 'Autorite regulation, societes\nnationales, enquetes menages', '2'),
            ('Fintech', 'Transactions P2P vs P2B vs B2B,\nfreqence d\'usage, panier moyen,\nadoption par tranche d\'age', 'GIM-UEMOA, operateurs, Global\nFindex, enquetes DATATYM', '2'),
            ('FMCG & Distribution', 'Panier moyen menage, frequence\nd\'achat, fidelite marque vs prix,\npart du vrac vs emballe', 'Nielsen, enquetes menages,\ngrandes surfaces', '3'),
            ('Agriculture', 'Demande interieure vs export,\nautosuffisance par culture,\nimportations alimentaires', 'FAO, douanes, ministere\nagriculture, statistiques commerce', '2'),
        ],
    },
    {
        'num': '04', 'titre': 'Couts, Prix & Rentabilite',
        'sous': 'La mecanique financiere reelle du secteur',
        'decision': 'Est-ce rentable ? Quel est le point mort ? Comment optimiser ?',
        'priorite': 'P1 — CRITIQUE',
        'questions': [
            ('Q16', 'Quelle est la structure de couts type d\'un acteur du secteur — quels postes representent plus de 15% des charges totales ?'),
            ('Q17', 'Quel est le prix de revient reel d\'une unite vendue ou d\'un service delivre, et comment se compare-t-il au prix affiche ?'),
            ('Q18', 'Quelle est la politique de prix du marche — y a-t-il un prix de reference implicite, des pratiques de discount, une pression tarifaire a la baisse ?'),
            ('Q19', 'Quel est le point mort typique d\'un acteur de taille moyenne — en volume, en revenus, en mois d\'exploitation ?'),
            ('Q20', 'Quelles sont les opportunites de reduction de couts que les acteurs leaders ont identifiees et exploitent deja ?'),
        ],
        'detenteurs': 'CFO et controleurs de gestion, commissaires aux comptes, associations sectorielles, fournisseurs du secteur, cabinets d\'audit (Deloitte, PwC, EY locaux).',
        'adaptations': [
            ('Banques & IMF', 'Coefficient d\'exploitation, cout du\nrisque, cout de la ressource, marge\nd\'intermediation', 'Rapports annuels, BCEAO,\ncommissaires aux comptes', '1'),
            ('Telecoms', 'Cout par Go, cout d\'acquisition\nabonne, CAPEX/OPEX ratio, cout\nde maintenance reseau', 'Rapports operateurs, ARTP,\nfournisseurs (Huawei, Ericsson)', '2'),
            ('Immobilier', 'Cout de construction au m2,\nmarge promoteur, frais de notaire,\ntaxe fonciere, rendement locatif', 'Promoteurs, architectes, ordres\ndes notaires, enquete terrain', '2'),
            ('Mines', 'All-in sustaining cost (AISC),\ncout d\'extraction par tonne,\nredevances et taxes minieres', 'Rapports operateurs, cadastre\nminier, EITI', '2'),
            ('Energie', 'Cout du kWh produit par source\n(thermique/solaire/hydro), tarif\nregule vs cout reel', 'Autorite regulation, operateurs,\nIRENA, rapports IPP', '2'),
            ('Fintech', 'Cout par transaction, cout\nd\'acquisition client (CAC), taux\nde commission, burn rate', 'Operateurs, investisseurs, BCEAO,\nrapports financiers', '2'),
            ('FMCG & Distribution', 'Marge distributeur, cout logistique\n(% du prix), cout d\'import (droits\n+ transport)', 'Importateurs, transitaires, grandes\nsurfaces, douanes', '3'),
            ('Agriculture', 'Cout de production par hectare/tonne,\ncout des intrants (engrais, semences),\nmarge producteur', 'Cooperatives, ministere agriculture,\nfournisseurs intrants', '2'),
        ],
    },
    {
        'num': '05', 'titre': 'Regulation, Conformite & Risques',
        'sous': 'Le cadre qui contraint et faconne le secteur',
        'decision': 'Quels risques je prends ? Qu\'est-ce qui va changer ? Combien coute la conformite ?',
        'priorite': 'P2 — IMPORTANT',
        'questions': [
            ('Q21', 'Quelles sont les evolutions reglementaires en cours ou prevues qui vont modifier les regles du jeu dans les 12-24 prochains mois ?'),
            ('Q22', 'Quel est le cout reel de la conformite — licences, reporting reglementaire, audits, fonds propres requis — en pourcentage du chiffre d\'affaires ?'),
            ('Q23', 'Quels sont les principaux risques operationnels du secteur — fraude, sinistres, defauts, interruptions — et comment les acteurs les gerent-ils ?'),
            ('Q24', 'Quelle est la qualite reelle de l\'application de la reglementation — y a-t-il un ecart entre ce qui est impose et ce qui est reellement controle ?'),
            ('Q25', 'Quelles sont les sanctions et litiges recents dans le secteur — que revelent-ils sur les zones de risque reel ?'),
        ],
        'detenteurs': 'Regulateurs sectoriels (BCEAO, BEAC, ARCEP), cabinets juridiques OHADA, directions conformite des grands acteurs, journaux officiels.',
        'adaptations': [
            ('Banques & IMF', 'Ratios prudentiels BCEAO (Bale II/III),\nlicences bancaires, reporting CRC,\nLAB/FT', 'BCEAO, Commission Bancaire,\ncabinets juridiques', '1'),
            ('Telecoms', 'Licences 3G/4G/5G, obligations\ncouverture, partage infrastructure,\ntaxes sur appels', 'ARTP/ARCEP, ministere telecoms,\ntextes officiels', '1'),
            ('Immobilier', 'Permis de construire, normes\nurbanistiques, fiscalite immobiliere,\ndroit OHADA', 'Mairies, cadastre, ordres\nprofessionnels, fiscalite', '2'),
            ('Mines', 'Code minier, redevances, obligations\nenvironnementales, contenu local,\nreporting EITI', 'Ministere mines, cadastre minier,\nEITI national', '1'),
            ('Energie', 'Tarification regulee, obligations\nde service public, normes\nenvironnementales, PPP', 'Autorite regulation, ministere\nenergie, textes PPP', '2'),
            ('Fintech', 'Instructions BCEAO (EME, PSP),\nKYC digital, plafonds transactions,\ninteroperabilite', 'BCEAO, GIM-UEMOA, cabinets\njuridiques fintech', '1'),
            ('FMCG & Distribution', 'Normes sanitaires (CEDEAO),\ndroits de douane, reglementation\nprix, etiquetage', 'Directions du commerce, douanes,\nnormes CEDEAO', '2'),
            ('Agriculture', 'Subventions, normes phytosanitaires,\nacces au foncier, politique agricole\nregionale', 'ECOWAP/CEDEAO, ministere\nagriculture, FAO', '2'),
        ],
    },
    {
        'num': '06', 'titre': 'Ressources Humaines & Capacites',
        'sous': 'Le capital humain et organisationnel du secteur',
        'decision': 'Vais-je trouver les competences ? A quel cout ? Quel est le turnover ?',
        'priorite': 'P3 — UTILE (complemente par DATATYM Talent)',
        'questions': [
            ('Q26', 'Quelle est la taille de la main d\'oeuvre du secteur — effectifs formels, niveau de qualification dominant, ratio cadres sur executants ?'),
            ('Q27', 'Quelles sont les grilles salariales reelles par fonction et niveau — ecarts entre grandes entreprises, PME et organisations publiques ?'),
            ('Q28', 'Quelles sont les competences en tension — profils introuvables, delais de recrutement anormalement longs, turnover eleve sur certains postes ?'),
            ('Q29', 'Quel est le niveau de digitalisation des equipes — adoption des outils digitaux, formation disponible, resistance au changement ?'),
            ('Q30', 'Quelle est la politique de formation et developpement — budget formation moyen, partenariats ecoles, programmes d\'apprentissage interne ?'),
        ],
        'detenteurs': 'DRH des grands acteurs, cabinets de recrutement, ministeres du Travail, resultats du Barometre DATATYM Talent.',
        'adaptations': [
            ('Banques & IMF', 'Effectifs par banque, grilles salariales\nbanquiers, turnover analystes/risk,\ndigitalisation agences', 'APBEF, rapports annuels, cabinets\nRH, DATATYM Talent', '2'),
            ('Telecoms', 'Effectifs operateurs, salaires\ningenieurs reseau/IT, turnover tech,\nformation digitale', 'Operateurs, cabinets RH, DATATYM\nTalent', '2'),
            ('Immobilier', 'Nombre d\'agents immobiliers,\nsalaires par fonction, formation\nprofessionnelle, turnover', 'Ordres professionnels, agences,\nenquete terrain', '3'),
            ('Mines', 'Effectifs par mine, ratio expat vs\nlocal, grilles operateurs/ingenieurs,\ncompetences QHSE', 'Operateurs, ministere travail,\nchambre des mines', '2'),
            ('Energie', 'Effectifs societes nationales,\ncompetences energies renouvelables,\nformation technique', 'Societes nationales, ministere\nenergie, ecoles ingenieurs', '3'),
            ('Fintech', 'Taille equipes par fintech, salaires\ndev/product/data, concurrence talent\navec les banques', 'Fintechs, cabinets RH, LinkedIn\ndata, DATATYM Talent', '2'),
            ('FMCG & Distribution', 'Effectifs distribution (formel),\nsalaires logistique/commercial,\nformation terrain', 'Grandes surfaces, distributeurs,\ncabinets RH', '3'),
            ('Agriculture', 'Main d\'oeuvre agricole (formelle\nvs informelle), mecanisation,\nformation technique', 'Ministere agriculture, cooperatives,\nOIT/BIT', '3'),
        ],
    },
    {
        'num': '07', 'titre': 'Perspectives & Decisions d\'Investissement',
        'sous': 'Ou va le secteur, qui investit, combien',
        'decision': 'Ou va l\'argent ? Qu\'est-ce qui va disruper ? Faut-il investir maintenant ?',
        'priorite': 'P2 — IMPORTANT',
        'questions': [
            ('Q31', 'Quels sont les projets d\'investissement majeurs en cours ou planifies dans le secteur — capex, expansion geographique, acquisitions ?'),
            ('Q32', 'Quels sont les postes budgetaires prioritaires pour l\'annee prochaine — ou les dirigeants vont-ils allouer le plus de ressources ?'),
            ('Q33', 'Quelles sont les technologies ou innovations que le secteur envisage d\'adopter dans les 24 prochains mois et a quels couts ?'),
            ('Q34', 'Quelle est la strategie d\'expansion regionale des acteurs leaders — quels pays ciblent-ils et avec quel modele d\'entree ?'),
            ('Q35', 'Quels sont les signaux avances d\'une rupture de marche — nouveaux entrants, changement de modele, disruption technologique imminente ?'),
        ],
        'detenteurs': 'DG et directeurs strategie, fonds d\'investissement, presse economique, conferences et salons sectoriels, rapports DFI (AFD, BAD, IFC).',
        'adaptations': [
            ('Banques & IMF', 'Nouvelles licences, fusion/acquisitions,\nexpansion sous-regionale, investissement\ndigital banking', 'BCEAO, presse eco, Jeune Afrique,\nrapports BAD/IFC', '1'),
            ('Telecoms', 'Deploiement 5G, fibre optique,\nacquisitions, mobile money,\npartenariats OTT', 'Operateurs, GSMA, presse,\nlevees de fonds', '1'),
            ('Immobilier', 'Projets immobiliers annonces,\nzones d\'amenagement, investissements\netrangers, PPP logement', 'Promoteurs, mairies, presse,\nbanques (credits immo)', '2'),
            ('Mines', 'Nouveaux permis, investissements\nexploration, projets d\'infrastructures\nminieres, M&A', 'Cadastre minier, presse, rapports\noperateurs, Africa Mining Summit', '2'),
            ('Energie', 'Projets solaires/eoliens, IPP\nannonces, electrification rurale,\nstockage batterie', 'Autorite regulation, DFI (BAD,\nAFD, IFC), presse', '2'),
            ('Fintech', 'Levees de fonds, nouvelles licences,\npartenariats banques, expansion\nregionale, M&A', 'Partech Africa, TechCrunch Africa,\nDisrupt Africa, BCEAO', '1'),
            ('FMCG & Distribution', 'Nouveaux entrants, expansion\nchaînes, e-commerce, logistique\ndu dernier kilometre', 'Presse eco, rapports investisseurs,\ngrandes surfaces', '3'),
            ('Agriculture', 'Investissements agro-industriels,\nprojet irrigation, transformation\nlocale, export', 'Ministere agriculture, BAD, FAO,\nagro-industriels', '2'),
        ],
    },
]

for dim in dimensions:
    doc.add_heading(f'DIMENSION {dim["num"]}', level=1)
    doc.add_heading(dim['titre'], level=2)
    note(dim['sous'])

    p = doc.add_paragraph()
    r = p.add_run(f'Priorite : {dim["priorite"]}')
    r.font.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xC4, 0x2B, 0x1E) if 'P1' in dim['priorite'] else RGBColor(0xB5, 0x89, 0x1A)

    p = doc.add_paragraph()
    r = p.add_run(f'Decision servie : ')
    r.font.bold = True; r.font.size = Pt(10)
    p.add_run(dim['decision']).font.size = Pt(10)

    doc.add_heading('Les 5 questions', level=3)
    for qid, qtxt in dim['questions']:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(f'{qid}  ')
        r.font.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x00, 0x8F, 0x82)
        run2 = p.add_run(qtxt)
        run2.font.size = Pt(9)

    note(f'Detenteurs types : {dim["detenteurs"]}')

    doc.add_heading('Adaptation par secteur', level=3)
    tbl(
        ['Secteur', 'Reformulation / KPIs cles', 'Detenteur principal', 'Diff.'],
        dim['adaptations'],
    )

    section_break()


# ════════════════════════════════════════════════════════════════
# MATRICE PAR PAYS
# ════════════════════════════════════════════════════════════════
doc.add_heading('Matrice de collecte par pays', level=1)

doc.add_paragraph(
    'Pour chaque pays cible, identifier les sources prioritaires par dimension. '
    'Les 3 premiers pays (P1) : Cote d\'Ivoire, Senegal, Cameroun.'
)

tbl(
    ['Source', 'Type', 'Dimensions\ncouvertes', 'Pays\ncouverts', 'Acces'],
    [
        ['BCEAO\n(bfrancophone)', 'Regulateur bancaire', 'D01, D02, D04, D05', '8 pays UEMOA', 'Public (rapports annuels)\n+ contact direct'],
        ['BEAC', 'Regulateur bancaire', 'D01, D02, D04, D05', '6 pays CEMAC', 'Public + contact'],
        ['ARTP / ARCEP\n(par pays)', 'Regulateur telecoms', 'D01, D02, D03, D05', 'Par pays', 'Public (rapports)\n+ demande formelle'],
        ['INS / ANSD\n(par pays)', 'Institut statistique', 'D01, D03, D06', 'Par pays', 'Public'],
        ['Cadastres miniers', 'Administration mines', 'D01, D02, D05, D07', 'Par pays', 'Semi-public\n+ contact'],
        ['Ordres des notaires', 'Profession reglementee', 'D01, D04, D05', 'Par pays', 'Contact direct\n+ enquete'],
        ['Agences immobilieres\npartenaires', 'Acteurs prives', 'D01, D03, D04', 'Par ville', 'Partenariat\n(incentive visibilite)'],
        ['Rapports annuels\nentreprises cotees', 'Donnees publiques', 'D01, D02, D04', 'BRVM, DSX', 'Public\n(sites bourse)'],
        ['Barometre DATATYM\nTalent', 'Collecte proprietaire', 'D06', '16+ pays', 'Interne\n(deja disponible)'],
        ['Presse eco specialisee\n(Jeune Afrique, Financial\nAfrik, etc.)', 'Secondaire', 'D02, D07', 'Panafricain', 'Public\n(veille)'],
    ]
)

section_break()


# ════════════════════════════════════════════════════════════════
# KPIs DE COLLECTE
# ════════════════════════════════════════════════════════════════
doc.add_heading('KPIs de collecte & Suivi', level=1)

doc.add_heading('KPIs par secteur/pays', level=2)
tbl(
    ['KPI', 'Definition', 'Cible Phase 1', 'Cible Phase 2'],
    [
        ['Couverture dimensionnelle', 'Nombre de dimensions couvertes / 7', '4/7 (D01,D02,D04,D05)', '7/7'],
        ['Couverture questions', 'Nombre de questions repondues / 35', '15/35', '30/35'],
        ['Score de maturite', 'Voir grille ci-dessus (0 a 4)', '2 (Emergent)', '3 (Operationnel)'],
        ['Nombre de sources', 'Sources uniques utilisees par secteur/pays', '5+', '10+'],
        ['Taux de triangulation', '% de reponses avec 3+ sources', '50%', '80%'],
        ['Fraicheur moyenne', 'Age moyen des donnees', '< 12 mois', '< 6 mois'],
        ['Delai de publication', 'Temps entre collecte et publication', '< 2 semaines', '< 1 semaine'],
    ]
)

doc.add_heading('Tableau de bord de suivi (modele)', level=2)
tbl(
    ['Secteur x Pays', 'D01', 'D02', 'D03', 'D04', 'D05', 'D06', 'D07', 'Score'],
    [
        ['Banques — CIV', '-', '-', '-', '-', '-', '-', '-', '0'],
        ['Banques — SEN', '-', '-', '-', '-', '-', '-', '-', '0'],
        ['Banques — CMR', '-', '-', '-', '-', '-', '-', '-', '0'],
        ['Telecoms — CIV', '-', '-', '-', '-', '-', '-', '-', '0'],
        ['Telecoms — SEN', '-', '-', '-', '-', '-', '-', '-', '0'],
        ['Immobilier — CIV', '-', '-', '-', '-', '-', '-', '-', '0'],
    ]
)
note('Remplir avec le nombre de questions couvertes par dimension (0-5). Score = somme / 35 x 4.')

section_break()


# ════════════════════════════════════════════════════════════════
# RECAPITULATIF 35 QUESTIONS
# ════════════════════════════════════════════════════════════════
doc.add_heading('Recapitulatif — Les 35 questions', level=1)
note('Check-list terrain avant chaque session de collecte.')

all_questions = []
for dim in dimensions:
    for qid, qtxt in dim['questions']:
        all_questions.append((qid, f'D{dim["num"]}', dim['priorite'].split(' — ')[0], qtxt))

# Split in 2 tables to avoid overflow
for start in range(0, len(all_questions), 18):
    batch = all_questions[start:start+18]
    tbl(
        ['#', 'Dim', 'Prio', 'Question'],
        [(q[0], q[1], q[2], q[3]) for q in batch],
    )


# ── FIN ────────────────────────────────────────────────────────
doc.add_paragraph()
note('Framework GETITHERE V2 — Confidentiel — Usage interne H&C Executive — datatym.ai — Avril 2026')

out = '/Users/Apple/Desktop/H&C/Enquetes/Suivi enquetes/enquetes/Framework_Questions_GETITHERE_V2.docx'
doc.save(out)
print(f'OK -> {out}')
