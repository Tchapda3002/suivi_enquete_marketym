"""
Génération du rapport Word - Suivi des baromètres Marketym
Focus : écart par pays vs quotas
Données 100% depuis la DB (pas besoin d'API QuestionPro)
"""
import os, sys
from collections import defaultdict
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv(os.path.join(os.path.dirname(__file__), "backend/.env"))

# Connexion DB directe
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "backend"))
from app.db import DirectClient
db = DirectClient(os.getenv("DATABASE_URL"))

# ─── Chargement des données depuis la DB ─────────────────────────────────
enquetes = db.table("enquetes").select("id, code, nom, taille_echantillon, survey_id").order("code").execute()
all_affs = db.table("affectations").select("id, enquete_id, completions_total, objectif_total, clics").execute()

aff_by_enq = defaultdict(list)
for a in all_affs.data:
    aff_by_enq[a["enquete_id"]].append(a)

BAROMETRES = []
for e in enquetes.data:
    affs = aff_by_enq.get(e["id"], [])
    comp = sum(a["completions_total"] or 0 for a in affs)
    obj = e["taille_echantillon"] or 1500
    BAROMETRES.append({
        "enquete_id": e["id"],
        "code": e["code"],
        "nom": e["nom"],
        "survey_id": e["survey_id"],
        "objectif_global": obj,
        "completions_global": comp,
        "manquant": obj - comp,
        "pct_progression": round(comp / max(obj, 1) * 100, 1),
        "nb_enqueteurs": len(affs),
        "total_clics": sum(a["clics"] or 0 for a in affs),
    })

# ─── Couleurs ─────────────────────────────────────────────────────────────
VERT    = RGBColor(0x05, 0x96, 0x69)
ROUGE   = RGBColor(0xDC, 0x26, 0x26)
ORANGE  = RGBColor(0xD9, 0x77, 0x06)
BLEU    = RGBColor(0x1D, 0x4E, 0xD8)
VIOLET  = RGBColor(0x7C, 0x3A, 0xED)
NOIR    = RGBColor(0x11, 0x18, 0x27)
GRIS    = RGBColor(0x6B, 0x72, 0x80)
BLANC   = RGBColor(0xFF, 0xFF, 0xFF)


# ─── Données pays depuis la DB ────────────────────────────────────────────
def get_pays_data(enquete_id, taille_echantillon):
    """Retourne liste de dicts {pays, pct_quota, obj, comp, ecart, pct_atteint}"""
    quotas = (db.table("quotas")
                .select("segment_value, pourcentage, answer_option_id")
                .eq("enquete_id", enquete_id)
                .is_("affectation_id", "null")
                .execute())

    ao_ids = [q["answer_option_id"] for q in quotas.data if q["answer_option_id"]]
    ao_map = {}
    if ao_ids:
        aos = db.table("answer_options").select("id, valeur").in_("id", ao_ids).execute()
        ao_map = {ao["id"]: ao["valeur"] for ao in aos.data}

    affs = db.table("affectations").select("id").eq("enquete_id", enquete_id).execute()
    aff_ids = [a["id"] for a in affs.data]
    rc_totals = defaultdict(int)
    if aff_ids:
        rc = db.table("response_counts").select("answer_option_id, count").in_("affectation_id", aff_ids).execute()
        for r in rc.data:
            rc_totals[r["answer_option_id"]] += r["count"]

    rows = []
    for q in sorted(quotas.data, key=lambda x: -(x["pourcentage"] or 0)):
        ao_id      = q["answer_option_id"]
        pays       = ao_map.get(ao_id, q["segment_value"] or "—")
        pct_quota  = float(q["pourcentage"] or 0)
        obj        = round(pct_quota / 100 * taille_echantillon)
        comp       = rc_totals.get(ao_id, 0)
        ecart      = obj - comp
        pct_atteint = round(comp / max(obj, 1) * 100, 1)
        rows.append({
            "pays": pays, "pct_quota": pct_quota,
            "obj": obj, "comp": comp,
            "ecart": ecart, "pct_atteint": pct_atteint,
        })
    return rows

# ─── Helpers Word ─────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=9, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(str(text))
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color or NOIR

def para(doc, text, bold=False, size=11, color=None,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=4, after=4, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p

def section_title(doc, text, level=1):
    sizes  = {1: 16, 2: 13, 3: 11}
    colors = {1: VERT, 2: NOIR, 3: NOIR}
    p = para(doc, text, bold=True, size=sizes[level],
             color=colors[level], before=10 if level == 1 else 8, after=4)
    if level == 1:
        pPr   = p._p.get_or_add_pPr()
        pBdr  = OxmlElement("w:pBdr")
        bot   = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "059669")
        pBdr.append(bot)
        pPr.append(pBdr)
    return p

def kpi_row(doc, items):
    """items = [(label, value, bg_hex, fg_color), ...]"""
    t = doc.add_table(rows=2, cols=len(items))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value, bg, fg) in enumerate(items):
        cv = t.cell(0, i)
        set_cell_bg(cv, bg)
        cv.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_text(cv, value, bold=True, size=18, color=BLANC,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cl = t.cell(1, i)
        set_cell_bg(cl, "F3F4F6")
        cell_text(cl, label, size=8, color=GRIS,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def progress_bar_text(pct, width=25):
    filled = int(pct / 100 * width)
    return "█" * filled + "░" * (width - filled)

# ─── Tableau pays ─────────────────────────────────────────────────────────
def add_pays_table(doc, rows, taille_echantillon):
    if not rows:
        para(doc, "Aucun quota pays défini.", color=GRIS, size=9)
        return

    total_obj  = sum(r["obj"]  for r in rows)
    total_comp = sum(r["comp"] for r in rows)
    total_ecart= total_obj - total_comp

    # Note introductive
    para(doc,
         f"Quotas définis pour {len(rows)} pays — objectif global : {taille_echantillon} répondants. "
         f"Completions totales sur pays avec quota : {total_comp} / {total_obj} "
         f"(écart global : {total_ecart:+d}).",
         size=9, color=GRIS, italic=True, before=0, after=6)

    hdrs = ["Pays", "Quota %", "Objectif", "Réalisé", "Écart", "% Atteint", "Progression visuelle"]
    col_widths = [Cm(3.8), Cm(1.6), Cm(1.6), Cm(1.6), Cm(1.6), Cm(1.8), Cm(5.0)]

    t = doc.add_table(rows=1 + len(rows), cols=len(hdrs))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (h, w) in enumerate(zip(hdrs, col_widths)):
        c = t.cell(0, i)
        c.width = w
        set_cell_bg(c, "1F2937")
        cell_text(c, h, bold=True, size=8, color=BLANC,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    for ri, r in enumerate(rows):
        bg = "F9FAFB" if ri % 2 == 0 else "FFFFFF"
        ecart   = r["ecart"]
        pct_att = r["pct_atteint"]
        bar     = progress_bar_text(min(pct_att, 100), width=18)

        # Couleur écart
        ecart_color = ROUGE if ecart > 0 else VERT

        # Couleur % atteint
        if pct_att >= 100:
            att_color = VERT
        elif pct_att >= 60:
            att_color = ORANGE
        else:
            att_color = ROUGE

        row_vals = [
            (r["pays"],              NOIR,       WD_ALIGN_PARAGRAPH.LEFT),
            (f"{r['pct_quota']}%",   GRIS,       WD_ALIGN_PARAGRAPH.CENTER),
            (str(r["obj"]),          NOIR,       WD_ALIGN_PARAGRAPH.CENTER),
            (str(r["comp"]),         VERT,       WD_ALIGN_PARAGRAPH.CENTER),
            (f"{ecart:+d}",          ecart_color,WD_ALIGN_PARAGRAPH.CENTER),
            (f"{pct_att}%",          att_color,  WD_ALIGN_PARAGRAPH.CENTER),
            (bar,                    VERT,       WD_ALIGN_PARAGRAPH.LEFT),
        ]
        for ci, (val, color, align) in enumerate(row_vals):
            c = t.cell(ri + 1, ci)
            set_cell_bg(c, bg)
            cell_text(c, val, size=8 if ci == 6 else 9,
                      color=color, align=align)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Pays prioritaires (écart > 0, non atteints)
    prioritaires = [r for r in rows if r["ecart"] > 30]
    prioritaires.sort(key=lambda x: -x["ecart"])
    if prioritaires:
        para(doc, "Pays nécessitant le plus d'effort :", bold=True,
             size=9, color=ROUGE, before=2, after=2)
        for r in prioritaires:
            para(doc,
                 f"  • {r['pays']} : il manque {r['ecart']} répondants "
                 f"({r['pct_atteint']}% atteint sur objectif {r['obj']})",
                 size=9, color=NOIR, before=0, after=1)

# ─── Tableau distribution ──────────────────────────────────────────────────
def add_distribution_table(doc, question_text, counts, total_resp):
    if not counts:
        para(doc, "Aucune donnée disponible.", size=9, color=GRIS)
        return

    para(doc, f"Question : {question_text}",
         size=9, color=GRIS, italic=True, before=2, after=4)

    total_ans = sum(counts.values())
    para(doc,
         f"{total_ans} répondants ont répondu à cette question "
         f"({round(total_ans/max(total_resp,1)*100,1)}% du total collecté).",
         size=9, color=GRIS, italic=True, before=0, after=4)

    hdrs = ["Modalité", "Réponses", "% des répondants", "Répartition visuelle"]
    col_widths = [Cm(5.5), Cm(2.0), Cm(2.5), Cm(7.0)]

    t = doc.add_table(rows=1 + len(counts), cols=len(hdrs))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (h, w) in enumerate(zip(hdrs, col_widths)):
        c = t.cell(0, i)
        c.width = w
        set_cell_bg(c, "374151")
        cell_text(c, h, bold=True, size=8, color=BLANC,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    max_val = max(counts.values(), default=1)
    for ri, (modalite, cnt) in enumerate(counts.items()):
        bg  = "F9FAFB" if ri % 2 == 0 else "FFFFFF"
        pct = round(cnt / max(total_ans, 1) * 100, 1)
        bar = progress_bar_text(cnt / max_val * 100, width=22)

        row_vals = [
            (modalite,   NOIR, WD_ALIGN_PARAGRAPH.LEFT),
            (str(cnt),   BLEU, WD_ALIGN_PARAGRAPH.CENTER),
            (f"{pct}%",  NOIR, WD_ALIGN_PARAGRAPH.CENTER),
            (bar,        VERT, WD_ALIGN_PARAGRAPH.LEFT),
        ]
        for ci, (val, color, align) in enumerate(row_vals):
            c = t.cell(ri + 1, ci)
            set_cell_bg(c, bg)
            cell_text(c, val, size=8 if ci == 3 else 9,
                      color=color, align=align)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ─── Génération ───────────────────────────────────────────────────────────
def generate():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ── PAGE DE TITRE ──────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MARKETYM")
    r.bold = True; r.font.size = Pt(30); r.font.color.rgb = VERT

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("RAPPORT DE SUIVI DES BAROMÈTRES")
    r2.bold = True; r2.font.size = Pt(20); r2.font.color.rgb = NOIR

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Généré le {datetime.now().strftime('%d %B %Y')}")
    r3.font.size = Pt(11); r3.font.color.rgb = GRIS

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(
        "Ce rapport présente pour chaque baromètre l'écart entre les quotas définis\n"
        "par pays et les completions actuelles, ainsi que le détail par enquêteur."
    )
    r4.font.size = Pt(10); r4.font.color.rgb = GRIS

    doc.add_paragraph()

    # KPIs globaux
    total_obj  = sum(b["objectif_global"]  for b in BAROMETRES)
    total_comp = sum(b["completions_global"] for b in BAROMETRES)
    total_man  = total_obj - total_comp
    total_pct  = round(total_comp / max(total_obj, 1) * 100, 1)
    kpi_row(doc, [
        ("Objectif global",    str(total_obj),  "1D4ED8", BLANC),
        ("Completions totales",str(total_comp), "059669", BLANC),
        ("Manquant",           str(total_man),  "DC2626", BLANC),
        ("Progression globale",f"{total_pct}%", "D97706", BLANC),
    ])

    doc.add_page_break()

    # ── SYNTHÈSE COMPARATIVE ──────────────────────────────────────────
    section_title(doc, "SYNTHÈSE COMPARATIVE DES 3 BAROMÈTRES", level=1)

    hdrs = ["Baromètre", "Objectif", "Réalisé", "Manquant", "Progression",
            "Barre de progression"]
    col_w = [Cm(4.5), Cm(1.8), Cm(1.8), Cm(1.8), Cm(2.0), Cm(6.0)]
    t = doc.add_table(rows=1 + len(BAROMETRES), cols=len(hdrs))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (h, w) in enumerate(zip(hdrs, col_w)):
        c = t.cell(0, i)
        c.width = w
        set_cell_bg(c, "059669")
        cell_text(c, h, bold=True, size=9, color=BLANC,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, b in enumerate(BAROMETRES):
        pct  = b["pct_progression"]
        bar  = progress_bar_text(pct, width=20)
        bg   = "F0FDF4" if ri % 2 == 0 else "FFFFFF"
        vals = [
            (b["nom"],                   NOIR,   WD_ALIGN_PARAGRAPH.LEFT),
            (str(b["objectif_global"]),  BLEU,   WD_ALIGN_PARAGRAPH.CENTER),
            (str(b["completions_global"]),VERT,  WD_ALIGN_PARAGRAPH.CENTER),
            (str(b["manquant"]),         ROUGE,  WD_ALIGN_PARAGRAPH.CENTER),
            (f"{pct}%",                  ORANGE, WD_ALIGN_PARAGRAPH.CENTER),
            (bar,                        VERT,   WD_ALIGN_PARAGRAPH.LEFT),
        ]
        for ci, (val, col, align) in enumerate(vals):
            c = t.cell(ri + 1, ci)
            set_cell_bg(c, bg)
            cell_text(c, val, size=8 if ci == 5 else 9,
                      color=col, align=align)

    doc.add_paragraph()
    para(doc,
         f"Total manquant : {total_man} completions, soit {round(total_man/max(total_obj,1)*100,1)}% "
         f"de l'objectif global restant à atteindre.",
         bold=True, size=10, color=ROUGE, before=4, after=8)

    doc.add_page_break()

    # ── UN CHAPITRE PAR BAROMÈTRE ─────────────────────────────────────
    for b_idx, b in enumerate(BAROMETRES, 1):

        section_title(doc, f"{b_idx}. {b['nom'].upper()}", level=1)

        pct  = b["pct_progression"]
        bar  = progress_bar_text(pct, width=30)
        kpi_row(doc, [
            ("Objectif",      str(b["objectif_global"]),  "1D4ED8", BLANC),
            ("Réalisé",       str(b["completions_global"]),"059669", BLANC),
            ("Manquant",      str(b["manquant"]),          "DC2626", BLANC),
            ("Progression",   f"{pct}%",                   "D97706", BLANC),
        ])
        para(doc, f"{bar}  {pct}%", size=10, color=VERT,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8)

        # ── ÉCART PAR PAYS ──────────────────────────────────────────
        section_title(doc, "Écart quota vs réalisé par pays", level=2)
        pays_rows = get_pays_data(b["enquete_id"], b["objectif_global"])
        add_pays_table(doc, pays_rows, b["objectif_global"])

        # ── DETAIL PAR ENQUETEUR ──────────────────────────────────────
        section_title(doc, "Détail par enquêteur", level=2)
        enq_affs = db.table("affectations").select("id, objectif_total, completions_total, clics, enqueteurs(nom, prenom)").eq("enquete_id", b["enquete_id"]).execute()
        if enq_affs.data:
            hdrs_e = ["Enquêteur", "Objectif", "Réalisé", "Clics", "Taux conv.", "Progression"]
            col_w_e = [Cm(4.5), Cm(1.8), Cm(1.8), Cm(1.5), Cm(2.0), Cm(5.5)]
            te = doc.add_table(rows=1 + len(enq_affs.data), cols=len(hdrs_e))
            te.style = "Table Grid"
            te.alignment = WD_TABLE_ALIGNMENT.LEFT
            for i, (h, w) in enumerate(zip(hdrs_e, col_w_e)):
                c = te.cell(0, i)
                c.width = w
                set_cell_bg(c, "374151")
                cell_text(c, h, bold=True, size=8, color=BLANC, align=WD_ALIGN_PARAGRAPH.CENTER)
            for ri, aff in enumerate(enq_affs.data):
                bg = "F9FAFB" if ri % 2 == 0 else "FFFFFF"
                enqr = aff.get("enqueteurs") or {}
                nom = f"{enqr.get('prenom', '')} {enqr.get('nom', '')}".strip() or "—"
                obj = aff["objectif_total"] or 0
                comp = aff["completions_total"] or 0
                clics = aff["clics"] or 0
                taux_conv = round(comp / max(clics, 1) * 100, 1) if clics > 0 else 0
                pct_e = round(comp / max(obj, 1) * 100, 1) if obj > 0 else 0
                bar_e = progress_bar_text(min(pct_e, 100), width=18)
                vals_e = [
                    (nom, NOIR, WD_ALIGN_PARAGRAPH.LEFT),
                    (str(obj), BLEU, WD_ALIGN_PARAGRAPH.CENTER),
                    (str(comp), VERT, WD_ALIGN_PARAGRAPH.CENTER),
                    (str(clics), NOIR, WD_ALIGN_PARAGRAPH.CENTER),
                    (f"{taux_conv}%", ORANGE if taux_conv < 50 else VERT, WD_ALIGN_PARAGRAPH.CENTER),
                    (f"{bar_e} {pct_e}%", VERT, WD_ALIGN_PARAGRAPH.LEFT),
                ]
                for ci, (val, color, align) in enumerate(vals_e):
                    c = te.cell(ri + 1, ci)
                    set_cell_bg(c, bg)
                    cell_text(c, val, size=8 if ci == 5 else 9, color=color, align=align)

        if b_idx < len(BAROMETRES):
            doc.add_page_break()

    # ── PIED ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    p_f = doc.add_paragraph()
    p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_f.add_run(
        "Marketym — Document confidentiel — "
        + datetime.now().strftime("%d/%m/%Y")
    )
    r_f.font.size = Pt(8); r_f.font.color.rgb = GRIS

    out = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        f"Rapport_Barometres_HC_{datetime.now().strftime('%Y%m%d')}.docx",
    )
    doc.save(out)
    print(f"Rapport généré : {out}")

if __name__ == "__main__":
    generate()
